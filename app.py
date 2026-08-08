import os
import re
import mimetypes
from collections import defaultdict
from flask import Flask, current_app, render_template, request, redirect, url_for, session, flash, send_file, jsonify, abort
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import LargeBinary
from werkzeug.utils import secure_filename
from datetime import datetime, date, timedelta
from zoneinfo import ZoneInfo as _ZoneInfo

# Fuseau horaire de Lubumbashi — CAT (UTC+2), sans heure d'été
_CAT = _ZoneInfo('Africa/Lubumbashi')

def now_cat() -> datetime:
    """Retourne l'heure actuelle de Lubumbashi (datetime naïf en heure locale CAT)."""
    return datetime.now(_CAT).replace(tzinfo=None)
import qrcode
from io import BytesIO
import base64
import json
import zipfile
import shutil
import tempfile
import subprocess
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

# Smoke-test python-docx at startup so a broken install fails immediately
# with a clear message instead of silently blowing up inside a letter route.
try:
    _docx_smoke = Document()
    _docx_smoke.add_paragraph()          # exercises the core XML machinery
    _p = _docx_smoke.paragraphs[0]
    _run = _p.add_run("test")
    _run.font.size = Pt(12)
    _run.font.color.rgb = RGBColor(0, 0, 0)
    _tc = OxmlElement('w:tc')
    _tc.set(qn('w:val'), 'single')
    del _docx_smoke, _p, _run, _tc
except Exception as _docx_err:
    raise ImportError(
        f"python-docx est installé mais cassé — la génération de lettres échouera : {_docx_err}"
    ) from _docx_err

import pdfplumber
import hashlib
from cryptography.fernet import Fernet, InvalidToken
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_CENTER, TA_LEFT

app = Flask(__name__)
_SESSION_SECRET_DEFAULT = 'sciences-sociales-unilu-2026-change-in-production'
_session_secret = (
    os.environ.get('SESSION_SECRET') or
    os.environ.get('SECRET_KEY') or
    _SESSION_SECRET_DEFAULT
)

# Fail fast in production when SESSION_SECRET is absent or still set to the
# development placeholder.  Without a real secret, each gunicorn worker would
# use a different key and every login would be invalidated between requests.
if os.environ.get('REPLIT_DEPLOYMENT') == '1':
    if not os.environ.get('SESSION_SECRET') or _session_secret == _SESSION_SECRET_DEFAULT:
        raise RuntimeError(
            "SESSION_SECRET n'est pas défini dans les secrets de production. "
            "Ajoutez SESSION_SECRET dans Replit Publishing → Secrets → Production "
            "avant de redéployer. Sans cette clé, les connexions DÉCANAT seront "
            "invalidées entre chaque requête."
        )

app.config['SECRET_KEY'] = _session_secret

# ── SMTP-password encryption at rest ────────────────────────────────────────
# A Fernet key is exactly 32 bytes encoded as URL-safe base64.  We derive it
# deterministically from SESSION_SECRET via SHA-256 so no extra secret needs
# to be stored, and it is stable across restarts.
def _fernet() -> Fernet:
    """Return a Fernet instance keyed from SESSION_SECRET."""
    raw = hashlib.sha256(_session_secret.encode()).digest()   # 32 bytes
    key = base64.urlsafe_b64encode(raw)                       # 44-char b64
    return Fernet(key)

_SMTP_ENC_PREFIX = b'enc1:'   # stored as the raw bytes prefix in the DB value

def _encrypt_smtp_password(plaintext: str) -> str:
    """Encrypt *plaintext* and return a string safe to store in AppConfig."""
    if not plaintext:
        return plaintext
    token = _fernet().encrypt(plaintext.encode())
    return (_SMTP_ENC_PREFIX + token).decode()

def _decrypt_smtp_password(stored: str) -> str:
    """Decrypt a value from AppConfig.  Plain-text values are returned as-is
    (migration path) — callers that detect this should re-encrypt and persist."""
    if not stored:
        return stored
    b = stored.encode()
    if not b.startswith(_SMTP_ENC_PREFIX):
        return stored   # plain-text legacy row — caller handles migration
    try:
        return _fernet().decrypt(b[len(_SMTP_ENC_PREFIX):]).decode()
    except (InvalidToken, Exception):
        return ''       # corrupted token → treat as empty to force re-entry

import os as _os

# La base de données est la seule source de vérité en production et dans le
# workspace lorsqu'une base Replit gérée est disponible. Le filesystem d'une
# application publiée peut être réinitialisé lors d'une republication, même
# avec une Reserved VM : SQLite reste uniquement le repli local hors Replit.
_configured_db_url = os.environ.get('DATABASE_URL', '').strip()
if _configured_db_url:
    if not _configured_db_url.startswith(('postgresql://', 'postgres://')):
        raise RuntimeError(
            "DATABASE_URL doit pointer vers une base PostgreSQL gérée. "
            "Aucune base SQLite locale ne sera utilisée avec DATABASE_URL."
        )
    _db_url = _configured_db_url
elif os.environ.get('REPLIT_DEPLOYMENT') == '1':
    raise RuntimeError(
        "DATABASE_URL PostgreSQL est absent en production. "
        "Reconnectez la base PostgreSQL gérée avant de publier."
    )
else:
    _db_url = 'sqlite:///esciales_unilu.db'
app.config['SQLALCHEMY_DATABASE_URI'] = _db_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_pre_ping': True,
}
app.config['UPLOAD_FOLDER'] = 'static/uploads'
app.config['QRCODE_FOLDER'] = 'static/qrcodes'
app.config['HORAIRES_FOLDER'] = 'static/horaires'
app.config['ACTUALITES_FOLDER'] = 'static/actualites'
app.config['COMMUNICATION_AUDIO_FOLDER'] = 'static/communication_audio'
app.config['GENERATEUR_UPLOADS'] = 'static/generateur_lettres/uploads'
app.config['GENERATEUR_GENERES'] = 'static/generateur_lettres/generes'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg'}
ALLOWED_AUDIO_EXTENSIONS = {'webm', 'ogg', 'mp4', 'm4a', 'wav', 'mpeg', 'mp3'}
ALLOWED_HORAIRE_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg'}
_DECANAT_PASSWORD_DEFAULT = 'DECANAT2026'
_PROF_PASSWORD_DEFAULT    = 'PROF2026'
DECANAT_PASSWORD = os.environ.get('DECANAT_PASSWORD', _DECANAT_PASSWORD_DEFAULT)
PROF_PASSWORD    = os.environ.get('PROF_PASSWORD',    _PROF_PASSWORD_DEFAULT)

# Avertir en production si les mots de passe sont absents ou encore aux valeurs par défaut.
# On ne bloque plus le démarrage pour ne pas empêcher le déploiement, mais l'avertissement
# est visible dans les logs de production.
import sys as _sys
if os.environ.get('REPLIT_DEPLOYMENT') == '1':
    if not os.environ.get('DECANAT_PASSWORD') or DECANAT_PASSWORD == _DECANAT_PASSWORD_DEFAULT:
        print(
            "⚠️  AVERTISSEMENT SÉCURITÉ : DECANAT_PASSWORD absent ou valeur par défaut. "
            "Ajoutez DECANAT_PASSWORD dans Publishing → Secrets → Production.",
            file=_sys.stderr
        )
    if not os.environ.get('PROF_PASSWORD') or PROF_PASSWORD == _PROF_PASSWORD_DEFAULT:
        print(
            "⚠️  AVERTISSEMENT SÉCURITÉ : PROF_PASSWORD absent ou valeur par défaut. "
            "Ajoutez PROF_PASSWORD dans Publishing → Secrets → Production.",
            file=_sys.stderr
        )

db = SQLAlchemy(app)

# Custom Jinja2 filter: parse JSON in templates
import json as _json
@app.template_filter('fromjson')
def fromjson_filter(s):
    return _json.loads(s) if s else {}

@app.template_filter('enumerate')
def enumerate_filter(iterable, start=0):
    return list(enumerate(iterable, start=start))

DEPARTEMENTS = ['RI', 'SPA POL', 'SPA SAM', 'Sociologie', 'Anthropologie']
PROMOTIONS = ['BAC 1', 'BAC 2', 'BAC 3', 'Master 1', 'Master 2']

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_horaire_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_HORAIRE_EXTENSIONS

_FILE_ASSET_FOLDERS = {
    'uploads': 'UPLOAD_FOLDER',
    'qrcodes': 'QRCODE_FOLDER',
    'horaires': 'HORAIRES_FOLDER',
    'actualites': 'ACTUALITES_FOLDER',
    'communication_audio': 'COMMUNICATION_AUDIO_FOLDER',
    'root': None,
    'recours_preuves': None,
    'recours_cartes': None,
}

def _asset_key(category, filename):
    """Build a stable, path-safe key for a file stored in PostgreSQL."""
    if category not in _FILE_ASSET_FOLDERS:
        raise ValueError('Catégorie de fichier inconnue')
    safe_name = secure_filename(os.path.basename(filename or ''))
    if not safe_name:
        raise ValueError('Nom de fichier invalide')
    return f'{category}/{safe_name}'

def _asset_local_path(category, filename):
    """Return the legacy local path used as a migration/read fallback."""
    safe_name = secure_filename(os.path.basename(filename or ''))
    if not safe_name:
        return None
    if category == 'root':
        return os.path.join(app.static_folder, safe_name)
    if category in ('recours_preuves', 'recours_cartes'):
        return os.path.join(app.static_folder, category, safe_name)
    folder_config = _FILE_ASSET_FOLDERS.get(category)
    if not folder_config:
        return None
    return os.path.join(app.config[folder_config], safe_name)

def _store_file_asset(category, filename, data, mime_type=None):
    """Upsert a file into the persistent database-backed file store."""
    if not data:
        return
    key = _asset_key(category, filename)
    asset = FileAsset.query.filter_by(storage_key=key).first()
    if asset is None:
        asset = FileAsset(storage_key=key)
        db.session.add(asset)
    asset.data = bytes(data)
    asset.size_bytes = len(data)
    asset.mime_type = mime_type or mimetypes.guess_type(filename)[0] or 'application/octet-stream'

def _store_local_file_asset(category, filename):
    """Copy an existing local file into PostgreSQL without deleting the source."""
    path = _asset_local_path(category, filename)
    if not path or not os.path.isfile(path):
        return False
    key = _asset_key(category, filename)
    if FileAsset.query.filter_by(storage_key=key).first():
        return False
    with open(path, 'rb') as source:
        _store_file_asset(category, filename, source.read())
    return True

def _save_uploaded_asset(file_storage, category, filename):
    """Save an upload locally and persist the same bytes in PostgreSQL."""
    data = file_storage.read()
    path = _asset_local_path(category, filename)
    if not path:
        raise ValueError('Chemin de stockage local invalide')
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, 'wb') as destination:
        destination.write(data)
    _store_file_asset(category, filename, data, getattr(file_storage, 'mimetype', None))

def _delete_file_asset(category, filename):
    """Remove a file from the persistent store after an explicit user deletion."""
    if not filename:
        return
    asset = FileAsset.query.filter_by(
        storage_key=_asset_key(category, filename)
    ).first()
    if asset:
        db.session.delete(asset)

def _send_persistent_or_local(category, filename, *, as_attachment=False,
                               download_name=None):
    """Return a persistent file response, with the legacy disk as fallback."""
    if not filename:
        return None
    try:
        asset = FileAsset.query.filter_by(
            storage_key=_asset_key(category, filename)
        ).first()
    except ValueError:
        asset = None
    if asset:
        return send_file(
            BytesIO(asset.data),
            mimetype=asset.mime_type,
            as_attachment=as_attachment,
            download_name=download_name or os.path.basename(filename),
        )
    local_path = _asset_local_path(category, filename)
    if local_path and os.path.isfile(local_path):
        return send_file(
            local_path,
            as_attachment=as_attachment,
            download_name=download_name or os.path.basename(filename),
        )
    return None

def _ensure_local_asset(category, filename):
    """Materialize a persistent file locally for libraries that need a path."""
    path = _asset_local_path(category, filename)
    if not path:
        return None
    if os.path.isfile(path):
        return path
    try:
        asset = FileAsset.query.filter_by(
            storage_key=_asset_key(category, filename)
        ).first()
    except ValueError:
        asset = None
    if not asset:
        return None
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, 'wb') as destination:
        destination.write(asset.data)
    return path

def _migrate_local_assets_to_db():
    """Import local uploads into PostgreSQL, without deleting local files."""
    categories = {
        'uploads': app.config['UPLOAD_FOLDER'],
        'qrcodes': app.config['QRCODE_FOLDER'],
        'horaires': app.config['HORAIRES_FOLDER'],
        'actualites': app.config['ACTUALITES_FOLDER'],
        'recours_preuves': os.path.join(app.static_folder, 'recours_preuves'),
        'recours_cartes': os.path.join(app.static_folder, 'recours_cartes'),
    }
    # The root folder contains a few referenced page images; do not copy all
    # framework/static assets into the database.
    root_candidates = {'faculte-facade.jpg'}
    try:
        page = PageContent.query.filter_by(page_name='a_propos').first()
        if page and page.image_principale:
            root_candidates.add(os.path.basename(page.image_principale))
    except Exception:
        pass

    migrated = 0
    for category, folder in categories.items():
        if not os.path.isdir(folder):
            continue
        for root, _dirs, files in os.walk(folder):
            for filename in files:
                if filename == '.gitkeep':
                    continue
                # Nested paths are not part of the current upload contract.
                if root != folder:
                    continue
                try:
                    if _store_local_file_asset(category, filename):
                        migrated += 1
                except Exception:
                    db.session.rollback()
                    raise
    for filename in root_candidates:
        try:
            if _store_local_file_asset('root', filename):
                migrated += 1
        except Exception:
            db.session.rollback()
            raise
    if migrated:
        db.session.commit()
        import logging
        logging.getLogger(__name__).info(
            "[files] %d fichier(s) importé(s) dans PostgreSQL.", migrated
        )
    return migrated

@app.route('/fichiers/<category>/<path:filename>')
def fichier_persistant(category, filename):
    """Serve a file from PostgreSQL, falling back to the legacy local copy."""
    try:
        key = _asset_key(category, filename)
    except ValueError:
        return 'Fichier introuvable', 404
    response = _send_persistent_or_local(category, filename)
    if response:
        return response
    return 'Fichier introuvable', 404

class Etudiant(db.Model):
    __tablename__ = 'etudiants'
    id = db.Column(db.Integer, primary_key=True)
    matricule = db.Column(db.String(60), unique=True, nullable=False)
    nom = db.Column(db.String(100), nullable=False)
    postnom = db.Column(db.String(100), nullable=False)
    prenom = db.Column(db.String(100), nullable=False)
    sexe = db.Column(db.String(10), nullable=False)
    telephone = db.Column(db.String(20), nullable=False)
    promotion = db.Column(db.String(20), nullable=False)
    departement = db.Column(db.String(50), nullable=False)
    photo = db.Column(db.String(200))
    qrcode_path = db.Column(db.String(200))
    date_inscription = db.Column(db.DateTime, default=now_cat)
    presences = db.relationship('Presence', backref='etudiant', lazy=True)

class Professeur(db.Model):
    __tablename__ = 'professeurs'
    id = db.Column(db.Integer, primary_key=True)
    matricule = db.Column(db.String(60), unique=True, nullable=False)
    nom = db.Column(db.String(100), nullable=False)
    postnom = db.Column(db.String(100), nullable=False)
    prenom = db.Column(db.String(100), nullable=False)
    telephone = db.Column(db.String(20))
    departement = db.Column(db.String(50))
    qrcode_path = db.Column(db.String(200))
    presences = db.relationship('Presence', backref='professeur', lazy=True)

class Cours(db.Model):
    __tablename__ = 'cours'
    id = db.Column(db.Integer, primary_key=True)
    code = db.Column(db.String(20), unique=True, nullable=False)
    nom = db.Column(db.String(200), nullable=False)
    departement = db.Column(db.String(50), nullable=False)
    promotion = db.Column(db.String(20), nullable=False)
    professeur_id = db.Column(db.Integer, db.ForeignKey('professeurs.id'))
    presences = db.relationship('Presence', backref='cours', lazy=True)

class Presence(db.Model):
    __tablename__ = 'presences'
    id = db.Column(db.Integer, primary_key=True)
    etudiant_id = db.Column(db.Integer, db.ForeignKey('etudiants.id'))
    professeur_id = db.Column(db.Integer, db.ForeignKey('professeurs.id'))
    cours_id = db.Column(db.Integer, db.ForeignKey('cours.id'))
    heure_entree = db.Column(db.DateTime, default=now_cat)
    date = db.Column(db.Date, default=lambda: now_cat().date())
    type_presence = db.Column(db.String(20))

class Horaire(db.Model):
    __tablename__ = 'horaires'
    id = db.Column(db.Integer, primary_key=True)
    departement = db.Column(db.String(50), nullable=False)
    promotion = db.Column(db.String(20), nullable=False)
    type_horaire = db.Column(db.String(20), nullable=False)
    fichier = db.Column(db.String(200), nullable=False)
    date_publication = db.Column(db.DateTime, default=now_cat)

class Actualite(db.Model):
    __tablename__ = 'actualites'
    id = db.Column(db.Integer, primary_key=True)
    titre = db.Column(db.String(200), nullable=False)
    description = db.Column(db.Text, nullable=False)
    image = db.Column(db.String(200))
    date_publication = db.Column(db.DateTime, default=now_cat)
    publie = db.Column(db.Boolean, default=True)
    type_publication = db.Column(db.String(20), default='actualite', nullable=False)
    epingle = db.Column(db.Boolean, default=False, nullable=False)


class ChatThread(db.Model):
    """Espace privé de communication d'un étudiant avec le DÉCANAT."""
    __tablename__ = 'chat_threads'
    id = db.Column(db.Integer, primary_key=True)
    matricule = db.Column(db.String(60), unique=True, nullable=False, index=True)
    nom_complet = db.Column(db.String(220), nullable=False)
    promotion = db.Column(db.String(100), nullable=False)
    actif = db.Column(db.Boolean, default=True, nullable=False)
    date_creation = db.Column(db.DateTime, default=now_cat, nullable=False)
    date_dernier_message = db.Column(db.DateTime, default=now_cat, nullable=False, index=True)
    messages = db.relationship(
        'ChatMessage', backref='thread', lazy=True,
        cascade='all, delete-orphan', order_by='ChatMessage.date_creation',
    )
    appels = db.relationship(
        'ChatCall', backref='thread', lazy=True,
        cascade='all, delete-orphan', order_by='ChatCall.date_creation',
    )


class ChatMessage(db.Model):
    __tablename__ = 'chat_messages'
    id = db.Column(db.Integer, primary_key=True)
    thread_id = db.Column(db.Integer, db.ForeignKey('chat_threads.id'), nullable=False, index=True)
    sender_role = db.Column(db.String(20), nullable=False)  # etudiant | decanat
    sender_name = db.Column(db.String(220), nullable=False)
    contenu = db.Column(db.Text, default='')
    audio_filename = db.Column(db.String(240))
    type_message = db.Column(db.String(30), default='texte', nullable=False)
    lu = db.Column(db.Boolean, default=False, nullable=False)
    date_creation = db.Column(db.DateTime, default=now_cat, nullable=False, index=True)


class ChatCall(db.Model):
    """Signalisation WebRTC persistée pour un appel audio étudiant-DÉCANAT."""
    __tablename__ = 'chat_calls'
    id = db.Column(db.Integer, primary_key=True)
    thread_id = db.Column(db.Integer, db.ForeignKey('chat_threads.id'), nullable=False, index=True)
    caller_role = db.Column(db.String(20), nullable=False)
    status = db.Column(db.String(20), default='pending', nullable=False)
    offer_json = db.Column(db.Text)
    answer_json = db.Column(db.Text)
    date_creation = db.Column(db.DateTime, default=now_cat, nullable=False)
    date_reponse = db.Column(db.DateTime)
    date_fin = db.Column(db.DateTime)


class PageContent(db.Model):
    __tablename__ = 'page_contents'
    id = db.Column(db.Integer, primary_key=True)
    page_name = db.Column(db.String(50), unique=True, nullable=False)
    content_json = db.Column(db.Text, nullable=False)
    image_principale = db.Column(db.String(200))
    date_modification = db.Column(db.DateTime, default=now_cat, onupdate=now_cat)

class FileAsset(db.Model):
    """Fichier téléversé conservé dans la base PostgreSQL persistante."""
    __tablename__ = 'file_assets'
    id = db.Column(db.Integer, primary_key=True)
    storage_key = db.Column(db.String(300), unique=True, nullable=False, index=True)
    mime_type = db.Column(db.String(120), nullable=False, default='application/octet-stream')
    data = db.Column(LargeBinary, nullable=False)
    size_bytes = db.Column(db.Integer, nullable=False, default=0)
    date_modification = db.Column(db.DateTime, default=now_cat, onupdate=now_cat)

class BulletinSession(db.Model):
    __tablename__ = 'bulletin_sessions'
    id = db.Column(db.Integer, primary_key=True)
    nom = db.Column(db.String(200))
    annee = db.Column(db.String(20))
    session_acad = db.Column(db.String(50))
    semestre = db.Column(db.String(50))
    promotion = db.Column(db.String(100))
    montant_fc   = db.Column(db.Integer, default=5000)
    departement  = db.Column(db.String(150), default='')  # département affiché sur le bulletin
    texte_intro  = db.Column(db.Text, default='')          # phrase d'introduction personnalisable
    type_grille  = db.Column(db.String(20), default='initial', nullable=False)
    # Les sessions historiques sont des bulletins initiaux. Une session de
    # recours est toujours séparée afin de ne jamais réinitialiser un paiement
    # déjà enregistré sur le bulletin initial.
    date_import = db.Column(db.DateTime, default=now_cat)
    bulletins = db.relationship('BulletinData', backref='bul_session', lazy=True,
                                cascade='all, delete-orphan')

class BulletinData(db.Model):
    __tablename__ = 'bulletin_data'
    id = db.Column(db.Integer, primary_key=True)
    session_id = db.Column(db.Integer, db.ForeignKey('bulletin_sessions.id'))
    matricule = db.Column(db.String(60), index=True)
    nom = db.Column(db.String(200))
    sexe = db.Column(db.String(10))
    data_json = db.Column(db.Text)             # JSON complet de l'étudiant
    numero_bulletin = db.Column(db.String(40), unique=True)
    paye = db.Column(db.Boolean, default=False)
    date_paiement = db.Column(db.DateTime)
    methode_paiement = db.Column(db.String(50))
    telephone_paiement = db.Column(db.String(20))
    reference_paiement = db.Column(db.String(50))
    montant_paye = db.Column(db.Integer, default=0)
    nb_telechargements = db.Column(db.Integer, default=0)
    date_dernier_telechargement = db.Column(db.DateTime)


class BulletinImportDraft(db.Model):
    """Import de grille en attente de validation des matricules."""
    __tablename__ = 'bulletin_import_drafts'
    id = db.Column(db.Integer, primary_key=True)
    token = db.Column(db.String(120), unique=True, nullable=False, index=True)
    promotion = db.Column(db.String(150), nullable=False)
    payload_json = db.Column(db.Text, nullable=False)
    date_creation = db.Column(db.DateTime, default=now_cat, nullable=False)


class PaiementAudit(db.Model):
    """Historique des corrections de paiement effectuées par le DÉCANAT."""
    __tablename__ = 'paiement_audits'
    id               = db.Column(db.Integer, primary_key=True)
    bulletin_id      = db.Column(db.Integer, db.ForeignKey('bulletin_data.id'), nullable=False, index=True)
    old_montant      = db.Column(db.Integer)
    new_montant      = db.Column(db.Integer)
    old_methode      = db.Column(db.String(50))
    new_methode      = db.Column(db.String(50))
    old_reference    = db.Column(db.String(50))
    new_reference    = db.Column(db.String(50))
    old_date_paiement= db.Column(db.DateTime)
    new_date_paiement= db.Column(db.DateTime)
    date_modification= db.Column(db.DateTime, default=now_cat, nullable=False)
    operator_name    = db.Column(db.String(100))   # nom de l'agent qui a effectué la correction
    # relationship for convenience
    bulletin         = db.relationship('BulletinData', backref=db.backref('audits', lazy=True,
                                       order_by='PaiementAudit.date_modification.desc()'))


class ListeIdentifiants(db.Model):
    """Liste officielle des matricules et mots de passe par promotion."""
    __tablename__ = 'liste_identifiants'
    id          = db.Column(db.Integer, primary_key=True)
    promotion   = db.Column(db.String(150), nullable=False, index=True)
    nom         = db.Column(db.String(200), nullable=False)
    nom_norm    = db.Column(db.String(200), nullable=False, index=True)  # nom normalisé pour matching
    matricule   = db.Column(db.String(60), nullable=False, index=True)
    mot_de_passe= db.Column(db.String(40))
    date_import = db.Column(db.DateTime, default=now_cat)


class RecuPaiement(db.Model):
    __tablename__ = 'recus_paiement'
    id = db.Column(db.Integer, primary_key=True)
    numero = db.Column(db.String(40), unique=True, nullable=False, index=True)
    code_qr = db.Column(db.String(80), unique=True, nullable=False, index=True)
    lot = db.Column(db.String(20))
    dept = db.Column(db.String(20))
    annee = db.Column(db.String(10))           # ex: "26"
    semestre = db.Column(db.String(10))        # ex: "S1", "S2"
    montant = db.Column(db.String(30), default='5000 CDF')
    montant_lettres = db.Column(db.String(200), default='Cinq mille Francs congolais')
    motif = db.Column(db.String(300), default='Bulletin des résultats')
    type_recu = db.Column(db.String(40), default='bulletin', nullable=False)
    annee_complete = db.Column(db.String(20), default='2025-2026')
    date_creation = db.Column(db.DateTime, default=now_cat)
    # Utilisation
    utilise = db.Column(db.Boolean, default=False)
    date_utilisation = db.Column(db.DateTime)
    matricule_etudiant = db.Column(db.String(60))
    nom_etudiant = db.Column(db.String(200))
    bulletin_id = db.Column(db.Integer, db.ForeignKey('bulletin_data.id'), nullable=True)
    # Présentation sur un bulletin déjà payé (reçu non consommé)
    tentative_bulletin_id = db.Column(db.Integer, db.ForeignKey('bulletin_data.id'), nullable=True)
    tentative_matricule = db.Column(db.String(60))
    tentative_nom = db.Column(db.String(200))
    date_tentative = db.Column(db.DateTime)
    tentative_revue = db.Column(db.Boolean, default=False, nullable=False, server_default='0')


class ReleveCommande(db.Model):
    """Commande étudiante d'un relevé de cotes après validation d'un reçu."""
    __tablename__ = 'releve_commandes'
    id = db.Column(db.Integer, primary_key=True)
    bulletin_id = db.Column(db.Integer, db.ForeignKey('bulletin_data.id'), nullable=False, index=True)
    recu_id = db.Column(db.Integer, db.ForeignKey('recus_paiement.id'), nullable=False, unique=True)
    matricule = db.Column(db.String(60), nullable=False, index=True)
    nom_etudiant = db.Column(db.String(200), nullable=False)
    statut = db.Column(db.String(30), nullable=False, default='soumise')
    date_commande = db.Column(db.DateTime, default=now_cat, nullable=False, index=True)
    date_traitement = db.Column(db.DateTime)
    note_decanat = db.Column(db.Text)
    bulletin = db.relationship('BulletinData', foreign_keys=[bulletin_id])
    recu = db.relationship('RecuPaiement', foreign_keys=[recu_id])


# Les valeurs historiques restent inchangées. Les deux nouvelles étapes sont
# ajoutées à côté d'elles afin de ne jamais réinterpréter les anciennes lignes.
_GRID_TYPES = ('initial', 'recours', 'session_2', 'recours_session_2')
_RECEIPT_TYPES = (
    'bulletin', 'recours', 'resultat_recours',
    'session_2', 'recours_session_2_soumission', 'recours_session_2',
    'releve',
)

_GRID_TYPE_LABELS = {
    'initial': 'Délibération 1ère session',
    'recours': 'Recours 1ère session',
    'session_2': 'Délibération 2ème session',
    'recours_session_2': 'Recours 2ème session',
}

_RECEIPT_TYPE_LABELS = {
    'bulletin': 'Résultats session 1',
    'recours': 'Soumission recours session 1',
    'resultat_recours': 'Résultats du recours session 1',
    'session_2': 'Résultats session 2',
    'recours_session_2_soumission': 'Soumission recours session 2',
    'recours_session_2': 'Résultats du recours session 2',
    'releve': 'Commande de relevé de cotes',
}

_RECEIPT_TO_GRID_TYPE = {
    'bulletin': 'initial',
    'resultat_recours': 'recours',
    'session_2': 'session_2',
    'recours_session_2': 'recours_session_2',
}

_RECEIPT_PREFIX_TO_TYPE = {
    'REL': 'releve',
    'RS2': 'recours_session_2_soumission',
    'RS': 'recours',
    'RR': 'resultat_recours',
    'S2': 'session_2',
    'R2': 'recours_session_2',
}


def _receipt_type_from_number(numero):
    """Retourne le type canonique déduit du préfixe imprimé du reçu.

    Le numéro imprimé est la source de vérité pour éviter qu'un reçu RR- ou
    R2- mal classé en base ne puisse débloquer une autre étape.
    """
    prefix = str(numero or '').strip().upper().split('-', 1)[0]
    if prefix in _RECEIPT_PREFIX_TO_TYPE:
        return _RECEIPT_PREFIX_TO_TYPE[prefix]
    if prefix.startswith('B') and prefix[1:].isdigit():
        return 'bulletin'
    return None


def _receipt_type_is_consistent(recu):
    """Vérifie que le type DB correspond au préfixe du reçu imprimé."""
    canonical = _receipt_type_from_number(getattr(recu, 'numero', ''))
    if canonical is None:
        return False
    return getattr(recu, 'type_recu', 'bulletin') == canonical


def _grid_type_label(value):
    return _GRID_TYPE_LABELS.get(value or 'initial', 'Délibération 1ère session')


def _receipt_type_label(value):
    return _RECEIPT_TYPE_LABELS.get(value or 'bulletin', 'Bulletin initial')


def _receipt_prefix(type_recu, lot='1'):
    if type_recu == 'releve':
        return 'REL'
    if type_recu == 'session_2':
        # B2- is already used by historical ordinary lot 2 receipts.
        return 'S2'
    if type_recu == 'recours_session_2':
        return 'R2'
    if type_recu == 'recours_session_2_soumission':
        return 'RS2'
    if type_recu == 'recours':
        return 'RS'
    if type_recu == 'resultat_recours':
        return 'RR'
    return f'B{lot}'


_RECEIPT_PAYMENT_DETAILS = {
    'bulletin': ('5000 FC', 'Cinq mille Francs congolais'),
    'recours': ('10 USD', 'Dix Dollars américains'),
    'resultat_recours': ('5000 FC', 'Cinq mille Francs congolais'),
    'session_2': ('5000 FC', 'Cinq mille Francs congolais'),
    'recours_session_2_soumission': ('10 USD', 'Dix Dollars américains'),
    'recours_session_2': ('5000 FC', 'Cinq mille Francs congolais'),
}


def _receipt_payment_details(type_recu):
    """Retourne le montant canonique imposé par la catégorie du reçu."""
    return _RECEIPT_PAYMENT_DETAILS.get(
        type_recu, _RECEIPT_PAYMENT_DETAILS['bulletin']
    )


def _receipt_pdf_title(type_recu):
    """Titre officiel imprimé sur le reçu selon sa catégorie."""
    return {
        'bulletin': 'PREUVE DE PAIEMENT DES RÉSULTATS SESSION 1',
        'recours': 'PREUVE DE PAIEMENT DU RECOURS SESSION 1',
        'resultat_recours': 'PREUVE DE PAIEMENT DES RÉSULTATS DU RECOURS SESSION 1',
        'session_2': 'PREUVE DE PAIEMENT DES RÉSULTATS SESSION 2',
        'recours_session_2_soumission':
            'PREUVE DE PAIEMENT DE LA SOUMISSION DU RECOURS SESSION 2',
        'recours_session_2':
            'PREUVE DE PAIEMENT DES RÉSULTATS DU RECOURS SESSION 2',
        'releve': 'PREUVE DE PAIEMENT DE LA COMMANDE DU RELEVÉ DE COTES',
    }.get(type_recu, 'PREUVE DE PAIEMENT DES RÉSULTATS SESSION 1')


class AppConfig(db.Model):
    """Paramètres de configuration de l'application (clé / valeur)."""
    __tablename__ = 'app_config'
    id    = db.Column(db.Integer, primary_key=True)
    key   = db.Column(db.String(80), unique=True, nullable=False, index=True)
    value = db.Column(db.Text, nullable=False, default='')


class SiteVisit(db.Model):
    """Visite publique de l'accueil, sans adresse IP ni donnée personnelle."""
    __tablename__ = 'site_visits'
    id          = db.Column(db.Integer, primary_key=True)
    date_visite = db.Column(db.DateTime, default=now_cat, nullable=False, index=True)


class ScanLog(db.Model):
    """Journal de toutes les tentatives de scan de reçus."""
    __tablename__ = 'scan_logs'
    id              = db.Column(db.Integer, primary_key=True)
    code            = db.Column(db.String(100), nullable=False, index=True)
    ip              = db.Column(db.String(50))
    date_scan       = db.Column(db.DateTime, default=now_cat, index=True)
    # 'ok' | 'invalide' | 'deja_utilise' | 'deja_paye'
    resultat        = db.Column(db.String(20), nullable=False, index=True)
    matricule       = db.Column(db.String(60))   # matricule saisi par l'étudiant (si POST)
    nom_etudiant    = db.Column(db.String(200))
    dismissed       = db.Column(db.Boolean, default=False, nullable=False, server_default='0')


class AdministrationAudit(db.Model):
    """Historique des opérations administratives sensibles."""
    __tablename__ = 'administration_audits'
    id            = db.Column(db.Integer, primary_key=True)
    action        = db.Column(db.String(80), nullable=False, index=True)
    operator_name = db.Column(db.String(100))
    details       = db.Column(db.Text)
    date_action   = db.Column(db.DateTime, default=now_cat, nullable=False, index=True)


class Recours(db.Model):
    """Formulaire de recours soumis par un étudiant."""
    __tablename__ = 'recours'
    id                = db.Column(db.Integer, primary_key=True)
    # Identité
    carte_etudiant_filename = db.Column(db.String(200))   # photo carte uploadée en ligne
    preuves_json      = db.Column(db.Text, default='{}')  # {claim_id: filename} preuves jointes
    nom               = db.Column(db.String(100), nullable=False)
    postnom           = db.Column(db.String(100), nullable=False)
    prenom            = db.Column(db.String(100), nullable=False)
    sexe              = db.Column(db.String(10))
    telephone         = db.Column(db.String(30))
    promotion         = db.Column(db.String(100), nullable=False, index=True)
    filiere           = db.Column(db.String(100), nullable=False)
    # Réclamations (JSON : liste de dicts {id, texte, detail})
    reclamations_json = db.Column(db.Text, nullable=False, default='[]')
    # Reçu de paiement du recours
    recu_id           = db.Column(db.Integer, db.ForeignKey('recus_paiement.id'), nullable=True)
    recu_numero       = db.Column(db.String(50))
    type_recu         = db.Column(db.String(40), nullable=False, default='recours')
    # Métadonnées
    date_soumission   = db.Column(db.DateTime, default=now_cat, index=True)
    statut            = db.Column(db.String(30), default='soumis')
    ip_soumission     = db.Column(db.String(45))
    # Relations
    recu = db.relationship('RecuPaiement', foreign_keys=[recu_id])


def _enregistrer_admin_audit(action, details=None):
    """Ajoute une trace d'une opération sensible à la transaction courante."""
    db.session.add(AdministrationAudit(
        action=action,
        operator_name=session.get('decanat_operator', 'Inconnu'),
        details=json.dumps(details or {}, ensure_ascii=False),
    ))


def _normaliser_recherche_recu(value):
    """Retourne une référence ou un code QR exploitable depuis une saisie libre."""
    value = str(value or '').strip()
    if not value:
        return ''
    # Les QR imprimés historiques peuvent contenir l'URL /scan/CODE.
    match = re.search(r'/scan/(?:scan/)?([A-Za-z0-9]+)', value, re.IGNORECASE)
    if match:
        return match.group(1).upper()
    try:
        payload = json.loads(value)
        if isinstance(payload, dict):
            value = payload.get('code_qr') or payload.get('code') or value
    except (TypeError, ValueError):
        pass
    return str(value).strip().upper()


def _chercher_recu(value):
    """Recherche un reçu par numéro imprimé, code QR ou URL de QR."""
    lookup = _normaliser_recherche_recu(value)
    if not lookup:
        return None, ''
    recu = RecuPaiement.query.filter(
        db.func.upper(RecuPaiement.numero) == lookup
    ).first()
    if recu:
        return recu, 'référence'
    recu = RecuPaiement.query.filter(
        db.func.upper(RecuPaiement.code_qr) == lookup
    ).first()
    return recu, 'QR' if recu else ''


def _description_utilisation_recu(recu):
    """Construit les informations métier affichées lors d'une vérification."""
    service = _receipt_type_label(getattr(recu, 'type_recu', 'bulletin'))
    session_concernee = getattr(recu, 'motif', None) or service
    bulletin = None
    if recu.bulletin_id:
        bulletin = db.session.get(BulletinData, recu.bulletin_id)
        if bulletin and bulletin.bul_session:
            session_concernee = _grid_type_label(
                getattr(bulletin.bul_session, 'type_grille', 'initial')
            )
            service = 'Consultation du bulletin'
    recours = Recours.query.filter_by(recu_id=recu.id).first()
    if recours:
        session_concernee = recours.type_recu or session_concernee
        service = 'Soumission d’un recours'
    tentatives = ScanLog.query.filter(
        db.func.upper(ScanLog.code) == str(recu.code_qr or '').upper()
    ).order_by(ScanLog.date_scan.desc()).limit(100).all()
    return {
        'service': service,
        'session': session_concernee,
        'bulletin': bulletin,
        'tentatives': tentatives,
        'nb_tentatives': len(tentatives),
    }


def _get_modele_docx() -> bytes | None:
    """Retourne le contenu binaire du modèle Word stocké en base ou None."""
    try:
        row = AppConfig.query.filter_by(key='modele_docx_b64').first()
        if row and row.value:
            return base64.b64decode(row.value)
    except Exception:
        pass
    return None


def _get_app_config_bool(key, default=True):
    """Lit un booléen de configuration stocké dans AppConfig."""
    try:
        row = AppConfig.query.filter_by(key=key).first()
        if row is None:
            return default
        return str(row.value).strip().lower() in ('1', 'true', 'yes', 'on')
    except Exception:
        return default


_RELEVE_ENTETE_DEFAULTS = {
    'initial': 'RELEVÉ DE COTES — PREMIÈRE SESSION',
    'recours': 'RELEVÉ DE COTES — RECOURS PREMIÈRE SESSION',
    'session_2': 'RELEVÉ DE COTES — DEUXIÈME SESSION',
    'recours_session_2': 'RELEVÉ DE COTES — RECOURS DEUXIÈME SESSION',
}

_RELEVE_MODELE_OFFICIEL = (
    'attached_assets/Relevé_de_cotes_BAC1_SAM_2024_1785839608716.docx'
)
_RELEVE_POINTILLES = '...............'
_RELEVE_PIED_OFFICIEL = [
    'Fait à Lubumbashi, le……/………/2024',
    'Le Secrétaire académique facultaire                                                                                                      Le Doyen de la Faculté',
    '    ',
    'NGUZA KARL IBOND Jean-Pierre                                                                                      MULUMBENI MUNYENGA Georges',
    '            Chef de Travaux                                                                                                                  Professeur',
]


def _releve_entetes():
    """Retourne les en-têtes configurables des quatre étapes de délibération."""
    result = {}
    for grid_type, default in _RELEVE_ENTETE_DEFAULTS.items():
        try:
            row = AppConfig.query.filter_by(key=f'releve_entete_{grid_type}').first()
            result[grid_type] = row.value.strip() if row and row.value.strip() else default
        except Exception:
            result[grid_type] = default
    return result


def _releve_type_label(grid_type):
    return {
        'initial': 'Première session',
        'recours': 'Recours première session',
        'session_2': 'Deuxième session',
        'recours_session_2': 'Recours deuxième session',
    }.get(grid_type or 'initial', 'Première session')


def _releve_modele_path():
    """Chemin du modèle officiel fourni par la faculté."""
    return os.path.join(current_app.root_path, _RELEVE_MODELE_OFFICIEL)


def _releve_num(value, decimals=2):
    """Nombre lisible dans les cellules du modèle officiel."""
    if value in (None, ''):
        return ''
    try:
        number = float(value)
        if decimals == 0:
            return str(int(round(number)))
        return f'{number:.{decimals}f}'.rstrip('0').rstrip('.').replace('.', ',')
    except (TypeError, ValueError):
        return str(value)


def _releve_valeur(value, fallback=_RELEVE_POINTILLES):
    return str(value).strip() if value not in (None, '') and str(value).strip() else fallback


def _releve_set_paragraph_text(paragraph, value):
    """Remplace le texte sans toucher au formatage ni aux cadres de la zone."""
    text_nodes = paragraph.xpath('.//w:t', namespaces=_RELEVE_XML_NS)
    if not text_nodes:
        return
    text_nodes[0].text = str(value if value is not None else '')
    for node in text_nodes[1:]:
        node.text = ''


def _releve_box_paragraphs(box):
    return box.xpath('./w:p', namespaces=_RELEVE_XML_NS)


def _releve_replace_textbox(box, values, start=0):
    paragraphs = _releve_box_paragraphs(box)
    for index, paragraph in enumerate(paragraphs):
        if index < start:
            continue
        value_index = index - start
        _releve_set_paragraph_text(
            paragraph, values[value_index] if value_index < len(values) else ''
        )


def _releve_official_values(etudiant, bs):
    """Prépare uniquement les valeurs prévues par le modèle Word officiel."""
    courses = etudiant.get('cours') or []
    # Le modèle officiel BAC1 contient exactement 18 lignes. Ne pas créer de ligne
    # supplémentaire ni modifier sa géométrie.
    courses = courses[:18]
    notes = [_releve_valeur(_releve_num(course.get('note'))) for course in courses]
    names = [_releve_valeur(course.get('name')) for course in courses]
    credits = [
        _releve_valeur(_releve_num(course.get('credit'), decimals=0))
        for course in courses
    ]
    weighted = []
    for course in courses:
        if not course.get('has_pond_col'):
            weighted.append(_RELEVE_POINTILLES)
            continue
        numerator = _releve_num(course.get('pondere'))
        denominator = _releve_num(course.get('max_pondere'), decimals=0)
        weighted.append(_releve_valeur(
            f'{numerator}/{denominator}' if denominator else numerator
        ))

    total = etudiant.get('total')
    pct = etudiant.get('pct')
    denominator = ''
    try:
        if total is not None and pct:
            denominator = _releve_num(float(total) * 100 / float(pct), decimals=0)
    except (TypeError, ValueError, ZeroDivisionError):
        denominator = ''
    total_value = _releve_valeur(
        f'{_releve_num(total, decimals=0)}/{denominator}'
        if denominator else _releve_num(total, decimals=0)
    )

    decision = _releve_valeur(etudiant.get('decision'))
    appreciation = _releve_valeur(etudiant.get('appre'))
    name = _releve_valeur(etudiant.get('nom'))
    lieu_naissance = _releve_valeur(
        etudiant.get('lieu_naissance') or etudiant.get('lieu')
    )
    date_naissance = _releve_valeur(etudiant.get('date_naissance'))
    credits_valides = _releve_num(etudiant.get('cr_val'), decimals=0)
    credits_non_valides = _releve_num(etudiant.get('cr_nval'), decimals=0)
    moyenne = _releve_num(etudiant.get('moyenne'))
    intro = (
        f'Monsieur, Madame, Mademoiselle {name}, né(e) à {lieu_naissance}, '
        f'le {date_naissance}, a obtenue à l’issue des évaluations du premier '
        f'et second semestres (Epreuves de rattrapage), de l’année académique '
        f'{bs.annee}, le résultat ci-dessous sur l’ensemble des UE (et ECUE) '
        f'prévues au programme de {bs.promotion} '
    )
    return {
        'intro': intro,
        'notes': notes,
        'names': names,
        'credits': credits,
        'weighted': weighted,
        'total': total_value,
        'pct': _releve_valeur(
            f'{_releve_num(pct, decimals=1)}%'
            if pct not in (None, '') else ''
        ),
        'moyenne': _releve_valeur(f'{moyenne}/20' if moyenne else ''),
        'credits_valides': _releve_valeur(credits_valides),
        'credits_non_valides': _releve_valeur(credits_non_valides),
        'decision': decision,
        'appreciation': appreciation,
        'year': str(bs.annee or ''),
    }


def _releve_pied_officiel(aperçu=False):
    """Retourne le pied officiel, y compris dans l'aperçu étudiant."""
    return list(_RELEVE_PIED_OFFICIEL)


def _releve_set_institution_line(paragraph, value):
    """Place l'Université et la Faculté sur une seule ligne du modèle."""
    _releve_set_paragraph_text(paragraph, value)
    for run in paragraph.xpath('./w:r', namespaces=_RELEVE_XML_NS):
        run_props = run.find('w:rPr', namespaces=_RELEVE_XML_NS)
        if run_props is None:
            run_props = OxmlElement('w:rPr')
            run.insert(0, run_props)
        for tag in ('w:sz', 'w:szCs'):
            size = run_props.find(tag, namespaces=_RELEVE_XML_NS)
            if size is None:
                size = OxmlElement(tag)
                run_props.append(size)
            size.set(qn('w:val'), '24')


def _generer_releve_cotes_modele(etudiant, bs, aperçu=False):
    """Remplit le modèle Word officiel sans en recréer la mise en page."""
    from lxml import etree

    template_path = _releve_modele_path()
    if not os.path.isfile(template_path):
        raise FileNotFoundError(f'Modèle officiel introuvable : {template_path}')

    values = _releve_official_values(etudiant, bs)
    with zipfile.ZipFile(template_path, 'r') as source:
        output = BytesIO()
        with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as target:
            for item in source.infolist():
                data = source.read(item.filename)
                if item.filename == 'word/document.xml':
                    root = etree.fromstring(data)
                    institution_line = (
                        'UNIVERSITE DE LUBUMBASHI — '
                        'FACULTE DES SCIENCES SOCIALES, POLITIQUES ET ADMINISTRATIVES.'
                    )
                    institution_seen = False
                    for paragraph in root.xpath('.//w:p', namespaces=_RELEVE_XML_NS):
                        paragraph_text = ''.join(
                            paragraph.xpath('.//w:t/text()', namespaces=_RELEVE_XML_NS)
                        ).strip()
                        if paragraph_text == 'UNIVERSITE DE LUBUMBASHI':
                            _releve_set_institution_line(paragraph, institution_line)
                            institution_seen = True
                        elif (
                            institution_seen
                            and paragraph_text
                            == 'FACULTE DES SCIENCES SOCIALES, POLITIQUES ET ADMINISTRATIVES.'
                        ):
                            _releve_set_paragraph_text(paragraph, '')
                            institution_seen = False
                    boxes = root.xpath('.//*[local-name()="txbxContent"]')
                    for box in boxes:
                        text = ''.join(box.xpath('.//w:t/text()', namespaces=_RELEVE_XML_NS))
                        paragraphs = _releve_box_paragraphs(box)
                        if 'a obtenue' in text or 'AHADI MUGOBE' in text:
                            _releve_replace_textbox(box, [values['intro']])
                        elif 'Introduction à la science administrative' in text:
                            _releve_replace_textbox(box, values['names'])
                        elif (
                            len(paragraphs) == 36
                            and not ''.join(
                                paragraphs[0].xpath(
                                    './/w:t/text()', namespaces=_RELEVE_XML_NS
                                )
                            )
                            and not ''.join(
                                paragraphs[1].xpath(
                                    './/w:t/text()', namespaces=_RELEVE_XML_NS
                                )
                            )
                        ):
                            _releve_replace_textbox(box, values['notes'], start=2)
                        elif (
                            len(paragraphs) == 37
                            and not ''.join(
                                paragraphs[0].xpath(
                                    './/w:t/text()', namespaces=_RELEVE_XML_NS
                                )
                            )
                            and '/' not in ''.join(
                                paragraphs[1].xpath(
                                    './/w:t/text()', namespaces=_RELEVE_XML_NS
                                )
                            )
                        ):
                            _releve_replace_textbox(box, values['credits'], start=1)
                        elif (
                            len(paragraphs) == 37
                            and '/' in ''.join(
                                paragraphs[1].xpath(
                                    './/w:t/text()', namespaces=_RELEVE_XML_NS
                                )
                            )
                        ):
                            _releve_replace_textbox(box, values['weighted'], start=1)
                        elif len(paragraphs) == 16 and '/' in text:
                            _releve_replace_textbox(box, [
                                values['total'],
                                values['pct'],
                                values['moyenne'],
                                values['credits_valides'],
                                values['credits_non_valides'],
                                values['appreciation'],
                                '',
                                values['decision'],
                            ])
                        elif text.startswith('Fait à Lubumbashi'):
                            _releve_replace_textbox(
                                box, _releve_pied_officiel(aperçu=aperçu)
                            )
                    # Le titre et les autres éléments statiques du modèle sont
                    # laissés intacts ; seuls les blocs de données et le pied
                    # officiel demandé sont remplacés.
                    data = etree.tostring(
                        root, xml_declaration=True, encoding='UTF-8', standalone='yes'
                    )
                target.writestr(item, data)
        output.seek(0)
        return output


def _rendre_releve_apercu_png(etudiant, bs):
    """Rend le modèle officiel en image sans exposer le DOCX à l'étudiant."""
    docx_buf = _generer_releve_cotes_modele(etudiant, bs, aperçu=True)
    with tempfile.TemporaryDirectory(prefix='releve-preview-') as workdir:
        docx_path = os.path.join(workdir, 'releve_officiel_apercu.docx')
        pdf_path = os.path.join(workdir, 'releve_officiel_apercu.pdf')
        with open(docx_path, 'wb') as handle:
            handle.write(docx_buf.read())

        env = os.environ.copy()
        env['HOME'] = workdir
        env['UserInstallation'] = f'file://{workdir}/lo-profile'
        subprocess.run(
            [
                'libreoffice', '--headless', '--convert-to', 'pdf',
                '--outdir', workdir, docx_path,
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=45,
            env=env,
        )
        if not os.path.isfile(pdf_path):
            raise RuntimeError('Le rendu PDF du relevé officiel est introuvable.')

        import fitz
        document = fitz.open(pdf_path)
        try:
            if not document.page_count:
                raise RuntimeError('Le relevé officiel ne contient aucune page.')
            page = document.load_page(0)
            pixmap = page.get_pixmap(
                matrix=fitz.Matrix(1.7, 1.7),
                alpha=False,
            )
            return BytesIO(pixmap.tobytes('png'))
        finally:
            document.close()


_RELEVE_XML_NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
}


def _enregistrer_visite_site():
    """Enregistre une ouverture de l'accueil et retourne le total courant."""
    try:
        db.session.add(SiteVisit())
        db.session.commit()
        return SiteVisit.query.count()
    except Exception:
        db.session.rollback()
        return None


def _save_modele_docx(data: bytes) -> None:
    """Sauvegarde le modèle Word en base de données (encodé en base64)."""
    encoded = base64.b64encode(data).decode('ascii')
    row = AppConfig.query.filter_by(key='modele_docx_b64').first()
    if row:
        row.value = encoded
    else:
        db.session.add(AppConfig(key='modele_docx_b64', value=encoded))
    db.session.commit()


def _get_alert_email():
    """Retourne l'adresse e-mail d'alerte : DB en priorité, puis DECANAT_EMAIL env var."""
    try:
        cfg = AppConfig.query.filter_by(key='alert_email').first()
        if cfg and cfg.value.strip():
            return cfg.value.strip()
    except Exception:
        pass
    return os.environ.get('DECANAT_EMAIL', '').strip()


def _normaliser_nom(nom):
    """Normalise un nom pour la comparaison : majuscules, sans accents, espaces simples."""
    import unicodedata
    nom = nom.upper().strip()
    nom = unicodedata.normalize('NFD', nom)
    nom = ''.join(c for c in nom if unicodedata.category(c) != 'Mn')
    nom = ' '.join(nom.split())
    return nom


def _zero_moyenne_counts(session_ids=None):
    """Retourne un dict {session_id: nb_bulletins_zero_moyenne}.

    Un bulletin est compté si son data_json contient moy_col_present == True
    ET moyenne == 0.0 (colonne présente mais valeur nulle → données incomplètes).
    Si session_ids est None, calcule pour toutes les sessions.
    """
    q = BulletinData.query
    if session_ids is not None:
        q = q.filter(BulletinData.session_id.in_(session_ids))
    counts = {}
    for bd in q.with_entities(BulletinData.session_id, BulletinData.data_json):
        try:
            etu = json.loads(bd.data_json) if bd.data_json else {}
        except Exception:
            etu = {}
        if etu.get('moy_col_present') and etu.get('moyenne', 1) == 0.0:
            counts[bd.session_id] = counts.get(bd.session_id, 0) + 1
    return counts


def _nb_suspects_nouveaux():
    """Nombre de scans suspects non-renvoyés (invalide / deja_utilise) depuis le dernier acquittement explicite."""
    try:
        cfg = AppConfig.query.filter_by(key='scans_last_acknowledged').first()
        if not cfg or not cfg.value.strip():
            return ScanLog.query.filter(
                ScanLog.resultat.in_(['invalide', 'deja_utilise']),
                ScanLog.dismissed != True,
            ).count()
        last_ack = datetime.fromisoformat(cfg.value.strip())
    except Exception:
        return ScanLog.query.filter(
            ScanLog.resultat.in_(['invalide', 'deja_utilise']),
            ScanLog.dismissed != True,
        ).count()
    return ScanLog.query.filter(
        ScanLog.resultat.in_(['invalide', 'deja_utilise']),
        ScanLog.dismissed != True,
        ScanLog.date_scan > last_ack,
    ).count()


def _construire_suivi_etudiants(sessions, promotion=None):
    """Construit le suivi indépendant des quatre étapes par matricule officiel.

    La clé de rapprochement est toujours le matricule, jamais le nom. Quand
    plusieurs imports existent pour une même étape, une ligne est considérée
    comme présente si elle existe dans l'une des sessions de cette étape et
    comme payée si l'un de ces bulletins est payé.
    """
    sessions = list(sessions or [])
    if promotion:
        sessions = [bs for bs in sessions if bs.promotion == promotion]

    grouped = defaultdict(lambda: defaultdict(list))
    names = {}
    promotions = set()
    for bs in sessions:
        grid_type = getattr(bs, 'type_grille', 'initial') or 'initial'
        if grid_type not in _GRID_TYPES:
            grid_type = 'initial'
        promotions.add(bs.promotion or '')
        for bd in bs.bulletins:
            matricule = (bd.matricule or '').strip().upper()
            if not matricule:
                continue
            grouped[matricule][grid_type].append(bd)
            names.setdefault(matricule, bd.nom or '')

    # La liste officielle permet aussi d'afficher les étudiants absents d'une
    # étape, sans les transformer en étudiants non-payés.
    official_rows = ListeIdentifiants.query
    if promotion:
        official_rows = official_rows.filter_by(promotion=promotion)
    else:
        official_rows = official_rows.filter(
            ListeIdentifiants.promotion.in_(list(promotions))
        ) if promotions else official_rows.filter(ListeIdentifiants.id == -1)
    for row in official_rows.all():
        matricule = (row.matricule or '').strip().upper()
        if matricule:
            grouped.setdefault(matricule, defaultdict(list))
            names.setdefault(matricule, row.nom or '')

    suivi = []
    for matricule, by_type in grouped.items():
        entry = {
            'matricule': matricule,
            'nom': names.get(matricule, ''),
            'promotion': promotion or '',
        }
        for grid_type in _GRID_TYPES:
            bulletins = by_type.get(grid_type, [])
            if not bulletins:
                entry[grid_type] = {'statut': 'absent', 'bulletin': None}
            else:
                paid = next((bd for bd in bulletins if bd.paye), None)
                latest = sorted(bulletins, key=lambda bd: bd.id, reverse=True)[0]
                entry[grid_type] = {
                    'statut': 'paye' if paid else 'attente',
                    'bulletin': paid or latest,
                }
        suivi.append(entry)

    return sorted(suivi, key=lambda item: (
        item['promotion'], item['nom'].upper(), item['matricule']
    ))


def _get_smtp_config():
    """Retourne (host, port, user, password, mode) depuis la DB ou les variables d'environnement.

    *mode* is one of: 'starttls' (default), 'ssl', 'none'.
    """
    def _cfg(key, env_key, default=''):
        try:
            row = AppConfig.query.filter_by(key=key).first()
            if row and row.value.strip():
                return row.value.strip()
        except Exception:
            pass
        return os.environ.get(env_key, default).strip()

    host  = _cfg('smtp_host',     'SMTP_HOST',     '')
    port  = _cfg('smtp_port',     'SMTP_PORT',     '587')
    user  = _cfg('smtp_user',     'SMTP_USER',     '')
    mode  = _cfg('smtp_mode',     'SMTP_MODE',     'starttls')
    if mode not in ('starttls', 'ssl', 'none'):
        mode = 'starttls'

    # ── SMTP password: decrypt from DB, fall back to env var ────────────────
    passw = ''
    try:
        row = AppConfig.query.filter_by(key='smtp_password').first()
        if row and row.value.strip():
            stored = row.value.strip()
            decrypted = _decrypt_smtp_password(stored)
            # Migration: plain-text row detected → re-encrypt and persist now
            if decrypted and not stored.encode().startswith(_SMTP_ENC_PREFIX):
                row.value = _encrypt_smtp_password(decrypted)
                try:
                    db.session.commit()
                except Exception:
                    db.session.rollback()
            passw = decrypted
    except Exception:
        pass
    if not passw:
        passw = os.environ.get('SMTP_PASSWORD', '').strip()

    try:
        port = int(port)
    except (ValueError, TypeError):
        port = 587
    return host, port, user, passw, mode


def _smtp_send(smtp_host, smtp_port, smtp_user, smtp_pass, smtp_mode, msg_obj, dest_email):
    """Ouvre une connexion SMTP selon le mode choisi et envoie *msg_obj*."""
    import smtplib
    if smtp_mode == 'ssl':
        with smtplib.SMTP_SSL(smtp_host, smtp_port, timeout=10) as srv:
            srv.login(smtp_user, smtp_pass)
            srv.sendmail(smtp_user, [dest_email], msg_obj.as_string())
    elif smtp_mode == 'none':
        with smtplib.SMTP(smtp_host, smtp_port, timeout=10) as srv:
            srv.ehlo()
            if smtp_pass:
                srv.login(smtp_user, smtp_pass)
            srv.sendmail(smtp_user, [dest_email], msg_obj.as_string())
    else:  # starttls (default)
        with smtplib.SMTP(smtp_host, smtp_port, timeout=10) as srv:
            srv.ehlo()
            srv.starttls()
            srv.login(smtp_user, smtp_pass)
            srv.sendmail(smtp_user, [dest_email], msg_obj.as_string())


def _send_scan_alert_email(code, resultat, ip, matricule=None, nom_etudiant=None):
    """Envoie un e-mail d'alerte au DÉCANAT en arrière-plan (silencieux si SMTP non configuré)."""
    smtp_host, smtp_port, smtp_user, smtp_pass, smtp_mode = _get_smtp_config()
    dest_email = _get_alert_email()

    if not (smtp_host and smtp_user and dest_email):
        return  # SMTP non configuré → pas d'envoi

    import threading, smtplib
    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart

    label  = 'Code QR inconnu' if resultat == 'invalide' else 'Reçu déjà utilisé'
    sujet  = f'⚠️ Alerte fraude – {label} détecté'
    horod  = now_cat().strftime('%d/%m/%Y à %H:%M:%S')

    rows = f"""
      <tr><td style="padding:.4rem .8rem;background:#f4f4f4;font-weight:600;">Résultat</td>
          <td style="padding:.4rem .8rem;">{label}</td></tr>
      <tr><td style="padding:.4rem .8rem;background:#f4f4f4;font-weight:600;">Code scanné</td>
          <td style="padding:.4rem .8rem;font-family:monospace;">{code}</td></tr>
      <tr><td style="padding:.4rem .8rem;background:#f4f4f4;font-weight:600;">Adresse IP</td>
          <td style="padding:.4rem .8rem;">{ip or '—'}</td></tr>
      <tr><td style="padding:.4rem .8rem;background:#f4f4f4;font-weight:600;">Date / Heure</td>
          <td style="padding:.4rem .8rem;">{horod}</td></tr>
    """
    if matricule:
        rows += f"""<tr><td style="padding:.4rem .8rem;background:#f4f4f4;font-weight:600;">Matricule</td>
          <td style="padding:.4rem .8rem;">{matricule}</td></tr>"""
    if nom_etudiant:
        rows += f"""<tr><td style="padding:.4rem .8rem;background:#f4f4f4;font-weight:600;">Étudiant</td>
          <td style="padding:.4rem .8rem;">{nom_etudiant}</td></tr>"""

    corps_html = f"""<html><body style="font-family:Arial,sans-serif;color:#222;max-width:560px;">
    <h2 style="color:#dc3545;margin-bottom:.5rem;">⚠️ Scan suspect détecté</h2>
    <p style="color:#555;margin-top:0;">Un scan inhabituel vient d'être enregistré sur le système de bulletins.</p>
    <table style="border-collapse:collapse;width:100%;">{rows}</table>
    <p style="color:#888;font-size:.85rem;margin-top:1.2rem;">
      Connectez-vous au tableau de bord pour consulter l'historique complet des scans suspects.
    </p>
    </body></html>"""

    def _send():
        try:
            msg = MIMEMultipart('alternative')
            msg['Subject'] = sujet
            msg['From']    = smtp_user
            msg['To']      = dest_email
            msg.attach(MIMEText(corps_html, 'html', 'utf-8'))
            _smtp_send(smtp_host, smtp_port, smtp_user, smtp_pass, smtp_mode, msg, dest_email)
        except Exception:
            pass  # Silencieux : ne pas bloquer la réponse au scanner

    threading.Thread(target=_send, daemon=True).start()


def _parse_liste_identifiants(file_path):
    """
    Parse un PDF de liste d'identifiants UNILU.
    3 étiquettes par ligne. Utilise les coordonnées x/y de chaque mot
    pour reconstruire exactement chaque étiquette (nom, matricule, mot de passe).
    Retourne (promotion, list_of_dict{nom, matricule, mot_de_passe}).
    """
    # ── Extraire tous les mots avec positions, page par page ──
    all_words = []
    promo = ''

    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            if page_num == 0:
                raw = page.extract_text() or ''
                m = re.search(r'Promotion\s*:\s*(.+?)(?:\n|Ann)', raw, re.IGNORECASE)
                if m:
                    p = m.group(1).strip()
                    p = re.sub(r'\s*/\s*FACULTE.*', '', p, flags=re.IGNORECASE).strip()
                    p = re.sub(r'\s*/\s*', ' ', p).strip()
                    promo = p
            for w in page.extract_words(x_tolerance=3, y_tolerance=3):
                all_words.append({
                    'text': w['text'],
                    'x0':   w['x0'],
                    'x1':   w['x1'],
                    'top':  w['top'],
                    'page': page_num,
                })

    if not all_words:
        return promo, []

    # ── Grouper les mots par ligne (top similaire, même page) ──
    Y_TOL = 4
    lines = []   # chaque ligne = liste de mots triés par x0
    for w in sorted(all_words, key=lambda w: (w['page'], w['top'], w['x0'])):
        placed = False
        for line in lines:
            if line[0]['page'] == w['page'] and abs(line[0]['top'] - w['top']) <= Y_TOL:
                line.append(w)
                placed = True
                break
        if not placed:
            lines.append([w])

    # ── Identifier les lignes "Matricule:" ──
    mat_lines = [ln for ln in lines
                 if any(w['text'].lower().startswith('matricule') for w in ln)]

    etudiants = []

    for mat_line in mat_lines:
        # Trouver les positions x des mots "Matricule:" → définissent les colonnes
        mat_anchors = sorted(
            [w for w in mat_line if w['text'].lower().startswith('matricule')],
            key=lambda w: w['x0']
        )
        n_cols = len(mat_anchors)
        if n_cols == 0:
            continue

        # Bornes x de chaque colonne : milieux entre les ancres + marges externes
        page_n = mat_anchors[0]['page']
        xs = [a['x0'] for a in mat_anchors]
        bounds = []
        for i, x in enumerate(xs):
            lo = (xs[i-1] + x) / 2 if i > 0 else 0
            hi = (x + xs[i+1]) / 2 if i < n_cols - 1 else 9999
            bounds.append((lo, hi))

        mat_top = mat_anchors[0]['top']

        # Trouver la ligne "Mot de passe:" juste en dessous
        pw_line = next(
            (ln for ln in lines
             if ln[0]['page'] == page_n
             and 2 < ln[0]['top'] - mat_top < 40
             and any('passe' in w['text'].lower() for w in ln)),
            None
        )

        # Trouver la ligne de noms juste au-dessus
        name_line = next(
            (ln for ln in reversed(lines)
             if ln[0]['page'] == page_n
             and 5 < mat_top - ln[0]['top'] < 60
             and all(not w['text'].lower().startswith(('matricule', 'promotion',
                     'ann', 'universit', 'page', 'mot'))
                     for w in ln)),
            None
        )

        for i, anchor in enumerate(mat_anchors):
            lo, hi = bounds[i]

            # ── Matricule value (mots à droite du label "Matricule:" sur la même ligne) ──
            mat_val_words = [w for w in mat_line
                             if w['x0'] > anchor['x1']
                             and lo <= w['x0'] < hi]
            mat_val = ' '.join(w['text'] for w in
                               sorted(mat_val_words, key=lambda w: w['x0'])).strip().upper()

            if not mat_val:
                continue

            # ── Mot de passe (mots dans cette colonne sur la ligne pw_line) ──
            pw_val = ''
            if pw_line:
                pw_words = [w for w in pw_line if lo <= w['x0'] < hi]
                pw_text  = ' '.join(w['text'] for w in sorted(pw_words, key=lambda w: w['x0']))
                m_pw = re.search(r'(\d{3,6})', pw_text)
                if m_pw:
                    pw_val = m_pw.group(1)

            # ── Nom (mots dans cette colonne sur la ligne name_line) ──
            nom = ''
            if name_line:
                name_words = [w for w in name_line if lo <= w['x0'] < hi]
                nom = ' '.join(w['text'] for w in
                               sorted(name_words, key=lambda w: w['x0'])).strip().upper()

            SKIP_NAMES = {'UNIVERSITE', 'UNIVERSITÉ', 'PROMOTION', 'ANNEE', 'ANNÉE',
                          'FACULTE', 'TOTAL', 'PAGE', 'ETIQUETTES'}
            if not nom or nom in SKIP_NAMES or len(nom) < 3:
                continue

            etudiants.append({
                'nom':          nom,
                'matricule':    mat_val,
                'mot_de_passe': pw_val,
            })

    return promo, etudiants


def generer_matricule(type_personne='etudiant'):
    import random
    
    if type_personne == 'etudiant':
        prefix = 'ETU'
        Model = Etudiant
    else:
        prefix = 'PROF'
        Model = Professeur
    
    while True:
        numero = random.randint(1000, 9999)
        matricule = f"{prefix}{numero}"
        
        existing = Model.query.filter_by(matricule=matricule).first()
        if not existing:
            return matricule

def generer_qrcode(data, matricule):
    qr = qrcode.QRCode(version=1, box_size=10, border=5)
    qr.add_data(json.dumps(data))
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    
    filename = f"{matricule}.png"
    filepath = os.path.join(app.config['QRCODE_FOLDER'], filename)
    img.save(filepath)
    with open(filepath, 'rb') as qr_file:
        _store_file_asset('qrcodes', filename, qr_file.read(), 'image/png')
    return filename

@app.route('/healthz')
def healthz():
    """Endpoint de health-check sans requête DB — utilisé par le probe autoscale."""
    return 'ok', 200

@app.route('/health')
def health():
    """Health-check avec vérification DB — utilisé par la page db_unavailable pour l'auto-retry."""
    try:
        db.session.execute(db.text('SELECT 1'))
        return jsonify(status='ok'), 200
    except Exception:
        return jsonify(status='unavailable'), 503

@app.route('/')
def index():
    total_visites = _enregistrer_visite_site()
    scanner_visible = _get_app_config_bool('home_scanner_visible', True)
    search_visible = _get_app_config_bool('home_search_visible', True)
    actualites_visible = _get_app_config_bool('home_actualites_visible', True)
    info_visible = _get_app_config_bool('home_info_visible', True)
    try:
        actualites = (
            Actualite.query.filter_by(publie=True)
            .filter_by(type_publication='actualite')
            .order_by(Actualite.epingle.desc(), Actualite.date_publication.desc())
            .limit(5)
            .all()
            if actualites_visible else []
        )
    except Exception:
        actualites = []
    return render_template(
        'index.html',
        actualites=actualites,
        scanner_visible=scanner_visible,
        search_visible=search_visible,
        info_visible=info_visible,
        total_visites=total_visites,
    )


@app.route('/communiques')
def communiques():
    """Page publique dédiée aux communiqués officiels de la faculté."""
    publications = (
        Actualite.query.filter_by(publie=True, type_publication='communique')
        .order_by(Actualite.epingle.desc(), Actualite.date_publication.desc())
        .all()
    )
    return render_template('communiques.html', publications=publications)


def _chat_enabled():
    return _get_app_config_bool('chat_enabled', True)


def _chat_registration_open():
    return _get_app_config_bool('chat_registration_open', True)


def _student_registration_open():
    """Indique si le formulaire principal d'inscription est ouvert."""
    return _get_app_config_bool('student_registration_open', True)


def _releves_public_accessible():
    """Indique si le portail public des relevés est ouvert aux étudiants."""
    return _get_app_config_bool('releves_public_access', True)


@app.context_processor
def inject_chat_widget_state():
    return {
        'chat_widget_enabled': _chat_enabled(),
        'releves_public_accessible': _releves_public_accessible(),
    }


def _normaliser_nom_chat(value):
    return re.sub(r'\s+', ' ', (value or '').strip()).casefold()


def _chat_thread_for_student():
    matricule = session.get('chat_matricule')
    if not matricule:
        return None
    return ChatThread.query.filter_by(matricule=matricule).first()


def _chat_payload(thread):
    return [{
        'id': message.id,
        'role': message.sender_role,
        'nom': message.sender_name,
        'contenu': message.contenu or '',
        'type': 'texte',
        'date': message.date_creation.strftime('%d/%m/%Y %H:%M'),
    } for message in thread.messages if message.type_message == 'texte']


@app.route('/chat', methods=['GET', 'POST'])
def chat():
    """Connexion et espace privé de communication étudiant-DÉCANAT."""
    if not _chat_enabled():
        return render_template(
            'chat.html', mode='disabled', promotions=PROMOTIONS,
            registration_open=_chat_registration_open(),
        )

    thread = _chat_thread_for_student()
    if request.method == 'POST':
        matricule = request.form.get('matricule', '').strip().upper()
        existing_thread = ChatThread.query.filter_by(matricule=matricule).first()
        if existing_thread:
            if existing_thread.actif:
                session['chat_matricule'] = matricule
                return redirect(url_for('chat'))
            flash('Votre espace de communication est momentanément désactivé par le DÉCANAT.', 'error')
            return render_template(
                'chat.html', mode='login', promotions=PROMOTIONS,
                registration_open=_chat_registration_open(),
            )

        if not _chat_registration_open():
            flash('Les nouvelles inscriptions de contact sont momentanément fermées par le DÉCANAT.', 'error')
            return render_template(
                'chat.html', mode='login', promotions=PROMOTIONS,
                registration_open=False,
            )

        identifiant = ListeIdentifiants.query.filter(
            db.func.upper(ListeIdentifiants.matricule) == matricule
        ).order_by(ListeIdentifiants.id.desc()).first()
        if not identifiant:
            flash('Ce CODE-ID FAC ne figure pas dans la liste officielle des identifiants publiée.', 'error')
            return render_template(
                'chat.html', mode='login', promotions=PROMOTIONS,
                registration_open=True,
            )

        thread = ChatThread(
            matricule=identifiant.matricule.strip().upper(),
            nom_complet=identifiant.nom.strip(),
            promotion=identifiant.promotion.strip(),
        )
        db.session.add(thread)
        db.session.commit()
        session['chat_matricule'] = thread.matricule
        return redirect(url_for('chat'))

    if thread is None:
        return render_template(
            'chat.html', mode='login', promotions=PROMOTIONS,
            registration_open=_chat_registration_open(),
        )
    if not thread.actif:
        session.pop('chat_matricule', None)
        return render_template(
            'chat.html', mode='disabled', promotions=PROMOTIONS,
            registration_open=_chat_registration_open(),
        )
    return render_template(
        'chat.html',
        mode='chat',
        thread=thread,
        messages=_chat_payload(thread),
        promotions=PROMOTIONS,
    )


@app.route('/chat/deconnexion')
def chat_deconnexion():
    session.pop('chat_matricule', None)
    return redirect(url_for('chat'))


@app.route('/chat/message', methods=['POST'])
def chat_message():
    if not _chat_enabled():
        return jsonify({'ok': False, 'error': 'Chat désactivé'}), 403
    thread = _chat_thread_for_student()
    if not thread or not thread.actif:
        return jsonify({'ok': False, 'error': 'Espace non connecté'}), 403
    contenu = request.form.get('contenu', '').strip()
    if not contenu or len(contenu) > 4000:
        return jsonify({'ok': False, 'error': 'Message vide ou trop long'}), 400
    db.session.add(ChatMessage(
        thread_id=thread.id,
        sender_role='etudiant',
        sender_name=thread.nom_complet,
        contenu=contenu,
        type_message='texte',
    ))
    thread.date_dernier_message = now_cat()
    db.session.commit()
    return jsonify({'ok': True})


@app.route('/chat/audio', methods=['POST'])
def chat_audio_upload():
    return jsonify({'ok': False, 'error': 'Les messages vocaux sont désactivés'}), 410


@app.route('/chat/messages')
def chat_messages():
    thread = _chat_thread_for_student()
    if not thread or not thread.actif:
        return jsonify({'ok': False}), 403
    return jsonify({'ok': True, 'messages': _chat_payload(thread)})


@app.route('/chat/audio/<path:filename>')
def chat_audio(filename):
    return 'Les messages vocaux sont désactivés', 410


@app.route('/chat/appel', methods=['POST'])
def chat_appel():
    return jsonify({'ok': False, 'error': 'Les appels audio sont désactivés'}), 410


@app.route('/chat/appel/<int:call_id>')
def chat_appel_statut(call_id):
    return jsonify({'ok': False, 'error': 'Les appels audio sont désactivés'}), 410


@app.route('/decanat/chat')
def decanat_chat():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    threads = ChatThread.query.order_by(
        ChatThread.date_dernier_message.desc()
    ).all()
    selected_id = request.args.get('thread', type=int)
    selected = ChatThread.query.get(selected_id) if selected_id else (threads[0] if threads else None)
    if selected:
        ChatMessage.query.filter_by(
            thread_id=selected.id, sender_role='etudiant', lu=False
        ).update({'lu': True}, synchronize_session=False)
        db.session.commit()
    return render_template(
        'decanat_chat.html',
        threads=threads,
        selected=selected,
        calls=(
            ChatCall.query.filter_by(thread_id=selected.id)
            .filter(ChatCall.status.in_(['pending', 'accepted']))
            .order_by(ChatCall.date_creation.desc()).all()
            if selected else []
        ),
        chat_enabled=_chat_enabled(),
    )


@app.route('/decanat/chat/messages/<int:thread_id>')
def decanat_chat_messages(thread_id):
    if not session.get('decanat_logged_in'):
        return jsonify({'ok': False}), 403
    thread = ChatThread.query.get_or_404(thread_id)
    return jsonify({
        'ok': True,
        'messages': _chat_payload(thread),
        'calls': [{
            'id': call.id,
            'status': call.status,
            'date': call.date_creation.strftime('%d/%m/%Y %H:%M'),
        } for call in thread.appels if call.status in ('pending', 'accepted')],
    })


@app.route('/decanat/chat/<int:thread_id>/message', methods=['POST'])
def decanat_chat_message(thread_id):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    thread = ChatThread.query.get_or_404(thread_id)
    contenu = request.form.get('contenu', '').strip()
    if contenu and len(contenu) <= 4000 and thread.actif:
        db.session.add(ChatMessage(
            thread_id=thread.id,
            sender_role='decanat',
            sender_name=session.get('decanat_operator', 'DÉCANAT'),
            contenu=contenu,
            type_message='texte',
        ))
        thread.date_dernier_message = now_cat()
        db.session.commit()
    return redirect(url_for('decanat_chat', thread=thread.id))


@app.route('/decanat/chat/<int:thread_id>/audio', methods=['POST'])
def decanat_chat_audio(thread_id):
    return jsonify({'ok': False, 'error': 'Les messages vocaux sont désactivés'}), 410


@app.route('/decanat/chat/<int:thread_id>/actif', methods=['POST'])
def decanat_chat_toggle(thread_id):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    thread = ChatThread.query.get_or_404(thread_id)
    thread.actif = not thread.actif
    db.session.commit()
    flash(
        'Espace chat activé.' if thread.actif else 'Espace chat désactivé.',
        'success',
    )
    return redirect(url_for('decanat_chat', thread=thread.id))


@app.route('/decanat/chat/appel/<int:call_id>/accepter', methods=['POST'])
def decanat_chat_accept_call(call_id):
    return jsonify({'ok': False, 'error': 'Les appels audio sont désactivés'}), 410


@app.route('/decanat/chat/appel/<int:call_id>/terminer', methods=['POST'])
def decanat_chat_end_call(call_id):
    return jsonify({'ok': False, 'error': 'Les appels audio sont désactivés'}), 410


@app.route('/decanat/chat/appel/<int:call_id>/offre')
def decanat_chat_call_offer(call_id):
    return jsonify({'ok': False, 'error': 'Les appels audio sont désactivés'}), 410

@app.route('/a-propos')
def a_propos():
    page_content = PageContent.query.filter_by(page_name='a_propos').first()
    content_data = json.loads(page_content.content_json) if page_content and page_content.content_json else {}
    image_facade = page_content.image_principale if page_content and page_content.image_principale else 'faculte-facade.jpg'
    return render_template('a_propos.html', content_data=content_data, image_facade=image_facade)

@app.route('/departements')
def departements():
    return render_template('departements.html')

@app.route('/recherche', methods=['GET', 'POST'])
def recherche():
    if request.method == 'POST':
        matricule = request.form.get('matricule', '').strip().upper()
        if matricule:
            etudiant = Etudiant.query.filter_by(matricule=matricule).first()
            if etudiant:
                return render_template('resultat_recherche.html', etudiant=etudiant)
            else:
                flash(f'Aucun étudiant trouvé avec le matricule: {matricule}', 'error')
        else:
            flash('Veuillez entrer un matricule', 'error')
    
    return redirect(url_for('index'))

@app.route('/inscription', methods=['GET', 'POST'])
def inscription():
    if not _student_registration_open():
        return render_template(
            'inscription.html',
            departements=DEPARTEMENTS,
            promotions=PROMOTIONS,
            registration_open=False,
        ), 403

    if request.method == 'POST':
        try:
            nom = request.form['nom']
            postnom = request.form['postnom']
            prenom = request.form['prenom']
            sexe = request.form['sexe']
            telephone = request.form['telephone']
            promotion = request.form['promotion']
            departement = request.form['departement']
            
            if 'photo' not in request.files:
                flash('Photo requise', 'error')
                return redirect(request.url)
            
            file = request.files['photo']
            if file.filename == '':
                flash('Aucune photo sélectionnée', 'error')
                return redirect(request.url)
            
            if file and allowed_file(file.filename):
                matricule = generer_matricule('etudiant')
                filename = secure_filename(f"{matricule}_{file.filename}")
                _save_uploaded_asset(file, 'uploads', filename)
                
                qr_data = {
                    'matricule': matricule,
                    'nom': nom,
                    'postnom': postnom,
                    'prenom': prenom,
                    'departement': departement,
                    'promotion': promotion,
                    'type': 'etudiant'
                }
                qrcode_filename = generer_qrcode(qr_data, matricule)
                
                etudiant = Etudiant(
                    matricule=matricule,
                    nom=nom,
                    postnom=postnom,
                    prenom=prenom,
                    sexe=sexe,
                    telephone=telephone,
                    promotion=promotion,
                    departement=departement,
                    photo=filename,
                    qrcode_path=qrcode_filename
                )
                
                db.session.add(etudiant)
                db.session.commit()
                
                flash(f'Inscription réussie! Votre matricule est: {matricule}', 'success')
                return redirect(url_for('confirmation', matricule=matricule))
            else:
                flash('Format de fichier non autorisé. Utilisez PNG, JPG ou JPEG', 'error')
                
        except Exception as e:
            db.session.rollback()
            flash(f'Erreur lors de l\'inscription: {str(e)}', 'error')
    
    return render_template('inscription.html', departements=DEPARTEMENTS, promotions=PROMOTIONS)

@app.route('/confirmation/<matricule>')
def confirmation(matricule):
    etudiant = Etudiant.query.filter_by(matricule=matricule).first_or_404()
    return render_template('confirmation.html', etudiant=etudiant)

@app.route('/telecharger/qrcode/<matricule>')
def telecharger_qrcode(matricule):
    etudiant = Etudiant.query.filter_by(matricule=matricule).first_or_404()
    if etudiant.qrcode_path:
        response = _send_persistent_or_local(
            'qrcodes',
            etudiant.qrcode_path,
            as_attachment=True,
            download_name=f'QRCode_{matricule}.png',
        )
        if response is None:
            flash('Le fichier QR code est introuvable sur le serveur. Réinscrivez l\'étudiant pour le régénérer.', 'error')
            return redirect(url_for('index'))
        return response
    else:
        flash('QR code introuvable', 'error')
        return redirect(url_for('index'))

@app.route('/telecharger/preuve_inscription/<matricule>')
def telecharger_preuve_inscription(matricule):
    etudiant = Etudiant.query.filter_by(matricule=matricule).first_or_404()
    
    output = BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    elements = []
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#6c5ce7'),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    header_style = ParagraphStyle(
        'Header',
        parent=styles['Normal'],
        fontSize=14,
        textColor=colors.HexColor('#2d3436'),
        spaceAfter=20,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    elements.append(Paragraph("🎓 UNIVERSITÉ DE LUBUMBASHI (UNILU)", header_style))
    elements.append(Paragraph("E-SCIALES UNILU + - Faculté des Sciences Sociales, Politiques et Administratives", styles['Normal']))
    elements.append(Spacer(1, 0.5*cm))
    
    elements.append(Paragraph("PREUVE D'INSCRIPTION FACULTAIRE", title_style))
    elements.append(Spacer(1, 0.5*cm))
    
    date_str = etudiant.date_inscription.strftime('%d/%m/%Y à %H:%M')
    elements.append(Paragraph(f"<b>Date d'inscription:</b> {date_str}", styles['Normal']))
    elements.append(Spacer(1, 0.5*cm))
    
    data = [
        ['CODE-ID FAC:', etudiant.matricule],
        ['Nom:', etudiant.nom],
        ['Postnom:', etudiant.postnom],
        ['Prénom:', etudiant.prenom],
        ['Sexe:', etudiant.sexe],
        ['Téléphone:', etudiant.telephone],
        ['Département:', etudiant.departement],
        ['Promotion:', etudiant.promotion],
    ]
    
    table = Table(data, colWidths=[5*cm, 10*cm])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e3e8f0')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    elements.append(table)
    elements.append(Spacer(1, 1*cm))
    
    if etudiant.qrcode_path:
        qrcode_path = os.path.join(app.config['QRCODE_FOLDER'], etudiant.qrcode_path)
        if os.path.exists(qrcode_path):
            elements.append(Paragraph("QR Code de l'étudiant:", styles['Heading3']))
            elements.append(Spacer(1, 0.3*cm))
            qr_img = Image(qrcode_path, width=4*cm, height=4*cm)
            elements.append(qr_img)
            elements.append(Spacer(1, 0.3*cm))
            elements.append(Paragraph("<i>Ce QR code est utilisé pour le contrôle de présence aux cours.</i>", styles['Normal']))
    
    elements.append(Spacer(1, 1.5*cm))
    
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.grey,
        alignment=TA_CENTER
    )
    elements.append(Paragraph("___________________________", footer_style))
    elements.append(Spacer(1, 0.2*cm))
    elements.append(Paragraph("E-SCIALES UNILU +", footer_style))
    elements.append(Paragraph(f"Document généré le {now_cat().strftime('%d/%m/%Y à %H:%M')}", footer_style))
    
    doc.build(elements)
    output.seek(0)
    
    filename = f'Preuve_Inscription_{matricule}.pdf'
    return send_file(output, download_name=filename, as_attachment=True, mimetype='application/pdf')

@app.route('/telecharger/carte_examen/<matricule>')
def telecharger_carte_examen(matricule):
    etudiant = Etudiant.query.filter_by(matricule=matricule).first_or_404()

    from reportlab.pdfgen import canvas as pdfcanvas
    from reportlab.lib.units import cm
    from PIL import Image as PILImage, ImageDraw, ImageFont

    # ── Chemins des polices DejaVu ──
    FONT_REG  = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
    FONT_BOLD = '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'
    FONT_OBL  = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'

    def fnt(path, size):
        try:
            return ImageFont.truetype(path, size)
        except:
            return ImageFont.load_default()

    # ── Construction de la carte propre ──
    tpl_path = os.path.join('static', 'carte_template.png')
    template = PILImage.open(tpl_path).convert('RGBA')
    TW, TH = template.size   # 963 × 612
    BODY_TOP = 128            # le bandeau bleu (0,127,255) se termine à y≈128

    # Étape 1 : carte blanche pure
    card = PILImage.new('RGBA', (TW, TH), (255, 255, 255, 255))

    # Étape 2 : coller le bandeau (header uniquement, y=0..BODY_TOP)
    header = template.crop((0, 0, TW, BODY_TOP)).convert('RGBA')
    card.paste(header, (0, 0), header)

    # Étape 3 : appliquer l'alpha du template (coins arrondis) comme masque final
    # On prend UNIQUEMENT l'alpha — les zones transparentes du template (coins) restent transparentes
    alpha_mask = template.split()[3]   # canal alpha 963×612

    draw = ImageDraw.Draw(card)

    # ── Couleurs (RGBA) ──
    BLANC  = (255, 255, 255, 255)
    MARINE = (13,  38,  102,  255)   # bleu marine UNILU
    GRIS   = (107, 107, 118,  255)   # gris labels italiques

    # Ligne de séparation bleue sous le bandeau
    draw.rectangle([0, BODY_TOP, TW, BODY_TOP + 6], fill=(13, 38, 102, 255))

    # Fond gris clair zone photo
    PH_X1 = 28
    PH_Y1 = BODY_TOP + 2
    PH_X2 = 225
    PH_Y2 = TH - 10
    draw.rectangle([PH_X1, PH_Y1, PH_X2, PH_Y2], fill=(228, 228, 228, 255))

    # ═══════════════════════════════════════════════════
    # 2. PHOTO étudiant
    # ═══════════════════════════════════════════════════
    if etudiant.photo:
        photo_path = os.path.join(app.config['UPLOAD_FOLDER'], etudiant.photo)
        if os.path.exists(photo_path):
            try:
                ph_img = PILImage.open(photo_path).convert('RGBA')
                ph_img = ph_img.resize((PH_X2 - PH_X1, PH_Y2 - PH_Y1), PILImage.LANCZOS)
                card.paste(ph_img, (PH_X1, PH_Y1), ph_img)
            except:
                pass
    draw.rectangle([PH_X1, PH_Y1, PH_X2, PH_Y2], outline=(160, 160, 165, 255), width=2)

    # ═══════════════════════════════════════════════════
    # 3. N° CODE-ID FAC  (grand, marine, bold)
    # ═══════════════════════════════════════════════════
    f_num = fnt(FONT_BOLD, 30)
    num_txt = f'N\u00b0  {etudiant.matricule}'
    draw.text((248, BODY_TOP + 15), num_txt, font=f_num, fill=MARINE)
    draw.line([(248, BODY_TOP + 58), (700, BODY_TOP + 58)], fill=(200, 200, 205, 255), width=1)

    # ═══════════════════════════════════════════════════
    # 4. INFORMATIONS (colonne centre)
    # ═══════════════════════════════════════════════════
    f_lbl = fnt(FONT_OBL, 17)    # label italique gris
    f_val = fnt(FONT_BOLD, 21)   # valeur gras marine
    f_dep = fnt(FONT_BOLD, 22)   # département plus grand

    IX = 248          # x de départ colonne infos
    iy = BODY_TOP + 70

    def put_label(txt, y):
        draw.text((IX, y), txt, font=f_lbl, fill=GRIS)

    def put_value(txt, y, size=21):
        f = fnt(FONT_BOLD, size)
        draw.text((IX, y), txt, font=f, fill=MARINE)

    # Nom
    put_label('Nom :', iy)
    iy += 25
    put_value(etudiant.nom.upper(), iy)
    iy += 36

    # Post-nom & Prénom
    put_label('Post-nom & Pr\u00e9nom', iy)
    iy += 25
    put_value(f'{etudiant.postnom.upper()} {etudiant.prenom}', iy, size=20)
    iy += 36

    # Hash + Département
    draw.text((IX, iy), '#', font=fnt(FONT_BOLD, 20), fill=MARINE)
    iy += 24
    draw.text((IX, iy), etudiant.departement.upper(), font=f_dep, fill=MARINE)
    iy += 36

    # CODE-ID FAC
    put_label('CODE-ID FAC', iy)
    iy += 25
    put_value(etudiant.matricule, iy)
    iy += 36

    # Promotion
    put_label('Promotion', iy)
    iy += 25
    put_value(etudiant.promotion.upper(), iy, size=19)

    # ═══════════════════════════════════════════════════
    # 5. QR CODE (colonne droite)
    # ═══════════════════════════════════════════════════
    QR_X  = 722
    QR_Y  = BODY_TOP + 10
    QR_SIZE = 210
    if etudiant.qrcode_path:
        qr_path = os.path.join(app.config['QRCODE_FOLDER'], etudiant.qrcode_path)
        if os.path.exists(qr_path):
            try:
                qr_img = PILImage.open(qr_path).convert('RGBA')
                qr_img = qr_img.resize((QR_SIZE, QR_SIZE), PILImage.LANCZOS)
                card.paste(qr_img, (QR_X, QR_Y), qr_img)
            except:
                pass
    draw.rectangle([QR_X, QR_Y, QR_X + QR_SIZE, QR_Y + QR_SIZE],
                   outline=(140, 140, 145, 255), width=1)

    # ═══════════════════════════════════════════════════
    # 6. SEXE
    # ═══════════════════════════════════════════════════
    sexe_cx = 827
    draw.text((sexe_cx - 22, BODY_TOP + 232), 'SEXE', font=fnt(FONT_BOLD, 16), fill=GRIS)
    sexe_lettre = 'M' if etudiant.sexe.upper().startswith('M') else 'F'
    f_sexe = fnt(FONT_BOLD, 52)
    bbox = draw.textbbox((0, 0), sexe_lettre, font=f_sexe)
    sw = bbox[2] - bbox[0]
    draw.text((sexe_cx - sw // 2, BODY_TOP + 252), sexe_lettre, font=f_sexe, fill=MARINE)

    # ═══════════════════════════════════════════════════
    # 7. BADGE DORÉ ÉTUDIANT + DATE DE DÉLIVRANCE
    # ═══════════════════════════════════════════════════
    mois_fr = ['janvier','f\u00e9vrier','mars','avril','mai','juin',
               'juillet','ao\u00fbt','septembre','octobre','novembre','d\u00e9cembre']
    now = now_cat()
    date_str = f"{now.day} {mois_fr[now.month-1]} {now.year}"

    f_date_lbl = fnt(FONT_OBL,  18)
    f_date_val = fnt(FONT_BOLD, 18)

    # Badge doré ÉTUDIANT
    bdg_x, bdg_w, bdg_h = 718, 213, 76
    bdg_y = TH - 90
    draw.rounded_rectangle(
        [bdg_x, bdg_y, bdg_x + bdg_w, bdg_y + bdg_h],
        radius=10,
        fill=(191, 153, 26, 255),
        outline=(127, 102, 5, 255),
        width=2
    )
    f_bdg_ico = fnt(FONT_BOLD, 22)
    f_bdg_txt = fnt(FONT_BOLD, 18)
    ico_bb = draw.textbbox((0, 0), '\u26c3', font=f_bdg_ico)
    draw.text((bdg_x + bdg_w // 2 - (ico_bb[2] - ico_bb[0]) // 2, bdg_y + 7),
              '\u26c3', font=f_bdg_ico, fill=(255, 255, 255, 255))
    txt_bb = draw.textbbox((0, 0), '\u00c9TUDIANT', font=f_bdg_txt)
    draw.text((bdg_x + bdg_w // 2 - (txt_bb[2] - txt_bb[0]) // 2, bdg_y + 44),
              '\u00c9TUDIANT', font=f_bdg_txt, fill=(255, 255, 255, 255))

    # Date de délivrance
    draw.text((35, TH - 42), 'Date de d\u00e9livrance :', font=f_date_lbl, fill=(80, 80, 85, 255))
    lbl_w = draw.textbbox((0, 0), 'Date de d\u00e9livrance :', font=f_date_lbl)[2]
    draw.text((35 + lbl_w + 8, TH - 42), date_str, font=f_date_val, fill=MARINE)

    # ═══════════════════════════════════════════════════
    # 8. APPLIQUER LE MASQUE ARRONDI + CONVERTIR EN PDF
    # ═══════════════════════════════════════════════════
    # Appliquer l'alpha du template pour les coins arrondis
    card.putalpha(alpha_mask)
    # Convertir RGBA → RGB pour PDF (fond blanc)
    card_rgb = PILImage.new('RGB', card.size, (255, 255, 255))
    card_rgb.paste(card, mask=alpha_mask)

    # Sauvegarder l'image en mémoire
    img_buf = BytesIO()
    card_rgb.save(img_buf, format='PNG')
    img_buf.seek(0)

    # Dimensions PDF = taille de la carte en cm à 96 DPI
    PDF_W = TW / 96 * 2.54 * cm
    PDF_H = TH / 96 * 2.54 * cm

    output = BytesIO()
    c = pdfcanvas.Canvas(output, pagesize=(PDF_W, PDF_H))
    from reportlab.lib.utils import ImageReader
    c.drawImage(ImageReader(img_buf), 0, 0, width=PDF_W, height=PDF_H)
    c.save()
    output.seek(0)

    filename = f'Carte_Etudiant_{matricule}.pdf'
    return send_file(output, download_name=filename, as_attachment=True, mimetype='application/pdf')

@app.route('/decanat/login', methods=['GET', 'POST'])
def decanat_login():
    if request.method == 'POST':
        password = request.form['password']
        if password == DECANAT_PASSWORD:
            session['decanat_logged_in'] = True
            operator = request.form.get('operator_name', '').strip()
            session['decanat_operator'] = operator or 'Inconnu'
            flash('Connexion réussie', 'success')
            return redirect(url_for('decanat_dashboard'))
        else:
            flash('Mot de passe incorrect', 'error')
    
    return render_template('decanat_login.html')

@app.route('/decanat/switch-operator', methods=['POST'])
def decanat_switch_operator():
    """Change le nom de l'opérateur courant sans déconnexion."""
    if not session.get('decanat_logged_in'):
        return jsonify({'ok': False, 'error': 'Non connecté'}), 403
    new_name = request.form.get('operator_name', '').strip()
    if not new_name:
        return jsonify({'ok': False, 'error': 'Nom vide'}), 400
    session['decanat_operator'] = new_name
    next_url = request.form.get('next') or request.referrer or url_for('decanat_dashboard')
    return redirect(next_url)

@app.route('/decanat/logout')
def decanat_logout():
    session.pop('decanat_logged_in', None)
    session.pop('decanat_operator', None)
    flash('Déconnexion réussie', 'success')
    return redirect(url_for('index'))

@app.route('/prof/login', methods=['GET', 'POST'])
def prof_login():
    if request.method == 'POST':
        password = request.form['password']
        if password == PROF_PASSWORD:
            session['prof_logged_in'] = True
            flash('Connexion réussie', 'success')
            return redirect(url_for('decanat_presences'))
        else:
            flash('Mot de passe incorrect', 'error')
    
    return render_template('prof_login.html')

@app.route('/prof/logout')
def prof_logout():
    session.pop('prof_logged_in', None)
    flash('Déconnexion réussie', 'success')
    return redirect(url_for('index'))

@app.route('/decanat/dashboard')
def decanat_dashboard():
    if not session.get('decanat_logged_in'):
        flash('Accès non autorisé', 'error')
        return redirect(url_for('decanat_login'))
    
    stats = {
        'total_etudiants': Etudiant.query.count(),
        'total_professeurs': Professeur.query.count(),
        'total_cours': Cours.query.count(),
        'total_visites': SiteVisit.query.count(),
    }
    debut_jour = datetime.combine(now_cat().date(), datetime.min.time())
    stats['visites_aujourd_hui'] = SiteVisit.query.filter(
        SiteVisit.date_visite >= debut_jour
    ).count()

    template_manquant = _get_modele_docx() is None
    home_scanner_visible = _get_app_config_bool('home_scanner_visible', True)
    home_search_visible = _get_app_config_bool('home_search_visible', True)
    home_actualites_visible = _get_app_config_bool('home_actualites_visible', True)
    home_info_visible = _get_app_config_bool('home_info_visible', True)
    chat_enabled = _get_app_config_bool('chat_enabled', True)
    chat_registration_open = _chat_registration_open()
    student_registration_open = _student_registration_open()
    releves_public_accessible = _releves_public_accessible()

    return render_template('decanat_dashboard.html', stats=stats,
                         departements=DEPARTEMENTS, promotions=PROMOTIONS,
                          template_manquant=template_manquant,
                          home_scanner_visible=home_scanner_visible,
                           home_actualites_visible=home_actualites_visible,
                           home_search_visible=home_search_visible,
                           home_info_visible=home_info_visible,
                           chat_enabled=chat_enabled,
                           chat_registration_open=chat_registration_open,
                           student_registration_open=student_registration_open,
                           releves_public_accessible=releves_public_accessible)


@app.route('/decanat/dashboard/visibility', methods=['POST'])
def decanat_dashboard_visibility():
    """Enregistre les éléments visibles sur la page d'accueil publique."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    def _upsert_visibility(key, visible):
        row = AppConfig.query.filter_by(key=key).first()
        value = '1' if visible else '0'
        if row:
            row.value = value
        else:
            db.session.add(AppConfig(key=key, value=value))

    try:
        _upsert_visibility(
            'home_scanner_visible',
            request.form.get('home_scanner_visible') == '1',
        )
        _upsert_visibility(
            'home_actualites_visible',
            request.form.get('home_actualites_visible') == '1',
        )
        _upsert_visibility(
            'home_search_visible',
            request.form.get('home_search_visible') == '1',
        )
        _upsert_visibility(
            'home_info_visible',
            request.form.get('home_info_visible') == '1',
        )
        _upsert_visibility(
            'chat_enabled',
            request.form.get('chat_enabled') == '1',
        )
        _upsert_visibility(
            'chat_registration_open',
            request.form.get('chat_registration_open') == '1',
        )
        _upsert_visibility(
            'student_registration_open',
            request.form.get('student_registration_open') == '1',
        )
        db.session.commit()
        flash('Visibilité de la page d’accueil mise à jour.', 'success')
    except Exception as exc:
        db.session.rollback()
        flash(f'Erreur lors de l’enregistrement de la visibilité : {exc}', 'error')
    return redirect(url_for('decanat_dashboard'))


@app.route('/decanat/releves-cotes/acces-public', methods=['POST'])
def decanat_releves_acces_public():
    """Active ou désactive l’accès étudiant au portail public des relevés."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    row = AppConfig.query.filter_by(key='releves_public_access').first()
    value = '1' if request.form.get('releves_public_access') == '1' else '0'
    if row:
        row.value = value
    else:
        db.session.add(AppConfig(key='releves_public_access', value=value))
    db.session.commit()
    flash(
        'Accès public aux relevés activé.' if value == '1'
        else 'Accès public aux relevés désactivé.',
        'success',
    )
    return redirect(url_for('decanat_releves_cotes'))

@app.route('/decanat/etudiants')
def decanat_etudiants():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    
    departement = request.args.get('departement', '')
    promotion = request.args.get('promotion', '')
    
    query = Etudiant.query
    if departement:
        query = query.filter_by(departement=departement)
    if promotion:
        query = query.filter_by(promotion=promotion)
    
    etudiants = query.order_by(Etudiant.nom).all()
    
    return render_template('decanat_etudiants.html', etudiants=etudiants,
                         departements=DEPARTEMENTS, promotions=PROMOTIONS,
                         selected_dept=departement, selected_promo=promotion)


@app.route('/decanat/etudiants/supprimer/<int:etudiant_id>', methods=['POST'])
def supprimer_etudiant(etudiant_id):
    """Supprime un candidat inscrit, avec ses traces de présence et fichiers."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    try:
        etudiant = Etudiant.query.get_or_404(etudiant_id)
        matricule = etudiant.matricule
        if etudiant.photo:
            local_path = _asset_local_path('uploads', etudiant.photo)
            if local_path and os.path.exists(local_path):
                os.remove(local_path)
            _delete_file_asset('uploads', etudiant.photo)
        if etudiant.qrcode_path:
            local_path = _asset_local_path('qrcodes', etudiant.qrcode_path)
            if local_path and os.path.exists(local_path):
                os.remove(local_path)
            _delete_file_asset('qrcodes', etudiant.qrcode_path)
        Presence.query.filter_by(etudiant_id=etudiant.id).delete(
            synchronize_session=False
        )
        db.session.delete(etudiant)
        db.session.commit()
        flash(f'Candidat {matricule} supprimé avec succès.', 'success')
    except Exception as exc:
        db.session.rollback()
        flash(f'Erreur lors de la suppression du candidat : {exc}', 'error')
    return redirect(url_for(
        'decanat_etudiants',
        departement=request.form.get('departement', ''),
        promotion=request.form.get('promotion', ''),
    ))

@app.route('/decanat/cours', methods=['GET', 'POST'])
def decanat_cours():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    
    if request.method == 'POST':
        try:
            code = request.form['code']
            nom = request.form['nom']
            departement = request.form['departement']
            promotion = request.form['promotion']
            
            cours = Cours(
                code=code,
                nom=nom,
                departement=departement,
                promotion=promotion
            )
            
            db.session.add(cours)
            db.session.commit()
            
            flash(f'Cours "{nom}" ajouté avec succès!', 'success')
            return redirect(url_for('decanat_cours'))
            
        except Exception as e:
            db.session.rollback()
            flash(f'Erreur lors de l\'ajout du cours: {str(e)}', 'error')
    
    cours_list = Cours.query.order_by(Cours.nom).all()
    return render_template('decanat_cours.html', cours_list=cours_list,
                         departements=DEPARTEMENTS, promotions=PROMOTIONS)

@app.route('/decanat/cours/supprimer/<int:cours_id>', methods=['POST'])
def supprimer_cours(cours_id):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    
    try:
        cours = Cours.query.get_or_404(cours_id)
        db.session.delete(cours)
        db.session.commit()
        flash('Cours supprimé avec succès', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erreur lors de la suppression: {str(e)}', 'error')
    
    return redirect(url_for('decanat_cours'))

@app.route('/decanat/cours/masse', methods=['GET', 'POST'])
def decanat_cours_masse():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    
    if request.method == 'POST':
        try:
            departement = request.form['departement']
            promotion = request.form['promotion']
            cours_liste = request.form['cours_liste']
            
            lignes = cours_liste.strip().split('\n')
            cours_ajoutes = 0
            
            import re
            
            for ligne in lignes:
                ligne = ligne.strip()
                if not ligne:
                    continue
                
                nom_cours = re.sub(r'^\d+[\.\)]\s*', '', ligne)
                nom_cours = re.sub(r'^[\-\*]\s*', '', nom_cours)
                nom_cours = nom_cours.strip()
                
                if not nom_cours:
                    continue
                
                existing_count = Cours.query.filter_by(
                    nom=nom_cours,
                    departement=departement,
                    promotion=promotion
                ).count()
                
                if existing_count > 0:
                    continue
                
                code_base = ''.join([c[0].upper() for c in nom_cours.split()[:3] if c])[:3]
                if not code_base:
                    code_base = 'CRS'
                
                counter = 1
                code = f"{code_base}{counter:03d}"
                while Cours.query.filter_by(code=code).first():
                    counter += 1
                    code = f"{code_base}{counter:03d}"
                
                cours = Cours(
                    code=code,
                    nom=nom_cours,
                    departement=departement,
                    promotion=promotion
                )
                
                db.session.add(cours)
                cours_ajoutes += 1
            
            db.session.commit()
            
            if cours_ajoutes > 0:
                flash(f'{cours_ajoutes} cours ajoutés avec succès!', 'success')
            else:
                flash('Aucun nouveau cours à ajouter (tous existent déjà)', 'warning')
            
            return redirect(url_for('decanat_cours'))
            
        except Exception as e:
            db.session.rollback()
            flash(f'Erreur lors de l\'ajout des cours: {str(e)}', 'error')
    
    return render_template('decanat_cours_masse.html',
                         departements=DEPARTEMENTS, promotions=PROMOTIONS)

@app.route('/decanat/horaires', methods=['GET', 'POST'])
def decanat_horaires():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    
    if request.method == 'POST':
        try:
            departement = request.form['departement']
            promotion = request.form['promotion']
            type_horaire = request.form['type_horaire']
            file = request.files['fichier']
            
            if file and allowed_horaire_file(file.filename):
                filename = secure_filename(f"{departement}_{promotion}_{type_horaire}_{file.filename}")
                _save_uploaded_asset(file, 'horaires', filename)
                
                existing = Horaire.query.filter_by(
                    departement=departement,
                    promotion=promotion,
                    type_horaire=type_horaire
                ).first()
                
                if existing:
                    old_file = os.path.join(app.config['HORAIRES_FOLDER'], existing.fichier)
                    if os.path.exists(old_file):
                        os.remove(old_file)
                    _delete_file_asset('horaires', existing.fichier)
                    existing.fichier = filename
                    existing.date_publication = now_cat()
                else:
                    horaire = Horaire(
                        departement=departement,
                        promotion=promotion,
                        type_horaire=type_horaire,
                        fichier=filename
                    )
                    db.session.add(horaire)
                
                db.session.commit()
                flash(f'Horaire {type_horaire} publié avec succès!', 'success')
            else:
                flash('Format de fichier non autorisé. Utilisez PDF, PNG, JPG ou JPEG', 'error')
        except Exception as e:
            db.session.rollback()
            flash(f'Erreur lors de la publication: {str(e)}', 'error')
        
        return redirect(url_for('decanat_horaires'))
    
    horaires = Horaire.query.order_by(Horaire.date_publication.desc()).all()
    return render_template('decanat_horaires.html', horaires=horaires,
                         departements=DEPARTEMENTS, promotions=PROMOTIONS)

@app.route('/decanat/horaires/supprimer/<int:horaire_id>', methods=['POST'])
def supprimer_horaire(horaire_id):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    
    try:
        horaire = Horaire.query.get_or_404(horaire_id)
        fichier_path = os.path.join(app.config['HORAIRES_FOLDER'], horaire.fichier)
        if os.path.exists(fichier_path):
            os.remove(fichier_path)
        _delete_file_asset('horaires', horaire.fichier)
        db.session.delete(horaire)
        db.session.commit()
        flash('Horaire supprimé avec succès', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erreur lors de la suppression: {str(e)}', 'error')
    
    return redirect(url_for('decanat_horaires'))

@app.route('/decanat/actualites', methods=['GET', 'POST'])
def decanat_actualites():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    
    if request.method == 'POST':
        try:
            titre = request.form['titre'].strip()
            description = request.form['description'].strip()
            type_publication = request.form.get('type_publication', 'actualite')
            if type_publication not in ('actualite', 'communique'):
                type_publication = 'actualite'
            actualite_id = request.form.get('actualite_id', type='int')
            file = request.files.get('image')

            actualite = Actualite.query.get(actualite_id) if actualite_id else None
            if actualite_id and not actualite:
                flash('Publication introuvable.', 'error')
                return redirect(url_for('decanat_actualites'))

            image_filename = actualite.image if actualite else None
            if file and file.filename and allowed_file(file.filename):
                image_filename = secure_filename(f"{now_cat().strftime('%Y%m%d_%H%M%S')}_{file.filename}")
                _save_uploaded_asset(file, 'actualites', image_filename)
                if actualite and actualite.image and actualite.image != image_filename:
                    old_path = _asset_local_path('actualites', actualite.image)
                    if old_path and os.path.exists(old_path):
                        os.remove(old_path)
                    _delete_file_asset('actualites', actualite.image)

            if actualite is None:
                actualite = Actualite(
                    titre=titre,
                    description=description,
                    image=image_filename,
                    type_publication=type_publication,
                )
                db.session.add(actualite)
                message = 'Publication créée avec succès.'
            else:
                actualite.titre = titre
                actualite.description = description
                actualite.type_publication = type_publication
                actualite.image = image_filename
                actualite.date_publication = now_cat()
                message = 'Publication modifiée avec succès.'
            db.session.commit()
            flash(message, 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Erreur lors de la publication: {str(e)}', 'error')
        
        return redirect(url_for('decanat_actualites'))
    
    actualites = Actualite.query.order_by(
        Actualite.epingle.desc(), Actualite.date_publication.desc()
    ).all()
    edit_id = request.args.get('modifier', type=int)
    publication_a_modifier = Actualite.query.get(edit_id) if edit_id else None
    return render_template(
        'decanat_actualites.html',
        actualites=actualites,
        publication_a_modifier=publication_a_modifier,
    )


@app.route('/decanat/actualites/<int:actualite_id>/epingle', methods=['POST'])
def epingler_actualite(actualite_id):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    try:
        actualite = Actualite.query.get_or_404(actualite_id)
        if not actualite.epingle:
            nb_epinglees = Actualite.query.filter_by(epingle=True).count()
            if nb_epinglees >= 4:
                flash('Vous pouvez épingler au maximum 4 publications.', 'warning')
                return redirect(url_for('decanat_actualites'))
        actualite.epingle = not actualite.epingle
        db.session.commit()
        flash(
            'Publication épinglée en tête.' if actualite.epingle
            else 'Publication retirée des publications épinglées.',
            'success',
        )
    except Exception as exc:
        db.session.rollback()
        flash(f'Erreur lors de la modification de l’épinglage : {exc}', 'error')
    return redirect(url_for('decanat_actualites'))


@app.route('/decanat/actualites/<int:actualite_id>/visibilite', methods=['POST'])
def basculer_visibilite_actualite(actualite_id):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    try:
        actualite = Actualite.query.get_or_404(actualite_id)
        actualite.publie = not actualite.publie
        db.session.commit()
        flash(
            'Publication masquée du site.' if not actualite.publie
            else 'Publication visible sur le site.',
            'success',
        )
    except Exception as exc:
        db.session.rollback()
        flash(f'Erreur lors de la modification de la visibilité : {exc}', 'error')
    return redirect(url_for('decanat_actualites'))

@app.route('/decanat/actualites/supprimer/<int:actualite_id>', methods=['POST'])
def supprimer_actualite(actualite_id):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    
    try:
        actualite = Actualite.query.get_or_404(actualite_id)
        if actualite.image:
            image_path = os.path.join(app.config['ACTUALITES_FOLDER'], actualite.image)
            if os.path.exists(image_path):
                os.remove(image_path)
            _delete_file_asset('actualites', actualite.image)
        db.session.delete(actualite)
        db.session.commit()
        flash('Actualité supprimée avec succès', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erreur lors de la suppression: {str(e)}', 'error')
    
    return redirect(url_for('decanat_actualites'))

@app.route('/decanat/migration-fichiers', methods=['GET', 'POST'])
def decanat_migration_fichiers():
    """Page de migration des fichiers : statut des fichiers référencés + import ZIP depuis dev."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    import_log = []
    import_errors = []

    if request.method == 'POST':
        action = request.form.get('action', 'import_zip')

        if action == 'import_zip':
            fichier_zip = request.files.get('fichier_zip')
            if not fichier_zip or fichier_zip.filename == '':
                flash('Aucun fichier ZIP sélectionné.', 'error')
                return redirect(request.url)

            if not fichier_zip.filename.lower().endswith('.zip'):
                flash('Seul le format ZIP est accepté.', 'error')
                return redirect(request.url)

            try:
                import tempfile
                zip_bytes = BytesIO(fichier_zip.read())
                with zipfile.ZipFile(zip_bytes, 'r') as zf:
                    # Répertoires autorisés → dossier de destination
                    ALLOWED_DIRS = {
                        'uploads':    app.config['UPLOAD_FOLDER'],
                        'qrcodes':    app.config['QRCODE_FOLDER'],
                        'horaires':   app.config['HORAIRES_FOLDER'],
                        'actualites': app.config['ACTUALITES_FOLDER'],
                    }
                    ALLOWED_EXT = {'.png', '.jpg', '.jpeg', '.pdf'}
                    skipped = 0
                    for member in zf.namelist():
                        # Ignorer répertoires et fichiers cachés
                        if member.endswith('/') or os.path.basename(member).startswith('.'):
                            continue
                        parts = member.split('/', 1)
                        if len(parts) != 2:
                            skipped += 1
                            continue
                        top_dir, rel_path = parts
                        if top_dir not in ALLOWED_DIRS:
                            skipped += 1
                            continue
                        # Vérifier l'extension
                        _, ext = os.path.splitext(rel_path.lower())
                        if ext not in ALLOWED_EXT:
                            skipped += 1
                            continue
                        dest_folder = ALLOWED_DIRS[top_dir]
                        dest_path = os.path.join(dest_folder, os.path.basename(rel_path))
                        # Sécurité : éviter path traversal
                        if not os.path.abspath(dest_path).startswith(os.path.abspath(dest_folder)):
                            import_errors.append(f'Chemin dangereux ignoré : {member}')
                            continue
                        os.makedirs(dest_folder, exist_ok=True)
                        with zf.open(member) as src, open(dest_path, 'wb') as dst:
                            file_bytes = src.read()
                            dst.write(file_bytes)
                        category = {
                            'uploads': 'uploads',
                            'qrcodes': 'qrcodes',
                            'horaires': 'horaires',
                            'actualites': 'actualites',
                        }[top_dir]
                        _store_file_asset(category, os.path.basename(rel_path), file_bytes)
                        import_log.append(f'{top_dir}/{os.path.basename(rel_path)}')

                msg = f'{len(import_log)} fichier(s) importé(s)'
                if skipped:
                    msg += f', {skipped} ignoré(s) (répertoire non reconnu ou extension non autorisée)'
                if import_errors:
                    msg += f', {len(import_errors)} erreur(s)'
                db.session.commit()
                flash(msg, 'success' if not import_errors else 'error')
            except zipfile.BadZipFile:
                flash('Le fichier n\'est pas un ZIP valide.', 'error')
            except Exception as e:
                flash(f'Erreur lors de l\'import : {str(e)}', 'error')

            return redirect(url_for('decanat_migration_fichiers'))

    # ── Calcul du statut des fichiers référencés en DB ──
    def _check_dir(folder, filenames):
        present, missing = [], []
        for f in filenames:
            if f:
                path = os.path.join(folder, f)
                (present if os.path.exists(path) else missing).append(f)
        return present, missing

    photos_present, photos_missing = _check_dir(
        app.config['UPLOAD_FOLDER'],
        [e.photo for e in Etudiant.query.with_entities(Etudiant.photo).all() if e.photo]
    )
    qrcodes_present, qrcodes_missing = _check_dir(
        app.config['QRCODE_FOLDER'],
        [e.qrcode_path for e in Etudiant.query.with_entities(Etudiant.qrcode_path).all() if e.qrcode_path]
    )
    horaires_present, horaires_missing = _check_dir(
        app.config['HORAIRES_FOLDER'],
        [h.fichier for h in Horaire.query.with_entities(Horaire.fichier).all() if h.fichier]
    )
    actualites_present, actualites_missing = _check_dir(
        app.config['ACTUALITES_FOLDER'],
        [a.image for a in Actualite.query.with_entities(Actualite.image).all() if a.image]
    )

    statut = {
        'photos':    {'present': len(photos_present),    'missing': photos_missing},
        'qrcodes':   {'present': len(qrcodes_present),   'missing': qrcodes_missing},
        'horaires':  {'present': len(horaires_present),  'missing': horaires_missing},
        'actualites':{'present': len(actualites_present),'missing': actualites_missing},
    }
    total_missing = sum(len(v['missing']) for v in statut.values())

    return render_template(
        'decanat_migration_fichiers.html',
        statut=statut,
        total_missing=total_missing,
        import_log=import_log,
        import_errors=import_errors,
    )


@app.route('/decanat/page-apropos', methods=['GET', 'POST'])
def decanat_page_apropos():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    
    if request.method == 'POST':
        try:
            page_content = PageContent.query.filter_by(page_name='a_propos').first()
            if not page_content:
                page_content = PageContent(page_name='a_propos', content_json='{}')
            
            content_data = {
                'mission_texte': request.form.get('mission_texte', ''),
                'lmd_texte': request.form.get('lmd_texte', ''),
                'cycle2_texte': request.form.get('cycle2_texte', ''),
                'cycle3_texte': request.form.get('cycle3_texte', ''),
                'conclusion_texte': request.form.get('conclusion_texte', ''),
                'adresse': request.form.get('adresse', ''),
                'telephone1': request.form.get('telephone1', ''),
                'telephone2': request.form.get('telephone2', ''),
                'motto': request.form.get('motto', ''),
                'doyenne': request.form.get('doyenne', ''),
                'vice_doyen_enseignement': request.form.get('vice_doyen_enseignement', ''),
                'vice_doyen_recherche': request.form.get('vice_doyen_recherche', ''),
                'vice_doyen_tutorat': request.form.get('vice_doyen_tutorat', ''),
                'secretaire_academique': request.form.get('secretaire_academique', ''),
                'secretaire_administratif': request.form.get('secretaire_administratif', ''),
                'appariteur': request.form.get('appariteur', ''),
                'comptable': request.form.get('comptable', ''),
                'intendant': request.form.get('intendant', ''),
                'chef_anthropologie': request.form.get('chef_anthropologie', ''),
                'chef_sociologie': request.form.get('chef_sociologie', ''),
                'chef_relations_int': request.form.get('chef_relations_int', ''),
                'chef_spa': request.form.get('chef_spa', ''),
                'directeur_cespac': request.form.get('directeur_cespac', ''),
                'directeur_cee': request.form.get('directeur_cee', '')
            }
            
            file = request.files.get('image_facade')
            if file and file.filename and allowed_file(file.filename):
                if page_content.image_principale:
                    old_path = os.path.join('static', page_content.image_principale)
                    if os.path.exists(old_path):
                        os.remove(old_path)
                
                filename = secure_filename(f"faculte-facade_{now_cat().strftime('%Y%m%d_%H%M%S')}.jpg")
                _save_uploaded_asset(file, 'root', filename)
                page_content.image_principale = filename
            
            page_content.content_json = json.dumps(content_data)
            page_content.date_modification = now_cat()
            
            db.session.add(page_content)
            db.session.commit()
            flash('Page À PROPOS mise à jour avec succès!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Erreur lors de la mise à jour: {str(e)}', 'error')
        
        return redirect(url_for('decanat_page_apropos'))
    
    page_content = PageContent.query.filter_by(page_name='a_propos').first()
    content_data = json.loads(page_content.content_json) if page_content and page_content.content_json else {}
    image_facade = (page_content.image_principale
                    if page_content and page_content.image_principale else None)
    
    return render_template('decanat_page_apropos.html', content_data=content_data, image_facade=image_facade)

@app.route('/decanat/page-departements', methods=['GET', 'POST'])
def decanat_page_departements():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    
    if request.method == 'POST':
        try:
            page_content = PageContent.query.filter_by(page_name='departements').first()
            if not page_content:
                page_content = PageContent(page_name='departements', content_json='{}')
            
            content_data = {
                'sociologie_description': request.form.get('sociologie_description', ''),
                'sociologie_formation': request.form.get('sociologie_formation', ''),
                'sociologie_debouches': request.form.get('sociologie_debouches', ''),
                'anthropologie_description': request.form.get('anthropologie_description', ''),
                'anthropologie_formation': request.form.get('anthropologie_formation', ''),
                'anthropologie_debouches': request.form.get('anthropologie_debouches', ''),
                'relations_int_description': request.form.get('relations_int_description', ''),
                'relations_int_formation': request.form.get('relations_int_formation', ''),
                'relations_int_debouches': request.form.get('relations_int_debouches', ''),
                'spa_description': request.form.get('spa_description', ''),
                'spa_formation': request.form.get('spa_formation', ''),
                'spa_debouches': request.form.get('spa_debouches', '')
            }
            
            page_content.content_json = json.dumps(content_data)
            page_content.date_modification = now_cat()
            
            db.session.add(page_content)
            db.session.commit()
            flash('Page Départements mise à jour avec succès!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Erreur lors de la mise à jour: {str(e)}', 'error')
        
        return redirect(url_for('decanat_page_departements'))
    
    page_content = PageContent.query.filter_by(page_name='departements').first()
    content_data = json.loads(page_content.content_json) if page_content and page_content.content_json else {}
    
    return render_template('decanat_page_departements.html', content_data=content_data)

@app.route('/horaires')
def consulter_horaires():
    departement = request.args.get('departement', '')
    promotion = request.args.get('promotion', '')
    
    query = Horaire.query
    if departement:
        query = query.filter_by(departement=departement)
    if promotion:
        query = query.filter_by(promotion=promotion)
    
    horaires = query.order_by(Horaire.date_publication.desc()).all()
    
    return render_template('consulter_horaires.html', horaires=horaires,
                         departements=DEPARTEMENTS, promotions=PROMOTIONS,
                         selected_dept=departement, selected_promo=promotion)

@app.route('/scanner')
def scanner_public():
    cours_list = Cours.query.all()
    return render_template('scanner_public.html', cours_list=cours_list)

@app.route('/presence/scanner')
def presence_scanner():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    
    cours_list = Cours.query.all()
    return render_template('presence_scanner.html', cours_list=cours_list)

@app.route('/presence/enregistrer', methods=['POST'])
def enregistrer_presence():
    try:
        data = request.get_json()
        qr_data = json.loads(data['qr_data'])
        cours_id = data.get('cours_id')
        
        matricule = qr_data['matricule']
        type_personne = qr_data['type']
        
        une_heure_avant = now_cat() - timedelta(hours=1)
        
        if type_personne == 'etudiant':
            etudiant = Etudiant.query.filter_by(matricule=matricule).first()
            if not etudiant:
                return jsonify({'success': False, 'message': 'Étudiant non trouvé'})
            
            presence_existante = Presence.query.filter_by(
                etudiant_id=etudiant.id,
                cours_id=cours_id,
                type_presence='etudiant'
            ).filter(Presence.heure_entree >= une_heure_avant).first()
            
            if presence_existante:
                temps_ecoule = now_cat() - presence_existante.heure_entree
                minutes_restantes = 60 - int(temps_ecoule.total_seconds() / 60)
                return jsonify({
                    'success': False, 
                    'message': f'Présence déjà enregistrée. Veuillez attendre {minutes_restantes} minutes avant de scanner à nouveau.'
                })
            
            presence = Presence(
                etudiant_id=etudiant.id,
                cours_id=cours_id,
                type_presence='etudiant'
            )
            db.session.add(presence)
            db.session.commit()
            
            return jsonify({
                'success': True, 
                'message': f'Présence enregistrée: {etudiant.nom} {etudiant.postnom} {etudiant.prenom}',
                'type': 'etudiant',
                'nom_complet': f'{etudiant.nom} {etudiant.postnom} {etudiant.prenom}'
            })
        
        elif type_personne == 'professeur':
            prof = Professeur.query.filter_by(matricule=matricule).first()
            if not prof:
                return jsonify({'success': False, 'message': 'Professeur non trouvé'})
            
            presence_existante = Presence.query.filter_by(
                professeur_id=prof.id,
                cours_id=cours_id,
                type_presence='professeur'
            ).filter(Presence.heure_entree >= une_heure_avant).first()
            
            if presence_existante:
                temps_ecoule = now_cat() - presence_existante.heure_entree
                minutes_restantes = 60 - int(temps_ecoule.total_seconds() / 60)
                return jsonify({
                    'success': False, 
                    'message': f'Présence déjà enregistrée. Veuillez attendre {minutes_restantes} minutes avant de scanner à nouveau.'
                })
            
            presence = Presence(
                professeur_id=prof.id,
                cours_id=cours_id,
                type_presence='professeur'
            )
            db.session.add(presence)
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': f'Présence enregistrée: Prof. {prof.nom} {prof.postnom}',
                'type': 'professeur',
                'nom_complet': f'{prof.nom} {prof.postnom} {prof.prenom}'
            })
        
    except Exception as e:
        return jsonify({'success': False, 'message': f'Erreur: {str(e)}'})

@app.route('/decanat/presences')
def decanat_presences():
    if not session.get('decanat_logged_in') and not session.get('prof_logged_in'):
        return redirect(url_for('prof_login'))
    
    cours_id = request.args.get('cours_id', type=int)
    date_str = request.args.get('date')
    
    query = Presence.query
    
    if cours_id:
        query = query.filter_by(cours_id=cours_id)
    if date_str:
        date_obj = datetime.strptime(date_str, '%Y-%m-%d').date()
        query = query.filter_by(date=date_obj)
    
    presences = query.order_by(Presence.heure_entree.desc()).all()
    cours_list = Cours.query.all()
    
    return render_template('decanat_presences.html', presences=presences, 
                         cours_list=cours_list, selected_cours=cours_id)

@app.route('/export/presences/<format>')
def export_presences(format):
    if not session.get('decanat_logged_in') and not session.get('prof_logged_in'):
        return redirect(url_for('prof_login'))
    
    cours_id = request.args.get('cours_id', type=int)
    plage = request.args.get('plage', '2jours')
    jours_personnalises = request.args.get('jours', type=int)
    date_debut_str = request.args.get('date_debut')
    date_fin_str = request.args.get('date_fin')
    
    if not cours_id:
        flash('Veuillez sélectionner un cours pour exporter les présences', 'error')
        return redirect(url_for('decanat_presences'))
    
    cours = Cours.query.get(cours_id)
    if not cours:
        flash('Cours introuvable', 'error')
        return redirect(url_for('decanat_presences'))
    
    if plage == 'dates_personnalisees':
        if not date_debut_str or not date_fin_str:
            flash('Veuillez sélectionner les deux dates', 'error')
            return redirect(url_for('decanat_presences'))
        
        try:
            date_debut = datetime.strptime(date_debut_str, '%Y-%m-%d')
            date_fin = datetime.strptime(date_fin_str, '%Y-%m-%d')
            
            if date_debut > date_fin:
                flash('La date de début doit être antérieure ou égale à la date de fin', 'error')
                return redirect(url_for('decanat_presences'))
            
            date_fin = date_fin.replace(hour=23, minute=59, second=59)
            
            presences = Presence.query.filter(
                Presence.cours_id == cours_id,
                Presence.heure_entree >= date_debut,
                Presence.heure_entree <= date_fin
            ).order_by(Presence.date.desc(), Presence.heure_entree.desc()).all()
            
            periode_desc = f"Du {date_debut.strftime('%d/%m/%Y')} au {date_fin.strftime('%d/%m/%Y')}"
            
        except ValueError:
            flash('Format de date invalide', 'error')
            return redirect(url_for('decanat_presences'))
    else:
        date_limite = now_cat()
        if plage == 'jours_personnalise':
            if jours_personnalises and jours_personnalises > 0:
                date_limite = date_limite - timedelta(days=jours_personnalises)
                periode_desc = f'{jours_personnalises} dernier{"s" if jours_personnalises > 1 else ""} jour{"s" if jours_personnalises > 1 else ""}'
            else:
                flash('Nombre de jours invalide', 'error')
                return redirect(url_for('decanat_presences'))
        elif plage == '2jours':
            date_limite = date_limite - timedelta(days=2)
            periode_desc = '2 derniers jours'
        elif plage == 'semaine':
            date_limite = date_limite - timedelta(days=7)
            periode_desc = 'Dernière semaine'
        elif plage == 'mois':
            date_limite = date_limite - timedelta(days=30)
            periode_desc = 'Dernier mois'
        else:
            periode_desc = plage
        
        presences = Presence.query.filter(
            Presence.cours_id == cours_id,
            Presence.heure_entree >= date_limite
        ).order_by(Presence.date.desc(), Presence.heure_entree.desc()).all()
    
    if format == 'excel':
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Présences"
        
        ws.append([f'Liste de présence - {cours.nom}'])
        ws.append([f'Période: {periode_desc}'])
        ws.append([f'Total: {len(presences)} présence(s)'])
        ws.append([])
        
        headers = ['Date', 'Heure', 'Type', 'CODE-ID FAC', 'Nom', 'Postnom', 'Prénom', 'Département']
        ws.append(headers)
        
        for presence in presences:
            if presence.etudiant:
                ws.append([
                    presence.date.strftime('%d/%m/%Y'),
                    presence.heure_entree.strftime('%H:%M:%S'),
                    'Étudiant',
                    presence.etudiant.matricule,
                    presence.etudiant.nom,
                    presence.etudiant.postnom,
                    presence.etudiant.prenom,
                    presence.etudiant.departement
                ])
            elif presence.professeur:
                ws.append([
                    presence.date.strftime('%d/%m/%Y'),
                    presence.heure_entree.strftime('%H:%M:%S'),
                    'Professeur',
                    presence.professeur.matricule,
                    presence.professeur.nom,
                    presence.professeur.postnom,
                    presence.professeur.prenom,
                    presence.professeur.departement or '-'
                ])
        
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        
        filename = f'presences_{cours.code}_{plage}.xlsx'.replace(' ', '_')
        return send_file(output, download_name=filename, as_attachment=True,
                        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    
    elif format == 'pdf':
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors
        
        output = BytesIO()
        doc = SimpleDocTemplate(output, pagesize=landscape(A4))
        elements = []
        
        styles = getSampleStyleSheet()
        title = Paragraph(f"Liste de présence - {cours.nom}", styles['Title'])
        subtitle = Paragraph(f"Période: {periode_desc} | Total: {len(presences)} présence(s)", 
                           styles['Normal'])
        elements.append(title)
        elements.append(subtitle)
        elements.append(Spacer(1, 20))
        
        data = [['Date', 'Heure', 'Type', 'CODE-ID', 'Nom', 'Postnom', 'Prénom', 'Dépt.']]
        
        for presence in presences:
            if presence.etudiant:
                data.append([
                    presence.date.strftime('%d/%m/%Y'),
                    presence.heure_entree.strftime('%H:%M'),
                    'Étud.',
                    presence.etudiant.matricule,
                    presence.etudiant.nom,
                    presence.etudiant.postnom,
                    presence.etudiant.prenom,
                    presence.etudiant.departement[:10]
                ])
            elif presence.professeur:
                data.append([
                    presence.date.strftime('%d/%m/%Y'),
                    presence.heure_entree.strftime('%H:%M'),
                    'Prof.',
                    presence.professeur.matricule,
                    presence.professeur.nom,
                    presence.professeur.postnom,
                    presence.professeur.prenom,
                    (presence.professeur.departement or '-')[:10]
                ])
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('FONTSIZE', (0, 1), (-1, -1), 8),
        ]))
        
        elements.append(table)
        doc.build(elements)
        
        output.seek(0)
        filename = f'presences_{cours.code}_{plage}.pdf'.replace(' ', '_')
        return send_file(output, download_name=filename, as_attachment=True,
                        mimetype='application/pdf')
    
    else:
        flash(f'Format non supporté: {format}', 'error')
        return redirect(url_for('decanat_presences'))

@app.route('/export/etudiants/<format>')
def export_etudiants(format):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    
    departement = request.args.get('departement', '')
    promotion = request.args.get('promotion', '')
    
    query = Etudiant.query
    if departement:
        query = query.filter_by(departement=departement)
    if promotion:
        query = query.filter_by(promotion=promotion)
    
    etudiants = query.order_by(Etudiant.nom).all()
    
    if format == 'excel':
        from openpyxl import Workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Etudiants"
        
        headers = ['Matricule', 'Nom', 'Postnom', 'Prénom', 'Sexe', 'Téléphone', 
                  'Département', 'Promotion', 'Date Inscription']
        ws.append(headers)
        
        for etudiant in etudiants:
            ws.append([
                etudiant.matricule,
                etudiant.nom,
                etudiant.postnom,
                etudiant.prenom,
                etudiant.sexe,
                etudiant.telephone,
                etudiant.departement,
                etudiant.promotion,
                etudiant.date_inscription.strftime('%Y-%m-%d %H:%M')
            ])
        
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        
        filename = f'etudiants_{departement}_{promotion}.xlsx'.replace(' ', '_')
        return send_file(output, download_name=filename, as_attachment=True,
                        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    
    elif format == 'pdf':
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors
        
        output = BytesIO()
        doc = SimpleDocTemplate(output, pagesize=landscape(A4))
        elements = []
        
        styles = getSampleStyleSheet()
        title = Paragraph(f"Liste des Étudiants - {departement or 'Tous départements'} - {promotion or 'Toutes promotions'}", 
                         styles['Title'])
        elements.append(title)
        
        data = [['Matricule', 'Nom', 'Postnom', 'Prénom', 'Sexe', 'Tél.', 'Dépt.', 'Promo']]
        
        for etudiant in etudiants:
            data.append([
                etudiant.matricule,
                etudiant.nom,
                etudiant.postnom,
                etudiant.prenom,
                etudiant.sexe,
                etudiant.telephone,
                etudiant.departement,
                etudiant.promotion
            ])
        
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        elements.append(table)
        doc.build(elements)
        output.seek(0)
        
        filename = f'etudiants_{departement}_{promotion}.pdf'.replace(' ', '_')
        return send_file(output, download_name=filename, as_attachment=True,
                        mimetype='application/pdf')

def reparer_fichier_word(fichier_corrompu_path, fichier_repare_path):
    """
    Répare un fichier Word corrompu en nettoyant les références NULL du fichier ZIP.
    """
    import tempfile
    import xml.etree.ElementTree as ET
    
    temp_dir = tempfile.mkdtemp()
    
    try:
        with zipfile.ZipFile(fichier_corrompu_path, 'r') as zip_ref:
            for member in zip_ref.namelist():
                if member != 'word/NULL' and not member.endswith('/'):
                    content = zip_ref.read(member)
                    output_path = os.path.join(temp_dir, member)
                    
                    output_dir = os.path.dirname(output_path)
                    if output_dir:
                        os.makedirs(output_dir, exist_ok=True)
                    
                    with open(output_path, 'wb') as f:
                        f.write(content)
        
        for root_dir, dirs, files in os.walk(temp_dir):
            for file in files:
                if file.endswith('.rels'):
                    rels_path = os.path.join(root_dir, file)
                    try:
                        tree = ET.parse(rels_path)
                        root = tree.getroot()
                        
                        namespace = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}
                        
                        rels_to_remove = []
                        for rel in root.findall('.//r:Relationship', namespace):
                            target = rel.get('Target')
                            if target and ('NULL' in target or target == 'NULL'):
                                rels_to_remove.append(rel)
                        
                        for rel in rels_to_remove:
                            root.remove(rel)
                        
                        tree.write(rels_path, encoding='utf-8', xml_declaration=True)
                    except:
                        pass
        
        with zipfile.ZipFile(fichier_repare_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root_dir, dirs, files in os.walk(temp_dir):
                for file in files:
                    file_path = os.path.join(root_dir, file)
                    arcname = os.path.relpath(file_path, temp_dir)
                    zipf.write(file_path, arcname)
        
        shutil.rmtree(temp_dir)
        
        try:
            Document(fichier_repare_path)
        except Exception as e:
            raise Exception(f"Fichier réparé invalide: {str(e)}")
        
        return True
        
    except Exception as e:
        shutil.rmtree(temp_dir, ignore_errors=True)
        raise Exception(f"Impossible de réparer le fichier Word: {str(e)}")

@app.route('/decanat/generateur-lettres', methods=['GET', 'POST'])
def decanat_generateur_lettres():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    
    if request.method == 'POST':
        fichier_excel = request.files.get('fichier_excel')
        modele_word = request.files.get('modele_word')

        if not fichier_excel:
            flash('Veuillez fournir le fichier Excel.', 'error')
            return redirect(url_for('decanat_generateur_lettres'))

        if not fichier_excel.filename.endswith(('.xlsx', '.xls')):
            flash('Le fichier Excel doit être au format .xlsx ou .xls', 'error')
            return redirect(url_for('decanat_generateur_lettres'))

        if modele_word and modele_word.filename and not modele_word.filename.endswith('.docx'):
            flash('Le modèle Word doit être au format .docx', 'error')
            return redirect(url_for('decanat_generateur_lettres'))

        os.makedirs(app.config['GENERATEUR_GENERES'], exist_ok=True)

        for file in os.listdir(app.config['GENERATEUR_GENERES']):
            file_path = os.path.join(app.config['GENERATEUR_GENERES'], file)
            try:
                os.remove(file_path)
            except:
                pass

        excel_path = os.path.join(app.config['GENERATEUR_UPLOADS'], 'base.xlsx')
        os.makedirs(app.config['GENERATEUR_UPLOADS'], exist_ok=True)
        fichier_excel.save(excel_path)

        # Save the new template to the database when one was actually uploaded.
        if modele_word and modele_word.filename:
            _save_modele_docx(modele_word.read())

        # Read the template from the persistent store (DB).
        modele_bytes = _get_modele_docx()

        # Early check: refuse to proceed if the template is still missing.
        if not modele_bytes:
            flash(
                "Le modèle de lettre n'a pas encore été configuré — "
                "veuillez téléverser le fichier modele.docx avant de générer les lettres.",
                'error'
            )
            return redirect(url_for('decanat_generateur_lettres'))

        # Write to a temp file so the generation function can open it normally.
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.write(modele_bytes)
            modele_path = tmp.name

        try:
            Document(modele_path)
        except KeyError as e:
            if 'word/NULL' in str(e):
                flash('Fichier Word corrompu détecté. Réparation automatique en cours...', 'warning')
                modele_backup_path = modele_path + '.bak'
                os.rename(modele_path, modele_backup_path)
                reparer_fichier_word(modele_backup_path, modele_path)
                # Re-save the repaired bytes to the DB so the fix persists.
                try:
                    with open(modele_path, 'rb') as _f:
                        _save_modele_docx(_f.read())
                except Exception:
                    pass
                try:
                    os.remove(modele_backup_path)
                except:
                    pass
            else:
                raise

        try:
            lettres_info = generer_lettres_charge_horaire(excel_path, modele_path)
            
            session['lettres_generees'] = lettres_info
            flash(f'{len(lettres_info)} lettre(s) générée(s) avec succès !', 'success')
            
            return render_template('decanat_generateur_lettres.html',
                                 lettres_generees=lettres_info,
                                 nombre_lettres=len(lettres_info),
                                 nombre_professeurs=len(lettres_info))
        
        except Exception as e:
            flash(f'Erreur lors de la génération des lettres: {str(e)}', 'error')
            return redirect(url_for('decanat_generateur_lettres'))
        finally:
            try:
                os.remove(modele_path)
            except Exception:
                pass
    
    template_manquant = _get_modele_docx() is None
    return render_template('decanat_generateur_lettres.html',
                           template_manquant=template_manquant)

def set_cell_border(cell, **kwargs):
    """
    Définit les bordures d'une cellule de tableau Word.
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            element = OxmlElement(tag)
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '12')  # Taille de bordure
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')  # Noir
            tcBorders.append(element)
    
    tcPr.append(tcBorders)

def set_cell_background(cell, fill):
    """
    Définit la couleur de fond d'une cellule de tableau Word.
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcVAlign = OxmlElement('w:shd')
    tcVAlign.set(qn('w:fill'), fill)
    tcPr.append(tcVAlign)

def generer_lettres_charge_horaire(excel_path, modele_path=None):
    # Guard: the Word template must exist before we do any heavy processing.
    if not modele_path or not os.path.exists(modele_path):
        raise FileNotFoundError(
            "Le modèle de lettre (modele.docx) est introuvable. "
            "Veuillez téléverser le modèle Word avant de générer les lettres."
        )

    wb = load_workbook(excel_path, data_only=True)
    ws = wb.active
    
    titulaires_data = {}
    
    for row in ws.iter_rows(min_row=1, values_only=True):
        if not row or len(row) < 10:
            continue
        
        titulaire = str(row[9]).strip() if row[9] and row[9] != 'None' else ""
        cours = str(row[1]).strip() if row[1] else ""
        sem = str(row[2]).strip() if row[2] else ""
        promotion = str(row[3]).strip() if row[3] else ""
        
        cmi = row[4] if row[4] and isinstance(row[4], (int, float)) else 0
        td = row[5] if row[5] and isinstance(row[5], (int, float)) else 0
        tp = row[6] if row[6] and isinstance(row[6], (int, float)) else 0
        heures_total = int(cmi + td + tp)
        
        credits = str(row[8]).strip() if row[8] else "0"
        
        if 'titulaire' in titulaire.lower() or 'professeur' in titulaire.lower() or 'cours' in cours.lower() or 'N°' in titulaire:
            continue
        
        if not titulaire or not cours or titulaire == 'None':
            continue
        
        if titulaire not in titulaires_data:
            titulaires_data[titulaire] = []
        
        titulaires_data[titulaire].append({
            'cours': cours,
            'sem': sem,
            'promotion': promotion,
            'cmi': int(cmi),
            'td': int(td),
            'tp': int(tp),
            'total': heures_total,
            'credits': credits
        })
    
    lettres_info = []
    
    for titulaire, cours_list in titulaires_data.items():
        doc = Document(modele_path)
        
        from datetime import datetime
        date_actuelle = now_cat().strftime('%d/%m/%Y')
        
        for paragraph in doc.paragraphs:
            if '…………….' in paragraph.text or '………………' in paragraph.text:
                if 'Lubumbashi' in paragraph.text:
                    paragraph.text = paragraph.text.replace('……………….', date_actuelle)
                    paragraph.text = paragraph.text.replace('……………..', date_actuelle)
                    paragraph.text = paragraph.text.replace('…………….', date_actuelle)
                    paragraph.text = paragraph.text.replace('……………', date_actuelle)
                    paragraph.text = paragraph.text.replace('…………', date_actuelle)
                else:
                    paragraph.text = paragraph.text.replace('……………………………………….', titulaire)
                    paragraph.text = paragraph.text.replace('………………………………………', titulaire)
                    paragraph.text = paragraph.text.replace('……………………………………', titulaire)
        
        if len(doc.tables) > 0:
            old_table = doc.tables[0]
            old_table_element = old_table._element
            old_table_parent = old_table_element.getparent()
            
            new_table = doc.add_table(rows=1, cols=9)
            
            try:
                new_table.style = 'Table Grid'
            except:
                pass
            
            new_table_element = new_table._element
            old_table_parent.replace(old_table_element, new_table_element)
            
            table = new_table
        else:
            table = doc.add_table(rows=1, cols=9)
            try:
                table.style = 'Table Grid'
            except:
                pass
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'N°'
        hdr_cells[1].text = 'Cours'
        hdr_cells[2].text = 'SEM'
        hdr_cells[3].text = 'Promotions'
        hdr_cells[4].text = 'CMI'
        hdr_cells[5].text = 'TD'
        hdr_cells[6].text = 'TP'
        hdr_cells[7].text = 'TOT'
        hdr_cells[8].text = 'CR'
        
        for i, cell in enumerate(hdr_cells):
            set_cell_background(cell, '4472C4')
            set_cell_border(cell, top=True, left=True, bottom=True, right=True)
            
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
        
        total_cmi = 0
        total_td = 0
        total_tp = 0
        total_heures = 0
        total_credits = 0
        
        for idx, cours_data in enumerate(cours_list, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = cours_data['cours']
            row_cells[2].text = cours_data['sem']
            row_cells[3].text = cours_data['promotion']
            row_cells[4].text = str(cours_data['cmi'])
            row_cells[5].text = str(cours_data['td'])
            row_cells[6].text = str(cours_data['tp'])
            row_cells[7].text = str(cours_data['total'])
            row_cells[8].text = cours_data['credits']
            
            for cell in row_cells:
                set_cell_border(cell, top=True, left=True, bottom=True, right=True)
            
            total_cmi += cours_data['cmi']
            total_td += cours_data['td']
            total_tp += cours_data['tp']
            total_heures += cours_data['total']
            
            try:
                credits_num = int(cours_data['credits']) if cours_data['credits'].isdigit() else 0
            except:
                credits_num = 0
            total_credits += credits_num
        
        row_cells = table.add_row().cells
        row_cells[0].text = ''
        row_cells[1].text = ''
        row_cells[2].text = ''
        row_cells[3].text = 'TOTAL'
        row_cells[4].text = str(total_cmi)
        row_cells[5].text = str(total_td)
        row_cells[6].text = str(total_tp)
        row_cells[7].text = str(total_heures)
        row_cells[8].text = str(total_credits)
        
        for cell in row_cells:
            set_cell_border(cell, top=True, left=True, bottom=True, right=True)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        filename_safe = titulaire.replace(' ', '_').replace('/', '_')
        output_filename = f'Lettre_{filename_safe}.docx'
        output_path = os.path.join(app.config['GENERATEUR_GENERES'], output_filename)
        
        doc.save(output_path)
        
        lettres_info.append({
            'professeur': titulaire,
            'nombre_cours': len(cours_list),
            'filename': output_filename
        })
    
    return lettres_info

@app.route('/decanat/telecharger-lettres/<format>')
def decanat_telecharger_lettres(format):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    
    generes_dir = app.config['GENERATEUR_GENERES']
    
    if format == 'zip':
        output = BytesIO()
        with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for filename in os.listdir(generes_dir):
                if filename.endswith('.docx'):
                    file_path = os.path.join(generes_dir, filename)
                    zipf.write(file_path, filename)
        
        output.seek(0)
        return send_file(output, download_name='lettres_charge_horaire.zip',
                        as_attachment=True, mimetype='application/zip')
    
    elif format == 'pdf':
        flash('La conversion en PDF nécessite LibreOffice. Veuillez télécharger les fichiers Word.', 'info')
        return redirect(url_for('decanat_generateur_lettres'))
    
    return redirect(url_for('decanat_generateur_lettres'))

@app.route('/decanat/telecharger-lettre/<filename>')
def decanat_telecharger_lettre_individuelle(filename):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    
    file_path = os.path.join(app.config['GENERATEUR_GENERES'], filename)
    
    if not os.path.exists(file_path):
        flash('Fichier introuvable', 'error')
        return redirect(url_for('decanat_generateur_lettres'))
    
    return send_file(file_path, as_attachment=True)

# ═══════════════════════════════════════════════════════════════════
#  BULLETINS ACADÉMIQUES
# ═══════════════════════════════════════════════════════════════════

def _parse_grille(file_path, ext):
    """Lit la grille de délibération XLS/XLSX et retourne (courses, students, meta).

    Gère les formats non standard : cellules fusionnées, colonnes cours décalées
    (start > col 4), lignes d'en-tête introuvables par simple scan col 0–3,
    et grilles sans colonne pondérée explicite (M1 AGIE, etc.).
    """
    import re as _re

    if ext == 'xls':
        import xlrd
        wb = xlrd.open_workbook(file_path)
        ws = wb.sheet_by_index(0)
        sheet_name = wb.sheet_names()[0]
        def gv(r, c):
            if r >= ws.nrows or c >= ws.ncols: return None
            v = ws.cell_value(r, c)
            return v if v != '' else None
        nrows = ws.nrows
    else:
        from openpyxl import load_workbook as _lw
        wb = _lw(file_path, data_only=True)
        ws = wb.active
        sheet_name = ws.title or ''
        # ── Résoudre les cellules fusionnées : propager la valeur du coin haut-gauche ──
        # Sans cela, toutes les cellules fusionnées non-pivot retournent None, ce qui
        # casse la détection des noms de cours et de l'en-tête.
        _merge_map = {}
        for mr in ws.merged_cells.ranges:
            top_val = ws.cell(row=mr.min_row, column=mr.min_col).value
            for _mr in range(mr.min_row, mr.max_row + 1):
                for _mc in range(mr.min_col, mr.max_col + 1):
                    if _mr != mr.min_row or _mc != mr.min_col:
                        _merge_map[(_mr, _mc)] = top_val
        def gv(r, c):
            key = (r + 1, c + 1)
            if key in _merge_map:
                return _merge_map[key]
            v = ws.cell(row=r + 1, column=c + 1).value
            return v
        nrows = ws.max_row

    # ── Détecter automatiquement la ligne d'en-tête (jusqu'à 15 lignes) ──
    # Accepte toute ligne qui contient ≥ 2 marqueurs parmi :
    #   N° / Matricule / NOM / SEXE / GENRE — dans les 15 premières colonnes.
    hdr_row = 0
    _NUM_KW  = ('N°', 'N', 'NO', 'NUM', 'NUMÉRO', 'NUMERO', 'N.', '#')
    _MAT_KW  = ('MATRICULE', 'MATR', 'MAT.', 'MAT')
    _NOM_KW  = ('NOM', 'NOM ET PRÉNOM', 'NOM COMPLET', 'NOMS')
    _SEX_KW  = ('SEXE', 'GENRE', 'S')
    for scan_r in range(min(15, nrows)):
        vals = [str(gv(scan_r, c) or '').strip().upper() for c in range(min(15, 150))]
        has_num = any(v in _NUM_KW for v in vals)
        has_mat = any(any(k in v for k in _MAT_KW) for v in vals)
        has_nom = any(any(v.startswith(k) for k in _NOM_KW) for v in vals)
        has_sex = any(v in _SEX_KW for v in vals)
        if sum([has_num, has_mat, has_nom, has_sex]) >= 2:
            hdr_row = scan_r
            break
        # Fallback : ligne dense (≥ 4 cellules remplies) qui contient au moins NOM
        non_empty = sum(1 for v in vals if v)
        if non_empty >= 4 and has_nom:
            hdr_row = scan_r
            break

    NCOLS = 150
    r0 = [gv(hdr_row,     c) for c in range(NCOLS)]  # noms des cours
    r1 = [gv(hdr_row + 1, c) for c in range(NCOLS)]  # crédits
    r2 = [gv(hdr_row + 2, c) for c in range(NCOLS)]  # max note / max pondéré
    data_start = hdr_row + 3                           # première ligne étudiant

    # ── Détecter les colonnes étudiant (N°, Matricule, Nom, Sexe) dynamiquement ──
    # Chercher uniquement dans les premières colonnes (avant les cours) pour éviter
    # les faux positifs (ex : "MAT" dans "diplomatique").
    col_num, col_mat, col_nom, col_sex = 0, 1, 2, 3
    for ci in range(min(10, NCOLS)):
        v = str(r0[ci] or '').strip().upper()
        if v in _NUM_KW:
            col_num = ci
        elif v in _MAT_KW or v.startswith('MATRICULE'):
            # Correspondance exacte ou préfixe strict (pas de sous-chaîne)
            col_mat = ci
        elif any(v.startswith(k) for k in _NOM_KW):
            col_nom = ci
        elif v in _SEX_KW:
            col_sex = ci

    # Les cours commencent juste après la dernière colonne d'info étudiant
    course_start_col = max(col_num, col_mat, col_nom, col_sex) + 1

    # ── Détecter le format "cours-au-dessus" ─────────────────────────────────
    # Certaines grilles placent les noms de cours AVANT la ligne d'en-tête :
    #   hdr_row - 2 : noms de cours
    #   hdr_row - 1 : crédits
    #   hdr_row     : N°/Matricule/Nom/Sexe + max notes (chiffres uniquement)
    #   hdr_row + 1 : premier étudiant
    # → Dans ce cas r0 (= hdr_row) ne contient que des chiffres dans les
    #   colonnes cours, pas de texte lisible.
    _course_has_text = any(
        bool(_re.search(r'[A-Za-zÀ-ÿ]{3,}', str(r0[c] or '')))
        for c in range(course_start_col, min(course_start_col + 12, NCOLS))
    )
    if not _course_has_text and hdr_row >= 2:
        # Les noms de cours et les colonnes de synthèse sont 2 lignes plus haut
        r0         = [gv(hdr_row - 2, c) for c in range(NCOLS)]
        r1         = [gv(hdr_row - 1, c) for c in range(NCOLS)]
        # r2 reste hdr_row : contient les max notes (20.0, 60.0…) pour chaque cours
        r2         = [gv(hdr_row,     c) for c in range(NCOLS)]
        data_start = hdr_row + 1

    # ── Détecter si les colonnes cours viennent par paires (note + pondéré) ──
    # Une colonne pondérée se reconnaît à son nom court commençant par P ou contenant POND.
    # On utilise des correspondances strictes pour éviter les faux positifs
    # (ex : « PT » dans « COMPTABILITÉ »).
    _POND_EXACT = {'PT', 'PTS', 'P', 'POND', 'PONDÉ', 'PONDERE', 'PONDERE', 'PONDU'}
    _POND_SUBSTR = ('PONDÉ', 'PONDE', 'PONDU')   # sous-chaînes longues uniquement
    def _is_pond_col(ci):
        name = str(r0[ci] or '').strip().upper()
        if not name:
            return False
        # Correspondance exacte pour noms courts (P, PT, PTS, POND…)
        if name in _POND_EXACT:
            return True
        # Px ou P<chiffres> : P1, P2, P10, P_1 …
        if name.startswith('P') and len(name) <= 5 and name[1:].replace('_', '').isdigit():
            return True
        # Sous-chaînes longues sans ambiguïté
        return any(k in name for k in _POND_SUBSTR)

    # ── Extraire les cours ──
    BAD_KW = ('TOTAL', 'PORCENT', 'POURCENT', 'MOYENNE', 'CREDIT', 'CRÉDIT',
              'APPRE', 'DECIS', 'NOMBR', 'PTS', '%', 'ECHEC', 'VALID',
              'MENTION', 'RÉSULTAT', 'RESULTAT', 'DECISION', 'DÉCISION')

    # Préfixe EC/UE suivi d'un séparateur → toujours un cours, jamais une synthèse
    # Ex : "EC5. Résultats…", "UE3. Crédits de…"  ne doivent PAS être filtrés par BAD_KW
    _re_course_code_with_desc = _re.compile(r'^P?(UE|EC)\s*\d+[\.\-\:\s]', _re.IGNORECASE)

    def _is_forced_course(col_name):
        """Vrai si le nom commence par un code EC/UE suivi d'un séparateur : toujours un cours."""
        return bool(_re_course_code_with_desc.match(col_name))

    def _is_bad_kw(col_name):
        if _is_forced_course(col_name):
            return False   # EC5. Résultats… n'est jamais une colonne de synthèse
        return any(k in col_name.upper() for k in BAD_KW)

    def _looks_like_course(ci):
        """Vrai si la colonne ci ressemble à un vrai cours (pas vide, pas BAD_KW, pas pond, pas code seul)."""
        n = str(r0[ci] or '').strip()
        if not n:
            return False
        if _is_bad_kw(n):          # déjà immunisé contre EC/UE+desc grâce à _is_forced_course
            return False
        if _is_pond_col(ci):
            return False
        nu = n.upper()
        if _re.match(r'^P?(UE|EC)\s*\d+$', nu):
            return False
        return True

    def _has_courses_ahead(from_col, lookahead=12):
        """Regarde si au moins un cours valide existe dans les prochaines colonnes."""
        for lc in range(from_col, min(from_col + lookahead, NCOLS)):
            if _looks_like_course(lc):
                return True
        return False

    courses = []
    c = course_start_col
    consecutive_empty = 0
    while c < NCOLS:
        name = str(r0[c] or '').strip()
        if not name:
            # Tolérer jusqu'à 3 cases vides consécutives (cellules fusionnées mal lues)
            consecutive_empty += 1
            if consecutive_empty > 3:
                break
            c += 1
            continue
        consecutive_empty = 0
        # Colonne de synthèse inter-groupes (TOTAL UE, CRÉDIT, …)
        # → seulement arrêter si aucun cours ne suit dans les 12 prochaines colonnes
        # Note : _is_bad_kw() est déjà faux pour tout nom commençant par EC/UE+séparateur
        if _is_bad_kw(name):
            if _has_courses_ahead(c + 1):
                c += 1   # colonne inter-groupes : on saute et on continue
                continue
            else:
                break    # vrai fin des cours
        # Colonne pondérée (P1, P2, POND…) → on la saute
        if _is_pond_col(c):
            c += 1
            continue
        # Labels de sous-groupes EXACTS (UE1, EC1, PEC1, PUE1… sans description) → on saute
        _name_up = name.strip().upper()
        if _re.match(r'^P?(UE|EC)\s*\d+$', _name_up):
            c += 1
            continue
        credit = float(r1[c] or 0) if r1[c] not in (None, '') else 0.0
        # Chercher la colonne pondérée associée (immédiatement après)
        if c + 1 < NCOLS and _is_pond_col(c + 1):
            pond_col = c + 1
            max_pondere = float(r2[pond_col]) if r2[pond_col] not in (None, '') else 0.0
        else:
            pond_col = None
            max_pondere = float(r2[c]) if r2[c] not in (None, '') else 0.0
        clean = _re.sub(r'^\s*\d+[\.\-]?\s*', '', name).strip()
        courses.append({
            'col':         c,
            'pond_col':    pond_col,  # None = pas de colonne pondérée dédiée
            'name':        clean,
            'credit':      credit,
            'max_pondere': max_pondere,
        })
        c += 1

    # ── Localiser les colonnes de synthèse ──
    col_map = {}
    for ci, v in enumerate(r0):
        s = str(v or '').upper()
        if 'ECHEC' in s:                                     col_map['echecs']  = ci
        elif 'TOTAL PONDERE' in s or 'TOTAL PONDÉRÉ' in s:  col_map['total']   = ci
        elif 'PORCENT' in s or 'POURCENT' in s:             col_map['pct']     = ci
        elif 'MOYENNE' in s:                                 col_map['moy']     = ci
        elif 'CREDITS VALID' in s or 'CRÉDITS VALID' in s:  col_map['cr_val']  = ci
        elif 'CREDITS NON'  in s or 'CRÉDITS NON'  in s:    col_map['cr_nval'] = ci
        elif 'APPRE' in s:                                   col_map['appre']   = ci
        elif 'DECISION' in s or 'DÉCISION' in s:            col_map['decision']= ci

    # ── Extraire la promo/département depuis le nom de feuille ──
    # Ex: "1er Semestre_BAC 1 RI"  ou  "S1_M1 AGIE"
    parts = sheet_name.split('_')
    semestre_meta   = parts[0].strip() if len(parts) > 0 else ''
    promo_dept_meta = parts[1].strip() if len(parts) > 1 else sheet_name.strip()

    meta = {
        'sheet':            sheet_name,
        'semestre':         semestre_meta,
        'promo_dept':       promo_dept_meta,
        'hdr_row':          hdr_row,            # pour debug
        'nb_cours':         len(courses),        # pour message flash
        'course_start_col': course_start_col,    # pour debug
        'col_num':          col_num,
        'col_mat':          col_mat,
        'col_nom':          col_nom,
        'col_sex':          col_sex,
        'col_map':          col_map,             # colonnes de synthèse détectées
    }

    # ── Lire les étudiants ──
    students = []
    for ri in range(data_start, nrows):
        row = [gv(ri, ci) for ci in range(NCOLS)]
        num = row[col_num]
        if not num:
            continue
        mat = str(row[col_mat] or '').strip()
        nom = str(row[col_nom] or '').strip()
        sex = str(row[col_sex] or '').strip()
        if not nom:
            continue

        cours_res = []
        for crs in courses:
            note_raw = row[crs['col']]
            if crs['pond_col'] is not None:
                pondere_raw = row[crs['pond_col']]
            else:
                pondere_raw = None
            note_val    = float(note_raw)    if note_raw    not in (None, '') else None
            pondere_val = float(pondere_raw) if pondere_raw not in (None, '') else 0.0
            if note_val is None:
                res = 'AB'
            elif note_val >= 10:
                res = 'ADM'
            else:
                res = 'REPR'
            cours_res.append({
                'name':        crs['name'],
                'credit':      crs['credit'],
                'note':        note_val,
                'pondere':     pondere_val,
                'max_pondere': crs['max_pondere'],
                'has_pond_col': crs['pond_col'] is not None,  # False = note-only format
                'result':      res,
            })

        def _fv(key, default=0):
            ci = col_map.get(key)
            return float(row[ci] or default) if ci is not None and row[ci] not in (None, '') else default

        students.append({
            'num':           int(float(num)),
            'matricule':     mat,
            'nom':           nom,
            'sexe':          sex,
            'cours':         cours_res,
            'nb_echecs':     int(_fv('echecs')),
            'total':         _fv('total'),
            'pct':           _fv('pct'),
            'moyenne':       _fv('moy'),
            'moy_col_present': 'moy' in col_map,  # True si la colonne Moyenne existait dans la grille
            'cr_val':        _fv('cr_val'),
            'cr_nval':       _fv('cr_nval'),
            'appre':         str(row[col_map['appre']]    if col_map.get('appre')    and row[col_map['appre']]    not in (None, '') else ''),
            'decision':      str(row[col_map['decision']] if col_map.get('decision') and row[col_map['decision']] not in (None, '') else ''),
        })

    return courses, students, meta


def _generer_bulletin_pdf(etudiant, promotion, annee, session_acad, semestre_label,
                          numero_bulletin=None, texte_intro=None, departement=None):
    """Génère un PDF bulletin pour un étudiant (modèle UNILU officiel). Retourne un BytesIO."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                     Paragraph, Spacer, HRFlowable)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    import datetime as _dt
    import re as _re_pdf

    # ── Détection du modèle AVANT création du doc (marges différentes) ──
    _dept_str  = (departement or '').strip()
    _promo_up  = (promotion or '').upper()
    _is_master = bool(_re_pdf.search(r'\bM(ASTER)?\s*[12]\b', _promo_up))
    _is_anthro = ("anthropologie" in _dept_str.lower()) and _is_master

    # Marges : Anthro compact pour tenir sur 1 page, standard sinon
    if _is_anthro:
        _mH, _mV, _mL, _mR = 1.2*cm, 0.6*cm, 1.2*cm, 1.2*cm
    else:
        _mH, _mV, _mL, _mR = 1.8*cm, 1.2*cm, 1.8*cm, 1.8*cm

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             topMargin=_mV, bottomMargin=_mV,
                             leftMargin=_mH, rightMargin=_mH)

    PAGE_W = A4[0] - _mH * 2   # largeur utile

    MARINE  = colors.Color(13 / 255, 38 / 255, 102 / 255)
    OR      = colors.Color(191 / 255, 153 / 255, 26 / 255)
    VERT    = colors.Color(0.0, 0.45, 0.1)
    ROUGE   = colors.Color(0.78, 0.06, 0.06)
    GRIS_L  = colors.Color(0.93, 0.93, 0.93)
    BG_ALT  = colors.Color(0.96, 0.97, 1.0)

    def P(txt, **kw):
        kw.setdefault('fontName', 'Helvetica')
        kw.setdefault('fontSize', 9)
        kw.setdefault('leading',  11)
        return Paragraph(txt, ParagraphStyle('_p', **kw))

    story = []

    # ══════════════════ 1. EN-TÊTE INSTITUTIONNEL ══════════════════
    from reportlab.platypus import Image as RLImage
    from reportlab.lib.utils import ImageReader
    from flask import current_app as _cur_app

    # _is_anthro / _dept_str / _re_pdf déjà définis avant la création du doc
    if _is_anthro:
        _logo_left  = os.path.join(_cur_app.root_path, 'static', 'logo-anthro-unilu.png')
        _logo_right = os.path.join(_cur_app.root_path, 'static', 'logo-anthro-dept.png')
    else:
        _logo_left  = os.path.join(_cur_app.root_path, 'static', 'logo-sciences-sociales-fac.jpg')
        _logo_right = os.path.join(_cur_app.root_path, 'static', 'logo-unilu-armoiries.png')

    logo_h = (1.5 if _is_anthro else 2.0) * cm

    def _logo(path, h=logo_h):
        try:
            ir = ImageReader(path)
            iw, ih = ir.getSize()
            ratio = iw / ih
            return RLImage(path, width=h * ratio, height=h)
        except Exception:
            return P('', fontSize=6)

    # Bloc texte central
    if _is_anthro:
        # Modèle Anthropologie :
        #   REPUBLIQUE DEMOCRATIQUE DU CONGO
        #   UNIVERSITÉ DE LUBUMBASHI
        #   Faculté des Sciences Sociales...
        #   Département d'Anthropologie          ← ligne propre
        #   BULLETIN DE RÉSULTATS ACADÉMIQUES    ← ligne propre, centrée en dessous
        _dept_label = _dept_str if _dept_str else "Département d'Anthropologie"
        _centre_block = [
            P('<b>REPUBLIQUE DEMOCRATIQUE DU CONGO</b>',
              fontName='Helvetica-Bold', fontSize=7, leading=9,
              alignment=TA_CENTER, textColor=MARINE),
            P('<b>UNIVERSITÉ DE LUBUMBASHI</b>',
              fontName='Helvetica-Bold', fontSize=10, leading=12,
              alignment=TA_CENTER, textColor=MARINE),
            P('Faculté des Sciences Sociales, Politiques et Administratives',
              fontSize=8, leading=10, alignment=TA_CENTER, textColor=MARINE),
            P(f'<b>{_dept_label}</b>',
              fontName='Helvetica-Bold', fontSize=8, leading=11,
              alignment=TA_CENTER, textColor=MARINE),
            P('<b>BULLETIN DES RÉSULTATS ACADÉMIQUES</b>',
              fontName='Helvetica-Bold', fontSize=9, leading=11,
              alignment=TA_CENTER, textColor=MARINE),
        ]
    else:
        # Modèle standard : toutes les lignes centrées
        _centre_block = [
            P('<b>UNIVERSITÉ DE LUBUMBASHI</b>',
              fontName='Helvetica-Bold', fontSize=10, leading=13,
              alignment=TA_CENTER, textColor=MARINE),
            P('Faculté des Sciences Sociales, Politiques et Administratives',
              fontSize=8, leading=11, alignment=TA_CENTER, textColor=MARINE),
        ]
        if _dept_str:
            _centre_block.append(
                P(f'<b>{_dept_str}</b>',
                  fontName='Helvetica-Bold', fontSize=8, leading=11,
                  alignment=TA_CENTER, textColor=MARINE)
            )
        _centre_block.append(
            P('<b>BULLETIN DES RÉSULTATS ACADÉMIQUES</b>',
              fontName='Helvetica-Bold', fontSize=10, leading=13,
              alignment=TA_CENTER, textColor=MARINE)
        )

    hdr_data = [[_logo(_logo_left), _centre_block, _logo(_logo_right)]]
    hdr_tbl = Table(hdr_data, colWidths=[PAGE_W * 0.15, PAGE_W * 0.70, PAGE_W * 0.15])
    hdr_tbl.setStyle(TableStyle([
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN',         (0, 0), (0, 0),   'LEFT'),
        ('ALIGN',         (2, 0), (2, 0),   'RIGHT'),
        ('LINEBELOW',     (0, 0), (-1, -1), 2.0, MARINE),
        ('TOPPADDING',    (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    story += [hdr_tbl, Spacer(1, 0.15 * cm if _is_anthro else 0.35 * cm)]

    # ══════════════════ 2. FICHE ÉTUDIANT (modèle officiel) ══════════════════
    # Colonnes : NOM/Matricule/Promotion | QR CODE (milieu) | Sexe/Bulletin n°
    nom_str  = etudiant.get('nom', '').upper()
    sexe_str = etudiant.get('sexe', '')
    bul_no   = str(numero_bulletin) if numero_bulletin else '—'
    mat_str  = etudiant.get('matricule', '') or '—'

    # ── Générer le QR code étudiant ──────────────────────────────────────────
    # Contenu : NOM, PROMOTION, POURCENTAGE, DECISION DU JURY
    _qr_pct = etudiant.get('pct', 0) or 0
    _qr_dec = etudiant.get('decision', '') or ''
    _qr_content = (
        f"NOM: {nom_str}\n"
        f"PROMOTION: {promotion}\n"
        f"POURCENTAGE: {_qr_pct:.1f}%\n"
        f"DECISION: {_qr_dec}"
    )
    try:
        _qr_obj = qrcode.QRCode(version=1, box_size=4, border=2,
                                 error_correction=qrcode.constants.ERROR_CORRECT_M)
        _qr_obj.add_data(_qr_content)
        _qr_obj.make(fit=True)
        _qr_pil  = _qr_obj.make_image(fill_color='black', back_color='white')
        _qr_buf  = BytesIO()
        _qr_pil.save(_qr_buf, format='PNG')
        _qr_buf.seek(0)
        _qr_size = 2.2 * cm
        _qr_cell = RLImage(_qr_buf, width=_qr_size, height=_qr_size)
    except Exception:
        _qr_cell = P('', fontSize=6)

    ident_data = [
        [
            P(f'<b>{nom_str}</b>', fontSize=11, fontName='Helvetica-Bold', textColor=MARINE),
            _qr_cell,                                                          # QR — span 3 lignes
            P(f'Sexe : <b>{sexe_str}</b>', fontSize=9, alignment=TA_CENTER),
        ],
        [
            P(f'Matricule : <b>{mat_str}</b>', fontSize=9, textColor=MARINE),
            '',                                                                 # cellule fusionnée
            P(f'Bulletin n° : <b>{bul_no}</b>', fontSize=9, alignment=TA_RIGHT),
        ],
        [
            P(f'<b>{promotion}</b>', fontSize=10, fontName='Helvetica-Bold',
              textColor=MARINE, leading=13),
            '',                                                                 # cellule fusionnée
            P('', fontSize=9),
        ],
    ]
    _pad = 2 if _is_anthro else 3
    ident_tbl = Table(ident_data, colWidths=[PAGE_W * 0.50, PAGE_W * 0.22, PAGE_W * 0.28])
    ident_tbl.setStyle(TableStyle([
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN',         (1, 0), (1, 2),   'CENTER'),
        ('TOPPADDING',    (0, 0), (-1, -1), _pad),
        ('BOTTOMPADDING', (0, 0), (-1, -1), _pad),
        ('LINEBELOW',     (0, 2), (-1, 2),  0.8, MARINE),
        ('SPAN',          (1, 0), (1, 2)),  # QR code occupe les 3 lignes de la colonne centrale
    ]))
    story += [ident_tbl, Spacer(1, 0.15 * cm if _is_anthro else 0.3 * cm)]

    # ══════════════════ 3. PHRASE INTRO ══════════════════
    _intro = texte_intro if texte_intro and texte_intro.strip() else (
        f'A obtenu à l\'issue des évaluations de la {session_acad} '
        f'de l\'année académique {annee}, '
        f'semestre {semestre_label}, les notes ci-dessous :'
    )
    _intro_fs = 8 if _is_anthro else 9
    story.append(P(_intro, fontSize=_intro_fs, leading=_intro_fs + 3))
    story.append(Spacer(1, 0.12 * cm if _is_anthro else 0.25 * cm))

    # ══════════════════ 4. TABLEAU DES COURS ══════════════════
    # Détecter si des colonnes pondérées étaient présentes dans la grille.
    # On utilise has_pond_col (stocké lors du parsing) comme signal principal.
    # Fallback sur pondere != 0 pour les bulletins importés avant l'ajout de ce champ.
    _cours_list = etudiant.get('cours', [])
    if any('has_pond_col' in c for c in _cours_list):
        _has_pondere = any(c.get('has_pond_col') for c in _cours_list)
    else:
        # Données anciennes : pas de has_pond_col — on teste les valeurs pondérées réelles
        _has_pondere = any((c.get('pondere') or 0) != 0 for c in _cours_list)

    if _is_anthro:
        # Modèle Anthropologie : N° | Intitulé | Note/20 | Note pondérée (pas de Crédit)
        th = [
            P('<b>N°</b>',                textColor=colors.white, alignment=TA_CENTER),
            P('<b>Intitulé du cours</b>',  textColor=colors.white),
            P('<b>Note sur 20</b>',        textColor=colors.white, alignment=TA_CENTER),
        ]
        if _has_pondere:
            th.append(P('<b>Note pondérée</b>', textColor=colors.white, alignment=TA_CENTER))
    else:
        th = [
            P('<b>Num</b>',               textColor=colors.white, alignment=TA_CENTER),
            P('<b>Intitulé du cours</b>',  textColor=colors.white),
            P('<b>Crédit</b>',             textColor=colors.white, alignment=TA_CENTER),
            P('<b>Note sur 20</b>',        textColor=colors.white, alignment=TA_CENTER),
        ]
        if _has_pondere:
            th.append(P('<b>Note pondérée</b>', textColor=colors.white, alignment=TA_CENTER))
    rows      = [th]
    row_styles = []

    _re_code_exact  = _re_pdf.compile(r'^P?(UE|EC)\s*\d+$',               _re_pdf.IGNORECASE)
    _re_ue_prefix   = _re_pdf.compile(r'^P?UE\s*\d+',                      _re_pdf.IGNORECASE)
    _re_ec_strip    = _re_pdf.compile(r'^EC\s*\d+[\.\-\:\s]+',             _re_pdf.IGNORECASE)
    _re_code_prefix = _re_pdf.compile(r'^P?(UE|EC)\s*\d+[\.\-\:]\s*',     _re_pdf.IGNORECASE)

    row_num = 0  # numéro de ligne affiché (on saute les lignes filtrées)
    for crs in etudiant.get('cours', []):
        raw_name = crs.get('name', '').strip()

        # Filtrer les labels de groupes UE (toutes promos) — codes exacts ou avec description
        if _re_ue_prefix.match(raw_name):   # UE1, UE1. Méthodes…, PUE3. Culture…
            continue
        if _re_code_exact.match(raw_name):  # EC1, PEC1 seuls (sans description)
            continue
        if _is_anthro:
            pass   # le bloc anthro ne filtre rien de plus ici

        # Nettoyer le préfixe EC1. / UE2. en début de nom
        if _is_anthro:
            clean_name = _re_ec_strip.sub('', raw_name).strip()
        else:
            clean_name = _re_code_prefix.sub('', raw_name).strip()

        row_num += 1
        ri = row_num
        bg = colors.white if row_num % 2 == 1 else BG_ALT
        row_styles.append(('BACKGROUND', (0, ri), (-1, ri), bg))

        note_val = crs.get('note')
        note_str = f'{note_val:.1f}' if note_val is not None else ''
        pond_str = f'{crs.get("pondere", 0):.0f}'

        # Note < 10 : rouge gras
        n_col = ROUGE if (note_val is not None and note_val < 10) else colors.black
        n_fn  = 'Helvetica-Bold' if (note_val is not None and note_val < 10) else 'Helvetica'

        _cfs = 7 if _is_anthro else 8   # taille police cours
        if _is_anthro:
            _row = [
                P(str(row_num), fontSize=_cfs, alignment=TA_CENTER),
                P(clean_name,   fontSize=_cfs, leading=_cfs + 2),
                P(f'<b>{note_str}</b>', fontSize=_cfs, fontName=n_fn,
                  textColor=n_col, alignment=TA_CENTER),
            ]
            if _has_pondere:
                _row.append(P(pond_str, fontSize=_cfs, alignment=TA_CENTER))
            rows.append(_row)
        else:
            _row = [
                P(str(row_num), alignment=TA_CENTER),
                P(clean_name,   leading=11, fontSize=8),
                P(str(int(crs.get('credit', 0))), alignment=TA_CENTER),
                P(f'<b>{note_str}</b>', fontName=n_fn, textColor=n_col,
                  alignment=TA_CENTER),
            ]
            if _has_pondere:
                _row.append(P(pond_str, alignment=TA_CENTER))
            rows.append(_row)

    # ── Notice si aucun cours n'a été chargé ──────────────────────────────────
    if row_num == 0:
        _notice_msg = "Aucun cours n'a pu être chargé pour ce bulletin."
        _n_cols = (3 if _is_anthro else 4) + (1 if _has_pondere else 0)
        _notice_row = [P(f'<i>{_notice_msg}</i>', fontSize=9,
                         textColor=ROUGE, alignment=TA_CENTER)]
        for _ in range(_n_cols - 1):
            _notice_row.append(P(''))
        rows.append(_notice_row)
        row_styles.append(('SPAN', (0, 1), (_n_cols - 1, 1)))
        row_styles.append(('ALIGN', (0, 1), (_n_cols - 1, 1), 'CENTER'))
        row_styles.append(('BACKGROUND', (0, 1), (_n_cols - 1, 1), colors.Color(1.0, 0.95, 0.95)))

    # Largeurs selon le modèle
    if _is_anthro:
        if _has_pondere:
            # N° | Cours | Note/20 | Pondérée
            W = [0.8*cm, PAGE_W - 0.8*cm - 2.0*cm - 2.4*cm, 2.0*cm, 2.4*cm]
        else:
            # N° | Cours | Note/20 (sans pondérée)
            W = [0.8*cm, PAGE_W - 0.8*cm - 2.4*cm, 2.4*cm]
    else:
        if _has_pondere:
            # N° | Cours | Crédit | Note/20 | Pondérée
            W = [1.0*cm, PAGE_W - 1.0*cm - 1.5*cm - 2.0*cm - 2.5*cm, 1.5*cm, 2.0*cm, 2.5*cm]
        else:
            # N° | Cours | Crédit | Note/20 (sans pondérée)
            W = [1.0*cm, PAGE_W - 1.0*cm - 1.5*cm - 2.5*cm, 1.5*cm, 2.5*cm]

    _cp = 1 if _is_anthro else 2   # cell padding cours
    cours_tbl = Table(rows, colWidths=W, repeatRows=1)
    cours_tbl.setStyle(TableStyle([
        ('BACKGROUND',    (0, 0), (-1, 0),  MARINE),
        ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
        ('BOX',           (0, 0), (-1, -1), 0.8, MARINE),
        ('INNERGRID',     (0, 0), (-1, -1), 0.3, colors.Color(0.78, 0.78, 0.78)),
        ('TOPPADDING',    (0, 0), (-1, -1), _cp),
        ('BOTTOMPADDING', (0, 0), (-1, -1), _cp),
        ('LEFTPADDING',   (0, 0), (-1, -1), 3 if _is_anthro else 4),
        ('RIGHTPADDING',  (0, 0), (-1, -1), 3 if _is_anthro else 4),
    ] + row_styles))
    story += [cours_tbl, Spacer(1, 0.15 * cm if _is_anthro else 0.35 * cm)]

    # ══════════════════ 5. SYNTHÈSE ══════════════════
    dec     = (etudiant.get('decision') or '').upper()
    dec_col = VERT if dec in ('ADM', 'COMP') else ROUGE
    dec_hex = dec_col.hexval() if hasattr(dec_col, 'hexval') else '#000000'
    total_p = etudiant.get('total', 0)
    moy     = etudiant.get('moyenne', 0)
    appre   = etudiant.get('appre', '')
    pct     = etudiant.get('pct', 0)
    cr_val  = etudiant.get('cr_val', 0)

    if _is_anthro:
        # Modèle Anthropologie : lignes verticales, comme dans l'ODT officiel
        syn_lines = [
            (f'TOTAL PONDERE ANNUEL',          f'{total_p:.3f}'),
            (f'MOYENNE PONDEREE ANNUELLE',      f'{moy:.12g}'),
            (f'APPRECIATION',                   appre),
            (f'DECISION DU JURY',               dec),
            (f'POURCENTAGE PONDERE ANNUEL',     f'{pct:.1f}'),
            (f'TOTAL CREDITS VALISES',          f'{cr_val:.0f}'),
        ]
        syn_rows = []
        for label, val in syn_lines:
            is_dec = label == 'DECISION DU JURY'
            v_para = P(
                f'<font color="{dec_hex}"><b>{val}</b></font>' if is_dec
                else f'<b>{val}</b>',
                fontSize=9, textColor=(dec_col if is_dec else colors.black))
            syn_rows.append([
                P(f'{label} :', fontSize=9, fontName='Helvetica-Bold'),
                v_para,
            ])
        syn_tbl = Table(syn_rows,
                        colWidths=[PAGE_W * 0.60, PAGE_W * 0.40])
        syn_tbl.setStyle(TableStyle([
            ('BOX',           (0, 0), (-1, -1), 1.0, MARINE),
            ('LINEBELOW',     (0, 0), (-1, -2), 0.3, colors.lightgrey),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING',    (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING',   (0, 0), (-1, -1), 6),
        ]))
    else:
        # Modèle standard : tableau 3×2
        syn_data = [
            [
                P(f'TOTAL PONDÉRÉ ANNUEL : <b>{total_p:.0f}</b>', fontSize=9),
                P(f'MOYENNE PONDÉRÉE ANNUELLE : <b>{moy:.2f}</b>', fontSize=9),
                P(f'APPRÉCIATION ANNUELLE : <b>{appre}</b>', fontSize=9,
                  textColor=MARINE),
            ],
            [
                P(f'POURCENTAGE PONDÉRÉ ANNUEL : <b>{pct:.1f}</b>', fontSize=9),
                P(f'TOTAL CRÉDITS VALIDÉS : <b>{cr_val:.1f}</b>', fontSize=9),
                P(f'<b>DÉCISION DU JURY : <font color="{dec_hex}">{dec}</font></b>',
                  fontName='Helvetica-Bold', fontSize=12, alignment=TA_CENTER,
                  textColor=dec_col),
            ],
        ]
        syn_tbl = Table(syn_data, colWidths=[PAGE_W * 0.34, PAGE_W * 0.34, PAGE_W * 0.32])
        syn_tbl.setStyle(TableStyle([
            ('BOX',           (0, 0), (-1, -1), 1.2, MARINE),
            ('INNERGRID',     (0, 0), (-1, -1), 0.3, colors.lightgrey),
            ('BACKGROUND',    (0, 0), (-1, 0),  colors.Color(0.93, 0.96, 1.0)),
            ('BACKGROUND',    (0, 1), (-1, 1),  colors.Color(0.88, 0.92, 1.0)),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING',    (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING',   (0, 0), (-1, -1), 6),
        ]))
    story += [syn_tbl, Spacer(1, 0.15 * cm if _is_anthro else 0.4 * cm)]

    # ══════════════════ 5b. NOTE DE BAS DE PAGE (sans pondération) ══════════════════
    if not _has_pondere:
        story.append(P(
            '<i>* Ce bulletin présente les notes brutes sans pondération.</i>',
            fontSize=7, textColor=colors.grey, leading=9,
        ))
        story.append(Spacer(1, 0.1 * cm))

    # ══════════════════ 6. LÉGENDE (standard seulement) ══════════════════
    if not _is_anthro:
        leg_data = [[
            P('<b>Légende :</b>', fontSize=8, textColor=MARINE),
            P('<b>ADM</b> : Admis (réussit sans échec)', fontSize=8),
            P('<b>COMP</b> : Compensation (réussit avec compensation)', fontSize=8),
            P('<b>AJ</b> : Ajourné (échoue pour accumulation d\'échecs)', fontSize=8),
            P('<b>DEF</b> : Défaillant (échoue pour manque de notes)', fontSize=8),
        ]]
        leg_tbl = Table(leg_data, colWidths=[
            PAGE_W * 0.10, PAGE_W * 0.20, PAGE_W * 0.25, PAGE_W * 0.25, PAGE_W * 0.20])
        leg_tbl.setStyle(TableStyle([
            ('BOX',           (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ('BACKGROUND',    (0, 0), (-1, -1), GRIS_L),
            ('VALIGN',        (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING',    (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
            ('LEFTPADDING',   (0, 0), (-1, -1), 4),
        ]))
        story += [leg_tbl, Spacer(1, 0.45 * cm)]

    # ══════════════════ 7. LIEU ET DATE ══════════════════
    today = _dt.date.today()
    date_str = today.strftime('%d/%m/%Y')
    story.append(P(
        f'Fait à Lubumbashi, le <b>{date_str}</b>',
        fontSize=9, alignment=TA_LEFT,
    ))
    story.append(Spacer(1, 0.15 * cm if _is_anthro else 0.3 * cm))

    # ══════════════════ 8. SIGNATURES ══════════════════
    story.append(P(
        '<b>POUR LE BUREAU DU JURY</b>',
        fontSize=9, fontName='Helvetica-Bold', alignment=TA_CENTER, textColor=MARINE,
    ))
    story.append(Spacer(1, 0.1 * cm if _is_anthro else 0.25 * cm))

    _sig_top = 18 if _is_anthro else 30
    sig_data = [[
        P('Secrétaires', fontSize=9, alignment=TA_CENTER),
        P('', fontSize=9),
        P('Président', fontSize=9, alignment=TA_CENTER),
    ]]
    sig_tbl = Table(sig_data, colWidths=[PAGE_W * 0.30, PAGE_W * 0.40, PAGE_W * 0.30])
    sig_tbl.setStyle(TableStyle([
        ('TOPPADDING',  (0, 0), (-1, -1), _sig_top),
        ('LINEABOVE',   (0, 0), (0, 0),   0.5, colors.grey),
        ('LINEABOVE',   (2, 0), (2, 0),   0.5, colors.grey),
    ]))
    story.append(sig_tbl)

    # ══════════════════ 9. QR CODE D'AUTHENTICITÉ ══════════════════
    if numero_bulletin:
        try:
            import qrcode as _qrcode
            qr_img = _qrcode.make(f'/bulletins/verifier/{numero_bulletin}')
            qr_buf_pdf = BytesIO()
            qr_img.save(qr_buf_pdf, format='PNG')
            qr_buf_pdf.seek(0)
            from reportlab.platypus import Image as RLImage
            from reportlab.lib.utils import ImageReader
            qr_row = [[
                RLImage(ImageReader(qr_buf_pdf), width=1.8*cm, height=1.8*cm),
                P(f'<b>N° Bulletin :</b> {numero_bulletin}<br/>'
                  f'<font size="7">Scannez pour vérifier l\'authenticité</font>',
                  fontSize=7, leading=10, textColor=colors.grey),
            ]]
            qr_tbl = Table(qr_row, colWidths=[2.2*cm, None])
            qr_tbl.setStyle(TableStyle([
                ('VALIGN',    (0, 0), (-1, -1), 'MIDDLE'),
                ('LINEABOVE', (0, 0), (-1, 0),  0.4, colors.lightgrey),
                ('TOPPADDING',(0, 0), (-1, -1),  5),
            ]))
            story += [Spacer(1, 0.4*cm), qr_tbl]
        except Exception:
            pass

    doc.build(story)
    buf.seek(0)
    return buf


def _charger_liste_identifiants_promotion(promotion):
    """Retourne uniquement la liste dont la promotion normalisée est exacte."""

    rows = ListeIdentifiants.query.order_by(ListeIdentifiants.nom).all()
    if not rows:
        return '', []

    cible = _normaliser_nom(promotion)
    matching_rows = [row for row in rows if _normaliser_nom(row.promotion) == cible]
    if not matching_rows:
        return '', []
    return matching_rows[0].promotion, matching_rows


def _preparer_attributions_matricules(students, liste_rows):
    """Prépare des associations conservatrices pour validation par le DÉCANAT.

    Une correspondance n'est automatique que si un seul matricule officiel
    correspond au nom complet ou exactement aux tokens nom/postnom. Les
    ressemblances approximatives restent à vérifier manuellement.
    """
    from collections import defaultdict

    by_full = defaultdict(list)
    by_tokens = defaultdict(list)
    for row in liste_rows:
        by_full[row.nom_norm].append(row)
        token_key = ' '.join(sorted(_tokens_nom(row.nom_norm)))
        if token_key:
            by_tokens[token_key].append(row)

    assignments = []
    for index, student in enumerate(students):
        nom = student.get('nom', '').strip()
        norm = _normaliser_nom(nom)
        token_key = ' '.join(sorted(_tokens_nom(norm)))
        candidates = list(by_full.get(norm, []))
        match_type = 'nom complet'

        if not candidates and token_key:
            candidates = list(by_tokens.get(token_key, []))
            match_type = 'nom/post-nom'

        # Si la grille contient un prénom supplémentaire, accepter uniquement
        # une liste dont tous les tokens sont présents, avec au moins 2 tokens.
        if not candidates and len(_tokens_nom(norm)) >= 2:
            grid_tokens = _tokens_nom(norm)
            subset_rows = []
            for row in liste_rows:
                list_tokens = _tokens_nom(row.nom_norm)
                if len(list_tokens) >= 2 and (
                    list_tokens.issubset(grid_tokens)
                    or grid_tokens.issubset(list_tokens)
                ):
                    subset_rows.append(row)
            if subset_rows:
                candidates = subset_rows
                match_type = 'nom/post-nom'

        unique = {}
        for row in candidates:
            unique[row.matricule] = row
        candidates = list(unique.values())

        item = {
            'index': index,
            'nom': nom,
            'matricule_grille': student.get('matricule', ''),
            'statut': 'introuvable',
            'type_correspondance': '',
            'matricule': '',
            'nom_liste': '',
            'candidats': [
                {'matricule': row.matricule, 'nom': row.nom}
                for row in candidates
            ],
        }
        if len(candidates) == 1:
            item.update({
                'statut': 'confirme',
                'type_correspondance': match_type,
                'matricule': candidates[0].matricule,
                'nom_liste': candidates[0].nom,
            })
        elif len(candidates) > 1:
            item['statut'] = 'ambigu'
        assignments.append(item)

    # Deux lignes de grille ne peuvent pas recevoir automatiquement le même
    # matricule : les deux doivent être revues par l'administrateur.
    auto_counts = defaultdict(list)
    for item in assignments:
        if item['statut'] == 'confirme':
            auto_counts[item['matricule']].append(item)
    for matricule, items in auto_counts.items():
        if len(items) > 1:
            for item in items:
                item['statut'] = 'doublon'
                item['candidats'] = [{
                    'matricule': matricule,
                    'nom': item['nom_liste'],
                }]
                item['matricule'] = ''
                item['nom_liste'] = ''
    return assignments


def _enregistrer_grille_validee(payload, matricules):
    """Enregistre une grille après validation de tous ses matricules officiels."""
    import random
    import string

    students = payload['students']
    liste_matricules = {row['matricule'] for row in payload['identifiants']}
    if len(matricules) != len(students):
        raise ValueError('Le nombre de matricules validés ne correspond pas aux étudiants.')
    if any(not matricule or matricule not in liste_matricules for matricule in matricules):
        raise ValueError('Un matricule validé ne provient pas de la liste officielle.')
    if len(set(matricules)) != len(matricules):
        raise ValueError('Un même matricule a été attribué à plusieurs étudiants.')

    type_grille = payload.get('type_grille', 'initial')
    if type_grille not in _GRID_TYPES:
        raise ValueError('Type de grille invalide.')

    # Une nouvelle étape est indépendante. Une nouvelle importation du même
    # type remplace uniquement ce même brouillon de session ; elle ne touche
    # jamais aux trois autres étapes ni à leurs paiements.
    existing_query = BulletinSession.query.filter_by(
        promotion=payload['promotion'],
        annee=payload['annee'],
        semestre=payload['semestre'],
        session_acad=payload['session_acad'],
        type_grille=type_grille,
    ).order_by(BulletinSession.date_import.desc())
    existing_bs = existing_query.first()
    # Ne jamais supprimer une session qui contient déjà un accès payé :
    # l'import d'une nouvelle grille doit préserver cet accès définitivement.
    if existing_bs and BulletinData.query.filter_by(
            session_id=existing_bs.id, paye=True).first():
        existing_bs = None

    if existing_bs:
        departement = payload.get('departement') or existing_bs.departement or ''
        texte_intro = payload.get('texte_intro') or existing_bs.texte_intro or ''
        BulletinData.query.filter_by(session_id=existing_bs.id).delete()
        existing_bs.nom = f"{payload['promotion']} – {payload['semestre']} {payload['annee']}"
        existing_bs.session_acad = payload['session_acad']
        existing_bs.montant_fc = payload['montant_fc']
        existing_bs.departement = departement
        existing_bs.texte_intro = texte_intro
        existing_bs.date_import = now_cat()
        db.session.flush()
        bs = existing_bs
    else:
        suffix = f" – {_grid_type_label(type_grille).upper()}" if type_grille != 'initial' else ''
        bs = BulletinSession(
            nom=f"{payload['promotion']} – {payload['semestre']} {payload['annee']}{suffix}",
            annee=payload['annee'],
            session_acad=payload['session_acad'],
            semestre=payload['semestre'],
            promotion=payload['promotion'],
            montant_fc=payload['montant_fc'],
            departement=payload.get('departement', ''),
            texte_intro=payload.get('texte_intro', ''),
            type_grille=type_grille,
        )
        db.session.add(bs)
        db.session.flush()

    for student, matricule in zip(students, matricules):
        saved = dict(student)
        saved['matricule'] = matricule
        suffix = ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))
        db.session.add(BulletinData(
            session_id=bs.id,
            matricule=matricule,
            nom=saved['nom'],
            sexe=saved['sexe'],
            data_json=json.dumps(saved, ensure_ascii=False),
            numero_bulletin=f"BUL-{bs.id:04d}-{suffix}",
        ))
    db.session.flush()
    return bs


def _est_session_recours(bs):
    """Retourne True pour l'une des deux sessions de recours."""
    return bool(bs and getattr(bs, 'type_grille', 'initial') in ('recours', 'recours_session_2'))


def _type_recu_attendu_pour_bulletin(bd):
    """Type de reçu autorisé pour le bulletin sélectionné."""
    grid_type = getattr(bd.bul_session, 'type_grille', 'initial')
    if grid_type == 'recours':
        return 'resultat_recours'
    if grid_type == 'session_2':
        return 'session_2'
    if grid_type == 'recours_session_2':
        return 'recours_session_2'
    return 'bulletin'


def _generer_zip_session(bs):
    import zipfile as _zipfile

    zip_buf = BytesIO()
    with _zipfile.ZipFile(zip_buf, 'w', _zipfile.ZIP_DEFLATED) as zf:
        for bd in BulletinData.query.filter_by(session_id=bs.id).all():
            etu = json.loads(bd.data_json)
            pdf_buf = _generer_bulletin_pdf(
                etu, bs.promotion, bs.annee, bs.session_acad, bs.semestre,
                numero_bulletin=bd.numero_bulletin,
                texte_intro=bs.texte_intro, departement=bs.departement)
            zf.writestr(f"Bulletin_{bd.nom.replace(' ', '_')[:35]}.pdf", pdf_buf.read())
    zip_buf.seek(0)
    return zip_buf, f"Bulletins_{bs.promotion.replace(' ', '_')}.zip"


@app.route('/decanat/bulletins', methods=['GET', 'POST'])
def decanat_bulletins():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    if request.method == 'POST':
        grille_file   = request.files.get('grille')
        annee         = request.form.get('annee_academique', '2025-2026').strip()
        session_acad  = request.form.get('session', '1ère session').strip()
        semestre_label= request.form.get('semestre', '1er Semestre').strip()
        promotion     = request.form.get('promotion', '').strip()
        montant_fc    = int(request.form.get('montant_fc', 5000))

        if not grille_file or grille_file.filename == '':
            flash('Veuillez sélectionner un fichier Excel (.xls ou .xlsx)', 'error')
            return redirect(request.url)

        ext = grille_file.filename.rsplit('.', 1)[-1].lower()
        if ext not in ('xls', 'xlsx'):
            flash('Format non supporté. Utilisez .xls ou .xlsx', 'error')
            return redirect(request.url)

        import tempfile
        with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{ext}') as tmp:
            grille_file.save(tmp.name)
            tmp_path = tmp.name

        try:
            courses, students, meta = _parse_grille(tmp_path, ext)
        except Exception as exc:
            try: os.unlink(tmp_path)
            except: pass
            flash(f'Erreur de lecture du fichier : {exc}', 'error')
            return redirect(request.url)

        # Si 0 cours ou 0 étudiants : garder le fichier pour diagnostic
        if not students or not courses:
            debug_dir = os.path.join(current_app.root_path, 'attached_assets', 'debug_grilles')
            os.makedirs(debug_dir, exist_ok=True)
            import shutil as _shutil
            debug_path = os.path.join(debug_dir, f'debug_grille_{int(os.path.getmtime(tmp_path))}.{ext}')
            _shutil.copy2(tmp_path, debug_path)
            os.unlink(tmp_path)
            if not students:
                flash(f'⚠️ Aucun étudiant trouvé dans la grille (ligne en-tête détectée : ligne {meta.get("hdr_row",0)+1}). '
                      f'Vérifiez que la grille respecte le format UNILU standard.', 'error')
            else:
                flash(f'⚠️ Aucun cours trouvé dans la grille ({len(students)} étudiants détectés). '
                      f'La ligne d\'en-tête a été trouvée en ligne {meta.get("hdr_row",0)+1}. '
                      f'Vérifiez que les cours commencent bien en colonne 5.', 'error')
            return redirect(request.url)

        os.unlink(tmp_path)

        # ── Avertir si aucune colonne pondérée n'a été détectée ──
        if courses and all(crs['pond_col'] is None for crs in courses):
            flash('⚠️ Aucune colonne pondérée détectée — les totaux pondérés seront 0.', 'warning')

        # ── Avertir si les colonnes de synthèse clés sont manquantes ──
        col_map = meta.get('col_map', {})
        missing_summary = [k for k in ('total', 'moy', 'decision') if k not in col_map]
        if missing_summary:
            labels = {'total': 'Total pondéré', 'moy': 'Moyenne', 'decision': 'Décision du jury'}
            missing_labels = ', '.join(labels[k] for k in missing_summary)
            flash(f'⚠️ Colonnes de synthèse non détectées ({missing_labels}) — '
                  f'totaux et décision du jury afficheront 0 / vide.', 'warning')

        # ── Avertir si des valeurs suspectes sont détectées ──
        nb_valeurs_suspectes = sum(
            1 for s in students
            if s.get('moyenne', 0) > 20 or s.get('moyenne', 0) < 0 or s.get('total', 0) < 0
        )
        if nb_valeurs_suspectes:
            flash(
                f'⚠️ {nb_valeurs_suspectes} étudiant(s) ont des valeurs suspectes '
                f'(moyenne > 20 ou valeur négative) — vérifiez l\'alignement des colonnes '
                f'avant de générer les bulletins.',
                'warning'
            )

        # ── Bloquer l'import si la colonne Moyenne est détectée mais toutes les valeurs sont 0 ──
        if 'moy' in col_map and students and all(s.get('moyenne', 0) == 0.0 for s in students):
            flash(
                '❌ La colonne Moyenne existe mais toutes les valeurs sont 0 — '
                'vérifiez que les résultats ont bien été saisis avant d\'importer.',
                'error'
            )
            return redirect(request.url)

        promo_final = promotion or meta.get('promo_dept', '')

        # ── Charger la liste officielle sans choisir silencieusement une autre promotion ──
        liste_promo, liste_rows = _charger_liste_identifiants_promotion(promo_final)
        if not liste_rows:
            flash(
                f'❌ Aucune liste d’identifiants unique ne correspond à « {promo_final} ». '
                'Importez d’abord la liste de cette promotion ou saisissez son nom exact.',
                'error')
            return redirect(request.url)

        # ── Préparer un brouillon avant toute écriture de bulletins ──
        assignments = _preparer_attributions_matricules(students, liste_rows)
        payload = {
            'students': students,
            'assignments': assignments,
            'identifiants': [
                {'matricule': row.matricule, 'nom': row.nom}
                for row in liste_rows
            ],
            'promotion': promo_final,
            'liste_promotion': liste_promo,
            'annee': annee,
            'session_acad': session_acad,
            'semestre': semestre_label,
            'montant_fc': montant_fc,
            'type_grille': request.form.get('type_grille', 'initial').strip().lower(),
            'departement': request.form.get('departement', '').strip(),
            'texte_intro': request.form.get('texte_intro', '').strip(),
            'action': request.form.get('action', 'generer'),
        }
        import secrets as _secrets
        token = _secrets.token_urlsafe(32)
        db.session.add(BulletinImportDraft(
            token=token,
            promotion=promo_final,
            payload_json=json.dumps(payload, ensure_ascii=False),
        ))
        db.session.commit()
        nb_confirmes = sum(a['statut'] == 'confirme' for a in assignments)
        nb_a_verifier = len(assignments) - nb_confirmes
        flash(
            f'Grille analysée : {nb_confirmes} correspondance(s) confirmée(s), '
            f'{nb_a_verifier} cas à vérifier. Aucun bulletin n’est encore enregistré.',
            'warning' if nb_a_verifier else 'success')
        return redirect(url_for('decanat_bulletins_import_review', token=token))

    sessions_list = BulletinSession.query.order_by(BulletinSession.date_import.desc()).all()
    return render_template('decanat_bulletins.html', sessions_list=sessions_list,
                           nb_nouveaux_suspects=_nb_suspects_nouveaux())


@app.route('/decanat/releves-cotes')
def decanat_releves_cotes():
    """Affiche les sessions disponibles pour les relevés officiels."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    sessions_list = BulletinSession.query.order_by(BulletinSession.date_import.desc()).all()
    releve_commandes = (
        ReleveCommande.query
        .order_by(ReleveCommande.date_commande.desc())
        .limit(100)
        .all()
    )
    return render_template(
        'decanat_releves_cotes.html',
        sessions_list=sessions_list,
        grid_type_labels={key: _releve_type_label(key) for key in _RELEVE_ENTETE_DEFAULTS},
        releve_commandes=releve_commandes,
    )


@app.route('/releves', methods=['GET', 'POST'])
def releves_portail():
    """Portail étudiant des relevés : accès uniquement par matricule."""
    if request.method == 'POST':
        matricule = request.form.get('matricule', '').strip().upper()
        if not matricule:
            flash('Veuillez saisir votre matricule.', 'error')
            return redirect(url_for('releves_portail'))
        if not _releves_public_accessible():
            return render_template(
                'releves_portail.html',
                mode='unavailable',
                matricule=matricule,
            ), 200
        resultats = BulletinData.query.filter(
            db.func.upper(BulletinData.matricule) == matricule
        ).join(BulletinSession).order_by(
            BulletinSession.date_import.desc()
        ).all()
        if not resultats:
            flash(f'Aucun relevé trouvé pour le matricule « {matricule} ».', 'error')
            return redirect(url_for('releves_portail'))
        session['releves_matricule'] = matricule
        if len(resultats) == 1:
            return redirect(url_for(
                'releves_preview',
                bid=resultats[0].id,
                matricule=matricule,
            ))
        return render_template(
            'releves_portail.html',
            mode='list',
            resultats=resultats,
            matricule=matricule,
        )
    return render_template('releves_portail.html', mode='search')


def _releve_bd_for_student(bid, matricule):
    """Retourne le relevé seulement si le matricule fourni correspond."""
    bd = BulletinData.query.get_or_404(bid)
    supplied = (matricule or '').strip().upper()
    if not supplied or supplied != (bd.matricule or '').strip().upper():
        abort(403, description='Matricule incorrect pour ce relevé.')
    return bd


@app.route('/releves/preview/<int:bid>')
def releves_preview(bid):
    if not _releves_public_accessible():
        return render_template('releves_portail.html', mode='disabled'), 403
    matricule = (
        request.args.get('matricule', '').strip().upper()
        or session.get('releves_matricule', '')
    )
    bd = _releve_bd_for_student(bid, matricule)
    bs = bd.bul_session
    if bs is None:
        abort(404, description='Ce relevé historique n’est plus associé à une session.')
    etu = json.loads(bd.data_json or '{}')
    commande = ReleveCommande.query.filter_by(bulletin_id=bd.id).first()
    return render_template(
        'releves_portail.html',
        mode='preview',
        bd=bd,
        bs=bs,
        etu=etu,
        matricule=matricule,
        commande=commande,
        commande_ok=request.args.get('commande') == 'ok',
    )


@app.route('/releves/preview/<int:bid>/image')
def releves_preview_image(bid):
    """Fournit uniquement le rendu image protégé du modèle officiel."""
    if not _releves_public_accessible():
        abort(403, description='L’accès public aux relevés est temporairement désactivé.')
    matricule = (
        request.args.get('matricule', '').strip().upper()
        or session.get('releves_matricule', '')
    )
    bd = _releve_bd_for_student(bid, matricule)
    bs = bd.bul_session
    if bs is None:
        abort(404, description='Ce relevé historique n’est plus associé à une session.')
    etu = json.loads(bd.data_json or '{}')
    image = _rendre_releve_apercu_png(etu, bs)
    image.seek(0)
    response = send_file(
        image,
        as_attachment=False,
        download_name=f'Releve_apercu_{bd.id}.png',
        mimetype='image/png',
    )
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Content-Disposition'] = 'inline'
    return response


@app.route('/releves/commander', methods=['POST'])
def releves_commander():
    """Consomme un reçu REL et enregistre une commande de relevé."""
    if not _releves_public_accessible():
        return render_template('releves_portail.html', mode='disabled'), 403
    matricule = request.form.get('matricule', '').strip().upper()
    try:
        bid = int(request.form.get('bid', ''))
    except (TypeError, ValueError):
        abort(400, description='Relevé non identifié.')
    bd = _releve_bd_for_student(bid, matricule)
    bs = bd.bul_session
    etu = json.loads(bd.data_json or '{}')
    existing = ReleveCommande.query.filter_by(bulletin_id=bd.id).first()
    if existing:
        return redirect(url_for(
            'releves_preview', bid=bd.id, matricule=matricule, commande='ok'
        ))

    raw_code = request.form.get('receipt_code', '').strip()
    lookup = _normaliser_recherche_recu(raw_code)
    recu = RecuPaiement.query.filter(
        db.func.upper(RecuPaiement.code_qr) == lookup
    ).with_for_update().first() if lookup else None
    if recu is None and lookup:
        recu = RecuPaiement.query.filter(
            db.func.upper(RecuPaiement.numero) == lookup
        ).with_for_update().first()
    error = None
    if not recu:
        error = 'Reçu non reconnu. Scannez un reçu officiel de relevé de cotes (REL-).'
    elif not _receipt_type_is_consistent(recu) or _receipt_type_from_number(recu.numero) != 'releve':
        error = 'Ce reçu ne correspond pas à une commande de relevé de cotes.'
    elif recu.utilise:
        error = 'Ce reçu a déjà été utilisé. Présentez un reçu REL- encore disponible.'
    if error:
        return render_template(
            'releves_portail.html',
            mode='preview', bd=bd, bs=bs, etu=etu, matricule=matricule,
            commande=None, commande_error=error,
        ), 400

    commande = ReleveCommande(
        bulletin_id=bd.id,
        recu_id=recu.id,
        matricule=matricule,
        nom_etudiant=bd.nom or '',
        statut='soumise',
    )
    recu.utilise = True
    recu.date_utilisation = now_cat()
    recu.matricule_etudiant = matricule
    recu.nom_etudiant = bd.nom
    db.session.add(commande)
    db.session.add(ScanLog(
        code=recu.code_qr,
        ip=request.headers.get('X-Forwarded-For', request.remote_addr or '').split(',')[0].strip(),
        resultat='releve_commande',
        matricule=matricule,
        nom_etudiant=bd.nom,
    ))
    _enregistrer_admin_audit(
        'commande_releve_cotes',
        {'bulletin_id': bd.id, 'recu_id': recu.id, 'matricule': matricule},
    )
    db.session.commit()
    return redirect(url_for(
        'releves_preview', bid=bd.id, matricule=matricule, commande='ok'
    ))


@app.route('/decanat/releves-cotes/commandes/<int:commande_id>/statut', methods=['POST'])
def decanat_releve_commande_statut(commande_id):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    commande = ReleveCommande.query.get_or_404(commande_id)
    statut = request.form.get('statut', 'soumise').strip()
    if statut not in {'soumise', 'en_traitement', 'prete', 'remise', 'annulee'}:
        statut = 'soumise'
    commande.statut = statut
    commande.note_decanat = request.form.get('note_decanat', '').strip() or None
    commande.date_traitement = now_cat()
    _enregistrer_admin_audit(
        'statut_commande_releve',
        {'commande_id': commande.id, 'statut': statut},
    )
    db.session.commit()
    return redirect(url_for('decanat_releves_cotes') + '#commandes-releves')


@app.route('/decanat/releves-cotes/session/<int:sid>/pdf/<int:bid>')
def decanat_releve_cotes_pdf(sid, bid):
    """Télécharge le relevé d'un étudiant d'une session donnée."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    bs = BulletinSession.query.get_or_404(sid)
    bd = BulletinData.query.filter_by(id=bid, session_id=sid).first_or_404()
    etu = json.loads(bd.data_json or '{}')
    docx_buf = _generer_releve_cotes_modele(etu, bs)
    _enregistrer_admin_audit(
        'generation_releve_cotes',
        {'session_id': sid, 'bulletin_id': bid, 'format': 'docx_officiel'},
    )
    db.session.commit()
    safe = re.sub(r'[^A-Za-z0-9_-]+', '_', bd.nom or 'Etudiant').strip('_')[:55]
    return send_file(
        docx_buf, as_attachment=True,
        download_name=f'Releve_de_cotes_{safe or "Etudiant"}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


@app.route('/decanat/releves-cotes/session/<int:sid>/apercu/<int:bid>')
def decanat_releve_cotes_apercu(sid, bid):
    """Ouvre le relevé officiel Word d'un étudiant."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    bs = BulletinSession.query.get_or_404(sid)
    bd = BulletinData.query.filter_by(id=bid, session_id=sid).first_or_404()
    etu = json.loads(bd.data_json or '{}')
    _enregistrer_admin_audit(
        'apercu_releve_cotes',
        {'session_id': sid, 'bulletin_id': bid, 'format': 'docx_officiel'},
    )
    db.session.commit()
    return send_file(
        _generer_releve_cotes_modele(etu, bs),
        as_attachment=False,
        download_name=f'Releve_de_cotes_{bd.id}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


@app.route('/decanat/releves-cotes/session/<int:sid>/zip')
def decanat_releves_cotes_zip(sid):
    """Télécharge tous les relevés d'une session dans une archive ZIP."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    bs = BulletinSession.query.get_or_404(sid)
    zip_buf = BytesIO()
    used_names = set()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as archive:
        for bd in BulletinData.query.filter_by(session_id=sid).order_by(BulletinData.nom).all():
            etu = json.loads(bd.data_json or '{}')
            docx_buf = _generer_releve_cotes_modele(etu, bs)
            base = re.sub(r'[^A-Za-z0-9_-]+', '_', bd.nom or f'Etudiant_{bd.id}').strip('_')[:55]
            base = base or f'Etudiant_{bd.id}'
            filename = f'Releve_de_cotes_{base}.docx'
            suffix = 2
            while filename in used_names:
                filename = f'Releve_de_cotes_{base}_{suffix}.docx'
                suffix += 1
            used_names.add(filename)
            archive.writestr(filename, docx_buf.read())
    _enregistrer_admin_audit(
        'generation_releves_cotes_lot',
        {
            'session_id': sid,
            'nombre_releves': len(used_names),
            'format': 'zip_docx_officiel',
        },
    )
    db.session.commit()
    zip_buf.seek(0)
    promo = re.sub(r'[^A-Za-z0-9_-]+', '_', bs.promotion or f'session_{sid}').strip('_')
    return send_file(
        zip_buf, as_attachment=True,
        download_name=f'Releves_de_cotes_officiels_{promo or sid}.zip',
        mimetype='application/zip',
    )


@app.route('/decanat/bulletins/import-review/<token>', methods=['GET', 'POST'])
def decanat_bulletins_import_review(token):
    """Revue obligatoire des matricules avant création des bulletins."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    draft = BulletinImportDraft.query.filter_by(token=token).first_or_404()
    try:
        payload = json.loads(draft.payload_json)
        assignments = payload['assignments']
        students = payload['students']
        identifiants = payload['identifiants']
    except (TypeError, ValueError, KeyError, json.JSONDecodeError):
        db.session.delete(draft)
        db.session.commit()
        flash('Brouillon d’import invalide ou incomplet.', 'error')
        return redirect(url_for('decanat_bulletins'))

    official_ids = {
        item['matricule'] for item in identifiants
        if item.get('matricule')
    }
    confirmes = sum(a.get('statut') == 'confirme' for a in assignments)
    a_verifier = len(assignments) - confirmes

    if request.method == 'GET':
        return render_template(
            'decanat_bulletins_import_review.html',
            draft=draft,
            payload=payload,
            assignments=assignments,
            identifiants=identifiants,
            confirmes=confirmes,
            a_verifier=a_verifier,
        )

    if request.form.get('action') == 'annuler':
        db.session.delete(draft)
        db.session.commit()
        flash('Import de la grille annulé. Aucun bulletin n’a été modifié.', 'info')
        return redirect(url_for('decanat_bulletins'))

    matricules = []
    erreurs = []
    for assignment in assignments:
        value = request.form.get(
            f"matricule_{assignment['index']}", ''
        ).strip().upper()
        matricules.append(value)
        if not value:
            erreurs.append(assignment['nom'])
        elif value not in official_ids:
            erreurs.append(f"{assignment['nom']} ({value})")

    if erreurs:
        flash(
            'Chaque étudiant doit recevoir un matricule de la liste officielle. '
            'À corriger : ' + ', '.join(erreurs[:8])
            + ('…' if len(erreurs) > 8 else ''),
            'error')
        return render_template(
            'decanat_bulletins_import_review.html',
            draft=draft,
            payload=payload,
            assignments=assignments,
            identifiants=identifiants,
            confirmes=confirmes,
            a_verifier=a_verifier,
            submitted=dict(request.form),
        )

    try:
        bs = _enregistrer_grille_validee(payload, matricules)
        db.session.delete(draft)
        db.session.commit()
    except ValueError as exc:
        db.session.rollback()
        flash(f'Validation refusée : {exc}', 'error')
        return render_template(
            'decanat_bulletins_import_review.html',
            draft=draft,
            payload=payload,
            assignments=assignments,
            identifiants=identifiants,
            confirmes=confirmes,
            a_verifier=a_verifier,
            submitted=dict(request.form),
        )
    except Exception:
        db.session.rollback()
        app.logger.exception('Validation import matricules échouée')
        flash('Erreur technique : aucun bulletin n’a été enregistré.', 'error')
        return render_template(
            'decanat_bulletins_import_review.html',
            draft=draft,
            payload=payload,
            assignments=assignments,
            identifiants=identifiants,
            confirmes=confirmes,
            a_verifier=a_verifier,
            submitted=dict(request.form),
        )

    flash(
        f'✅ {len(students)} bulletin(s) enregistré(s) avec les matricules validés.',
        'success')
    if payload.get('action') == 'generer_zip':
        zip_buf, filename = _generer_zip_session(bs)
        return send_file(
            zip_buf, as_attachment=True, download_name=filename,
            mimetype='application/zip')
    return redirect(url_for('decanat_bulletins_sessions'))


@app.route('/decanat/listes-identifiants', methods=['GET', 'POST'])
def decanat_listes_identifiants():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    if request.method == 'POST':
        fichiers  = request.files.getlist('fichiers')
        promo_man = request.form.get('promotion_manuelle', '').strip()

        import tempfile as _tmp
        total_importes = 0
        erreurs = []

        for f in fichiers:
            if not f or f.filename == '':
                continue
            if not f.filename.lower().endswith('.pdf'):
                erreurs.append(f'{f.filename} : format non supporté (PDF uniquement)')
                continue
            with _tmp.NamedTemporaryFile(delete=False, suffix='.pdf') as t:
                f.save(t.name)
                tmp_path = t.name
            try:
                promo_pdf, etudiants = _parse_liste_identifiants(tmp_path)
            except Exception as exc:
                erreurs.append(f'{f.filename} : erreur de lecture — {exc}')
                try: os.unlink(tmp_path)
                except: pass
                continue
            finally:
                try: os.unlink(tmp_path)
                except: pass

            promo_finale = promo_man or promo_pdf
            if not promo_finale:
                erreurs.append(f'{f.filename} : promotion introuvable dans le PDF')
                continue
            if not etudiants:
                erreurs.append(f'{f.filename} : aucun étudiant trouvé')
                continue

            # Supprimer les anciens enregistrements pour cette promotion si demandé
            if request.form.get('remplacer') == '1':
                ListeIdentifiants.query.filter_by(promotion=promo_finale).delete()

            for etu in etudiants:
                nom_norm = _normaliser_nom(etu['nom'])
                # Ne pas dupliquer
                existe = ListeIdentifiants.query.filter_by(
                    promotion=promo_finale,
                    nom_norm=nom_norm,
                    matricule=etu['matricule'],
                ).first()
                if not existe:
                    db.session.add(ListeIdentifiants(
                        promotion=promo_finale,
                        nom=etu['nom'],
                        nom_norm=nom_norm,
                        matricule=etu['matricule'],
                        mot_de_passe=etu.get('mot_de_passe', ''),
                    ))
                    total_importes += 1

        db.session.commit()

        if erreurs:
            for e in erreurs:
                flash(f'⚠️ {e}', 'warning')
        if total_importes:
            flash(f'✅ {total_importes} identifiant(s) importé(s) avec succès.', 'success')
        return redirect(url_for('decanat_listes_identifiants'))

    # GET — liste groupée par promotion
    from sqlalchemy import func as _func
    promos = db.session.query(
        ListeIdentifiants.promotion,
        _func.count(ListeIdentifiants.id).label('nb')
    ).group_by(ListeIdentifiants.promotion).order_by(ListeIdentifiants.promotion).all()

    promo_sel = request.args.get('promo', '')
    etudiants_promo = []
    if promo_sel:
        etudiants_promo = ListeIdentifiants.query.filter_by(
            promotion=promo_sel
        ).order_by(ListeIdentifiants.nom).all()

    return render_template('decanat_listes.html',
                           promos=promos,
                           promo_sel=promo_sel,
                           etudiants_promo=etudiants_promo)


@app.route('/decanat/listes-identifiants/supprimer', methods=['POST'])
def decanat_listes_supprimer():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    promo = request.form.get('promotion', '')
    if promo:
        nb = ListeIdentifiants.query.filter_by(promotion=promo).delete()
        db.session.commit()
        flash(f'🗑️ {nb} identifiant(s) supprimé(s) pour « {promo} ».', 'success')
    return redirect(url_for('decanat_listes_identifiants'))


@app.route('/decanat/bulletins/diagnostique', methods=['GET', 'POST'])
def decanat_bulletins_diagnostique():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    result = None
    if request.method == 'POST':
        f = request.files.get('grille')
        if not f or f.filename == '':
            flash('Sélectionnez un fichier.', 'error')
            return redirect(request.url)
        ext = f.filename.rsplit('.', 1)[-1].lower()
        if ext not in ('xls', 'xlsx'):
            flash('Format non supporté.', 'error')
            return redirect(request.url)
        import tempfile as _tf
        with _tf.NamedTemporaryFile(delete=False, suffix=f'.{ext}') as tmp:
            f.save(tmp.name)
            tmp_path = tmp.name
        try:
            # Lire les 8 premières lignes brutes
            raw_rows = []
            if ext == 'xls':
                import xlrd
                wb = xlrd.open_workbook(tmp_path)
                ws = wb.sheet_by_index(0)
                for r in range(min(10, ws.nrows)):
                    raw_rows.append([ws.cell_value(r, c) for c in range(min(20, ws.ncols))])
            else:
                from openpyxl import load_workbook as _lw
                wb = _lw(tmp_path, data_only=True)
                ws = wb.active
                for r in range(1, min(11, ws.max_row + 1)):
                    raw_rows.append([ws.cell(row=r, column=c).value for c in range(1, min(21, ws.max_column + 1))])

            courses, students, meta = _parse_grille(tmp_path, ext)
            os.unlink(tmp_path)
            # Préparer un aperçu : 3 premiers étudiants avec cours tronqués
            sample = []
            for s in students[:3]:
                moy   = round(s['moyenne'], 2)
                total = s.get('total', 0)
                sample.append({
                    'nom': s['nom'],
                    'matricule': s['matricule'],
                    'moyenne': moy,
                    'total': total,
                    'decision': s['decision'],
                    'nb_cours': len(s['cours']),
                    'cours3': [f"{c['name'][:30]} ({c['note']})" for c in s['cours'][:3]],
                    'moy_suspect':   moy > 20 or moy < 0,
                    'total_suspect': total < 0,
                })
            _col_map = meta.get('col_map', {})
            _summary_keys = ('total', 'moy', 'pct', 'cr_val', 'cr_nval', 'echecs', 'appre', 'decision')
            _summary_labels = {
                'total':   'Total pondéré',
                'moy':     'Moyenne',
                'pct':     'Pourcentage',
                'cr_val':  'Crédits validés',
                'cr_nval': 'Crédits non validés',
                'echecs':  'Nb échecs',
                'appre':   'Appréciation',
                'decision':'Décision',
            }
            summary_cols_found   = [_summary_labels[k] for k in _summary_keys if k in _col_map]
            summary_cols_missing = [_summary_labels[k] for k in ('total', 'moy', 'decision') if k not in _col_map]
            nb_valeurs_suspectes_diag = sum(
                1 for s in students
                if s.get('moyenne', 0) > 20 or s.get('moyenne', 0) < 0 or s.get('total', 0) < 0
            )
            all_moy_zero = bool(
                'moy' in _col_map
                and students
                and all(s.get('moyenne', 0) == 0.0 for s in students)
            )
            result = {
                'meta': meta,
                'nb_cours': len(courses),
                'nb_etudiants': len(students),
                'cours_sample': [c['name'] for c in courses[:6]],
                'students_sample': sample,
                'raw_rows': raw_rows,
                'no_pondere': bool(courses and all(crs['pond_col'] is None for crs in courses)),
                'summary_cols_found':   summary_cols_found,
                'summary_cols_missing': summary_cols_missing,
                'nb_valeurs_suspectes': nb_valeurs_suspectes_diag,
                'all_moy_zero': all_moy_zero,
            }
        except Exception as exc:
            try: os.unlink(tmp_path)
            except: pass
            flash(f'Erreur : {exc}', 'error')
    return render_template('decanat_bulletins_diagnostique.html', result=result)


@app.route('/decanat/bulletins/sessions')
def decanat_bulletins_sessions():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    sessions_list = BulletinSession.query.order_by(BulletinSession.date_import.desc()).all()
    zero_moy = _zero_moyenne_counts([bs.id for bs in sessions_list])
    return render_template('decanat_bulletins.html', sessions_list=sessions_list,
                           zero_moy=zero_moy,
                           nb_nouveaux_suspects=_nb_suspects_nouveaux())


@app.route('/decanat/bulletins/dashboard')
def decanat_bulletins_dashboard():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    # ── Filtres GET ──
    filtre_dept  = request.args.get('dept',  '').strip()
    filtre_promo = request.args.get('promo', '').strip()
    filtre_date  = request.args.get('date',  '').strip()   # YYYY-MM-DD
    filtre_date_mode = request.args.get('date_mode', '').strip()
    filtre_date_debut = request.args.get('date_debut', '').strip()
    filtre_date_fin = request.args.get('date_fin', '').strip()
    filtre_session = request.args.get('session', '').strip().lower()

    # Totaux globaux
    total_bulletins  = BulletinData.query.count()
    total_payes      = BulletinData.query.filter_by(paye=True).count()
    total_telecharge = db.session.query(db.func.sum(BulletinData.nb_telechargements)).scalar() or 0
    recettes         = db.session.query(db.func.sum(BulletinData.montant_paye)).scalar() or 0
    _np_q = (
        db.session.query(BulletinData)
        .join(BulletinSession, BulletinData.session_id == BulletinSession.id)
        .filter(BulletinData.paye == False)
    )
    if filtre_dept:
        _np_q = _np_q.filter(BulletinSession.departement == filtre_dept)
    if filtre_promo:
        _np_q = _np_q.filter(BulletinSession.promotion == filtre_promo)
    non_payes = _np_q.order_by(BulletinSession.promotion, BulletinData.nom).all()

    # ── Tableau détaillé par département / promotion ──
    sessions_list = BulletinSession.query.order_by(BulletinSession.date_import.desc()).all()
    suivi_etudiants = _construire_suivi_etudiants(sessions_list, filtre_promo or None)

    # Listes distinctes pour les filtres
    depts_dispo  = sorted({bs.departement for bs in sessions_list if bs.departement})
    promos_dispo = sorted({bs.promotion   for bs in sessions_list if bs.promotion})

    # Appliquer filtres sur les sessions
    sessions_filtrees = sessions_list
    if filtre_dept:
        sessions_filtrees = [s for s in sessions_filtrees if s.departement == filtre_dept]
    if filtre_promo:
        sessions_filtrees = [s for s in sessions_filtrees if s.promotion == filtre_promo]
    sessions_filtrees = [
        s for s in sessions_filtrees
        if _session_correspond_au_filtre(s, filtre_session)
    ]

    # Construire les lignes de stats (avec filtre date optionnel)
    today_str = now_cat().strftime('%Y-%m-%d')
    if not filtre_date_mode:
        filtre_date_mode = 'exact' if filtre_date else 'today'
    date_filter = _normaliser_filtres_date_paiements(
        filtre_date_mode, filtre_date, filtre_date_debut, filtre_date_fin,
        today_str=today_str,
    )
    date_cible = date_filter['label']

    stats_rows = []
    for bs in sessions_filtrees:
        bds = bs.bulletins
        total_s   = len(bds)
        payes_s   = [b for b in bds if b.paye]
        nb_payes  = len(payes_s)

        # Payés dans la période sélectionnée
        nb_payes_date = sum(
            1 for b in payes_s
            if _paiement_dans_periode(b.date_paiement, date_filter)
        )
        recettes_s = sum(b.montant_paye or 0 for b in payes_s)

        stats_rows.append({
            'bs':             bs,
            'total':          total_s,
            'payes':          nb_payes,
            'payes_date':     nb_payes_date,
            'en_attente':     total_s - nb_payes,
            'recettes':       recettes_s,
            'date_cible':     date_cible,
        })

    # ── Filtres scans suspects ──
    filtre_scan_type        = request.args.get('scan_type',       '').strip()
    filtre_scan_date_debut  = request.args.get('scan_date_debut', '').strip()
    filtre_scan_date_fin    = request.args.get('scan_date_fin',   '').strip()
    show_dismissed          = request.args.get('show_dismissed',  '') == '1'

    sq = ScanLog.query.filter(ScanLog.resultat.in_(['invalide', 'deja_utilise']))

    if not show_dismissed:
        sq = sq.filter(ScanLog.dismissed != True)

    if filtre_scan_type in ('invalide', 'deja_utilise'):
        sq = sq.filter(ScanLog.resultat == filtre_scan_type)

    if filtre_scan_date_debut:
        try:
            sq = sq.filter(ScanLog.date_scan >= datetime.strptime(filtre_scan_date_debut, '%Y-%m-%d'))
        except ValueError:
            pass

    if filtre_scan_date_fin:
        try:
            sq = sq.filter(ScanLog.date_scan < datetime.strptime(filtre_scan_date_fin, '%Y-%m-%d') + timedelta(days=1))
        except ValueError:
            pass

    scans_suspects = sq.order_by(ScanLog.date_scan.desc()).limit(50).all()
    nb_suspects_total = sq.count()

    # Badge "non lus" : compter les suspects depuis le dernier acquittement explicite
    nb_nouveaux_suspects = _nb_suspects_nouveaux()

    # Adresse e-mail d'alerte actuelle (DB ou env var)
    alert_email_current = _get_alert_email()

    # Paramètres SMTP actuels (DB ou env var) – le mot de passe n'est jamais affiché
    smtp_host_current, smtp_port_current, smtp_user_current, _, smtp_mode_current = _get_smtp_config()

    return render_template('decanat_bulletins_dashboard.html',
        total_bulletins=total_bulletins,
        total_payes=total_payes,
        total_telecharge=total_telecharge,
        recettes=recettes,
        non_payes=non_payes,
        sessions_list=sessions_list,
        stats_rows=stats_rows,
        suivi_etudiants=suivi_etudiants,
        depts_dispo=depts_dispo,
        promos_dispo=promos_dispo,
        filtre_dept=filtre_dept,
        filtre_promo=filtre_promo,
        filtre_date=filtre_date,
        filtre_date_mode=filtre_date_mode,
        filtre_date_debut=filtre_date_debut,
        filtre_date_fin=filtre_date_fin,
        filtre_session=filtre_session,
        date_filter=date_filter,
        today_str=today_str,
        scans_suspects=scans_suspects,
        nb_suspects_total=nb_suspects_total,
        nb_nouveaux_suspects=nb_nouveaux_suspects,
        filtre_scan_type=filtre_scan_type,
        filtre_scan_date_debut=filtre_scan_date_debut,
        filtre_scan_date_fin=filtre_scan_date_fin,
        show_dismissed=show_dismissed,
        alert_email_current=alert_email_current,
        smtp_host_current=smtp_host_current,
        smtp_port_current=smtp_port_current,
        smtp_user_current=smtp_user_current,
        smtp_mode_current=smtp_mode_current,
    )


@app.route('/decanat/bulletins/scans/<int:scan_id>/dismiss', methods=['POST'])
def decanat_scan_dismiss(scan_id):
    """Renvoi individuel d'un scan suspect (le cacher de la vue par défaut)."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    scan = ScanLog.query.get_or_404(scan_id)
    try:
        scan.dismissed = True
        db.session.commit()
        flash('Scan renvoyé.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erreur : {e}', 'error')
    # Preserve current filter state on redirect
    return redirect(url_for('decanat_bulletins_dashboard',
        dept=request.form.get('dept', ''),
        promo=request.form.get('promo', ''),
        date=request.form.get('date', ''),
        scan_type=request.form.get('scan_type', ''),
        scan_date_debut=request.form.get('scan_date_debut', ''),
        scan_date_fin=request.form.get('scan_date_fin', ''),
        show_dismissed=request.form.get('show_dismissed', ''),
    ))


@app.route('/decanat/bulletins/scans/acknowledge', methods=['POST'])
def decanat_scans_acknowledge():
    """Marque tous les scans suspects actuels comme vus (acquittement explicite)."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    try:
        now_str = now_cat().isoformat()
        cfg = AppConfig.query.filter_by(key='scans_last_acknowledged').first()
        if cfg:
            cfg.value = now_str
        else:
            cfg = AppConfig(key='scans_last_acknowledged', value=now_str)
            db.session.add(cfg)
        db.session.commit()
        flash('Tous les scans suspects ont été marqués comme vus.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erreur lors de l\'acquittement : {e}', 'error')
    return redirect(url_for('decanat_bulletins_dashboard'))


@app.route('/decanat/bulletins/dashboard/settings', methods=['POST'])
def decanat_dashboard_settings():
    """Enregistre les paramètres de configuration du tableau de bord (e-mail d'alerte, SMTP, etc.)."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    def _upsert(key, value):
        cfg = AppConfig.query.filter_by(key=key).first()
        if cfg:
            cfg.value = value
        else:
            cfg = AppConfig(key=key, value=value)
            db.session.add(cfg)

    try:
        # ── Adresse e-mail d'alerte ──
        new_email = request.form.get('alert_email', '').strip()
        _upsert('alert_email', new_email)

        # ── Paramètres SMTP ──
        smtp_host = request.form.get('smtp_host', '').strip()
        smtp_port = request.form.get('smtp_port', '').strip()
        smtp_user = request.form.get('smtp_user', '').strip()
        smtp_pass = request.form.get('smtp_password', '').strip()
        smtp_mode = request.form.get('smtp_mode', 'starttls').strip()
        if smtp_mode not in ('starttls', 'ssl', 'none'):
            smtp_mode = 'starttls'

        _upsert('smtp_host', smtp_host)
        _upsert('smtp_port', smtp_port)
        _upsert('smtp_user', smtp_user)
        _upsert('smtp_mode', smtp_mode)
        # Only overwrite password if a new one was submitted (blank = keep existing)
        # Encrypt before storing so the DB never holds plain text.
        if smtp_pass:
            _upsert('smtp_password', _encrypt_smtp_password(smtp_pass))

        db.session.commit()
        flash('Paramètres des alertes enregistrés avec succès.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erreur lors de la sauvegarde : {e}', 'error')

    return redirect(url_for('decanat_bulletins_dashboard'))


@app.route('/decanat/bulletins/dashboard/test-smtp', methods=['POST'])
def decanat_test_smtp():
    """Envoie un e-mail de test pour vérifier la configuration SMTP."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    smtp_host, smtp_port, smtp_user, smtp_pass, smtp_mode = _get_smtp_config()
    dest_email = _get_alert_email()

    if not (smtp_host and smtp_user and dest_email):
        flash('Configuration SMTP incomplète : veuillez renseigner l\'hôte, l\'utilisateur et l\'adresse de destination.', 'error')
        return redirect(url_for('decanat_bulletins_dashboard'))

    from email.mime.text import MIMEText
    from email.mime.multipart import MIMEMultipart

    try:
        corps_html = """<html><body style="font-family:Arial,sans-serif;color:#222;max-width:560px;">
        <h2 style="color:#198754;">✅ Test de configuration SMTP</h2>
        <p>Cet e-mail confirme que la configuration SMTP du tableau de bord DÉCANAT est correcte.</p>
        <p style="color:#888;font-size:.85rem;">Envoyé depuis le tableau de bord des bulletins académiques.</p>
        </body></html>"""

        msg = MIMEMultipart('alternative')
        msg['Subject'] = '✅ Test SMTP – Configuration e-mail DÉCANAT'
        msg['From']    = smtp_user
        msg['To']      = dest_email
        msg.attach(MIMEText(corps_html, 'html', 'utf-8'))

        _smtp_send(smtp_host, smtp_port, smtp_user, smtp_pass, smtp_mode, msg, dest_email)

        flash(f'E-mail de test envoyé avec succès à {dest_email}.', 'success')
    except Exception as e:
        flash(f'Échec de l\'envoi : {e}', 'error')

    return redirect(url_for('decanat_bulletins_dashboard'))


@app.route('/decanat/bulletins/scans-export-csv')
def decanat_scans_export_csv():
    """Exporte l'historique complet des tentatives de scan en CSV (filtrable).

    Le fichier est généré en streaming (par blocs de 500 lignes) afin d'éviter
    tout timeout ou débordement mémoire sur des journaux de plusieurs dizaines
    de milliers d'entrées.
    """
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    import csv
    from io import StringIO
    from flask import stream_with_context, Response

    # ── Filtres GET ──
    date_debut  = request.args.get('date_debut', '').strip()   # YYYY-MM-DD
    date_fin    = request.args.get('date_fin',   '').strip()   # YYYY-MM-DD
    filtre_type = request.args.get('type',       '').strip()   # '' | 'invalide' | 'deja_utilise' | 'ok'

    q = ScanLog.query.order_by(ScanLog.date_scan.desc())

    if filtre_type:
        q = q.filter(ScanLog.resultat == filtre_type)

    if date_debut:
        try:
            dt_debut = datetime.strptime(date_debut, '%Y-%m-%d')
            q = q.filter(ScanLog.date_scan >= dt_debut)
        except ValueError:
            pass

    if date_fin:
        try:
            dt_fin = datetime.strptime(date_fin, '%Y-%m-%d') + timedelta(days=1)
            q = q.filter(ScanLog.date_scan < dt_fin)
        except ValueError:
            pass

    LABELS = {
        'ok':           'OK',
        'invalide':     'Code inconnu',
        'deja_utilise': 'Déjà utilisé',
        'deja_paye':    'Déjà payé',
    }

    def generate():
        # En-tête avec BOM UTF-8 pour une ouverture correcte dans Excel
        si = StringIO()
        writer = csv.writer(si, delimiter=';', quoting=csv.QUOTE_MINIMAL)
        writer.writerow(['Date/Heure', 'Code scanné', 'Résultat', 'Matricule', 'Nom étudiant', 'IP source'])
        yield si.getvalue().encode('utf-8-sig')

        # Itérer par blocs de 500 lignes pour ne jamais tout charger en mémoire
        for log in q.yield_per(500):
            si = StringIO()
            writer = csv.writer(si, delimiter=';', quoting=csv.QUOTE_MINIMAL)
            writer.writerow([
                log.date_scan.strftime('%d/%m/%Y %H:%M:%S'),
                log.code,
                LABELS.get(log.resultat, log.resultat),
                log.matricule or '',
                log.nom_etudiant or '',
                log.ip or '',
            ])
            yield si.getvalue().encode('utf-8')

    filename = f"scans_suspects_{now_cat().strftime('%Y%m%d_%H%M%S')}.csv"
    headers = {
        'Content-Disposition': f'attachment; filename="{filename}"',
        'Content-Type': 'text/csv; charset=utf-8',
        'X-Accel-Buffering': 'no',   # désactive le buffering Nginx si présent
    }
    return Response(stream_with_context(generate()), headers=headers)


@app.route('/decanat/bulletins/non-payes-export-csv')
def decanat_non_payes_export_csv():
    """Exporte la liste des étudiants n'ayant pas encore payé en CSV.

    Paramètres GET optionnels :
      sid   – filtre par session (id exact) — prioritaire sur dept/promo
      dept  – filtre par département (valeur exacte)
      promo – filtre par promotion   (valeur exacte)
    """
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    import csv
    from io import StringIO

    filtre_sid   = request.args.get('sid',   '').strip()
    filtre_dept  = request.args.get('dept',  '').strip()
    filtre_promo = request.args.get('promo', '').strip()

    q = (
        db.session.query(BulletinData, BulletinSession)
        .join(BulletinSession, BulletinData.session_id == BulletinSession.id)
        .filter(BulletinData.paye == False)
    )
    if filtre_sid:
        try:
            q = q.filter(BulletinSession.id == int(filtre_sid))
        except ValueError:
            pass
    else:
        if filtre_dept:
            q = q.filter(BulletinSession.departement == filtre_dept)
        if filtre_promo:
            q = q.filter(BulletinSession.promotion == filtre_promo)

    rows = q.order_by(BulletinSession.promotion, BulletinData.nom).all()

    si = StringIO()
    writer = csv.writer(si, delimiter=';', quoting=csv.QUOTE_MINIMAL)
    writer.writerow([
        'Matricule', 'Nom', 'Département', 'Promotion',
        'Session', 'Montant dû (FC)'
    ])

    for bd, bs in rows:
        writer.writerow([
            bd.matricule or '',
            bd.nom or '',
            bs.departement or '',
            bs.promotion or '',
            bs.nom or '',
            bs.montant_fc or 0,
        ])

    output = si.getvalue().encode('utf-8-sig')  # BOM pour Excel
    buf = BytesIO(output)
    buf.seek(0)

    # Nom de fichier reflétant les filtres actifs
    parts = ['non_payes']
    if filtre_sid and rows:
        # Use session name when exporting a single session
        session_name = rows[0][1].nom or filtre_sid
        parts.append(re.sub(r'[^A-Za-z0-9]', '_', session_name))
    else:
        if filtre_dept:
            parts.append(re.sub(r'[^A-Za-z0-9]', '_', filtre_dept))
        if filtre_promo:
            parts.append(re.sub(r'[^A-Za-z0-9]', '_', filtre_promo))
    parts.append(now_cat().strftime('%Y%m%d'))
    filename = '_'.join(parts) + '.csv'

    return send_file(buf, mimetype='text/csv',
                     as_attachment=True, download_name=filename)


@app.route('/decanat/bulletins/zip-filtre')
def decanat_bulletins_zip_filtre():
    """Télécharger le ZIP des bulletins filtrés par département et/ou promotion."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    filtre_dept  = request.args.get('dept',  '').strip()
    filtre_promo = request.args.get('promo', '').strip()
    sid_param    = request.args.get('sid',   '').strip()   # session unique si fourni

    # Trouver les sessions correspondantes
    q = BulletinSession.query
    if sid_param:
        q = q.filter_by(id=int(sid_param))
    else:
        if filtre_dept:
            q = q.filter(BulletinSession.departement == filtre_dept)
        if filtre_promo:
            q = q.filter(BulletinSession.promotion == filtre_promo)

    sessions_sel = q.all()
    if not sessions_sel:
        flash('Aucune session trouvée avec ces critères.', 'error')
        return redirect(url_for('decanat_bulletins_dashboard'))

    zip_buf = BytesIO()
    total_gen = 0
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for bs in sessions_sel:
            bds = BulletinData.query.filter_by(session_id=bs.id).all()
            folder = bs.promotion.replace(' ', '_')[:30]
            for bd in bds:
                etu = json.loads(bd.data_json)
                pdf_buf = _generer_bulletin_pdf(
                    etu, bs.promotion, bs.annee,
                    bs.session_acad, bs.semestre,
                    numero_bulletin=bd.numero_bulletin,
                    texte_intro=bs.texte_intro,
                    departement=bs.departement,
                )
                safe = bd.nom.replace(' ', '_')[:35]
                zf.writestr(f'{folder}/Bulletin_{safe}.pdf', pdf_buf.read())
                total_gen += 1

    zip_buf.seek(0)
    label = filtre_promo or filtre_dept or 'Tous'
    return send_file(zip_buf, as_attachment=True,
                     download_name=f'Bulletins_{label.replace(" ","_")}.zip',
                     mimetype='application/zip')


# ─────────────────────────────────────────────────────────────────────────────
#  HELPER : récupérer les étudiants payés selon les filtres
# ─────────────────────────────────────────────────────────────────────────────
def _normaliser_filtres_date_paiements(
    mode='', date_precise='', date_debut='', date_fin='', today_str=None
):
    """Normalise les modes de date du dashboard et du rapport de paiements."""
    today_str = today_str or now_cat().strftime('%Y-%m-%d')
    modes = {'today', 'all', 'exact', 'range'}
    if mode not in modes:
        mode = 'exact' if date_precise else 'all'

    selection_mode = mode
    if mode == 'today':
        date_precise = today_str
        effective_mode = 'exact'
    else:
        effective_mode = mode

    def valid_date(value):
        try:
            datetime.strptime(value, '%Y-%m-%d')
            return value
        except (TypeError, ValueError):
            return ''

    date_precise = valid_date(date_precise)
    date_debut = valid_date(date_debut)
    date_fin = valid_date(date_fin)

    if effective_mode == 'exact':
        debut = fin = date_precise
    elif effective_mode == 'range':
        debut, fin = date_debut, date_fin
    else:
        debut = fin = ''

    def display(value):
        return datetime.strptime(value, '%Y-%m-%d').strftime('%d/%m/%Y')

    if not debut and not fin:
        label = 'Toutes les dates'
    elif debut and fin:
        label = (
            display(debut) if debut == fin
            else f'Du {display(debut)} au {display(fin)}'
        )
    elif debut:
        label = f'À partir du {display(debut)}'
    else:
        label = f"Jusqu'au {display(fin)}"

    return {
        'mode': selection_mode,
        'debut': debut,
        'fin': fin,
        'label': label,
        'slug': f'{debut or "toutes"}_{fin or ""}'.strip('_'),
    }


def _paiement_dans_periode(date_paiement, date_filter):
    if not date_paiement:
        return False
    paiement = date_paiement.strftime('%Y-%m-%d')
    return (
        (not date_filter['debut'] or paiement >= date_filter['debut'])
        and (not date_filter['fin'] or paiement <= date_filter['fin'])
    )


def _paiement_est_recours(bulletin_session):
    """Classe les bulletins issus d'une grille de recours séparément."""
    return getattr(bulletin_session, 'type_grille', 'initial') in (
        'recours', 'recours_session_2'
    )


def _rapport_stage_key(bulletin_session):
    """Retourne l'étape exacte d'un paiement de bulletin."""
    return getattr(bulletin_session, 'type_grille', 'initial') or 'initial'


_RAPPORT_STAGE_DETAILS = {
    'initial': {
        'label': 'Paiements des bulletins de 1ère session',
        'short_label': '1ère session',
        'order': 0,
        'recours': False,
    },
    'session_2': {
        'label': 'Paiements des bulletins de 2ème session',
        'short_label': '2ème session',
        'order': 1,
        'recours': False,
    },
    'recours': {
        'label': 'Paiements des bulletins de recours de 1ère session',
        'short_label': 'Recours 1ère session',
        'order': 2,
        'recours': True,
    },
    'recours_session_2': {
        'label': 'Paiements des bulletins de recours de 2ème session',
        'short_label': 'Recours 2ème session',
        'order': 3,
        'recours': True,
    },
}


def _session_correspond_au_filtre(bulletin_session, filtre_session=''):
    """Applique le filtre global de session au rapport de paiements."""
    type_grille = getattr(bulletin_session, 'type_grille', 'initial') or 'initial'
    if filtre_session == 'recours_1':
        return type_grille == 'recours'
    if filtre_session in ('initial', 'session_2', 'recours_session_2'):
        return type_grille == filtre_session
    if filtre_session == 'recours':
        return type_grille in ('recours', 'recours_session_2')
    if filtre_session == 'speciale':
        return type_grille in ('initial', 'session_2')
    return True


def _get_rapport_payes(
    filtre_dept='', filtre_promo='', filtre_date='',
    filtre_date_mode='', filtre_date_debut='', filtre_date_fin='',
    filtre_session=''
):
    """Retourne (lignes, meta) pour le rapport de paiements."""
    date_filter = _normaliser_filtres_date_paiements(
        filtre_date_mode, filtre_date, filtre_date_debut, filtre_date_fin
    )

    q = BulletinSession.query
    if filtre_dept:
        q = q.filter(BulletinSession.departement == filtre_dept)
    if filtre_promo:
        q = q.filter(BulletinSession.promotion == filtre_promo)
    sessions_sel = q.all()
    sessions_sel = [
        bs for bs in sessions_sel
        if _session_correspond_au_filtre(bs, filtre_session)
    ]

    lignes = []
    total_montant = 0
    for bs in sessions_sel:
        for bd in bs.bulletins:
            if not bd.paye:
                continue
            if not _paiement_dans_periode(bd.date_paiement, date_filter):
                continue
            lignes.append({
                'nom':        bd.nom,
                'matricule':  bd.matricule,
                'promotion':  bs.promotion,
                'departement': bs.departement or '—',
                'stage_key':  _rapport_stage_key(bs),
                'est_recours': _paiement_est_recours(bs),
                'date_paiement': bd.date_paiement.strftime('%d/%m/%Y %H:%M') if bd.date_paiement else '—',
                'montant':    bd.montant_paye or 0,
                'methode':    bd.methode_paiement or '—',
                'reference':  bd.reference_paiement or '—',
            })
            total_montant += bd.montant_paye or 0

    # tri : département > promotion > nom
    lignes.sort(key=lambda r: (r['departement'], r['promotion'], r['nom']))
    lignes_recours = [ligne for ligne in lignes if ligne['est_recours']]
    lignes_ordinaires = [ligne for ligne in lignes if not ligne['est_recours']]

    meta = {
        'filtre_dept':  filtre_dept  or 'Tous',
        'filtre_promo': filtre_promo or 'Toutes',
        'filtre_session': {
            'speciale': 'Spéciale',
            'recours': 'Recours',
            'recours_1': 'Recours 1ère session',
            **{
                key: details['short_label']
                for key, details in _RAPPORT_STAGE_DETAILS.items()
            },
        }.get(filtre_session, 'Toutes les sessions'),
        'date_cible':   date_filter['label'],
        'date_slug':    date_filter['slug'],
        'nb_payes':     len(lignes),
        'total_montant': total_montant,
        'nb_recours':   len(lignes_recours),
        'total_recours': sum(ligne['montant'] for ligne in lignes_recours),
        'nb_ordinaires': len(lignes_ordinaires),
        'total_ordinaires': sum(ligne['montant'] for ligne in lignes_ordinaires),
        'genere_le':    now_cat().strftime('%d/%m/%Y à %H:%M'),
    }
    return lignes, meta


@app.route('/decanat/bulletins/rapport-paiements')
def decanat_bulletins_rapport_preview():
    """Prévisualisation HTML du rapport de paiements avant impression PDF."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    filtre_dept  = request.args.get('dept',  '').strip()
    filtre_promo = request.args.get('promo', '').strip()
    filtre_date  = request.args.get('date',  '').strip()
    filtre_date_mode = request.args.get('date_mode', '').strip()
    filtre_date_debut = request.args.get('date_debut', '').strip()
    filtre_date_fin = request.args.get('date_fin', '').strip()
    filtre_session = request.args.get('session', '').strip().lower()

    lignes, meta = _get_rapport_payes(
        filtre_dept, filtre_promo, filtre_date,
        filtre_date_mode, filtre_date_debut, filtre_date_fin, filtre_session,
    )

    # Regrouper séparément les quatre étapes de paiement.
    groupes = {}
    for l in lignes:
        key = (l['stage_key'], l['departement'], l['promotion'])
        groupes.setdefault(key, []).append(l)
    sections = []
    for stage_key, details in sorted(
        _RAPPORT_STAGE_DETAILS.items(),
        key=lambda item: item[1]['order'],
    ):
        stage_groupes = {
            (dept, promo): groupe
            for (key, dept, promo), groupe in groupes.items()
            if key == stage_key
        }
        stage_lignes = [
            ligne for ligne in lignes if ligne['stage_key'] == stage_key
        ]
        sections.append({
            'key': stage_key,
            'label': details['label'],
            'groupes': stage_groupes,
            'nb': len(stage_lignes),
            'total': sum(ligne['montant'] for ligne in stage_lignes),
            'recours': details['recours'],
        })

    sessions_list = BulletinSession.query.all()
    depts_dispo   = sorted({bs.departement for bs in sessions_list if bs.departement})
    promos_dispo  = sorted({bs.promotion   for bs in sessions_list if bs.promotion})

    return render_template('decanat_rapport_paiements.html',
        lignes=lignes, meta=meta, groupes=groupes, sections=sections,
        filtre_dept=filtre_dept, filtre_promo=filtre_promo, filtre_date=filtre_date,
        filtre_date_mode=filtre_date_mode or ('exact' if filtre_date else 'all'),
        filtre_date_debut=filtre_date_debut, filtre_date_fin=filtre_date_fin,
        filtre_session=filtre_session,
        depts_dispo=depts_dispo, promos_dispo=promos_dispo,
    )


@app.route('/decanat/bulletins/rapport-paiements/pdf')
def decanat_bulletins_rapport_pdf():
    """Génère et télécharge le PDF du rapport de paiements."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    filtre_dept  = request.args.get('dept',  '').strip()
    filtre_promo = request.args.get('promo', '').strip()
    filtre_date  = request.args.get('date',  '').strip()
    filtre_date_mode = request.args.get('date_mode', '').strip()
    filtre_date_debut = request.args.get('date_debut', '').strip()
    filtre_date_fin = request.args.get('date_fin', '').strip()
    filtre_session = request.args.get('session', '').strip().lower()

    lignes, meta = _get_rapport_payes(
        filtre_dept, filtre_promo, filtre_date,
        filtre_date_mode, filtre_date_debut, filtre_date_fin, filtre_session,
    )

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                     Paragraph, Spacer, HRFlowable)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

    MARINE = colors.Color(13/255, 38/255, 102/255)
    OR     = colors.Color(191/255, 153/255, 26/255)
    VERT   = colors.Color(0.0, 0.45, 0.1)
    GRIS_L = colors.Color(0.93, 0.93, 0.93)

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             topMargin=1.5*cm, bottomMargin=1.5*cm,
                             leftMargin=1.8*cm, rightMargin=1.8*cm)
    PAGE_W = A4[0] - 1.8*cm*2

    def P(txt, **kw):
        kw.setdefault('fontName', 'Helvetica')
        kw.setdefault('fontSize', 9)
        kw.setdefault('leading', 11)
        return Paragraph(txt, ParagraphStyle('_p', **kw))

    story = []

    # ── En-tête ──
    story.append(P('<b>UNIVERSITÉ DE LUBUMBASHI</b>',
                   fontName='Helvetica-Bold', fontSize=13, leading=16,
                   alignment=TA_CENTER, textColor=MARINE))
    story.append(P('Faculté des Sciences Sociales, Politiques et Administratives',
                   fontSize=9, leading=12, alignment=TA_CENTER, textColor=MARINE))
    story.append(Spacer(1, 0.2*cm))
    story.append(HRFlowable(width=PAGE_W, thickness=2, color=MARINE))
    story.append(Spacer(1, 0.3*cm))

    # ── Titre du rapport ──
    story.append(P('<b>RAPPORT DE PAIEMENTS — BULLETINS ACADÉMIQUES</b>',
                   fontName='Helvetica-Bold', fontSize=12, leading=15,
                   alignment=TA_CENTER, textColor=MARINE))
    story.append(Spacer(1, 0.25*cm))

    # ── Bloc de synthèse ──
    date_label = meta['date_cible']
    syn_data = [[
        P(f"<b>Département :</b> {meta['filtre_dept']}", fontSize=9),
        P(f"<b>Promotion :</b> {meta['filtre_promo']}", fontSize=9),
        P(f"<b>Date :</b> {date_label}", fontSize=9),
    ],[
        P(f"<b>Étudiants payés :</b> <font color='#0d7a2e'><b>{meta['nb_payes']}</b></font>",
          fontSize=10),
        P(f"<b>Total recettes :</b> <font color='#0d7a2e'><b>{meta['total_montant']:,} FC</b></font>",
          fontSize=10),
        P(f"<b>Généré le :</b> {meta['genere_le']}", fontSize=9, textColor=colors.grey),
    ]]
    syn_tbl = Table(syn_data, colWidths=[PAGE_W*0.34, PAGE_W*0.33, PAGE_W*0.33])
    syn_tbl.setStyle(TableStyle([
        ('BOX',           (0,0), (-1,-1), 0.8, MARINE),
        ('INNERGRID',     (0,0), (-1,-1), 0.3, colors.lightgrey),
        ('BACKGROUND',    (0,0), (-1, 0), colors.Color(0.93,0.96,1.0)),
        ('BACKGROUND',    (0,1), (-1, 1), colors.Color(0.87,0.95,0.90)),
        ('TOPPADDING',    (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING',   (0,0), (-1,-1), 6),
        ('VALIGN',        (0,0), (-1,-1), 'MIDDLE'),
    ]))
    story += [syn_tbl, Spacer(1, 0.5*cm)]

    if not lignes:
        story.append(P('Aucun paiement enregistré pour ces critères.',
                       alignment=TA_CENTER, textColor=colors.grey, fontSize=10))
    else:
        # Grouper par étape exacte, puis département > promotion.
        from itertools import groupby
        groupby_key = lambda r: (
            _RAPPORT_STAGE_DETAILS.get(
                r['stage_key'], _RAPPORT_STAGE_DETAILS['initial']
            )['order'],
            r['departement'],
            r['promotion'],
        )
        lignes_sorted = sorted(lignes, key=groupby_key)

        num = 1
        previous_section = None
        for (section_order, dept, promo), groupe in groupby(
            lignes_sorted, key=groupby_key
        ):
            groupe = list(groupe)
            sous_total = sum(r['montant'] for r in groupe)

            if section_order != previous_section:
                stage_key = next(
                    (
                        key for key, details in _RAPPORT_STAGE_DETAILS.items()
                        if details['order'] == section_order
                    ),
                    'initial',
                )
                stage_details = _RAPPORT_STAGE_DETAILS[stage_key]
                section_label = (
                    '🎓 ' if stage_details['recours'] else '📘 '
                ) + stage_details['label'].upper()
                section_count = sum(
                    1 for ligne in lignes
                    if ligne['stage_key'] == stage_key
                )
                section_total = sum(
                    ligne['montant'] for ligne in lignes
                    if ligne['stage_key'] == stage_key
                )
                if stage_details['recours']:
                    section_color = colors.Color(1.0, 0.95, 0.80)
                    section_text = colors.Color(0.50, 0.31, 0.0)
                else:
                    section_color = colors.Color(0.91, 0.95, 1.0)
                    section_text = MARINE
                section_hdr = Table([[
                    P(
                        f'<b>{section_label}</b> — {section_count} étudiant(s)'
                        f' · {section_total:,} FC',
                        fontName='Helvetica-Bold', fontSize=10,
                        textColor=section_text,
                    )
                ]], colWidths=[PAGE_W])
                section_hdr.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), section_color),
                    ('BOX', (0, 0), (-1, -1), 0.7, section_text),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('LEFTPADDING', (0, 0), (-1, -1), 8),
                ]))
                story += [section_hdr, Spacer(1, 0.2*cm)]
                previous_section = section_order

            # Titre du groupe
            story.append(P(f'<b>{dept} — {promo}</b>',
                           fontName='Helvetica-Bold', fontSize=9,
                           textColor=colors.white,
                           backColor=MARINE))
            # on ne peut pas mettre backColor sur Paragraph directement ;
            # on utilise un mini-tableau pour le fond coloré
            hdr_grp = Table([[
                P(f'<b>{dept} — {promo}</b>',
                  fontName='Helvetica-Bold', fontSize=9, textColor=colors.white),
            ]], colWidths=[PAGE_W])
            hdr_grp.setStyle(TableStyle([
                ('BACKGROUND',  (0,0),(-1,-1), MARINE),
                ('TOPPADDING',  (0,0),(-1,-1), 4),
                ('BOTTOMPADDING',(0,0),(-1,-1), 4),
                ('LEFTPADDING', (0,0),(-1,-1), 6),
            ]))
            story.append(hdr_grp)

            # En-tête colonnes
            th = [
                P('<b>N°</b>',         textColor=colors.white, alignment=TA_CENTER, fontSize=8),
                P('<b>Nom complet</b>', textColor=colors.white, fontSize=8),
                P('<b>Matricule</b>',   textColor=colors.white, fontSize=8),
                P('<b>Date paiement</b>',textColor=colors.white, alignment=TA_CENTER, fontSize=8),
                P('<b>Méthode</b>',     textColor=colors.white, alignment=TA_CENTER, fontSize=8),
                P('<b>Montant (FC)</b>',textColor=colors.white, alignment=TA_RIGHT, fontSize=8),
            ]
            rows = [th]
            row_styles = []

            for i, r in enumerate(groupe):
                bg = colors.white if i % 2 == 0 else colors.Color(0.95,0.97,1.0)
                row_styles.append(('BACKGROUND', (0, i+1), (-1, i+1), bg))
                rows.append([
                    P(str(num), fontSize=8, alignment=TA_CENTER),
                    P(r['nom'], fontSize=8),
                    P(r['matricule'], fontSize=8),
                    P(r['date_paiement'], fontSize=7, alignment=TA_CENTER),
                    P(r['methode'], fontSize=7, alignment=TA_CENTER),
                    P(f"<b>{r['montant']:,}</b>", fontSize=8, alignment=TA_RIGHT),
                ])
                num += 1

            # Ligne sous-total du groupe
            rows.append([
                P('', fontSize=8),
                P(f'<b>Sous-total : {len(groupe)} étudiant(s)</b>',
                  fontName='Helvetica-Bold', fontSize=8, textColor=MARINE),
                P('', fontSize=8), P('', fontSize=8), P('', fontSize=8),
                P(f'<b>{sous_total:,}</b>',
                  fontName='Helvetica-Bold', fontSize=8, textColor=VERT,
                  alignment=TA_RIGHT),
            ])
            row_styles.append(('BACKGROUND', (0, len(rows)-1), (-1, len(rows)-1),
                                colors.Color(0.90, 0.96, 0.90)))

            W = [0.7*cm, PAGE_W-0.7*cm-2.8*cm-3.0*cm-1.8*cm-2.2*cm, 2.8*cm, 3.0*cm, 1.8*cm, 2.2*cm]
            tbl = Table(rows, colWidths=W, repeatRows=1)
            tbl.setStyle(TableStyle([
                ('BACKGROUND',    (0,0), (-1, 0),  colors.Color(0.20,0.35,0.60)),
                ('VALIGN',        (0,0), (-1,-1),  'MIDDLE'),
                ('BOX',           (0,0), (-1,-1),  0.6, MARINE),
                ('INNERGRID',     (0,0), (-1,-1),  0.3, colors.Color(0.80,0.80,0.80)),
                ('TOPPADDING',    (0,0), (-1,-1),  2),
                ('BOTTOMPADDING', (0,0), (-1,-1),  2),
                ('LEFTPADDING',   (0,0), (-1,-1),  4),
                ('RIGHTPADDING',  (0,0), (-1,-1),  4),
            ] + row_styles))
            story += [tbl, Spacer(1, 0.3*cm)]

        # ── TOTAL GÉNÉRAL ──
        story.append(HRFlowable(width=PAGE_W, thickness=1.5, color=MARINE))
        story.append(Spacer(1, 0.15*cm))
        total_tbl = Table([[
            P(f'<b>TOTAL GÉNÉRAL : {meta["nb_payes"]} étudiants payés</b>',
              fontName='Helvetica-Bold', fontSize=10, textColor=MARINE),
            P(f'<b>{meta["total_montant"]:,} FC</b>',
              fontName='Helvetica-Bold', fontSize=11, textColor=VERT,
              alignment=TA_RIGHT),
        ]], colWidths=[PAGE_W*0.65, PAGE_W*0.35])
        total_tbl.setStyle(TableStyle([
            ('VALIGN',  (0,0),(-1,-1),'MIDDLE'),
            ('TOPPADDING',(0,0),(-1,-1),4),
            ('BOTTOMPADDING',(0,0),(-1,-1),4),
        ]))
        story.append(total_tbl)

    doc.build(story)
    buf.seek(0)
    fname = f'Rapport_paiements_{meta["date_slug"]}.pdf'
    return send_file(buf, as_attachment=True,
                     download_name=fname, mimetype='application/pdf')


@app.route('/decanat/bulletins/session/<int:sid>')
def decanat_bulletins_session_detail(sid):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    bs = BulletinSession.query.get_or_404(sid)
    bulletins = BulletinData.query.filter_by(session_id=sid).order_by(BulletinData.nom).all()
    # Audit trail : toutes les corrections de paiement pour cette session
    bulletin_ids = [bd.id for bd in bulletins]
    audits = (PaiementAudit.query
              .filter(PaiementAudit.bulletin_id.in_(bulletin_ids))
              .order_by(PaiementAudit.date_modification.desc())
              .limit(200)
              .all()) if bulletin_ids else []
    # Créer un dict nom_etudiant par bulletin_id pour l'affichage
    nom_par_id = {bd.id: bd.nom for bd in bulletins}
    # Compter les bulletins avec moyenne == 0.0 et colonne présente
    nb_zero_moyenne = sum(
        1 for bd in bulletins
        if (lambda e: e.get('moy_col_present') and e.get('moyenne', 1) == 0.0)(
            json.loads(bd.data_json) if bd.data_json else {}
        )
    )
    return render_template('decanat_bulletins_session.html',
                           bs=bs, bulletins=bulletins,
                           audits=audits, nom_par_id=nom_par_id,
                           nb_zero_moyenne=nb_zero_moyenne)


@app.route('/decanat/bulletins/session/<int:sid>/ajouter-manuel', methods=['POST'])
def decanat_bulletins_ajouter_manuel(sid):
    """Ajoute un étudiant absent de la grille après validation de son reçu."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    bs = BulletinSession.query.get_or_404(sid)
    nom = request.form.get('nom', '').strip()
    matricule = request.form.get('matricule', '').strip().upper()
    numero_recu = request.form.get('numero_recu', '').strip().upper()
    date_str = request.form.get('date_paiement', '').strip()
    departement_form = request.form.get('departement', '').strip()
    promotion_form = request.form.get('promotion', '').strip()
    annee_form = request.form.get('annee', '').strip()
    session_form = request.form.get('session', '').strip()
    montant_form = request.form.get('montant', '').strip()

    if not nom or not matricule or not numero_recu:
        flash('Nom complet, matricule et numéro du reçu sont obligatoires.', 'error')
        return redirect(url_for('decanat_bulletins_session_detail', sid=sid))

    try:
        date_paiement = (
            datetime.strptime(date_str, '%Y-%m-%d') if date_str else now_cat()
        )
    except ValueError:
        flash('La date de paiement est invalide.', 'error')
        return redirect(url_for('decanat_bulletins_session_detail', sid=sid))

    try:
        if departement_form != (bs.departement or ''):
            raise ValueError('Le département envoyé ne correspond pas à cette session.')
        if promotion_form != (bs.promotion or ''):
            raise ValueError('La promotion envoyée ne correspond pas à cette session.')
        if annee_form != (bs.annee or ''):
            raise ValueError("L'année académique envoyée ne correspond pas à cette session.")
        if session_form != (bs.session_acad or ''):
            raise ValueError('La session académique envoyée ne correspond pas à cette session.')
        try:
            montant_act = int(montant_form)
        except (TypeError, ValueError):
            raise ValueError('Le montant du paiement est invalide.')
        if montant_act <= 0 or montant_act != (bs.montant_fc or 5000):
            raise ValueError(
                f'Le montant doit être exactement celui de la session : '
                f'{bs.montant_fc or 5000:,} FC.'
            )

        # Le verrou est pris avant tout contrôle puis conservé jusqu'au commit :
        # deux opérateurs ne peuvent pas attribuer simultanément le même reçu.
        recu = RecuPaiement.query.filter(
            db.or_(
                db.func.upper(RecuPaiement.numero) == numero_recu,
                db.func.upper(RecuPaiement.code_qr) == numero_recu,
            )
        ).with_for_update().first()
        if recu:
            origine_recu = 'référence'
        else:
            recu, origine_recu = _chercher_recu(numero_recu)
        if not recu:
            raise ValueError('Ce numéro de reçu ne correspond à aucun reçu officiel.')
        if not _receipt_type_is_consistent(recu):
            raise ValueError('Le reçu est incohérent : son préfixe ne correspond pas à son type.')

        expected_type = {
            'initial': 'bulletin',
            'recours': 'resultat_recours',
            'session_2': 'session_2',
            'recours_session_2': 'recours_session_2',
        }.get(getattr(bs, 'type_grille', 'initial'), 'bulletin')
        if _receipt_type_from_number(recu.numero) != expected_type:
            raise ValueError(
                f'Ce reçu ({_receipt_type_label(expected_type)}) ne correspond pas '
                f'à la session « {_grid_type_label(bs.type_grille)} ».'
            )
        if recu.utilise:
            raise ValueError(
                f'Ce reçu a déjà été utilisé par {recu.nom_etudiant or "un autre étudiant"}.'
            )
        if BulletinData.query.filter_by(session_id=sid, matricule=matricule).first():
            raise ValueError('Ce matricule existe déjà dans cette session.')

        # Le verrou PostgreSQL est complété par une mise à jour conditionnelle :
        # cela reste atomique avec SQLite, où FOR UPDATE n'est pas appliqué.
        reservee = RecuPaiement.query.filter(
            RecuPaiement.id == recu.id,
            RecuPaiement.utilise == False,  # noqa: E712
        ).update(
            {RecuPaiement.utilise: True},
            synchronize_session=False,
        )
        if reservee != 1:
            raise ValueError('Ce reçu vient d’être utilisé par un autre opérateur.')

        # Une ligne minimale est conservée pour que le paiement soit visible
        # immédiatement, même si la grille de résultats n'était pas disponible.
        import uuid as _uuid
        data = {
            'num': 0,
            'matricule': matricule,
            'nom': nom,
            'sexe': '',
            'cours': [],
            'manuel_paiement': True,
            'message_manuel': 'Résultat absent lors de la validation du paiement.',
        }
        bd = BulletinData(
            session_id=sid,
            matricule=matricule,
            nom=nom,
            sexe='',
            data_json=json.dumps(data, ensure_ascii=False),
            numero_bulletin=f'BUL-{sid:04d}-MAN-{_uuid.uuid4().hex[:8].upper()}',
            paye=True,
            date_paiement=date_paiement,
            methode_paiement='Validation manuelle',
            reference_paiement=recu.numero,
            montant_paye=montant_act,
        )
        db.session.add(bd)
        db.session.flush()

        recu.utilise = True
        recu.date_utilisation = date_paiement
        recu.matricule_etudiant = matricule
        recu.nom_etudiant = nom
        recu.bulletin_id = bd.id
        db.session.add(PaiementAudit(
            bulletin_id=bd.id,
            old_montant=0,
            new_montant=montant_act,
            old_methode=None,
            new_methode='Validation manuelle',
            old_reference=None,
            new_reference=recu.numero,
            old_date_paiement=None,
            new_date_paiement=date_paiement,
            operator_name=session.get('decanat_operator', 'Inconnu'),
        ))
        _enregistrer_admin_audit(
            'ajout_paiement_bulletin_manuel',
            {
                'bulletin_id': bd.id,
                'session_id': sid,
                'nom': nom,
                'matricule': matricule,
                'recu_id': recu.id,
                'reference': recu.numero,
            },
        )
        db.session.commit()
        flash(
            f'✅ Paiement Bulletin validé manuellement pour {nom}. '
            f'Reçu {recu.numero} consommé.',
            'success',
        )
    except ValueError as exc:
        db.session.rollback()
        flash(str(exc), 'error')
    except Exception as exc:
        db.session.rollback()
        app.logger.exception('Ajout manuel de paiement bulletin échoué')
        flash(f'Erreur lors de la validation manuelle : {exc}', 'error')
    return redirect(url_for('decanat_bulletins_session_detail', sid=sid))


@app.route('/decanat/bulletins/session/<int:sid>/activer-tout', methods=['POST'])
def decanat_bulletins_activer_tout(sid):
    """Active le paiement de TOUS les bulletins non-payés de la session."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    bs = BulletinSession.query.get_or_404(sid)
    date_str = request.form.get('date_paiement', '').strip()
    try:
        date_act = datetime.strptime(date_str, '%Y-%m-%d') if date_str else now_cat()
    except ValueError:
        date_act = now_cat()
    montant_act = bs.montant_fc or 5000
    operateur   = session.get('decanat_operator', 'Inconnu')
    non_payes   = BulletinData.query.filter_by(session_id=sid, paye=False).all()
    nb = 0
    for bd in non_payes:
        audit = PaiementAudit(
            bulletin_id=bd.id,
            old_montant=bd.montant_paye, new_montant=montant_act,
            old_methode=bd.methode_paiement, new_methode='Manuel',
            old_reference=bd.reference_paiement, new_reference=None,
            old_date_paiement=bd.date_paiement, new_date_paiement=date_act,
            operator_name=operateur,
        )
        db.session.add(audit)
        bd.paye = True
        bd.montant_paye = montant_act
        bd.methode_paiement = 'Manuel'
        bd.date_paiement = date_act
        nb += 1
    db.session.commit()
    flash(f'{nb} bulletin(s) activé(s) — {montant_act:,} FC chacun.', 'success')
    return redirect(url_for('decanat_bulletins_session_detail', sid=sid))


def _tokens_nom(nom: str) -> set:
    """Retourne un ensemble de tokens normalisés (sans accents, majuscules) pour le matching par intersection."""
    import unicodedata
    nfkd = unicodedata.normalize('NFKD', nom.upper())
    sans_accents = ''.join(c for c in nfkd if not unicodedata.combining(c))
    return set(sans_accents.split())


def _extraire_noms_pdf_colonnes(buf: bytes) -> list:
    """
    Extrait les noms d'étudiants d'un PDF tabulaire (rapport de paiement,
    liste de recours, liste d'émargement…) en utilisant les positions X/Y
    pour identifier la colonne NOM et recoller les noms coupés sur 2 lignes.
    """
    import fitz, re, unicodedata

    # Textes à ignorer même s'ils sont dans la colonne nom
    def _est_non_nom(t):
        tu = t.strip().upper()
        return (
            t.startswith('•') or
            tu in ('N°', 'N', 'NO', 'NUM', 'NOM', 'COMPLET', 'NOM COMPLET',
                   'PRENOM', 'PRÉNOM', 'NOMS', 'DATE', 'MONTANT', 'MATRICULE') or
            re.match(r'^[Tt][ée]l', t) or
            re.match(r'^[A-Z]{1,3}-[A-Z]{2,3}-\d', t) or   # codes reçu RS-SC-…
            re.match(r'^\d{2}/\d{2}/\d{4}', t) or           # dates
            re.match(r'^\d+$', t) or                          # numéros seuls
            re.match(r'^N[°o]\s*$', tu) or                   # "N°" seul
            len(tu) <= 2 or                                   # trop court
            any(kw in tu for kw in (
                'FACULTÉ', 'UNIVERSITÉ', 'LISTE DES', 'SECOND SEMESTRE',
                'PREMIER SEMESTRE', 'NOM COMPLET', 'FILIÈRE', 'RÉCLAMATION',
                'REÇU N°', 'TOTAL PROMOTION', 'GÉNÉRÉ LE', 'CARTE D',
                'REÇU DE', 'PIÈCE D', 'APPLIQUÉE', 'APPLIQUE',
                'DÉPARTEMENT', 'PROMOTION :', 'MÉTHODE',
                'RÉFÉRENCE', 'STATUT', 'DÉCISION',
            ))
        )

    doc = fitz.open(stream=buf, filetype='pdf')
    noms = []

    for page in doc:
        blocks = page.get_text("dict")["blocks"]

        # Collecter (x0, y0, texte)
        items = []
        for b in blocks:
            if b["type"] != 0:
                continue
            for line in b["lines"]:
                text = " ".join(s["text"] for s in line["spans"]).strip()
                if text:
                    items.append((line["bbox"][0], line["bbox"][1], text))

        if not items:
            continue

        # ── Trouver l'X de la colonne NOM ────────────────────────────────
        nom_col_x = None
        for x0, y0, text in items:
            tu = text.upper()
            if ('NOM' in tu and ('COMPLET' in tu or 'PRÉNOM' in tu or 'PRENOM' in tu or 'ETUDIANT' in tu)):
                nom_col_x = x0
                break
        if nom_col_x is None:
            # Fallback : chercher la 2e colonne d'une ligne numérotée
            for x0, y0, text in sorted(items, key=lambda i: i[1]):
                if re.match(r'^\d+$', text.strip()) and 1 <= int(text.strip()) <= 500:
                    row_y = y0
                    candidats = [(x, y, t) for x, y, t in items
                                 if abs(y - row_y) < 6 and x > x0 + 5]
                    if candidats:
                        nom_col_x = min(candidats, key=lambda i: i[0])[0]
                        break

        if nom_col_x is None:
            continue

        # ── Filtrer les items dans la colonne nom (tolérance ±35 pts) ────
        col_items = [
            (x0, y0, text) for x0, y0, text in items
            if abs(x0 - nom_col_x) <= 35 and not _est_non_nom(text)
        ]
        col_items.sort(key=lambda t: t[1])   # trier par Y croissant

        # ── Regrouper les lignes proches en Y (nom sur 2 lignes) ─────────
        groupes = []
        cur_parts = []
        cur_y = None
        SEUIL_MEME_NOM = 22   # points : < ce seuil = même étudiant

        for x0, y0, text in col_items:
            if cur_y is None:
                cur_y = y0
                cur_parts = [text]
            elif y0 - cur_y <= SEUIL_MEME_NOM:
                cur_parts.append(text)
                cur_y = y0
            else:
                groupes.append(' '.join(cur_parts))
                cur_parts = [text]
                cur_y = y0
        if cur_parts:
            groupes.append(' '.join(cur_parts))

        # ── Nettoyer et ajouter ───────────────────────────────────────────
        for nom in groupes:
            nom = nom.strip()
            if len(nom) >= 4 and len(nom.split()) >= 1:
                noms.append(nom)

    doc.close()
    return noms


@app.route('/decanat/bulletins/session/<int:sid>/activer-par-liste', methods=['POST'])
def decanat_bulletins_activer_par_liste(sid):
    """
    Active les bulletins correspondant à une liste de noms.
    La liste peut venir d'un textarea (un nom par ligne) ou d'un PDF uploadé.
    """
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    bs = BulletinSession.query.get_or_404(sid)
    date_str = request.form.get('date_paiement', '').strip()
    try:
        date_act = datetime.strptime(date_str, '%Y-%m-%d') if date_str else now_cat()
    except ValueError:
        date_act = now_cat()
    montant_act = bs.montant_fc or 5000
    operateur   = session.get('decanat_operator', 'Inconnu')

    # ── Récupérer la liste de noms ──────────────────────────────────────────
    noms_bruts = []

    SKIP_MOTS = {
        'FACULTÉ', 'UNIVERSITÉ', 'LISTE', 'SECOND', 'PREMIER', 'NOM', 'COMPLET',
        'NOM COMPLET', 'FILIÈRE', 'RÉCLAMATIONS', 'REÇU', 'DATE', 'TOTAL',
        'N°', 'PROMOTION', 'DÉPARTEMENT', 'MATRICULE', 'MONTANT', 'MÉTHODE',
        'RÉFÉRENCE', 'SEXE', 'SESSION', 'SEMESTRE', 'ANNEE', 'ANNÉE',
        'OBSERVATION', 'REMARQUE', 'PAIEMENT', 'STATUT', 'DECISION',
    }

    fichier = request.files.get('fichier_liste')
    ext = fichier.filename.rsplit('.', 1)[-1].lower() if fichier and fichier.filename else ''

    # 1. Fichier Excel (.xlsx ou .xls)
    if fichier and ext in ('xlsx', 'xls'):
        try:
            buf = fichier.read()
            import io as _io
            if ext == 'xlsx':
                import openpyxl as _openpyxl
                wb = _openpyxl.load_workbook(_io.BytesIO(buf), data_only=True)
                ws = wb.active
                rows = [[str(c.value).strip() if c.value is not None else '' for c in row]
                        for row in ws.iter_rows()]
            else:
                import xlrd as _xlrd
                wb = _xlrd.open_workbook(file_contents=buf)
                ws = wb.sheet_by_index(0)
                rows = [[str(ws.cell(r, c).value).strip() for c in range(ws.ncols)]
                        for r in range(ws.nrows)]

            # Trouver la colonne "nom" : chercher header contenant NOM ou ETUDIANT
            col_nom = None
            for ri, row in enumerate(rows[:5]):
                for ci, val in enumerate(row):
                    v = val.upper().replace('É', 'E').replace('È', 'E')
                    if 'NOM' in v or 'ETUDIANT' in v or 'PRENOM' in v:
                        col_nom = ci
                        break
                if col_nom is not None:
                    break

            if col_nom is not None:
                # Extraire uniquement la colonne des noms (en sautant l'en-tête)
                for row in rows:
                    val = row[col_nom].strip() if col_nom < len(row) else ''
                    if (val and val.upper() not in SKIP_MOTS
                            and len(val) >= 4 and not val.replace(' ', '').isdigit()):
                        noms_bruts.append(val)
            else:
                # Pas d'en-tête trouvé → scanner toutes les cellules ressemblant à des noms
                for row in rows:
                    for val in row:
                        val = val.strip()
                        mots = val.split()
                        if (len(mots) >= 2 and len(val) >= 5
                                and not val.replace(' ', '').isdigit()
                                and val.upper() not in SKIP_MOTS):
                            noms_bruts.append(val)
        except Exception as e:
            flash(f'Erreur lecture Excel : {e}', 'error')
            return redirect(url_for('decanat_bulletins_session_detail', sid=sid))

    # 2. Fichier PDF uploadé
    elif fichier and ext == 'pdf':
        try:
            import fitz as _fitz, re as _re
            buf = fichier.read()
            noms_bruts.extend(_extraire_noms_pdf_colonnes(buf))
        except Exception as e:
            flash(f'Erreur lecture PDF : {e}', 'error')
            return redirect(url_for('decanat_bulletins_session_detail', sid=sid))

    # 3. Textarea (un nom par ligne)
    texte_liste = request.form.get('noms_liste', '').strip()
    if texte_liste:
        for ligne in texte_liste.splitlines():
            ligne = ligne.strip()
            if ligne:
                noms_bruts.append(ligne)

    if not noms_bruts:
        flash('Aucun nom trouvé dans la liste fournie.', 'error')
        return redirect(url_for('decanat_bulletins_session_detail', sid=sid))

    # ── Normaliser les noms de la liste ─────────────────────────────────────
    tokens_liste = [_tokens_nom(n) for n in noms_bruts if n]
    tokens_liste = [t for t in tokens_liste if len(t) >= 1]

    # ── Chercher les bulletins correspondants ───────────────────────────────
    tous = BulletinData.query.filter_by(session_id=sid).all()
    actives, deja_payes, non_trouves_noms = [], [], list(noms_bruts)

    for bd in tous:
        tok_bd = _tokens_nom(bd.nom)
        for i, tok_liste in enumerate(tokens_liste):
            # Correspondance si au moins 2 tokens communs (ou tous les tokens courts)
            communs = tok_bd & tok_liste
            seuil = 2 if (len(tok_bd) >= 3 and len(tok_liste) >= 3) else 1
            if len(communs) >= seuil:
                if bd.paye:
                    deja_payes.append(bd.nom)
                else:
                    audit = PaiementAudit(
                        bulletin_id=bd.id,
                        old_montant=bd.montant_paye, new_montant=montant_act,
                        old_methode=bd.methode_paiement, new_methode='Manuel',
                        old_reference=bd.reference_paiement, new_reference=None,
                        old_date_paiement=bd.date_paiement, new_date_paiement=date_act,
                        operator_name=operateur,
                    )
                    db.session.add(audit)
                    bd.paye = True
                    bd.montant_paye = montant_act
                    bd.methode_paiement = 'Manuel'
                    bd.date_paiement = date_act
                    actives.append(bd.nom)
                # Retirer de la liste des non-trouvés
                if noms_bruts[i] in non_trouves_noms:
                    non_trouves_noms.remove(noms_bruts[i])
                break

    db.session.commit()

    if actives:
        flash(f'✅ {len(actives)} bulletin(s) activé(s) : {", ".join(actives[:5])}{"…" if len(actives)>5 else ""}', 'success')
    if deja_payes:
        flash(f'ℹ️ {len(deja_payes)} déjà payé(s) : {", ".join(deja_payes[:3])}', 'info')
    if non_trouves_noms:
        flash(f'⚠️ {len(non_trouves_noms)} nom(s) non trouvé(s) dans cette session : {", ".join(non_trouves_noms[:5])}', 'warning')
    if not actives and not deja_payes:
        flash('Aucun bulletin correspondant trouvé dans cette session.', 'warning')

    return redirect(url_for('decanat_bulletins_session_detail', sid=sid))


@app.route('/decanat/bulletins/session/<int:sid>/supprimer', methods=['POST'])
def decanat_bulletins_supprimer_session(sid):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    bs = BulletinSession.query.get_or_404(sid)
    # Supprimer les audits en premier : bulletin_id est NOT NULL → ne peut pas être mis à NULL
    bul_ids = [b.id for b in bs.bulletins]
    if bul_ids:
        PaiementAudit.query.filter(
            PaiementAudit.bulletin_id.in_(bul_ids)
        ).delete(synchronize_session=False)
    db.session.delete(bs)
    db.session.commit()
    flash('Session supprimée.', 'success')
    return redirect(url_for('decanat_bulletins_sessions'))


@app.route('/decanat/bulletins/session/<int:sid>/departement', methods=['POST'])
def decanat_bulletins_update_departement(sid):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    bs = BulletinSession.query.get_or_404(sid)
    bs.departement = request.form.get('departement', '').strip()
    db.session.commit()
    flash('Département mis à jour.', 'success')
    return redirect(url_for('decanat_bulletins_session_detail', sid=sid))


@app.route('/decanat/bulletins/session/<int:sid>/intro', methods=['POST'])
def decanat_bulletins_update_intro(sid):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    bs = BulletinSession.query.get_or_404(sid)
    bs.texte_intro = request.form.get('texte_intro', '').strip()
    db.session.commit()
    flash('Texte d\'introduction mis à jour.', 'success')
    return redirect(url_for('decanat_bulletins_session_detail', sid=sid))


@app.route('/decanat/bulletins/session/<int:sid>/audits.csv')
def decanat_bulletins_audits_csv(sid):
    """Export CSV des corrections de paiement pour une session.

    Généré en streaming (blocs de 500 lignes) pour éviter tout timeout ou
    débordement mémoire sur des historiques volumineux.
    """
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    import csv
    from io import StringIO
    from flask import stream_with_context, Response

    bs = BulletinSession.query.get_or_404(sid)
    bulletins = BulletinData.query.filter_by(session_id=sid).with_entities(
        BulletinData.id, BulletinData.nom).all()
    nom_par_id = {bd.id: bd.nom for bd in bulletins}
    bulletin_ids = list(nom_par_id.keys())

    q = (PaiementAudit.query
         .filter(PaiementAudit.bulletin_id.in_(bulletin_ids))
         .order_by(PaiementAudit.date_modification.desc())) if bulletin_ids else None

    def generate():
        si = StringIO()
        writer = csv.writer(si, delimiter=';', quoting=csv.QUOTE_MINIMAL)
        writer.writerow([
            'Date correction', 'Opérateur', 'Étudiant',
            'Montant avant (FC)', 'Montant après (FC)',
            'Méthode avant', 'Méthode après',
            'Référence avant', 'Référence après',
            'Date paiement avant', 'Date paiement après',
        ])
        yield si.getvalue().encode('utf-8-sig')

        if q is None:
            return

        for a in q.yield_per(500):
            si = StringIO()
            writer = csv.writer(si, delimiter=';', quoting=csv.QUOTE_MINIMAL)
            writer.writerow([
                a.date_modification.strftime('%d/%m/%Y %H:%M'),
                a.operator_name or '',
                nom_par_id.get(a.bulletin_id, ''),
                a.old_montant if a.old_montant is not None else '',
                a.new_montant if a.new_montant is not None else '',
                a.old_methode or '',
                a.new_methode or '',
                a.old_reference or '',
                a.new_reference or '',
                a.old_date_paiement.strftime('%d/%m/%Y') if a.old_date_paiement else '',
                a.new_date_paiement.strftime('%d/%m/%Y') if a.new_date_paiement else '',
            ])
            yield si.getvalue().encode('utf-8')

    fname = f'corrections_paiements_{bs.nom.replace(" ", "_")}_{bs.annee}.csv'
    headers = {
        'Content-Disposition': f'attachment; filename="{fname}"',
        'Content-Type': 'text/csv; charset=utf-8',
        'X-Accel-Buffering': 'no',
    }
    return Response(stream_with_context(generate()), headers=headers)


@app.route('/decanat/bulletins/audits.csv')
def decanat_bulletins_audits_all_csv():
    """Export CSV de tout l'historique des corrections de paiement, toutes sessions confondues.

    Paramètres GET optionnels :
      - date_from : date de début au format YYYY-MM-DD (incluse)
      - date_to   : date de fin   au format YYYY-MM-DD (incluse, jusqu'à 23:59:59)

    Généré en streaming (blocs de 500 lignes) pour éviter tout timeout ou
    débordement mémoire sur des historiques volumineux.
    """
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    import csv
    from io import StringIO
    from flask import stream_with_context, Response

    # Filtres de date optionnels
    date_from_str = request.args.get('date_from', '').strip()
    date_to_str   = request.args.get('date_to',   '').strip()
    date_from_dt  = None
    date_to_dt    = None
    try:
        if date_from_str:
            date_from_dt = datetime.strptime(date_from_str, '%Y-%m-%d')
    except ValueError:
        date_from_str = ''
    try:
        if date_to_str:
            # inclure toute la journée de fin
            date_to_dt = datetime.strptime(date_to_str, '%Y-%m-%d').replace(
                hour=23, minute=59, second=59)
    except ValueError:
        date_to_str = ''

    # Les métadonnées sessions/bulletins sont petites — on les charge une fois
    # en mémoire pour les lookups ; seuls les audits sont streamés.
    sessions_map = {bs.id: bs for bs in BulletinSession.query.all()}
    bulletins = BulletinData.query.with_entities(
        BulletinData.id, BulletinData.nom, BulletinData.session_id).all()
    nom_par_id     = {bd.id: bd.nom        for bd in bulletins}
    session_par_id = {bd.id: bd.session_id for bd in bulletins}

    q = PaiementAudit.query
    if date_from_dt:
        q = q.filter(PaiementAudit.date_modification >= date_from_dt)
    if date_to_dt:
        q = q.filter(PaiementAudit.date_modification <= date_to_dt)
    q = q.order_by(PaiementAudit.date_modification.desc())

    def generate():
        si = StringIO()
        writer = csv.writer(si, delimiter=';', quoting=csv.QUOTE_MINIMAL)
        writer.writerow([
            'Session',
            'Date correction', 'Opérateur', 'Étudiant',
            'Montant avant (FC)', 'Montant après (FC)',
            'Méthode avant', 'Méthode après',
            'Référence avant', 'Référence après',
            'Date paiement avant', 'Date paiement après',
        ])
        yield si.getvalue().encode('utf-8-sig')

        for a in q.yield_per(500):
            sid_audit = session_par_id.get(a.bulletin_id)
            bs = sessions_map.get(sid_audit)
            session_label = f"{bs.nom} ({bs.annee})" if bs else ''
            si = StringIO()
            writer = csv.writer(si, delimiter=';', quoting=csv.QUOTE_MINIMAL)
            writer.writerow([
                session_label,
                a.date_modification.strftime('%d/%m/%Y %H:%M'),
                a.operator_name or '',
                nom_par_id.get(a.bulletin_id, ''),
                a.old_montant if a.old_montant is not None else '',
                a.new_montant if a.new_montant is not None else '',
                a.old_methode or '',
                a.new_methode or '',
                a.old_reference or '',
                a.new_reference or '',
                a.old_date_paiement.strftime('%d/%m/%Y') if a.old_date_paiement else '',
                a.new_date_paiement.strftime('%d/%m/%Y') if a.new_date_paiement else '',
            ])
            yield si.getvalue().encode('utf-8')

    headers = {
        'Content-Disposition': 'attachment; filename="corrections_paiements_toutes_sessions.csv"',
        'Content-Type': 'text/csv; charset=utf-8',
        'X-Accel-Buffering': 'no',
    }
    return Response(stream_with_context(generate()), headers=headers)


@app.route('/decanat/bulletins/session/<int:sid>/montant', methods=['POST'])
def decanat_bulletins_update_montant(sid):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    bs = BulletinSession.query.get_or_404(sid)
    try:
        bs.montant_fc = int(request.form.get('montant_fc', bs.montant_fc))
        db.session.commit()
        flash('Montant mis à jour.', 'success')
    except Exception:
        flash('Valeur invalide.', 'error')
    return redirect(url_for('decanat_bulletins_session_detail', sid=sid))


@app.route('/decanat/bulletins/session/<int:sid>/zip')
def decanat_bulletins_zip(sid):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    bs = BulletinSession.query.get_or_404(sid)
    bulletins = BulletinData.query.filter_by(session_id=sid).all()

    zip_buf = BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for bd in bulletins:
            etu = json.loads(bd.data_json)
            pdf_buf = _generer_bulletin_pdf(etu, bs.promotion, bs.annee,
                                            bs.session_acad, bs.semestre,
                                            numero_bulletin=bd.numero_bulletin,
                                            texte_intro=bs.texte_intro,
                                            departement=bs.departement)
            safe = bd.nom.replace(' ', '_')[:35]
            zf.writestr(f'Bulletin_{safe}.pdf', pdf_buf.read())

    zip_buf.seek(0)
    return send_file(zip_buf, as_attachment=True,
                     download_name=f'Bulletins_{bs.promotion.replace(" ", "_")}.zip',
                     mimetype='application/zip')


@app.route('/decanat/bulletins/bulletin/<int:bid>/activer_paiement', methods=['POST'])
def decanat_bulletins_activer_paiement(bid):
    """Activer le paiement d'un bulletin en un seul clic (sans formulaire)."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    bd = BulletinData.query.get_or_404(bid)
    sid = bd.session_id
    try:
        date_str = request.form.get('date_paiement', '').strip()
        if date_str:
            date_act = datetime.strptime(date_str, '%Y-%m-%d')
        else:
            date_act = now_cat()
        montant_act = bd.bul_session.montant_fc if bd.bul_session else 5000
        audit = PaiementAudit(
            bulletin_id       = bd.id,
            old_montant       = bd.montant_paye,
            new_montant       = montant_act,
            old_methode       = bd.methode_paiement,
            new_methode       = 'Manuel',
            old_reference     = bd.reference_paiement,
            new_reference     = None,
            old_date_paiement = bd.date_paiement,
            new_date_paiement = date_act,
            operator_name     = session.get('decanat_operator', 'Inconnu'),
        )
        db.session.add(audit)
        bd.paye             = True
        bd.montant_paye     = montant_act
        bd.methode_paiement = 'Manuel'
        bd.date_paiement    = date_act
        db.session.commit()
        flash(f'Paiement de {bd.nom} activé ({montant_act:,} FC).', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erreur : {e}', 'error')
    return redirect(url_for('decanat_bulletins_session_detail', sid=sid))


@app.route('/decanat/bulletins/bulletin/<int:bid>/paiement', methods=['POST'])
def decanat_bulletins_update_paiement(bid):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    bd = BulletinData.query.get_or_404(bid)
    sid = bd.session_id
    try:
        montant_str = request.form.get('montant_paye', '').strip()
        new_montant   = int(montant_str) if montant_str else 0
        new_methode   = request.form.get('methode_paiement', '').strip() or None
        new_reference = request.form.get('reference_paiement', '').strip() or None
        date_str = request.form.get('date_paiement', '').strip()
        new_date = datetime.strptime(date_str, '%Y-%m-%d') if date_str else None

        # ── Enregistrer l'audit avant toute modification ──
        audit = PaiementAudit(
            bulletin_id       = bd.id,
            old_montant       = bd.montant_paye,
            new_montant       = new_montant,
            old_methode       = bd.methode_paiement,
            new_methode       = new_methode,
            old_reference     = bd.reference_paiement,
            new_reference     = new_reference,
            old_date_paiement = bd.date_paiement,
            new_date_paiement = new_date,
            operator_name     = session.get('decanat_operator', 'Inconnu'),
        )
        db.session.add(audit)

        # ── Appliquer les nouvelles valeurs ──
        bd.montant_paye      = new_montant
        bd.methode_paiement  = new_methode
        bd.reference_paiement= new_reference
        bd.date_paiement     = new_date
        # marquer payé si montant > 0, sinon remettre en attente
        bd.paye = bd.montant_paye > 0
        db.session.commit()
        flash(f'Paiement de {bd.nom} mis à jour.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erreur lors de la mise à jour : {e}', 'error')
    return redirect(url_for('decanat_bulletins_session_detail', sid=sid))


@app.route('/decanat/bulletins/bulletin/<int:bid>/pdf')
def decanat_bulletins_pdf_admin(bid):
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    bd = BulletinData.query.get_or_404(bid)
    bs = bd.bul_session
    etu = json.loads(bd.data_json)
    pdf_buf = _generer_bulletin_pdf(etu, bs.promotion, bs.annee,
                                    bs.session_acad, bs.semestre,
                                    numero_bulletin=bd.numero_bulletin,
                                    texte_intro=bs.texte_intro,
                                    departement=bs.departement)
    safe = bd.nom.replace(' ', '_')[:35]
    return send_file(pdf_buf, as_attachment=True,
                     download_name=f'Bulletin_{safe}.pdf',
                     mimetype='application/pdf')


# ═══════════════════════════════════════════════════════════════════
#  PORTAIL PUBLIC BULLETINS (étudiant)
# ═══════════════════════════════════════════════════════════════════

@app.route('/bulletins')
def bulletins_portail():
    return render_template('bulletin_portail.html', mode='search')


@app.route('/bulletins/rechercher', methods=['POST'])
def bulletins_rechercher():
    matricule = request.form.get('matricule', '').strip().upper()
    if not matricule:
        flash('Veuillez saisir votre matricule.', 'error')
        return redirect(url_for('bulletins_portail'))

    resultats = BulletinData.query.filter(
        db.func.upper(BulletinData.matricule) == matricule
    ).all()

    if not resultats:
        flash(f'Aucun bulletin trouvé pour le matricule « {matricule} ».', 'error')
        return redirect(url_for('bulletins_portail'))

    if len(resultats) == 1:
        return redirect(url_for('bulletins_result', bid=resultats[0].id))

    # Plusieurs bulletins (plusieurs sessions)
    return render_template('bulletin_portail.html', mode='list',
                           resultats=resultats, matricule=matricule)


@app.route('/bulletins/result/<int:bid>')
def bulletins_result(bid):
    bd = BulletinData.query.get_or_404(bid)
    bs = bd.bul_session
    etu = json.loads(bd.data_json) if bd.paye else None
    return render_template('bulletin_portail.html', mode='result',
                           bd=bd, bs=bs, etu=etu)


@app.route('/bulletins/payer/<int:bid>', methods=['POST', 'GET'])
def bulletins_payer(bid):
    """Conservé pour compatibilité — redirige vers le portail bulletin."""
    return redirect(url_for('bulletins_result', bid=bid))


@app.route('/bulletins/apercu/<int:bid>')
def bulletins_apercu(bid):
    """Affiche le bulletin PDF inline dans le navigateur (prévisualisation)."""
    bd = BulletinData.query.get_or_404(bid)
    if not bd.paye:
        return "Bulletin verrouillé", 403
    bs  = bd.bul_session
    etu = json.loads(bd.data_json)

    # Bloquer la prévisualisation si la moyenne est 0 alors que la colonne existait
    if etu.get('moy_col_present') and etu.get('moyenne', 0) == 0.0:
        return "Bulletin temporairement indisponible : les résultats n'ont pas encore été saisis. Contactez le DÉCANAT.", 403

    pdf_buf = _generer_bulletin_pdf(etu, bs.promotion, bs.annee,
                                    bs.session_acad, bs.semestre,
                                    numero_bulletin=bd.numero_bulletin,
                                    texte_intro=bs.texte_intro,
                                    departement=bs.departement)
    return send_file(pdf_buf, as_attachment=False, mimetype='application/pdf')


@app.route('/decanat/bulletins/bulletin/<int:bid>/apercu')
def decanat_bulletins_apercu(bid):
    """Admin : affiche le bulletin PDF inline (prévisualisation)."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    bd = BulletinData.query.get_or_404(bid)
    bs  = bd.bul_session
    etu = json.loads(bd.data_json)
    pdf_buf = _generer_bulletin_pdf(etu, bs.promotion, bs.annee,
                                    bs.session_acad, bs.semestre,
                                    numero_bulletin=bd.numero_bulletin,
                                    texte_intro=bs.texte_intro,
                                    departement=bs.departement)
    return send_file(pdf_buf, as_attachment=False, mimetype='application/pdf')


@app.route('/bulletins/telecharger/<int:bid>')
def bulletins_telecharger(bid):
    bd = BulletinData.query.get_or_404(bid)
    if not bd.paye:
        flash('Veuillez régler les frais de consultation pour télécharger votre bulletin.', 'error')
        return redirect(url_for('bulletins_result', bid=bid))

    bs  = bd.bul_session
    etu = json.loads(bd.data_json)

    # Bloquer le téléchargement si la colonne Moyenne était présente dans la grille
    # mais que la valeur est 0 — les résultats n'ont pas été saisis.
    if etu.get('moy_col_present') and etu.get('moyenne', 0) == 0.0:
        flash(
            '⛔ Votre bulletin ne peut pas être téléchargé pour l\'instant : '
            'la moyenne affichée est 0.00, ce qui indique que les résultats '
            'n\'ont pas encore été saisis dans le système. '
            'Contactez le DÉCANAT pour plus d\'informations.',
            'error'
        )
        return redirect(url_for('bulletins_result', bid=bid))

    pdf_buf = _generer_bulletin_pdf(etu, bs.promotion, bs.annee,
                                    bs.session_acad, bs.semestre,
                                    numero_bulletin=bd.numero_bulletin,
                                    texte_intro=bs.texte_intro,
                                    departement=bs.departement)

    bd.nb_telechargements += 1
    bd.date_dernier_telechargement = now_cat()
    db.session.commit()

    safe = bd.nom.replace(' ', '_')[:35]
    return send_file(pdf_buf, as_attachment=True,
                     download_name=f'Bulletin_{safe}.pdf',
                     mimetype='application/pdf')


@app.route('/bulletins/verifier/<numero>')
def bulletins_verifier_qr(numero):
    bd = BulletinData.query.filter_by(numero_bulletin=numero).first()
    if not bd:
        return render_template('bulletin_portail.html', mode='verif_fail', numero=numero)
    bs  = bd.bul_session
    etu = json.loads(bd.data_json)
    return render_template('bulletin_portail.html', mode='verif_ok',
                           bd=bd, bs=bs, etu=etu)


# ═══════════════════════════════════════════════════════════════════
#  GÉNÉRATION DE REÇUS PHYSIQUES
# ═══════════════════════════════════════════════════════════════════

def _generer_recus_pdf(recus, type_recu="bulletin"):
    """Génère un PDF avec 4 reçus par page A4 selon le modèle officiel UNILU."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.utils import ImageReader
    import qrcode as _qrcode

    buf = BytesIO()
    W, H = A4   # 595 × 842 pts

    COLS = 2     # 2 colonnes × 2 rangées = 4 reçus par page
    ROWS = 2
    RW = W / COLS   # largeur d'un reçu  ≈ 105 mm
    RH = H / ROWS   # hauteur d'un reçu  ≈ 149 mm

    MARINE = colors.Color(13/255, 38/255, 102/255)
    OR     = colors.Color(191/255, 153/255, 26/255)
    LGRAY  = colors.Color(0.6, 0.6, 0.6)

    c = rl_canvas.Canvas(buf, pagesize=A4)
    # positions : ligne du haut (rangée 1) puis ligne du bas (rangée 0)
    positions = [
        (0,   RH), (RW,  RH),   # rangée haute
        (0,   0),  (RW,  0),    # rangée basse
    ]

    def field_line(c, lx, ly, label, dots_width, value=''):
        """Dessine une ligne de formulaire : label + pointillés + valeur."""
        c.setFont('Helvetica-Bold', 9)
        c.setFillColor(colors.black)
        c.drawString(lx, ly, label)
        lbl_w = c.stringWidth(label, 'Helvetica-Bold', 9)
        c.setFont('Helvetica', 8)
        c.setFillColor(LGRAY)
        dots = '.' * int(dots_width / c.stringWidth('.', 'Helvetica', 8))
        c.drawString(lx + lbl_w + 1*mm, ly, dots)
        if value:
            c.setFillColor(MARINE)
            c.setFont('Helvetica-Bold', 9)
            c.drawString(lx + lbl_w + 2*mm, ly, value)

    def draw_recu(c, x, y, recu):
        foot_h = 18*mm   # hauteur du pied (défini ici pour être accessible partout)

        # ─── Bordure extérieure ───
        c.setStrokeColor(MARINE)
        c.setLineWidth(1.2)
        c.rect(x + 3*mm, y + 3*mm, RW - 6*mm, RH - 6*mm, fill=0)

        # ─── Tirets de découpe ───
        dl = 6*mm
        c.setStrokeColor(LGRAY)
        c.setLineWidth(0.35)
        c.setDash(2, 2)
        for sx, sy, ex, ey in [
            (x, y+RH, x+dl, y+RH), (x+RW-dl, y+RH, x+RW, y+RH),
            (x, y, x+dl, y), (x+RW-dl, y, x+RW, y),
            (x, y, x, y+dl), (x, y+RH-dl, x, y+RH),
            (x+RW, y, x+RW, y+dl), (x+RW, y+RH-dl, x+RW, y+RH),
        ]:
            c.line(sx, sy, ex, ey)
        c.setDash()

        # ─── Bandeau UNILU ───
        hdr_h = 24*mm
        hy = y + RH - 3*mm - hdr_h
        c.setFillColor(MARINE)
        c.rect(x + 3*mm, hy, RW - 6*mm, hdr_h, fill=1, stroke=0)

        c.setFillColor(colors.white)
        c.setFont('Helvetica-Bold', 13)
        c.drawCentredString(x + RW/2, hy + hdr_h - 7*mm, 'UNIVERSITE DE LUBUMBASHI')
        c.setFont('Helvetica', 9)
        c.drawCentredString(x + RW/2, hy + hdr_h - 13*mm,
                            'Faculté des Sciences Sociales, Politiques et Administratives')
        c.setFillColor(OR)
        _titre_recu = _receipt_pdf_title(type_recu)
        # Les titres RR/R2 sont longs : les ajuster dans le bandeau au lieu
        # de les laisser déborder du reçu ou être coupés à l'impression.
        title_max_width = RW - 12*mm
        title_font = 8.2
        title_lines = [_titre_recu]
        if c.stringWidth(_titre_recu, 'Helvetica-Bold', title_font) > title_max_width:
            title_font = 7.2
        if c.stringWidth(_titre_recu, 'Helvetica-Bold', title_font) > title_max_width:
            title_lines = []
            current_line = ''
            for word in _titre_recu.split():
                candidate = f'{current_line} {word}'.strip()
                if (current_line and
                        c.stringWidth(candidate, 'Helvetica-Bold', title_font)
                        > title_max_width):
                    title_lines.append(current_line)
                    current_line = word
                else:
                    current_line = candidate
            if current_line:
                title_lines.append(current_line)
        c.setFont('Helvetica-Bold', title_font)
        if len(title_lines) == 1:
            c.drawCentredString(x + RW/2, hy + 3.5*mm, title_lines[0])
        else:
            line_gap = 3.2*mm
            first_y = hy + 3.5*mm + line_gap
            for index, line in enumerate(title_lines[:2]):
                c.drawCentredString(
                    x + RW/2, first_y - index * line_gap, line
                )

        # ─── Numéro du reçu + ligne or ───
        num_y = hy - 12*mm
        c.setFillColor(MARINE)
        c.setFont('Helvetica-Bold', 15)
        c.drawCentredString(x + RW/2, num_y, f'Preuve de paiement : {recu.numero}')
        c.setStrokeColor(OR)
        c.setLineWidth(1.2)
        c.line(x + 8*mm, num_y - 2*mm, x + RW - 8*mm, num_y - 2*mm)

        # ─── Champs de formulaire ───
        lx = x + 6*mm
        fw = RW - 12*mm   # largeur totale disponible pour les champs
        line_h = 9*mm

        # QR code à droite des champs — on réserve 38mm à droite
        qr_size = 38*mm
        qr_x = x + RW - 3*mm - qr_size
        qr_y_pos = y + 3*mm + foot_h + 2*mm   # juste au-dessus du pied

        try:
            scan_url = (f'/recours/scan/{recu.code_qr}'
                        if type_recu in ('recours', 'recours_session_2_soumission')
                        else f'/scan/{recu.code_qr}')
            qr_img = _qrcode.make(scan_url)
            qr_buf = BytesIO()
            qr_img.save(qr_buf, format='PNG')
            qr_buf.seek(0)
            c.drawImage(ImageReader(qr_buf), qr_x, qr_y_pos, qr_size, qr_size)
            # légende sous QR
            c.setFont('Helvetica', 5.5)
            c.setFillColor(LGRAY)
            c.drawCentredString(qr_x + qr_size/2, qr_y_pos - 3*mm, 'Scanner pour accéder')
            c.drawCentredString(qr_x + qr_size/2, qr_y_pos - 6*mm, 'au bulletin (1 seule fois)')
        except Exception:
            pass

        field_w = fw - qr_size - 6*mm   # largeur champs (sans la zone QR)

        cur_y = num_y - 8*mm
        # Mme, M.
        c.setFont('Helvetica-Bold', 9)
        c.setFillColor(colors.black)
        c.drawString(lx, cur_y, 'Mme, M.')
        c.setFont('Helvetica', 8)
        c.setFillColor(LGRAY)
        dots1 = '.' * int((field_w - 16*mm) / c.stringWidth('.', 'Helvetica', 8))
        c.drawString(lx + 16*mm, cur_y, dots1)
        c.setStrokeColor(colors.Color(0.85,0.85,0.85))
        c.setLineWidth(0.3)
        c.line(lx, cur_y - 1.5*mm, lx + field_w, cur_y - 1.5*mm)

        cur_y -= line_h
        # Promotion
        c.setFont('Helvetica-Bold', 9)
        c.setFillColor(colors.black)
        c.drawString(lx, cur_y, 'Promotion :')
        c.setFont('Helvetica', 8)
        c.setFillColor(LGRAY)
        dots2 = '.' * int((field_w - 25*mm) / c.stringWidth('.', 'Helvetica', 8))
        c.drawString(lx + 25*mm, cur_y, dots2)
        c.setStrokeColor(colors.Color(0.85,0.85,0.85))
        c.line(lx, cur_y - 1.5*mm, lx + field_w, cur_y - 1.5*mm)

        cur_y -= line_h
        # Montant (fond bleu clair) — label sur une ligne, valeur sur la même à droite
        box_h = 8*mm
        c.setFillColor(colors.Color(0.92, 0.95, 1.0))
        c.rect(lx - 1*mm, cur_y - 1.5*mm, field_w + qr_size + 5*mm, box_h, fill=1, stroke=0)
        c.setFont('Helvetica-Bold', 9)
        c.setFillColor(colors.black)
        c.drawString(lx, cur_y + 1.5*mm, 'Montant :')
        c.setFont('Helvetica-Bold', 12)
        c.setFillColor(MARINE)
        montant_txt = recu.montant or '5000 CDF'
        c.drawString(lx + 26*mm, cur_y + 1.5*mm, montant_txt)

        cur_y -= line_h
        # En lettres
        c.setFont('Helvetica-Bold', 9)
        c.setFillColor(colors.black)
        c.drawString(lx, cur_y, 'En lettres :')
        c.setFont('Helvetica-Oblique', 9)
        c.setFillColor(MARINE)
        c.drawString(lx + 26*mm, cur_y, recu.montant_lettres or 'Cinq mille Francs congolais')
        c.setStrokeColor(colors.Color(0.85,0.85,0.85))
        c.setLineWidth(0.3)
        c.line(lx, cur_y - 1.5*mm, lx + field_w, cur_y - 1.5*mm)

        cur_y -= line_h
        # Session — affiche le type sélectionné (Second semestre / Recours…)
        session_val = (recu.motif or 'Second semestre').strip()
        c.setFont('Helvetica-Bold', 9)
        c.setFillColor(colors.black)
        c.drawString(lx, cur_y, 'Session :')
        c.setFont('Helvetica-Bold', 9)
        c.setFillColor(MARINE)
        sess_max_w = fw - 22*mm
        orig_sess = session_val
        while c.stringWidth(session_val, 'Helvetica-Bold', 9) > sess_max_w and len(session_val) > 4:
            session_val = session_val[:-1]
        if session_val != orig_sess:
            session_val = session_val.rstrip() + '…'
        c.drawString(lx + 22*mm, cur_y, session_val)
        c.setStrokeColor(colors.Color(0.85, 0.85, 0.85))
        c.setLineWidth(0.3)
        c.line(lx, cur_y - 1.5*mm, lx + field_w, cur_y - 1.5*mm)
        c.setStrokeColor(colors.Color(0.85,0.85,0.85))
        c.line(lx, cur_y - 1.5*mm, lx + field_w, cur_y - 1.5*mm)

        cur_y -= line_h
        # Date
        c.setFont('Helvetica-Bold', 9)
        c.setFillColor(colors.black)
        c.drawString(lx, cur_y, 'Fait à Lubumbashi, le')
        c.setFont('Helvetica', 8)
        c.setFillColor(LGRAY)
        c.drawString(lx + 44*mm, cur_y, '.' * 20)
        c.setStrokeColor(colors.Color(0.85,0.85,0.85))
        c.line(lx, cur_y - 1.5*mm, lx + field_w, cur_y - 1.5*mm)

        # ─── Pied : Sceau | Comptable ───
        foot_y = y + 3*mm
        c.setFillColor(colors.Color(0.94, 0.96, 1.0))
        c.rect(x + 3*mm, foot_y, RW - 6*mm, foot_h, fill=1, stroke=0)
        c.setStrokeColor(MARINE)
        c.setLineWidth(0.4)
        c.line(x + 3*mm, foot_y + foot_h, x + RW - 3*mm, foot_y + foot_h)

        mid_x = x + RW/2
        c.setFont('Helvetica-Bold', 9)
        c.setFillColor(MARINE)
        c.drawCentredString(x + RW/4, foot_y + foot_h - 6*mm, 'Sceau de la faculté')
        c.drawCentredString(x + 3*RW/4, foot_y + foot_h - 6*mm, 'Comptable')
        c.setStrokeColor(MARINE)
        c.setLineWidth(0.4)
        c.line(mid_x, foot_y + 2*mm, mid_x, foot_y + foot_h - 2*mm)

    page_recus = []
    for i, recu in enumerate(recus):
        page_recus.append(recu)
        if len(page_recus) == 4 or i == len(recus) - 1:
            # Lignes de coupe (pointillées) : 1 verticale centrale + 1 horizontale centrale
            c.setStrokeColor(LGRAY)
            c.setLineWidth(0.4)
            c.setDash(3, 4)
            c.line(RW, 0, RW, H)    # ligne verticale centrale
            c.line(0, H/2, W, H/2)  # ligne horizontale centrale
            c.setDash()

            for j, recu in enumerate(page_recus):
                px, py = positions[j]
                draw_recu(c, px, py, recu)

            c.showPage()
            page_recus = []

    c.save()
    buf.seek(0)
    return buf


def _extraire_recus_pdf_import(pdf_bytes: bytes) -> list[dict]:
    """Extrait les numéros et les QR existants d'un PDF de reçus imprimés.

    Les QR ne sont pas régénérés : ils sont décodés depuis les images embarquées
    du PDF afin que les reçus déjà vendus deviennent immédiatement utilisables.
    """
    import re as _re
    import fitz as _fitz
    import cv2 as _cv2
    import numpy as _np

    numero_re = _re.compile(
        r'\b((?:REL|RS2|S2|R2|B\d+|RS|RR)-[A-Z0-9]+-\d{3}-[A-Z0-9]+-\d{2,4})\b',
        _re.IGNORECASE,
    )
    qr_detector = _cv2.QRCodeDetector()

    def _decode_qr(image_bytes: bytes) -> str | None:
        image = _cv2.imdecode(
            _np.frombuffer(image_bytes, dtype=_np.uint8),
            _cv2.IMREAD_GRAYSCALE,
        )
        if image is None:
            return None
        variants = []
        for quarter_turns in range(4):
            rotated = _np.rot90(image, quarter_turns)
            variants.append(rotated)
            variants.append(_cv2.resize(
                rotated, None, fx=2, fy=2, interpolation=_cv2.INTER_NEAREST
            ))
        for variant in variants:
            value, _points, _straight = qr_detector.detectAndDecode(variant)
            if value:
                value = value.strip()
                match = _re.search(r'(?:^|/)(?:scan)/([A-Z0-9_-]+)', value, _re.I)
                if match:
                    return match.group(1).upper()
        return None

    def _code_metadata(numero: str) -> dict:
        match = _re.match(
            r'^(?P<prefix>REL|RS2|S2|R2|RS|RR|B(?P<lot>\d+))-(?P<dept>[A-Z0-9]+)-'
            r'(?P<seq>\d+)-(?P<sem>[A-Z0-9]+)-(?P<annee>\d{2,4})$',
            numero.upper(),
        )
        if not match:
            raise ValueError(f'Numéro de reçu non reconnu : {numero}')
        annee = match.group('annee')
        if len(annee) == 2:
            fin = 2000 + int(annee)
            annee_complete = f'{fin - 1}-{fin}'
        else:
            annee_complete = annee
        semestre = match.group('sem')
        prefix = match.group('prefix')
        if prefix == 'REL':
            motif = 'Commande de relevé de cotes'
        elif prefix == 'RS':
            motif = 'Paiement de soumission du recours'
        elif prefix == 'RS2':
            motif = 'Paiement de soumission du recours de 2ème session'
        elif prefix == 'RR':
            motif = 'Consultation du résultat de recours'
        elif prefix == 'R2':
            motif = 'Consultation du recours de 2ème session'
        elif prefix == 'S2':
            motif = 'Bulletin de 2ème session'
        elif semestre == 'S1':
            motif = 'Premier semestre'
        elif semestre == 'S2':
            motif = 'Second semestre'
        else:
            motif = semestre
        return {
            'numero': numero.upper(),
            'lot': match.group('lot') or '0',
            'type_recu': (
                'releve' if match.group('prefix') == 'REL'
                else 'recours' if match.group('prefix') == 'RS'
                else 'recours_session_2_soumission'
                if match.group('prefix') == 'RS2'
                else 'resultat_recours' if match.group('prefix') == 'RR'
                else 'recours_session_2' if match.group('prefix') == 'R2'
                else 'session_2' if match.group('prefix') == 'S2'
                else 'bulletin'
            ),
            'dept': match.group('dept').upper(),
            'annee': annee[-2:],
            'semestre': semestre,
            'annee_complete': annee_complete,
            'motif': motif,
        }

    doc = _fitz.open(stream=pdf_bytes, filetype='pdf')
    records = []
    try:
        for page_number, page in enumerate(doc, start=1):
            text_spans = []
            for block in page.get_text('dict').get('blocks', []):
                if block.get('type') != 0:
                    continue
                for line in block.get('lines', []):
                    for span in line.get('spans', []):
                        match = numero_re.search(span.get('text', ''))
                        if match:
                            x0, y0, x1, y1 = span['bbox']
                            text_spans.append({
                                'numero': match.group(1).upper(),
                                'cx': (x0 + x1) / 2,
                                'cy': (y0 + y1) / 2,
                                'page': page_number,
                            })

            image_candidates = []
            for image_info in page.get_images(full=True):
                xref = image_info[0]
                rects = page.get_image_rects(xref)
                if not rects:
                    continue
                extracted = doc.extract_image(xref)
                code_qr = _decode_qr(extracted['image'])
                if not code_qr:
                    raise ValueError(
                        f'QR illisible à la page {page_number}; '
                        'aucun reçu n’a été importé.'
                    )
                rect = rects[0]
                image_candidates.append({
                    'code_qr': code_qr,
                    'cx': (rect.x0 + rect.x1) / 2,
                    'cy': (rect.y0 + rect.y1) / 2,
                    'page': page_number,
                })

            if len(text_spans) != len(image_candidates):
                raise ValueError(
                    f'Page {page_number}: {len(text_spans)} numéro(s) et '
                    f'{len(image_candidates)} QR détecté(s). '
                    'Le PDF n’a pas été importé.'
                )

            assigned = set()
            for image in image_candidates:
                available = [
                    (index, span) for index, span in enumerate(text_spans)
                    if index not in assigned
                ]
                if not available:
                    raise ValueError(f'Association numéro/QR impossible page {page_number}.')
                index, span = min(
                    available,
                    key=lambda item: abs(item[1]['cx'] - image['cx'])
                    + abs(item[1]['cy'] - image['cy']),
                )
                assigned.add(index)
                metadata = _code_metadata(span['numero'])
                metadata['code_qr'] = image['code_qr']
                metadata['page'] = page_number
                records.append(metadata)
    finally:
        doc.close()

    if not records:
        raise ValueError('Aucun reçu détecté dans le PDF.')
    numeros = [r['numero'] for r in records]
    codes = [r['code_qr'] for r in records]
    if len(set(numeros)) != len(numeros):
        raise ValueError('Le PDF contient des numéros de reçu en double.')
    if len(set(codes)) != len(codes):
        raise ValueError('Le PDF contient des QR codes en double.')
    return sorted(records, key=lambda r: (r['dept'], r['lot'], r['numero']))


@app.route('/decanat/recus', methods=['GET', 'POST'])
def decanat_recus():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    if request.method == 'POST':
        dept            = request.form.get('dept', 'RI').strip().upper()
        lot             = request.form.get('lot', '1').strip()
        annee           = request.form.get('annee', '26').strip()
        semestre        = request.form.get('semestre', 'S2').strip()
        annee_complete  = request.form.get('annee_complete', '2025-2026').strip()
        montant         = request.form.get('montant', '5000 CDF').strip()
        montant_lettres = request.form.get('montant_lettres', 'Cinq mille Francs congolais').strip()
        type_session    = request.form.get('type_session', 'Second semestre').strip()
        motif           = type_session   # stocké dans motif, affiché comme "Session" sur le reçu
        type_recu       = request.form.get('type_recu', 'bulletin').strip()
        if type_recu not in _RECEIPT_TYPES:
            type_recu = 'bulletin'
        quantite        = min(int(request.form.get('quantite', 20)), 500)
        action          = request.form.get('action', 'generer')

        if type_recu != 'releve':
            montant, montant_lettres = _receipt_payment_details(type_recu)
        elif not montant:
            flash(
                'Indiquez le montant du reçu de relevé avant de générer le lot.',
                'error',
            )
            return redirect(url_for('decanat_recus') + '#form_releve')
        motif = {
            'bulletin': 'Résultats session 1',
            'recours': 'Soumission recours session 1',
            'resultat_recours': 'Résultats du recours session 1',
            'session_2': 'Résultats session 2',
            'recours_session_2_soumission': 'Soumission recours session 2',
            'recours_session_2': 'Résultats du recours session 2',
            'releve': 'Commande de relevé de cotes',
        }.get(type_recu, type_session)

        import uuid as _uuid

        # Calculer le dernier numéro séquentiel pour ce lot/dept/semestre.
        # On extrait le MAX réel du 3ᵉ segment de chaque numero existant
        # (ex. "B1-RI-037-S2-26" → 37 ou "RS-RI-037-S2-26" → 37)
        # pour ne jamais entrer en collision.
        import re as _re
        if type_recu != 'bulletin':
            prefix = _receipt_prefix(type_recu)
            prefix_like = f'{prefix}-{dept}-'
            _pat_seq    = _re.compile(rf'^{prefix}-[^-]+-(\d+)-')
        else:
            prefix_like = f'B{lot}-{dept}-'
            _pat_seq    = _re.compile(r'^B\d+-[^-]+-(\d+)-')
        existing_nums = RecuPaiement.query.filter(
            RecuPaiement.numero.like(f'{prefix_like}%')
        ).with_entities(RecuPaiement.numero).all()
        existing_seqs = []
        for (num,) in existing_nums:
            m = _pat_seq.match(num)
            if m:
                existing_seqs.append(int(m.group(1)))
        seq_start = (max(existing_seqs) + 1) if existing_seqs else 1

        # Ensemble des numeros déjà en base (vérification finale)
        existing_numeros = {num for (num,) in existing_nums}

        # Pré-générer des QR codes uniques sans toucher à la session
        # (évite l'autoflush sur des objets en attente).
        used_qr_in_db = {
            r.code_qr for r in RecuPaiement.query.filter(
                RecuPaiement.code_qr != None
            ).with_entities(RecuPaiement.code_qr).all()
        }
        used_qr_this_batch: set = set()

        def _gen_unique_qr() -> str:
            while True:
                q = _uuid.uuid4().hex[:16].upper()
                if q not in used_qr_in_db and q not in used_qr_this_batch:
                    used_qr_this_batch.add(q)
                    return q

        nouveaux = []
        seq = seq_start
        created = 0
        while created < quantite:
            if type_recu != 'bulletin':
                prefix = _receipt_prefix(type_recu)
                numero = f'{prefix}-{dept}-{seq:03d}-{semestre}-{annee}'
            else:
                numero = f'B{lot}-{dept}-{seq:03d}-{semestre}-{annee}'
            # Sauter tout numéro déjà présent en base (trous hérités d'anciennes tentatives)
            if numero in existing_numeros:
                seq += 1
                continue
            code_qr = _gen_unique_qr()
            r = RecuPaiement(
                numero=numero, code_qr=code_qr,
                lot=lot if type_recu == 'bulletin' else '0',
                dept=dept, annee=annee,
                semestre=semestre, annee_complete=annee_complete,
                montant=montant, montant_lettres=montant_lettres,
                motif=motif, type_recu=type_recu,
            )
            db.session.add(r)
            nouveaux.append(r)
            seq += 1
            created += 1

        db.session.flush()

        if action in ('generer_pdf', 'generer_recours_pdf'):
            db.session.commit()
            pdf_buf = _generer_recus_pdf(nouveaux, type_recu=type_recu)
            if type_recu == 'recours':
                fname = f'Recus_Recours_{dept}_{semestre}.pdf'
            elif type_recu == 'resultat_recours':
                fname = f'Recus_Resultat_Recours_{dept}_{semestre}.pdf'
            elif type_recu == 'session_2':
                fname = f'Recus_2e_Session_{dept}_{semestre}.pdf'
            elif type_recu == 'recours_session_2':
                fname = f'Recus_Recours_2e_Session_{dept}_{semestre}.pdf'
            elif type_recu == 'recours_session_2_soumission':
                fname = f'Recus_Soumission_Recours_2e_Session_{dept}_{semestre}.pdf'
            elif type_recu == 'releve':
                fname = f'Recus_Releves_Cotes_{dept}_{semestre}.pdf'
            else:
                fname = f'Recus_{dept}_Lot{lot}_{semestre}.pdf'
            return send_file(pdf_buf, as_attachment=True,
                             download_name=fname,
                             mimetype='application/pdf')
        else:
            db.session.commit()
            if type_recu == 'recours':
                flash(f'✅ {quantite} reçus de soumission de recours générés (RS-{dept}-xxx-{semestre}-{annee}).', 'success')
            elif type_recu == 'resultat_recours':
                flash(f'✅ {quantite} reçus de consultation de recours générés (RR-{dept}-xxx-{semestre}-{annee}).', 'success')
            elif type_recu == 'session_2':
                flash(f'✅ {quantite} reçus de 2ème session générés (S2-{dept}-xxx-{semestre}-{annee}).', 'success')
            elif type_recu == 'recours_session_2':
                flash(f'✅ {quantite} reçus de recours de 2ème session générés (R2-{dept}-xxx-{semestre}-{annee}).', 'success')
            elif type_recu == 'recours_session_2_soumission':
                flash(f'✅ {quantite} reçus de soumission de recours 2ème session générés (RS2-{dept}-xxx-{semestre}-{annee}).', 'success')
            else:
                flash(f'✅ {quantite} reçus générés (Lot {lot} – {dept} – {semestre}).', 'success')
            return redirect(url_for('decanat_recus'))

    # Stats
    total        = RecuPaiement.query.count()
    total_utilise= RecuPaiement.query.filter_by(utilise=True).count()
    total_deja_paye = RecuPaiement.query.filter(
        RecuPaiement.tentative_bulletin_id != None,
        RecuPaiement.utilise == False
    ).count()
    total_deja_paye_en_attente = RecuPaiement.query.filter(
        RecuPaiement.tentative_bulletin_id != None,
        RecuPaiement.utilise == False,
        RecuPaiement.tentative_revue == False
    ).count()
    total_deja_paye_revus = total_deja_paye - total_deja_paye_en_attente
    lots         = db.session.query(
        RecuPaiement.lot, RecuPaiement.dept, RecuPaiement.annee,
        db.func.count(RecuPaiement.id).label('nb'),
        db.func.sum(db.cast(RecuPaiement.utilise, db.Integer)).label('nb_util')
    ).group_by(RecuPaiement.lot, RecuPaiement.dept, RecuPaiement.annee).all()

    recus_recents = RecuPaiement.query.filter_by(utilise=True)\
        .order_by(RecuPaiement.date_utilisation.desc()).limit(20).all()

    recus_deja_paye = RecuPaiement.query.filter(
        RecuPaiement.tentative_bulletin_id != None,
        RecuPaiement.utilise == False
    ).order_by(RecuPaiement.tentative_revue.asc(),
               RecuPaiement.date_tentative.desc()).limit(50).all()

    return render_template('decanat_recus.html',
        total=total, total_utilise=total_utilise,
        total_deja_paye=total_deja_paye,
        total_deja_paye_en_attente=total_deja_paye_en_attente,
        total_deja_paye_revus=total_deja_paye_revus,
        lots=lots, recus_recents=recus_recents,
        recus_deja_paye=recus_deja_paye)


@app.route('/decanat/recus/import-pdf', methods=['POST'])
def decanat_recus_import_pdf():
    """Reconnaît comme officiels des reçus papier déjà imprimés/vendus."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    upload = request.files.get('recus_pdf')
    if not upload or not upload.filename:
        flash('Veuillez sélectionner un PDF de reçus.', 'error')
        return redirect(url_for('decanat_recus') + '#import-pdf-recus')
    if not upload.filename.lower().endswith('.pdf'):
        flash('Seuls les fichiers PDF sont acceptés.', 'error')
        return redirect(url_for('decanat_recus') + '#import-pdf-recus')

    try:
        records = _extraire_recus_pdf_import(upload.read())
        numeros = {r['numero'] for r in records}
        codes = {r['code_qr'] for r in records}
        existing_by_numero = {
            r.numero: r for r in RecuPaiement.query.filter(
                RecuPaiement.numero.in_(numeros)
            ).all()
        }
        existing_by_code = {
            r.code_qr: r for r in RecuPaiement.query.filter(
                RecuPaiement.code_qr.in_(codes)
            ).all()
        }

        conflicts = []
        duplicates = 0
        to_create = []
        for record in records:
            by_numero = existing_by_numero.get(record['numero'])
            by_code = existing_by_code.get(record['code_qr'])
            if by_numero or by_code:
                if by_numero and by_numero.code_qr == record['code_qr']:
                    duplicates += 1
                    continue
                conflicts.append(record['numero'])
                continue
            to_create.append(record)

        if conflicts:
            raise ValueError(
                'Conflit détecté pour les reçus : '
                + ', '.join(conflicts[:8])
                + ('…' if len(conflicts) > 8 else '')
                + '. Aucun reçu n’a été importé.'
            )

        for record in to_create:
            record_type = record.get('type_recu', 'bulletin')
            if record_type == 'releve':
                # Le montant du reçu REL est celui imprimé par le PDF importé ;
                # l'extraction historique ne le déduit pas du type.
                record_montant, record_montant_lettres = '', ''
            else:
                record_montant, record_montant_lettres = _receipt_payment_details(record_type)
            db.session.add(RecuPaiement(
                numero=record['numero'],
                code_qr=record['code_qr'],
                lot=record['lot'],
                dept=record['dept'],
                annee=record['annee'],
                semestre=record['semestre'],
                annee_complete=record['annee_complete'],
                montant=record_montant,
                montant_lettres=record_montant_lettres,
                motif=record['motif'],
                type_recu=record_type,
                utilise=False,
                tentative_revue=False,
            ))
        db.session.commit()
        flash(
            f'✅ {len(to_create)} reçu(s) importé(s) et activé(s) comme '
            f'reçus officiels à usage unique.'
            + (f' {duplicates} déjà présent(s) ignoré(s).' if duplicates else ''),
            'success',
        )
    except ValueError as exc:
        db.session.rollback()
        flash(f'❌ Import annulé : {exc}', 'error')
    except Exception as exc:
        db.session.rollback()
        app.logger.exception('Import PDF reçus échoué')
        flash(f'❌ Import annulé : erreur technique ({exc}).', 'error')
    return redirect(url_for('decanat_recus') + '#import-pdf-recus')


@app.route('/decanat/recus/lot-pdf', methods=['POST'])
def decanat_recus_lot_pdf():
    """Regénère le PDF d'un lot existant."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    lot  = request.form.get('lot')
    dept = request.form.get('dept')
    recus = RecuPaiement.query.filter_by(lot=lot, dept=dept)\
        .order_by(RecuPaiement.id).all()
    if not recus:
        flash('Aucun reçu trouvé pour ce lot.', 'error')
        return redirect(url_for('decanat_recus'))
    pdf_buf = _generer_recus_pdf(recus)
    return send_file(pdf_buf, as_attachment=True,
                     download_name=f'Recus_{dept}_Lot{lot}.pdf',
                     mimetype='application/pdf')


@app.route('/decanat/recus/supprimer-lot', methods=['POST'])
def decanat_recus_supprimer_lot():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    lot  = request.form.get('lot')
    dept = request.form.get('dept')
    RecuPaiement.query.filter_by(lot=lot, dept=dept, utilise=False).delete()
    db.session.commit()
    flash('Reçus non utilisés supprimés.', 'success')
    return redirect(url_for('decanat_recus'))


@app.route('/decanat/recus/marquer-revue/<int:recu_id>', methods=['POST'])
def decanat_recus_marquer_revue(recu_id):
    """Marque (ou démarque) un reçu 'déjà payé' comme revu par le DÉCANAT."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    recu = RecuPaiement.query.get_or_404(recu_id)
    recu.tentative_revue = not recu.tentative_revue
    db.session.commit()
    flash(
        '✅ Reçu marqué comme revu.' if recu.tentative_revue
        else 'ℹ️ Marquage supprimé — reçu de nouveau en attente de revue.',
        'success'
    )
    return redirect(url_for('decanat_recus') + '#section-deja-paye')


@app.route('/decanat/recus/verifier', methods=['GET', 'POST'])
def decanat_recus_verifier():
    """Vérifie un reçu par référence imprimée, code QR ou URL de QR."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    valeur = (
        request.form.get('recherche', '').strip()
        if request.method == 'POST'
        else request.args.get('q', '').strip()
    )
    recu = None
    origine = ''
    if valeur:
        recu, origine = _chercher_recu(valeur)
        _enregistrer_admin_audit(
            'verification_recu',
            {
                'recherche': valeur[:120],
                'trouve': bool(recu),
                'origine': origine,
                'recu_id': recu.id if recu else None,
            },
        )
        db.session.commit()
        if not recu:
            flash('Aucun reçu officiel ne correspond à cette référence ou ce QR.', 'error')

    utilisation = _description_utilisation_recu(recu) if recu else None
    return render_template(
        'decanat_verifier_recu.html',
        recherche=valeur,
        recu=recu,
        origine=origine,
        utilisation=utilisation,
    )


# ── Route publique : scan QR d'un reçu ──────────────────────────────

@app.route('/scan/<code>', methods=['GET', 'POST'])
def scan_recu(code):
    ip_client = request.headers.get('X-Forwarded-For', request.remote_addr or '').split(',')[0].strip()
    # Verrouiller la ligne pendant le scan : deux téléphones qui présentent
    # exactement le même reçu au même instant ne doivent jamais pouvoir le
    # consommer tous les deux. PostgreSQL applique ce verrou nativement ;
    # SQLite l'ignore mais conserve le comportement historique local.
    recu = RecuPaiement.query.filter_by(code_qr=code.upper()).with_for_update().first()

    if not recu:
        db.session.add(ScanLog(code=code, ip=ip_client, resultat='invalide'))
        db.session.commit()
        _send_scan_alert_email(code, 'invalide', ip_client)
        return render_template('scan_recu.html', mode='invalid', code=code)

    # Le préfixe imprimé est canonique : il empêche un reçu RR- de passer pour
    # un reçu de 2ème session à cause d'un type_recu historique mal enregistré.
    canonical_type = _receipt_type_from_number(recu.numero)
    if not _receipt_type_is_consistent(recu):
        db.session.add(ScanLog(code=code, ip=ip_client, resultat='invalide'))
        db.session.commit()
        return render_template(
            'scan_recu.html',
            mode='type_mismatch',
            recu=recu,
            error='Ce reçu est incohérent : son numéro ne correspond pas à son type. '
                  'Présentez un reçu officiel de l’étape demandée au DÉCANAT.',
        )

    # Quand le scan est lancé depuis une page de bulletin verrouillée, garder
    # le bulletin demandé dans le contexte. Sans ce lien, un reçu RR- présenté
    # depuis la page de 2ème session pouvait débloquer le recours de 1ère
    # session du même matricule, car la recherche ne connaissait que le type
    # du reçu et le matricule.
    target_bd = None
    target_bid = request.args.get('bid', '').strip()
    if target_bid:
        try:
            target_bd = db.session.get(BulletinData, int(target_bid))
        except (TypeError, ValueError):
            target_bd = None
        if target_bd is None:
            return render_template(
                'scan_recu.html',
                mode='type_mismatch',
                recu=recu,
                error='Le bulletin demandé n’existe plus. Recommencez depuis le portail.',
            )
        expected_type = _type_recu_attendu_pour_bulletin(target_bd)
        if expected_type != canonical_type:
            expected_label = _RECEIPT_TYPE_LABELS.get(expected_type, 'reçu approprié')
            expected_prefix = _receipt_prefix(expected_type) + '-'
            return render_template(
                'scan_recu.html',
                mode='type_mismatch',
                recu=recu,
                error=f'Ce reçu ne correspond pas à cette étape. '
                      f'Cette page nécessite un reçu {expected_prefix} '
                      f'({expected_label}).',
            )

    # Les reçus RS sont réservés à la soumission d'un recours. Ils ne peuvent
    # jamais être utilisés pour débloquer un bulletin. Les quatre reçus de
    # consultation sont, eux, strictement liés à leur étape.
    if canonical_type == 'recours':
        db.session.add(ScanLog(code=code, ip=ip_client, resultat='invalide'))
        db.session.commit()
        return render_template(
            'scan_recu.html',
            mode='type_mismatch',
            recu=recu,
            error='Ce reçu est réservé à la soumission d’un recours. '
                  'Il ne peut pas débloquer un bulletin.',
        )

    if recu.utilise:
        db.session.add(ScanLog(code=code, ip=ip_client, resultat='deja_utilise',
                               matricule=recu.matricule_etudiant,
                               nom_etudiant=recu.nom_etudiant))
        db.session.commit()
        _send_scan_alert_email(code, 'deja_utilise', ip_client,
                               recu.matricule_etudiant, recu.nom_etudiant)
        return render_template('scan_recu.html', mode='deja_utilise', recu=recu)

    if request.method == 'POST':
        matricule = request.form.get('matricule', '').strip().upper()
        if not matricule:
            return render_template('scan_recu.html', mode='scan_ok', recu=recu,
                                   error='Veuillez saisir votre matricule.')

        type_attendu = canonical_type
        if type_attendu not in _RECEIPT_TO_GRID_TYPE:
            return render_template(
                'scan_recu.html',
                mode='type_mismatch',
                recu=recu,
                error='Type de reçu non autorisé pour cette consultation.',
            )

        # Un reçu initial ne peut sélectionner qu'un bulletin initial. De la
        # même manière, un reçu RR ne peut sélectionner qu'un résultat de
        # recours. Cette séparation protège les deux paiements indépendants.
        type_grille = _RECEIPT_TO_GRID_TYPE[type_attendu]
        if target_bd is not None:
            if (
                target_bd.matricule.upper() != matricule
                or getattr(target_bd.bul_session, 'type_grille', 'initial') != type_grille
            ):
                return render_template(
                    'scan_recu.html',
                    mode='scan_ok',
                    recu=recu,
                    error='Le matricule saisi ne correspond pas à cette étape de bulletin.',
                )
            bulletins = [target_bd]
        else:
            bulletins = BulletinData.query.filter(
                db.func.upper(BulletinData.matricule) == matricule,
                BulletinData.session_id.in_(
                    db.session.query(BulletinSession.id).filter(
                        BulletinSession.type_grille == type_grille
                    )
                ),
            ).all()

        if not bulletins:
            label = _RECEIPT_TYPE_LABELS.get(type_attendu, 'bulletin')
            return render_template('scan_recu.html', mode='scan_ok', recu=recu,
                                   error=f'Aucun {label} trouvé pour le matricule « {matricule} ».')

        # Si plusieurs bulletins, prendre le plus récent non payé, sinon le plus récent
        bd = next((b for b in sorted(bulletins, key=lambda x: x.id, reverse=True) if not b.paye), None)
        if not bd:
            bd = sorted(bulletins, key=lambda x: x.id, reverse=True)[0]

        # Si le bulletin est déjà payé, ne pas consommer le reçu
        if bd.paye:
            # Mémoriser la tentative sur le reçu (sans le marquer utilisé)
            if not recu.tentative_bulletin_id:
                recu.tentative_bulletin_id = bd.id
                recu.date_tentative = now_cat()
            recu.tentative_matricule = matricule
            recu.tentative_nom = bd.nom
            db.session.add(ScanLog(code=code, ip=ip_client, resultat='deja_paye',
                                   matricule=matricule, nom_etudiant=bd.nom))
            db.session.commit()
            return render_template('scan_recu.html', mode='deja_paye', recu=recu,
                                   bd=bd, matricule=matricule)

        # Marquer le reçu comme utilisé
        recu.utilise = True
        recu.date_utilisation = now_cat()
        recu.matricule_etudiant = matricule
        recu.nom_etudiant = bd.nom
        recu.bulletin_id = bd.id

        # Débloquer le bulletin — chaque reçu scanné vaut 5000 FC
        bd.paye = True
        bd.date_paiement = now_cat()
        bd.methode_paiement = 'Reçu papier'
        bd.reference_paiement = recu.numero
        montant_session = bd.bul_session.montant_fc if bd.bul_session else 5000
        bd.montant_paye = montant_session if montant_session else 5000

        db.session.add(ScanLog(code=code, ip=ip_client, resultat='ok',
                               matricule=matricule, nom_etudiant=bd.nom))
        db.session.commit()
        return redirect(url_for('bulletins_result', bid=bd.id))

    return render_template('scan_recu.html', mode='scan_ok', recu=recu)


def _run_migrations():
    """Add columns that were introduced after the initial schema creation."""
    from sqlalchemy import inspect as sa_inspect, text
    inspector = sa_inspect(db.engine)

    # ── recus_paiement migrations ──
    existing_cols = {c['name'] for c in inspector.get_columns('recus_paiement')}
    migrations = [
        ("type_recu", "VARCHAR(40) NOT NULL DEFAULT 'bulletin'"),
        ("tentative_bulletin_id", "INTEGER REFERENCES bulletin_data(id)"),
        ("tentative_matricule",   "VARCHAR(30)"),
        ("tentative_nom",         "VARCHAR(200)"),
        ("date_tentative",        "DATETIME"),
        ("tentative_revue",       "BOOLEAN NOT NULL DEFAULT 0"),
    ]
    for col_name, col_def in migrations:
        if col_name not in existing_cols:
            db.session.execute(
                text(f"ALTER TABLE recus_paiement ADD COLUMN {col_name} {col_def}")
            )

    # ── bulletin_sessions migrations ──
    if inspector.has_table('bulletin_sessions'):
        session_cols = {c['name'] for c in inspector.get_columns('bulletin_sessions')}
        if 'type_grille' not in session_cols:
            db.session.execute(text(
                "ALTER TABLE bulletin_sessions "
                "ADD COLUMN type_grille VARCHAR(20) NOT NULL DEFAULT 'initial'"
            ))

    # ── paiement_audits migrations ──
    audit_cols = {c['name'] for c in inspector.get_columns('paiement_audits')}
    audit_migrations = [
        ("operator_name", "VARCHAR(100)"),
    ]
    for col_name, col_def in audit_migrations:
        if col_name not in audit_cols:
            db.session.execute(
                text(f"ALTER TABLE paiement_audits ADD COLUMN {col_name} {col_def}")
            )

    # ── scan_logs migrations ──
    scan_cols = {c['name'] for c in inspector.get_columns('scan_logs')}
    scan_migrations = [
        ("dismissed", "BOOLEAN NOT NULL DEFAULT 0"),
    ]
    for col_name, col_def in scan_migrations:
        if col_name not in scan_cols:
            db.session.execute(
                text(f"ALTER TABLE scan_logs ADD COLUMN {col_name} {col_def}")
            )

    # La table des commandes est créée par create_all pour les nouvelles
    # installations ; cette vérification garde le démarrage explicite sur les
    # bases historiques.
    if not inspector.has_table('releve_commandes'):
        ReleveCommande.__table__.create(bind=db.engine, checkfirst=True)

    # ── recours migrations ──
    if inspector.has_table('recours'):
        recours_cols = {c['name'] for c in inspector.get_columns('recours')}
        recours_migrations = [
            ("carte_etudiant_filename", "VARCHAR(200)"),
            ("preuves_json",            "TEXT DEFAULT '{}'"),
            ("type_recu",               "VARCHAR(40) NOT NULL DEFAULT 'recours'"),
        ]
        for col_name, col_def in recours_migrations:
            if col_name not in recours_cols:
                db.session.execute(
                    text(f"ALTER TABLE recours ADD COLUMN {col_name} {col_def}")
                )

    # ── actualites migrations ──
    if inspector.has_table('actualites'):
        actualites_cols = {c['name'] for c in inspector.get_columns('actualites')}
        actualites_migrations = [
            ("type_publication", "VARCHAR(20) NOT NULL DEFAULT 'actualite'"),
            ("epingle", "BOOLEAN NOT NULL DEFAULT FALSE"),
        ]
        for col_name, col_def in actualites_migrations:
            if col_name not in actualites_cols:
                db.session.execute(
                    text(f"ALTER TABLE actualites ADD COLUMN {col_name} {col_def}")
                )

    db.session.commit()

    # Les reçus RS existants ont été créés avant la distinction explicite des
    # usages. Les rattacher une seule fois à la soumission de recours permet
    # de préserver leur fonctionnement actuel sans les rendre utilisables
    # pour débloquer un résultat.
    try:
        db.session.execute(text(
            "UPDATE recus_paiement SET type_recu = 'recours' "
            "WHERE numero LIKE 'RS-%' AND (type_recu IS NULL OR type_recu = 'bulletin')"
        ))
        db.session.execute(text(
            "UPDATE recus_paiement SET type_recu = 'resultat_recours' "
            "WHERE numero LIKE 'RR-%' AND (type_recu IS NULL OR type_recu = 'bulletin')"
        ))
        db.session.execute(text(
            "UPDATE recus_paiement SET type_recu = 'session_2' "
            "WHERE numero LIKE 'S2-%' AND (type_recu IS NULL OR type_recu = 'bulletin')"
        ))
        db.session.execute(text(
            "UPDATE recus_paiement SET type_recu = 'recours_session_2' "
            "WHERE numero LIKE 'R2-%' AND (type_recu IS NULL OR type_recu = 'bulletin')"
        ))
        db.session.execute(text(
            "UPDATE recus_paiement SET type_recu = 'recours_session_2_soumission' "
            "WHERE numero LIKE 'RS2-%' AND (type_recu IS NULL OR type_recu = 'bulletin')"
        ))
        db.session.commit()
    except Exception:
        db.session.rollback()

    # ── backfill moy_col_present for bulletins imported before the flag existed ──
    # Heuristic (session-level):
    #   • ALL students in the session have moyenne == 0.0
    #   • AND at least one has a non-zero total
    # → the Moyenne column was present in the grille but contained no values
    #   → moy_col_present = True  (download block will apply)
    # Otherwise → moy_col_present = False (column was absent; no block needed)
    try:
        # Collect rows that are missing the flag
        rows_to_check = BulletinData.query.with_entities(
            BulletinData.id, BulletinData.session_id, BulletinData.data_json
        ).all()

        # Build per-session lists, keeping only rows without the flag
        from collections import defaultdict
        session_rows = defaultdict(list)   # session_id → list of (id, data_json_str)
        for row in rows_to_check:
            try:
                etu = json.loads(row.data_json) if row.data_json else {}
            except Exception:
                etu = {}
            if 'moy_col_present' not in etu:
                session_rows[row.session_id].append((row.id, row.data_json, etu))

        if session_rows:
            # For each session that has unflagged rows, we need the full session picture
            # (all rows, not just unflagged) to apply the all-zero test correctly.
            all_by_session = defaultdict(list)
            for row in rows_to_check:
                try:
                    etu = json.loads(row.data_json) if row.data_json else {}
                except Exception:
                    etu = {}
                all_by_session[row.session_id].append(etu)

            for sess_id, unflagged in session_rows.items():
                all_etu = all_by_session[sess_id]
                all_zero    = all(e.get('moyenne', 0) == 0.0 for e in all_etu)
                any_nonzero_total = any(e.get('total', 0) != 0 for e in all_etu)
                inferred = bool(all_zero and any_nonzero_total)

                for (row_id, data_json_str, etu) in unflagged:
                    etu['moy_col_present'] = inferred
                    BulletinData.query.filter_by(id=row_id).update(
                        {'data_json': json.dumps(etu, ensure_ascii=False)}
                    )

            db.session.commit()
    except Exception as _backfill_err:
        import logging
        logging.getLogger(__name__).warning(
            "moy_col_present backfill skipped: %s", _backfill_err
        )
        db.session.rollback()


# ─── Backup / Restore ────────────────────────────────────────────────────────

def _row_to_dict(row):
    """Convert a SQLAlchemy model instance to a plain dict (datetime/date → ISO string)."""
    d = {}
    for col in row.__table__.columns:
        val = getattr(row, col.name)
        if isinstance(val, datetime):
            val = val.isoformat()
        elif isinstance(val, date):
            val = val.isoformat()
        d[col.name] = val
    return d


@app.route('/decanat/backup')
def decanat_backup():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    _BACKUP_TABLES = [
        ('etudiants',          Etudiant),
        ('professeurs',        Professeur),
        ('cours',              Cours),
        ('presences',          Presence),
        ('horaires',           Horaire),
        ('actualites',         Actualite),
        ('page_contents',      PageContent),
        ('liste_identifiants', ListeIdentifiants),
        ('bulletin_sessions',  BulletinSession),
        ('bulletin_data',      BulletinData),
        ('paiement_audits',    PaiementAudit),
        ('administration_audits', AdministrationAudit),
        ('recus_paiement',     RecuPaiement),
        ('releve_commandes',   ReleveCommande),
        ('recours',            Recours),
        ('app_config',         AppConfig),
        ('scan_logs',          ScanLog),
    ]

    backup = {
        'version': 1,
        'exported_at': now_cat().isoformat(),
        'tables': {},
        'files': [],
    }
    for tname, model in _BACKUP_TABLES:
        backup['tables'][tname] = [_row_to_dict(r) for r in model.query.all()]
    backup['files'] = [
        {
            'storage_key': asset.storage_key,
            'mime_type': asset.mime_type,
            'size_bytes': asset.size_bytes,
            'data_b64': base64.b64encode(asset.data).decode('ascii'),
        }
        for asset in FileAsset.query.all()
    ]

    filename = f"sauvegarde_esciales_{now_cat().strftime('%Y%m%d_%H%M%S')}.json"
    data = json.dumps(backup, ensure_ascii=False, indent=2)
    from flask import Response
    return Response(
        data,
        mimetype='application/json',
        headers={'Content-Disposition': f'attachment; filename="{filename}"'},
    )


@app.route('/decanat/restore', methods=['GET', 'POST'])
def decanat_restore():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    if request.method == 'GET':
        return render_template('decanat_restore.html', counts=None)

    # ── parse & preflight ────────────────────────────────────────────────────
    f = request.files.get('backup_file')
    if not f or not f.filename.lower().endswith('.json'):
        flash('Veuillez choisir un fichier de sauvegarde .json valide.', 'error')
        return render_template('decanat_restore.html', counts=None)

    try:
        backup = json.loads(f.read().decode('utf-8'))
    except Exception:
        flash('Fichier invalide — impossible de lire le JSON.', 'error')
        return render_template('decanat_restore.html', counts=None)

    if not isinstance(backup, dict) or 'tables' not in backup:
        flash('Format non reconnu (clé "tables" manquante).', 'error')
        return render_template('decanat_restore.html', counts=None)

    if backup.get('version', 1) != 1:
        flash(f'Version de sauvegarde non supportée ({backup.get("version")}).', 'error')
        return render_template('decanat_restore.html', counts=None)

    td = backup.get('tables', {})
    if not isinstance(td, dict):
        flash('Format invalide : "tables" doit être un objet JSON.', 'error')
        return render_template('decanat_restore.html', counts=None)

    # ── helpers ──────────────────────────────────────────────────────────────
    def _dt(s):
        if not s:
            return None
        try:
            return datetime.fromisoformat(s)
        except Exception:
            return None

    def _d(s):
        if not s:
            return None
        try:
            return date.fromisoformat(s)
        except Exception:
            return None

    def _find_or_insert(model_cls, natural_key_kw, create_fn, backup_id=None):
        """Return (obj, was_inserted).

        Look up by natural key first.  If found, return existing row (no insert).
        If not found, create with create_fn().  Try to preserve backup_id; if that
        slot is already taken (cross-server ID collision), let SQLAlchemy
        auto-assign — the row is still inserted, never silently dropped.
        Flush so obj.id is populated before returning.
        """
        existing = model_cls.query.filter_by(**natural_key_kw).first()
        if existing:
            return existing, False
        obj = create_fn()
        if backup_id is not None and not model_cls.query.get(backup_id):
            obj.id = backup_id           # slot is free: preserve original ID
        db.session.add(obj)
        db.session.flush()               # populate obj.id immediately
        return obj, True

    def _insert_preserve_id(model_cls, create_fn, backup_id=None):
        """Insert a row for tables with no natural unique key.

        Tries to preserve backup_id.  If that ID is already taken (ID collision
        in a non-empty target DB), falls back to a new auto-assigned ID — the
        row is ALWAYS inserted, never silently dropped.  Flush before returning.
        """
        obj = create_fn()
        if backup_id is not None and not model_cls.query.get(backup_id):
            obj.id = backup_id
        db.session.add(obj)
        db.session.flush()
        return obj

    counts = {}

    try:
        # ── 1. etudiants (natural key: matricule) ────────────────────────────
        n = 0
        etudiant_id_map: dict = {}       # backup_id → live_id
        for r in td.get('etudiants', []):
            obj, inserted = _find_or_insert(
                Etudiant, {'matricule': r['matricule']},
                lambda r=r: Etudiant(
                    matricule=r['matricule'], nom=r['nom'], postnom=r['postnom'],
                    prenom=r['prenom'], sexe=r['sexe'], telephone=r['telephone'],
                    promotion=r['promotion'], departement=r['departement'],
                    photo=r.get('photo'), qrcode_path=r.get('qrcode_path'),
                    date_inscription=_dt(r.get('date_inscription')) or now_cat(),
                ),
                backup_id=r.get('id'),
            )
            if r.get('id') is not None:
                etudiant_id_map[r['id']] = obj.id
            if inserted:
                n += 1
        counts['Étudiants'] = n

        # ── 2. professeurs (natural key: matricule) ──────────────────────────
        n = 0
        professeur_id_map: dict = {}
        for r in td.get('professeurs', []):
            obj, inserted = _find_or_insert(
                Professeur, {'matricule': r['matricule']},
                lambda r=r: Professeur(
                    matricule=r['matricule'], nom=r['nom'], postnom=r['postnom'],
                    prenom=r['prenom'], telephone=r.get('telephone'),
                    departement=r.get('departement'), qrcode_path=r.get('qrcode_path'),
                ),
                backup_id=r.get('id'),
            )
            if r.get('id') is not None:
                professeur_id_map[r['id']] = obj.id
            if inserted:
                n += 1
        counts['Professeurs'] = n

        # ── 3. cours (natural key: code; FK: professeur_id) ──────────────────
        n = 0
        cours_id_map: dict = {}
        for r in td.get('cours', []):
            # remap professeur_id using the map built above
            prof_id = r.get('professeur_id')
            if prof_id is not None:
                prof_id = professeur_id_map.get(prof_id, prof_id)
            obj, inserted = _find_or_insert(
                Cours, {'code': r['code']},
                lambda r=r, pid=prof_id: Cours(
                    code=r['code'], nom=r['nom'],
                    departement=r['departement'], promotion=r['promotion'],
                    professeur_id=pid,
                ),
                backup_id=r.get('id'),
            )
            if r.get('id') is not None:
                cours_id_map[r['id']] = obj.id
            if inserted:
                n += 1
        counts['Cours'] = n

        # ── 4. presences (no natural key; FK: etudiant, professeur, cours) ─────
        n = 0
        for r in td.get('presences', []):
            etu_id  = etudiant_id_map.get(r.get('etudiant_id'),    r.get('etudiant_id'))
            prof_id = professeur_id_map.get(r.get('professeur_id'), r.get('professeur_id'))
            crs_id  = cours_id_map.get(r.get('cours_id'),           r.get('cours_id'))
            _insert_preserve_id(
                Presence,
                lambda r=r, ei=etu_id, pi=prof_id, ci=crs_id: Presence(
                    etudiant_id=ei, professeur_id=pi, cours_id=ci,
                    heure_entree=_dt(r.get('heure_entree')),
                    date=_d(r.get('date')), type_presence=r.get('type_presence'),
                ),
                backup_id=r.get('id'),
            )
            n += 1
        counts['Présences'] = n

        # ── 5. horaires (no natural key) ─────────────────────────────────────
        n = 0
        for r in td.get('horaires', []):
            _insert_preserve_id(
                Horaire,
                lambda r=r: Horaire(
                    departement=r['departement'], promotion=r['promotion'],
                    type_horaire=r['type_horaire'], fichier=r['fichier'],
                    date_publication=_dt(r.get('date_publication')),
                ),
                backup_id=r.get('id'),
            )
            n += 1
        counts['Horaires'] = n

        # ── 6. actualites (no natural key) ───────────────────────────────────
        n = 0
        for r in td.get('actualites', []):
            _insert_preserve_id(
                Actualite,
                lambda r=r: Actualite(
                    titre=r['titre'], description=r['description'],
                    image=r.get('image'),
                    date_publication=_dt(r.get('date_publication')),
                    publie=r.get('publie', True),
                ),
                backup_id=r.get('id'),
            )
            n += 1
        counts['Actualités'] = n

        # ── 7. page_contents (natural key: page_name) ────────────────────────
        n = 0
        for r in td.get('page_contents', []):
            _, inserted = _find_or_insert(
                PageContent, {'page_name': r['page_name']},
                lambda r=r: PageContent(
                    page_name=r['page_name'], content_json=r['content_json'],
                    image_principale=r.get('image_principale'),
                    date_modification=_dt(r.get('date_modification')),
                ),
                backup_id=r.get('id'),
            )
            if inserted:
                n += 1
        counts['Contenus pages'] = n

        # ── 8. liste_identifiants (natural key: promotion+matricule) ─────────
        n = 0
        for r in td.get('liste_identifiants', []):
            _, inserted = _find_or_insert(
                ListeIdentifiants,
                {'promotion': r['promotion'], 'matricule': r['matricule']},
                lambda r=r: ListeIdentifiants(
                    promotion=r['promotion'], nom=r['nom'],
                    nom_norm=r.get('nom_norm') or _normaliser_nom(r['nom']),
                    matricule=r['matricule'], mot_de_passe=r.get('mot_de_passe'),
                    date_import=_dt(r.get('date_import')),
                ),
                backup_id=r.get('id'),
            )
            if inserted:
                n += 1
        counts['Identifiants'] = n

        # ── 9. bulletin_sessions (composite natural key: annee+session_acad+
        #        semestre+promotion; builds session_id_map for downstream FK) ──
        n = 0
        session_id_map: dict = {}        # backup_id → live_id
        for r in td.get('bulletin_sessions', []):
            obj, inserted = _find_or_insert(
                BulletinSession,
                {
                    'annee':        r.get('annee', ''),
                    'session_acad': r.get('session_acad', ''),
                    'semestre':     r.get('semestre', ''),
                    'promotion':    r.get('promotion', ''),
                    'type_grille':  r.get('type_grille', 'initial'),
                },
                lambda r=r: BulletinSession(
                    nom=r.get('nom'), annee=r.get('annee'),
                    session_acad=r.get('session_acad'), semestre=r.get('semestre'),
                    promotion=r.get('promotion'), montant_fc=r.get('montant_fc', 5000),
                    departement=r.get('departement', ''),
                    texte_intro=r.get('texte_intro', ''),
                    type_grille=r.get('type_grille', 'initial'),
                    date_import=_dt(r.get('date_import')),
                ),
                backup_id=r.get('id'),
            )
            if r.get('id') is not None:
                session_id_map[r['id']] = obj.id
            if inserted:
                n += 1
        counts['Sessions bulletins'] = n

        # ── 10. bulletin_data (natural key: numero_bulletin; FK: session_id) ──
        n = 0
        bulletin_id_map: dict = {}       # backup_id → live_id
        for r in td.get('bulletin_data', []):
            # remap session_id through the map built in step 9
            sess_id = session_id_map.get(r.get('session_id'), r.get('session_id'))
            obj, inserted = _find_or_insert(
                BulletinData, {'numero_bulletin': r['numero_bulletin']},
                lambda r=r, sid=sess_id: BulletinData(
                    session_id=sid, matricule=r.get('matricule'),
                    nom=r.get('nom'), sexe=r.get('sexe'), data_json=r.get('data_json'),
                    numero_bulletin=r['numero_bulletin'],
                    paye=r.get('paye', False),
                    date_paiement=_dt(r.get('date_paiement')),
                    methode_paiement=r.get('methode_paiement'),
                    telephone_paiement=r.get('telephone_paiement'),
                    reference_paiement=r.get('reference_paiement'),
                    montant_paye=r.get('montant_paye', 0),
                    nb_telechargements=r.get('nb_telechargements', 0),
                    date_dernier_telechargement=_dt(r.get('date_dernier_telechargement')),
                ),
                backup_id=r.get('id'),
            )
            if r.get('id') is not None:
                bulletin_id_map[r['id']] = obj.id
            if inserted:
                n += 1
        counts['Bulletins'] = n

        # ── 11. paiement_audits (no natural key; FK: bulletin_id) ───────────
        n = 0
        for r in td.get('paiement_audits', []):
            bul_id = bulletin_id_map.get(r['bulletin_id'], r['bulletin_id'])
            _insert_preserve_id(
                PaiementAudit,
                lambda r=r, b=bul_id: PaiementAudit(
                    bulletin_id=b,
                    old_montant=r.get('old_montant'), new_montant=r.get('new_montant'),
                    old_methode=r.get('old_methode'), new_methode=r.get('new_methode'),
                    old_reference=r.get('old_reference'), new_reference=r.get('new_reference'),
                    old_date_paiement=_dt(r.get('old_date_paiement')),
                    new_date_paiement=_dt(r.get('new_date_paiement')),
                    date_modification=_dt(r.get('date_modification')) or now_cat(),
                    operator_name=r.get('operator_name'),
                ),
                backup_id=r.get('id'),
            )
            n += 1
        counts['Audits paiements'] = n

        # ── 12. recus_paiement (natural key: numero; FK: bulletin_id x2) ─────
        n = 0
        for r in td.get('recus_paiement', []):
            bul_id  = bulletin_id_map.get(r.get('bulletin_id'),            r.get('bulletin_id'))
            tent_id = bulletin_id_map.get(r.get('tentative_bulletin_id'),  r.get('tentative_bulletin_id'))
            _, inserted = _find_or_insert(
                RecuPaiement, {'numero': r['numero']},
                lambda r=r, b=bul_id, t=tent_id: RecuPaiement(
                    numero=r['numero'], code_qr=r['code_qr'],
                    lot=r.get('lot'), dept=r.get('dept'),
                    annee=r.get('annee'), semestre=r.get('semestre'),
                    montant=r.get('montant', '5000 CDF'),
                    montant_lettres=r.get('montant_lettres', 'Cinq mille Francs congolais'),
                    motif=r.get('motif', 'Bulletin des résultats'),
                    type_recu=r.get('type_recu', 'bulletin'),
                    annee_complete=r.get('annee_complete', '2025-2026'),
                    date_creation=_dt(r.get('date_creation')),
                    utilise=r.get('utilise', False),
                    date_utilisation=_dt(r.get('date_utilisation')),
                    matricule_etudiant=r.get('matricule_etudiant'),
                    nom_etudiant=r.get('nom_etudiant'),
                    bulletin_id=b, tentative_bulletin_id=t,
                    tentative_matricule=r.get('tentative_matricule'),
                    tentative_nom=r.get('tentative_nom'),
                    date_tentative=_dt(r.get('date_tentative')),
                    tentative_revue=r.get('tentative_revue', False),
                ),
                backup_id=r.get('id'),
            )
            if inserted:
                n += 1
        counts['Reçus'] = n

        # ── 13. app_config (natural key: key) ────────────────────────────────
        n = 0
        for r in td.get('app_config', []):
            _, inserted = _find_or_insert(
                AppConfig, {'key': r['key']},
                lambda r=r: AppConfig(key=r['key'], value=r.get('value', '')),
                backup_id=r.get('id'),
            )
            if inserted:
                n += 1
        counts['Configuration'] = n

        # ── 14. scan_logs (no natural key) ───────────────────────────────────
        n = 0
        for r in td.get('scan_logs', []):
            _insert_preserve_id(
                ScanLog,
                lambda r=r: ScanLog(
                    code=r['code'], ip=r.get('ip'),
                    date_scan=_dt(r.get('date_scan')),
                    resultat=r['resultat'], matricule=r.get('matricule'),
                    nom_etudiant=r.get('nom_etudiant'),
                    dismissed=r.get('dismissed', False),
                ),
                backup_id=r.get('id'),
            )
            n += 1
        counts['Scans'] = n

        db.session.commit()

        total = sum(counts.values())
        details = ', '.join(f"{v} {k}" for k, v in counts.items() if v > 0)
        flash(
            f'Restauration terminée — {total} enregistrement(s) importé(s)'
            + (f' : {details}' if details else '') + '.',
            'success',
        )

    except Exception as _restore_err:
        db.session.rollback()
        flash(f'Erreur lors de la restauration : {_restore_err}', 'error')
        counts = {}

    return render_template('decanat_restore.html', counts=counts)


# ═══════════════════════════════════════════════════════════════════
#  RECOURS — portail public + gestion DÉCANAT
# ═══════════════════════════════════════════════════════════════════

_RECOURS_PROMOTIONS = ['LICENCE 1', 'LICENCE 2', 'LICENCE 3', 'MASTER 1', 'MASTER 2']
_RECOURS_FILIERES   = [
    'Anthropologie', 'Relations Internationales', 'Sciences du travail',
    'Sociologie', 'Science Politique', 'Science Administrative et Management',
]
_RECOURS_MONTANT_USD = 10
_RECOURS_DEPARTEMENT_PAR_FILIERE = {
    'Anthropologie': 'Anthropologie',
    'Relations Internationales': 'RI',
    'Sciences du travail': 'Sociologie',
    'Sociologie': 'Sociologie',
    'Science Politique': 'SPA POL',
    'Science Administrative et Management': 'SPA SAM',
}
_RECOURS_CLAIMS = [
    ('omission',       "Omission de mon nom alors que j'ai payé tous les frais",      None),
    ('orthographe',    "Mauvaise orthographe de nom",                                  'nom_correct'),
    ('pct',            "Erreur dans le calcul du pourcentage annuel",                  None),
    ('sommation',      "Erreur dans la sommation des cotes",                           None),
    ('credits',        "Erreur dans le calcul des crédits validés",                    None),
    ('decision',       "Erreur dans la décision du jury",                              None),
    ('transcription',  "Erreur dans la transcription des cotes dans certains cours",   'cours_concernes'),
    ('manque_cote',    "Manque de cote dans les cours ci-après",                       'cours_manquants'),
    ('memoire',        "Manque de cote de mémoire alors que j'ai défendu",             None),
    ('stage',          "Manque de cote de stage",                                      None),
    ('dispense',       "Manque de cote des cours de dispense",                         'cours_dispense'),
    ('compensation',   "La compensation prévue par le règlement n'a pas été appliquée", None),
    ('autre',          "Autre motif",                                                   'autre_detail'),
]

# Claims qui acceptent une photo-preuve jointe depuis le formulaire (optionnel)
# claim_id → (nom_du_champ_fichier, libellé_bouton)
_RECOURS_CLAIM_FILES = {
    'omission': ('preuve_omission', "Carte d'étudiant"),
    'memoire':  ('preuve_memoire',  'Reçu de dépôt de mémoire'),
    'stage':    ('preuve_stage',    'Reçu des frais de stage'),
}


_RECOURS_CHRONO_DUREES = (1, 24, 48, 72)


def _get_recours_chronos():
    """Retourne l'état des chronos configurés par promotion."""
    try:
        cfg = AppConfig.query.filter_by(key='recours_promo_chronos').first()
        raw = json.loads(cfg.value) if (cfg and cfg.value) else {}
    except Exception:
        raw = {}

    now = now_cat()
    result = {}
    for promo in _RECOURS_PROMOTIONS:
        item = raw.get(promo)
        if not isinstance(item, dict) or not item.get('ends_at'):
            result[promo] = {
                'has_timer': False,
                'hours': None,
                'started_at': None,
                'ends_at': None,
                'remaining_seconds': None,
                'expired': False,
            }
            continue

        try:
            ends_at = datetime.fromisoformat(str(item['ends_at']))
            started_at = (
                datetime.fromisoformat(str(item['started_at']))
                if item.get('started_at') else None
            )
            hours = int(item.get('hours', 0))
            remaining = max(0, int((ends_at - now).total_seconds()))
            result[promo] = {
                'has_timer': True,
                'hours': hours,
                'started_at': started_at.isoformat() if started_at else None,
                'ends_at': ends_at.isoformat(),
                'remaining_seconds': remaining,
                'expired': remaining == 0,
            }
        except (TypeError, ValueError, OverflowError):
            result[promo] = {
                'has_timer': False,
                'hours': None,
                'started_at': None,
                'ends_at': None,
                'remaining_seconds': None,
                'expired': False,
            }
    return result


def _get_recours_actif():
    """Retourne {promo: {filiere: bool}} — True = accepte les soumissions. Défaut : tout actif.
    Seules les valeurs False sont stockées ; l'absence signifie True."""
    try:
        cfg = AppConfig.query.filter_by(key='recours_promo_actif').first()
        raw = json.loads(cfg.value) if (cfg and cfg.value) else {}
    except Exception:
        raw = {}
    chronos = _get_recours_chronos()
    result = {}
    for p in _RECOURS_PROMOTIONS:
        val = raw.get(p, {})
        # Compatibilité ancien format : la valeur était un bool (True/False par promo entière)
        if isinstance(val, bool):
            val = {}   # on ignore l'ancien état — tout est rouvert par défaut
        chrono_ouvert = (
            not chronos[p]['has_timer']
            or chronos[p]['remaining_seconds'] > 0
        )
        result[p] = {
            f: bool(val.get(f, True)) and chrono_ouvert
            for f in _RECOURS_FILIERES
        }
    return result


def _set_recours_promo_actif(promo, filiere, actif):
    """Active ou désactive les soumissions pour une (promotion, filière)."""
    try:
        cfg = AppConfig.query.filter_by(key='recours_promo_actif').first()
        raw = json.loads(cfg.value) if (cfg and cfg.value) else {}
    except Exception:
        raw = {}
    # Compatibilité : purger les anciennes valeurs bool par promo entière
    for k in list(raw.keys()):
        if isinstance(raw[k], bool):
            del raw[k]
    raw.setdefault(promo, {})[filiere] = bool(actif)
    # Nettoyer : si True (valeur par défaut), on peut supprimer pour alléger le JSON
    if raw[promo][filiere]:
        raw[promo].pop(filiere, None)
    if not raw[promo]:
        raw.pop(promo, None)
    cfg2 = AppConfig.query.filter_by(key='recours_promo_actif').first()
    if not cfg2:
        cfg2 = AppConfig(key='recours_promo_actif', value='{}')
        db.session.add(cfg2)
    cfg2.value = json.dumps(raw)
    db.session.commit()


def _set_recours_chrono(promo, hours):
    """Démarre ou redémarre le chrono d'une promotion."""
    if promo not in _RECOURS_PROMOTIONS or hours not in _RECOURS_CHRONO_DUREES:
        raise ValueError('Promotion ou durée de chrono invalide.')

    started_at = now_cat()
    ends_at = started_at + timedelta(hours=hours)
    try:
        cfg = AppConfig.query.filter_by(key='recours_promo_chronos').first()
        raw = json.loads(cfg.value) if (cfg and cfg.value) else {}
    except Exception:
        raw = {}
        cfg = None

    raw[promo] = {
        'hours': hours,
        'started_at': started_at.isoformat(),
        'ends_at': ends_at.isoformat(),
    }
    if not cfg:
        cfg = AppConfig(key='recours_promo_chronos', value='{}')
        db.session.add(cfg)
    cfg.value = json.dumps(raw)
    db.session.commit()


def _remove_recours_chrono(promo):
    """Retire le chrono d'une promotion sans modifier ses ouvertures manuelles."""
    if promo not in _RECOURS_PROMOTIONS:
        raise ValueError('Promotion invalide.')
    cfg = AppConfig.query.filter_by(key='recours_promo_chronos').first()
    if not cfg:
        return
    try:
        raw = json.loads(cfg.value) if cfg.value else {}
    except Exception:
        raw = {}
    if promo in raw:
        raw.pop(promo, None)
        cfg.value = json.dumps(raw)
        db.session.commit()


def _generer_recours_pdf(recours_list, promotion=None, include_attachments=True,
                         output_path=None):
    """Génère un PDF ReportLab pour une liste de recours (optionnellement filtrée par promotion)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable, PageBreak,
                                    Image as RLImage)
    from reportlab.lib.utils import ImageReader
    from io import BytesIO
    from xml.sax.saxutils import escape as _xml_escape
    import os as _ospdf
    from PIL import Image as _PILImage

    def _pdf_text(value, fallback='—'):
        """Échappe les données saisies avant leur insertion dans Paragraph."""
        text = fallback if value is None or value == '' else str(value)
        return _xml_escape(text)

    output = output_path or BytesIO()
    doc = SimpleDocTemplate(output, pagesize=A4,
                            topMargin=1.5*cm, bottomMargin=1.5*cm,
                            leftMargin=2*cm, rightMargin=2*cm)
    styles = getSampleStyleSheet()
    marine = colors.HexColor('#0d2666')
    gold   = colors.HexColor('#bfa01a')

    title_style  = ParagraphStyle('RecoursTitle', parent=styles['Heading1'],
                                   textColor=marine, fontSize=13, spaceAfter=4)
    sub_style    = ParagraphStyle('RecoursSub', parent=styles['Normal'],
                                   fontSize=9, textColor=colors.HexColor('#444444'), spaceAfter=8)
    cell_style   = ParagraphStyle('RecoursCell', parent=styles['Normal'],
                                   fontSize=8, leading=11, textColor=colors.HexColor('#111111'))
    # Réclamations : texte sombre lisible
    claim_style  = ParagraphStyle('RecoursClaim', parent=styles['Normal'],
                                   fontSize=7, textColor=colors.HexColor('#1a1a1a'), leading=10)
    # Style en-tête blanc
    hdr_style    = ParagraphStyle('RecoursHdr', parent=styles['Normal'],
                                   fontSize=8, leading=11,
                                   textColor=colors.white, fontName='Helvetica-Bold')
    annex_style  = ParagraphStyle('RecoursAnnex', parent=styles['Normal'],
                                   fontSize=8, textColor=marine, spaceAfter=2)

    story = []

    # Grouper par promotion
    by_promo = {}
    for r in recours_list:
        by_promo.setdefault(r.promotion, []).append(r)

    promos_order = [p for p in _RECOURS_PROMOTIONS if p in by_promo]
    for p in sorted(by_promo.keys()):
        if p not in promos_order:
            promos_order.append(p)

    for idx, promo in enumerate(promos_order):
        items = by_promo[promo]
        if idx > 0:
            story.append(PageBreak())

        story.append(Paragraph(
            "FACULTÉ DES SCIENCES SOCIALES — UNIVERSITÉ DE LUBUMBASHI", sub_style))
        story.append(Paragraph(
            f"LISTE DES RECOURS — {promo.upper()}", title_style))
        story.append(Paragraph(
            f"Second Semestre 2025-2026 &nbsp;&nbsp;·&nbsp;&nbsp; {len(items)} recours soumis",
            sub_style))
        story.append(HRFlowable(width='100%', thickness=1, color=gold, spaceAfter=8))

        # ── En-tête tableau (texte blanc) ──
        header = [
            Paragraph('N°', hdr_style),
            Paragraph('NOM COMPLET', hdr_style),
            Paragraph('Filière', hdr_style),
            Paragraph('Réclamations', hdr_style),
            Paragraph('Reçu N°', hdr_style),
            Paragraph('Date', hdr_style),
        ]
        rows = [header]

        claims_labels = {cid: label for cid, label, _ in _RECOURS_CLAIMS}

        for i, r in enumerate(items, 1):
            try:
                claims = json.loads(r.reclamations_json)
            except Exception:
                claims = []

            claims_text = []
            for c in claims:
                cid    = c.get('id', '')
                label  = claims_labels.get(cid, cid)
                detail_value = c.get('detail', '')
                detail = str(detail_value).strip() if detail_value is not None else ''
                line   = f"• {_pdf_text(label)}"
                if detail:
                    line += f" : {_pdf_text(detail)}"
                claims_text.append(line)

            date_str = r.date_soumission.strftime('%d/%m/%Y %H:%M') if r.date_soumission else '—'
            nom_complet = _pdf_text(
                f"{r.nom} {r.postnom}" + (f" {r.prenom}" if r.prenom else '')
            )

            rows.append([
                Paragraph(str(i), cell_style),
                Paragraph(f"<b>{nom_complet}</b><br/><font size='7' color='grey'>"
                          f"Tél : {_pdf_text(r.telephone)}</font>", cell_style),
                Paragraph(_pdf_text(r.filiere), cell_style),
                Paragraph('<br/>'.join(claims_text) or '—', claim_style),
                Paragraph(_pdf_text(r.recu_numero), cell_style),
                Paragraph(_pdf_text(date_str), cell_style),
            ])

        col_widths = [0.8*cm, 3.8*cm, 3.5*cm, 7.0*cm, 2.4*cm, 2.5*cm]
        IMG_W, IMG_H = 4.5*cm, 3.5*cm
        GRID_COLOR   = colors.HexColor('#c8d0e0')
        cap_style    = ParagraphStyle('RCap', parent=styles['Normal'],
                                      fontSize=6, textColor=colors.grey, alignment=1)
        _pdf_image_buffers = []

        def _collect_imgs(r):
            """Retourne [(label, filepath)] pour toutes les images d'un recours."""
            if not include_attachments:
                return []
            imgs = []
            try:
                preuves = json.loads(r.preuves_json or '{}')
            except Exception:
                preuves = {}
            for cid, fname in preuves.items():
                lbl   = _RECOURS_CLAIM_FILES.get(cid, (None, cid))[1]
                fpath = _ensure_local_asset('recours_preuves', fname)
                if fpath:
                    imgs.append((lbl, fpath))
            if r.carte_etudiant_filename:
                fpath = _ensure_local_asset('recours_cartes',
                                            r.carte_etudiant_filename)
                if fpath:
                    imgs.append(("Carte d'étudiant", fpath))
            return imgs

        def _img_grid_table(all_imgs):
            """Construit un mini-tableau de miniatures pour un recours."""
            cells = []
            for lbl, fpath in all_imgs:
                try:
                    # Ne pas intégrer les originaux parfois très lourds :
                    # une miniature JPEG conserve la lisibilité et réduit
                    # fortement la mémoire et le temps de génération.
                    with _PILImage.open(fpath) as source:
                        image = source.convert('RGB')
                        image.thumbnail((1400, 1000), _PILImage.Resampling.LANCZOS)
                        image_buffer = BytesIO()
                        image.save(
                            image_buffer,
                            format='JPEG',
                            quality=78,
                            optimize=True,
                        )
                    image_buffer.seek(0)
                    _pdf_image_buffers.append(image_buffer)
                    img_flow = RLImage(
                        image_buffer, width=IMG_W, height=IMG_H, kind='proportional'
                    )
                    cap = Paragraph(f"<i>{_pdf_text(lbl)}</i>", cap_style)
                    cells.append((img_flow, cap))
                except Exception:
                    pass
            if not cells:
                return None
            grid_rows = []
            for start in range(0, len(cells), 3):
                chunk = cells[start:start+3]
                pad   = 3 - len(chunk)
                grid_rows.append([c[0] for c in chunk] + [''] * pad)
                grid_rows.append([c[1] for c in chunk] + [''] * pad)
            grid = Table(grid_rows, colWidths=[(IMG_W + 0.3*cm)] * 3)
            grid.setStyle(TableStyle([
                ('ALIGN',          (0,0), (-1,-1), 'CENTER'),
                ('VALIGN',         (0,0), (-1,-1), 'MIDDLE'),
                ('TOPPADDING',     (0,0), (-1,-1), 2),
                ('BOTTOMPADDING',  (0,0), (-1,-1), 2),
                ('LEFTPADDING',    (0,0), (-1,-1), 3),
                ('RIGHTPADDING',   (0,0), (-1,-1), 3),
            ]))
            return grid

        # ── En-tête du tableau (une seule fois par promotion) ──────────
        hdr_tbl = Table([rows[1:][0:1][0:1] if len(rows) > 1 else header],
                        colWidths=col_widths)
        # Correction : rows[0] == header (seule ligne ajoutée jusqu'ici)
        hdr_tbl = Table([rows[0]], colWidths=col_widths)
        hdr_tbl.setStyle(TableStyle([
            ('BACKGROUND',    (0,0), (-1,0), marine),
            ('GRID',          (0,0), (-1,0), 0.4, GRID_COLOR),
            ('VALIGN',        (0,0), (-1,0), 'MIDDLE'),
            ('LEFTPADDING',   (0,0), (-1,0), 4),
            ('RIGHTPADDING',  (0,0), (-1,0), 4),
            ('TOPPADDING',    (0,0), (-1,0), 4),
            ('BOTTOMPADDING', (0,0), (-1,0), 4),
        ]))
        story.append(hdr_tbl)

        # ── Une ligne de données + pièces jointes par candidat ────────
        for i, (data_row, r) in enumerate(zip(rows[1:], items), 1):
            bg = colors.white if i % 2 == 1 else colors.HexColor('#f0f4ff')
            row_tbl = Table([data_row], colWidths=col_widths)
            row_tbl.setStyle(TableStyle([
                ('BACKGROUND',    (0,0), (-1,0), bg),
                ('GRID',          (0,0), (-1,0), 0.4, GRID_COLOR),
                ('VALIGN',        (0,0), (-1,0), 'TOP'),
                ('LEFTPADDING',   (0,0), (-1,0), 4),
                ('RIGHTPADDING',  (0,0), (-1,0), 4),
                ('TOPPADDING',    (0,0), (-1,0), 4),
                ('BOTTOMPADDING', (0,0), (-1,0), 4),
            ]))
            story.append(row_tbl)

            # Pièces jointes de ce candidat (immédiatement sous sa ligne)
            all_imgs = _collect_imgs(r)
            if all_imgs:
                grid = _img_grid_table(all_imgs)
                if grid:
                    outer = Table([[grid]], colWidths=[sum(col_widths)])
                    outer.setStyle(TableStyle([
                        ('BACKGROUND',    (0,0), (-1,-1), colors.HexColor('#fafbff')),
                        ('BOX',           (0,0), (-1,-1), 0.4, GRID_COLOR),
                        ('LEFTPADDING',   (0,0), (-1,-1), 6),
                        ('RIGHTPADDING',  (0,0), (-1,-1), 6),
                        ('TOPPADDING',    (0,0), (-1,-1), 4),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
                    ]))
                    story.append(outer)

        story.append(Spacer(1, 0.4*cm))
        story.append(Paragraph(
            f"Total promotion {promo} : <b>{len(items)}</b> recours — "
            f"Généré le {now_cat().strftime('%d/%m/%Y à %H:%M')}",
            sub_style))

    doc.build(story)
    if output_path:
        return output_path
    output.seek(0)
    return output


@app.route('/recours')
def recours_form():
    """Formulaire de recours public."""
    recours_actif = _get_recours_actif()
    recours_chronos = _get_recours_chronos()
    recours_promo_ouvert = {
        promo: any(recours_actif[promo].values())
        for promo in _RECOURS_PROMOTIONS
    }
    return render_template('recours_form.html',
        promotions=_RECOURS_PROMOTIONS,
        filieres=_RECOURS_FILIERES,
        claims=_RECOURS_CLAIMS,
        claim_files=_RECOURS_CLAIM_FILES,
        recours_actif=recours_actif,
        recours_chronos=recours_chronos,
        recours_promo_ouvert=recours_promo_ouvert,
    )


@app.route('/recours/valider', methods=['POST'])
def recours_valider():
    """Valide le formulaire et stocke les données en session Flask avant le scan du reçu."""
    import os as _osv, uuid as _uuidv
    from werkzeug.utils import secure_filename as _sfnv

    nom      = request.form.get('nom',      '').strip().upper()
    postnom  = request.form.get('postnom',  '').strip().upper()
    prenom   = request.form.get('prenom',   '').strip().upper()
    sexe     = request.form.get('sexe',     '').strip()
    telephone= request.form.get('telephone','').strip()
    promotion= request.form.get('promotion','').strip()
    filiere  = request.form.get('filiere',  '').strip()

    errors = []
    if not nom:     errors.append("Le NOM est obligatoire.")
    if not postnom: errors.append("Le POSTNOM est obligatoire.")
    # PRÉNOM est optionnel
    promo_ok   = promotion and promotion in _RECOURS_PROMOTIONS
    filiere_ok = filiere   and filiere   in _RECOURS_FILIERES
    if not promo_ok:
        errors.append("Choisissez une promotion valide.")
    if not filiere_ok:
        errors.append("Choisissez une filière valide.")
    if promo_ok and filiere_ok:
        if not _get_recours_actif().get(promotion, {}).get(filiere, True):
            errors.append(f"Les recours sont temporairement fermés pour {promotion} — {filiere}. "
                          f"Contactez le DÉCANAT.")

    # Récupérer les réclamations cochées
    selected_claims = request.form.getlist('claims')
    if not selected_claims:
        errors.append("Cochez au moins une réclamation.")

    # Valider les champs texte obligatoires (cochés + champ de détail présent)
    claims_labels = {cid: label for cid, label, _ in _RECOURS_CLAIMS}
    for cid in selected_claims:
        detail_key = next((fk for c_id, _, fk in _RECOURS_CLAIMS if c_id == cid and fk), None)
        if detail_key:
            val = request.form.get(detail_key, '').strip()
            if not val:
                label = claims_labels.get(cid, cid)
                errors.append(f'Précisez votre réclamation : « {label} ».')

    if errors:
        for e in errors:
            flash(e, 'error')
        return redirect(url_for('recours_form'))

    # Construire la liste de réclamations avec détails
    claims_data = []
    for cid in selected_claims:
        if cid not in claims_labels:
            continue
        detail_key = next((fk for c_id, _, fk in _RECOURS_CLAIMS if c_id == cid and fk), None)
        detail = request.form.get(detail_key, '').strip() if detail_key else ''
        claims_data.append({'id': cid, 'texte': claims_labels[cid], 'detail': detail})

    # Sauvegarder les photos-preuves jointes (optionnelles)
    preuves = {}
    preuves_dir = _osv.path.join(app.static_folder, 'recours_preuves')
    _osv.makedirs(preuves_dir, exist_ok=True)
    for cid, (field_name, _label) in _RECOURS_CLAIM_FILES.items():
        if cid in selected_claims:
            f = request.files.get(field_name)
            if f and f.filename:
                ext = _osv.path.splitext(_sfnv(f.filename))[1].lower()
                if ext in ('.jpg', '.jpeg', '.png', '.webp', '.gif'):
                    fname = f"preuve_{cid}_{int(now_cat().timestamp())}_{_uuidv.uuid4().hex[:8]}{ext}"
                    _save_uploaded_asset(f, 'recours_preuves', fname)
                    preuves[cid] = fname
    if preuves:
        # The draft may spend time in the session before the QR confirmation.
        # Commit the file bytes now so a restart cannot lose the attachment.
        db.session.commit()

    # Stocker en session Flask (pas en DB — pas encore payé)
    session['recours_draft'] = {
        'nom': nom, 'postnom': postnom, 'prenom': prenom,
        'sexe': sexe, 'telephone': telephone,
        'promotion': promotion, 'filiere': filiere,
        'reclamations': claims_data,
        'preuves': preuves,
    }
    return redirect(url_for('recours_scanner'))


@app.route('/recours/scanner')
def recours_scanner():
    """Page de scan du reçu de paiement du recours."""
    draft = session.get('recours_draft')
    if not draft:
        flash("Remplissez d'abord le formulaire de recours.", 'error')
        return redirect(url_for('recours_form'))
    return render_template('recours_scanner.html', draft=draft, mode='scan')


@app.route('/recours/scan/<code>', methods=['GET', 'POST'])
@app.route('/recours/scan/<path:code>', methods=['GET', 'POST'])
@app.route('/recours/scan/scan/<code>', methods=['GET', 'POST'])
def recours_scan(code):
    """Valide le reçu QR et soumet le recours."""
    # Compatibilité avec les QR déjà imprimés et avec l'ancien scanner :
    # certains transmettaient "/scan/CODE" comme valeur du paramètre.
    from urllib.parse import unquote
    code = unquote(str(code or '')).strip()
    path_match = re.search(r'(?:^|/)(?:recours/)?scan/([^/?#\s]+)', code, re.IGNORECASE)
    if path_match:
        code = path_match.group(1)

    draft = session.get('recours_draft')
    if not draft:
        flash("Session expirée — remplissez à nouveau le formulaire.", 'error')
        return redirect(url_for('recours_form'))

    draft_actif = _get_recours_actif().get(draft.get('promotion', ''), {})
    if not draft_actif.get(draft.get('filiere', ''), False):
        return render_template(
            'recours_scanner.html',
            draft=draft,
            mode='error',
            error='Le délai de recours de votre promotion est terminé. '
                  'Les soumissions sont maintenant fermées.',
        )

    recu = RecuPaiement.query.filter_by(code_qr=code).first()

    if not recu:
        return render_template('recours_scanner.html', draft=draft,
                               mode='error', error="Ce reçu n'est pas reconnu dans notre système.")

    # Les reçus RR/R2 sont réservés à la consultation d'un résultat.
    # Seuls RS/RS2 peuvent confirmer une nouvelle soumission.
    submission_types = ('recours', 'recours_session_2_soumission')
    if (not _receipt_type_is_consistent(recu)
            or getattr(recu, 'type_recu', 'bulletin') not in submission_types):
        return render_template(
            'recours_scanner.html',
            draft=draft,
            mode='error',
            error="Ce reçu n'est pas un reçu de soumission de recours. "
                  "Utilisez un reçu RS- ou RS2- remis pour les frais de recours.",
        )

    if recu.utilise:
        return render_template('recours_scanner.html', draft=draft,
                               mode='error', error=f"Ce reçu ({recu.numero}) a déjà été utilisé.")

    if request.method == 'GET':
        # Montrer la confirmation avant soumission
        return render_template('recours_scanner.html', draft=draft,
                               mode='confirm', recu=recu, code=code)

    # POST — enregistrer le recours
    ip_client = (request.headers.get('X-Forwarded-For') or request.remote_addr or '').split(',')[0].strip()
    try:
        import os as _os, uuid as _uuid2
        from werkzeug.utils import secure_filename as _sfn

        # ── Gestion optionnelle de la photo de carte d'étudiant ──
        carte_filename = None
        carte_file = request.files.get('carte_etudiant')
        if carte_file and carte_file.filename:
            _ext = _os.path.splitext(_sfn(carte_file.filename))[1].lower()
            if _ext in ('.jpg', '.jpeg', '.png', '.webp', '.gif'):
                carte_dir = _os.path.join(app.static_folder, 'recours_cartes')
                _os.makedirs(carte_dir, exist_ok=True)
                carte_filename = f"carte_{int(now_cat().timestamp())}_{_uuid2.uuid4().hex[:8]}{_ext}"
                _save_uploaded_asset(carte_file, 'recours_cartes', carte_filename)
                db.session.commit()

        prenom_val = draft.get('prenom') or ''
        r = Recours(
            nom        = draft['nom'],
            postnom    = draft['postnom'],
            prenom     = prenom_val,
            sexe       = draft.get('sexe'),
            telephone  = draft.get('telephone'),
            promotion  = draft['promotion'],
            filiere    = draft['filiere'],
            reclamations_json = json.dumps(draft['reclamations'], ensure_ascii=False),
            preuves_json      = json.dumps(draft.get('preuves', {}), ensure_ascii=False),
            recu_id    = recu.id,
            recu_numero= recu.numero,
            type_recu  = recu.type_recu,
            ip_soumission = ip_client,
            carte_etudiant_filename = carte_filename,
        )
        db.session.add(r)
        # Marquer le reçu comme utilisé (recours)
        recu.utilise           = True
        recu.date_utilisation  = now_cat()
        recu.matricule_etudiant= draft['nom'] + ' ' + draft['postnom']
        recu.nom_etudiant      = (draft['nom'] + ' ' + draft['postnom']
                                  + (' ' + prenom_val if prenom_val else ''))
        db.session.commit()
        session.pop('recours_draft', None)
        flash('Votre recours a été soumis avec succès. Le DÉCANAT en prendra connaissance.', 'success')
        return redirect(url_for('recours_success'))
    except Exception as e:
        db.session.rollback()
        return render_template('recours_scanner.html', draft=draft,
                               mode='error', error=f"Erreur lors de la soumission : {e}")


@app.route('/recours/succes')
def recours_success():
    return render_template('recours_scanner.html', draft=None, mode='success')


# ── DÉCANAT : gestion des recours ──────────────────────────────────

@app.route('/decanat/recours')
def decanat_recours():
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    filtre_promo    = request.args.get('promo',    '').strip()
    filtre_filiere  = request.args.get('filiere',  '').strip()

    q = Recours.query
    if filtre_promo:
        q = q.filter(Recours.promotion == filtre_promo)
    if filtre_filiere:
        q = q.filter(Recours.filiere == filtre_filiere)
    all_recours = q.order_by(Recours.promotion, Recours.date_soumission).all()

    # Grouper par promotion
    by_promo = {}
    for r in all_recours:
        by_promo.setdefault(r.promotion, []).append(r)

    # Reçus utilisés sans bulletin associé (paiements non comptabilisés)
    orphan_recus = RecuPaiement.query.filter(
        RecuPaiement.utilise == True,
        RecuPaiement.bulletin_id == None,
        ~RecuPaiement.numero.like('RS-%'),   # exclure les reçus de recours
    ).order_by(RecuPaiement.date_utilisation.desc()).all()

    promos_dispo  = sorted({r.promotion for r in Recours.query.all()})
    claims_labels = {cid: label for cid, label, _ in _RECOURS_CLAIMS}
    recours_actif = _get_recours_actif()
    recours_chronos = _get_recours_chronos()
    recours_promo_ouvert = {
        promo: any(recours_actif[promo].values())
        for promo in _RECOURS_PROMOTIONS
    }
    total_recours = len(all_recours)
    total_montant_usd = total_recours * _RECOURS_MONTANT_USD

    par_departement = {}
    for recours in all_recours:
        departement = _RECOURS_DEPARTEMENT_PAR_FILIERE.get(
            recours.filiere, recours.filiere
        )
        ligne = par_departement.setdefault(
            departement, {'recours': 0, 'montant_usd': 0, 'filieres': set()}
        )
        ligne['recours'] += 1
        ligne['montant_usd'] += _RECOURS_MONTANT_USD
        ligne['filieres'].add(recours.filiere)
    par_departement = [
        {
            **ligne,
            'departement': departement,
            'filieres': ', '.join(sorted(ligne['filieres'])),
        }
        for departement, ligne in sorted(par_departement.items())
    ]

    par_promotion = []
    for promotion in _RECOURS_PROMOTIONS:
        nombre = sum(1 for recours in all_recours if recours.promotion == promotion)
        if nombre:
            par_promotion.append({
                'promotion': promotion,
                'recours': nombre,
                'montant_usd': nombre * _RECOURS_MONTANT_USD,
            })

    return render_template('decanat_recours.html',
        by_promo=by_promo,
        filtre_promo=filtre_promo,
        filtre_filiere=filtre_filiere,
        promos_dispo=promos_dispo,
        claims_labels=claims_labels,
        orphan_recus=orphan_recus,
        promotions=_RECOURS_PROMOTIONS,
        filieres=_RECOURS_FILIERES,
        recours_actif=recours_actif,
        recours_chronos=recours_chronos,
        chrono_durees=_RECOURS_CHRONO_DUREES,
        recours_promo_ouvert=recours_promo_ouvert,
        total_recours=total_recours,
        total_montant_usd=total_montant_usd,
        par_departement=par_departement,
        par_promotion=par_promotion,
        recours_montant_usd=_RECOURS_MONTANT_USD,
    )


@app.route('/decanat/recours/pdf')
def decanat_recours_pdf():
    """PDF des recours, groupés par promotion, filtrable par filière."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    filtre_promo   = request.args.get('promo',   '').strip()
    filtre_filiere = request.args.get('filiere', '').strip()
    q = Recours.query
    if filtre_promo:
        q = q.filter(Recours.promotion == filtre_promo)
    if filtre_filiere:
        q = q.filter(Recours.filiere == filtre_filiere)
    all_recours = q.order_by(Recours.promotion, Recours.date_soumission).all()
    if not all_recours:
        flash('Aucun recours à exporter (filtre trop restrictif ?).', 'error')
        return redirect(url_for('decanat_recours'))
    temp_path = None
    try:
        # Écrire directement sur disque : les gros exports (30 Mo ou plus)
        # ne doivent pas être entièrement conservés en mémoire du worker.
        with tempfile.NamedTemporaryFile(
            prefix='recours_export_', suffix='.pdf', delete=False
        ) as temp_file:
            temp_path = temp_file.name
        _generer_recours_pdf(all_recours, output_path=temp_path)
    except Exception:
        # Un recours peut contenir du texte ou une pièce jointe mal formée.
        # Fournir au minimum la liste PDF sans images plutôt qu'une erreur 500.
        app.logger.exception("Échec de génération du PDF des recours avec pièces jointes")
        try:
            if temp_path and os.path.exists(temp_path):
                os.unlink(temp_path)
            with tempfile.NamedTemporaryFile(
                prefix='recours_export_', suffix='.pdf', delete=False
            ) as temp_file:
                temp_path = temp_file.name
            _generer_recours_pdf(
                all_recours,
                include_attachments=False,
                output_path=temp_path,
            )
        except Exception:
            app.logger.exception("Échec de génération du PDF des recours sans pièces jointes")
            if temp_path and os.path.exists(temp_path):
                os.unlink(temp_path)
            flash("Le PDF des recours n'a pas pu être généré. Réessayez ou contactez l'administrateur.", 'error')
            return redirect(url_for('decanat_recours', promo=filtre_promo, filiere=filtre_filiere))
    parts = []
    if filtre_promo:   parts.append(filtre_promo.replace(' ', '_'))
    if filtre_filiere: parts.append(filtre_filiere.replace(' ', '_'))
    fname = f"Recours_{'_'.join(parts) or 'Tous'}.pdf"
    response = send_file(
        temp_path,
        as_attachment=True,
        download_name=fname,
        mimetype='application/pdf',
        max_age=0,
    )

    def _remove_export_file(path=temp_path):
        try:
            if path and os.path.exists(path):
                os.unlink(path)
        except OSError:
            app.logger.warning("Impossible de supprimer le PDF temporaire %s", path)

    response.call_on_close(_remove_export_file)
    return response


@app.route('/decanat/recours/lier-bulletin', methods=['POST'])
def decanat_recours_lier_bulletin():
    """Lie manuellement un reçu orphelin à un bulletin (paiement non comptabilisé)."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    recu_id  = request.form.get('recu_id', type=int)
    matricule = request.form.get('matricule', '').strip()
    if not recu_id or not matricule:
        flash('Reçu ou matricule manquant.', 'error')
        return redirect(url_for('decanat_recours'))
    recu = RecuPaiement.query.get_or_404(recu_id)
    bd = BulletinData.query.filter(
        db.func.lower(BulletinData.matricule) == matricule.lower()
    ).first()
    if not bd:
        flash(f'Aucun bulletin trouvé pour le matricule « {matricule} ».', 'error')
        return redirect(url_for('decanat_recours'))
    try:
        recu.bulletin_id = bd.id
        recu.matricule_etudiant = bd.matricule
        recu.nom_etudiant = bd.nom
        if not bd.paye:
            bd.paye = True
            bd.date_paiement = recu.date_utilisation or now_cat()
            bd.methode_paiement = 'Reçu papier'
            bd.reference_paiement = recu.numero
            bd.montant_paye = recu.montant or (bd.bul_session.montant_fc if bd.bul_session else 5000)
        db.session.commit()
        flash(f'Reçu {recu.numero} lié au bulletin de {bd.nom} et paiement enregistré.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erreur : {e}', 'error')
    return redirect(url_for('decanat_recours'))


@app.route('/decanat/recours/toggle-promo', methods=['POST'])
def decanat_recours_toggle_promo():
    """Active ou désactive les soumissions pour une (promotion, filière)."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    promo   = request.form.get('promo',   '').strip()
    filiere = request.form.get('filiere', '').strip()
    if promo not in _RECOURS_PROMOTIONS or filiere not in _RECOURS_FILIERES:
        flash('Promotion ou filière invalide.', 'error')
        return redirect(url_for('decanat_recours'))
    current   = _get_recours_actif()
    new_state = not current.get(promo, {}).get(filiere, True)
    _set_recours_promo_actif(promo, filiere, new_state)
    etat = "✅ ouverts" if new_state else "🔒 fermés"
    flash(f'Recours {etat} pour {promo} — {filiere}.', 'success')
    return redirect(url_for('decanat_recours'))


@app.route('/decanat/recours/chrono', methods=['POST'])
def decanat_recours_chrono():
    """Démarre ou arrête le chrono de fermeture d'une promotion."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))

    promo = request.form.get('promo', '').strip()
    action = request.form.get('action', 'start').strip()
    try:
        if action == 'stop':
            _remove_recours_chrono(promo)
            flash(
                f'Chrono retiré pour {promo}. Les ouvertures manuelles restent inchangées.',
                'success',
            )
        else:
            hours = request.form.get('hours', type=int)
            _set_recours_chrono(promo, hours)
            flash(
                f'Chrono de {hours} heure(s) démarré pour {promo}. '
                'À expiration, toutes ses filières seront fermées.',
                'success',
            )
    except ValueError as exc:
        flash(str(exc), 'error')
    except Exception as exc:
        db.session.rollback()
        flash(f'Erreur lors de la configuration du chrono : {exc}', 'error')
    return redirect(url_for('decanat_recours'))


@app.route('/decanat/recours/<int:rid>/supprimer', methods=['POST'])
def decanat_recours_supprimer(rid):
    """Supprime un recours et libère le reçu associé."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    r = Recours.query.get_or_404(rid)
    nom_display = f"{r.nom} {r.postnom}"
    try:
        # Libérer le reçu associé (le remettre disponible)
        if r.recu_id:
            recu = RecuPaiement.query.get(r.recu_id)
            if recu:
                recu.utilise           = False
                recu.date_utilisation  = None
                recu.matricule_etudiant= None
                recu.nom_etudiant      = None
        db.session.delete(r)
        db.session.commit()
        flash(f'Recours de {nom_display} supprimé. Le reçu est à nouveau disponible.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Erreur lors de la suppression : {e}', 'error')
    return redirect(url_for('decanat_recours'))


@app.route('/decanat/recours/supprimer-tous', methods=['POST'])
def decanat_recours_supprimer_tous():
    """Supprime toutes les soumissions de recours après confirmation serveur."""
    if not session.get('decanat_logged_in'):
        return redirect(url_for('decanat_login'))
    if request.form.get('confirmation') != 'SUPPRIMER':
        flash('Suppression annulée : confirmation invalide.', 'error')
        return redirect(url_for('decanat_recours'))

    recours = Recours.query.all()
    count = len(recours)
    released_receipts = 0
    try:
        for item in recours:
            # Les pièces jointes appartiennent à la demande supprimée :
            # retirer aussi leur copie persistante pour éviter les fichiers orphelins.
            try:
                preuves = json.loads(item.preuves_json or '{}')
            except (TypeError, ValueError):
                preuves = {}
            for filename in list(preuves.values()) + [item.carte_etudiant_filename]:
                if filename:
                    _delete_file_asset('recours_preuves', filename)
                    _delete_file_asset('recours_cartes', filename)
            if item.recu_id:
                recu = db.session.get(RecuPaiement, item.recu_id)
                if recu:
                    recu.utilise = False
                    recu.date_utilisation = None
                    recu.matricule_etudiant = None
                    recu.nom_etudiant = None
                    recu.bulletin_id = None
                    released_receipts += 1
            db.session.delete(item)
        _enregistrer_admin_audit(
            'suppression_tous_recours',
            {'nombre_recours': count, 'recus_liberes': released_receipts},
        )
        db.session.commit()
        flash(
            f'{count} recours supprimé(s). {released_receipts} reçu(s) de recours libéré(s).',
            'success',
        )
    except Exception as exc:
        db.session.rollback()
        app.logger.exception('Suppression globale des recours échouée')
        flash(f'Erreur : {exc}', 'error')
    return redirect(url_for('decanat_recours'))


# Initialisation au démarrage — s'exécute sous gunicorn ET python3 app.py
def _auto_seed_if_empty():
    """Restaure depuis instance/seed_backup.json si toutes les tables critiques sont vides.
    S'exécute une seule fois au premier démarrage sur un Reserved VM vierge."""
    import os as _ose, json as _jse, logging as _log
    _logger = _log.getLogger(__name__)

    # PostgreSQL géré est la source de vérité permanente. Le seed JSON est
    # uniquement un filet de sécurité pour une ancienne base SQLite vierge ;
    # ne jamais l'exécuter sur PostgreSQL, où deux workers pourraient tenter
    # de restaurer simultanément des identifiants déjà présents.
    if os.environ.get('DATABASE_URL', '').strip():
        return

    seed_path = os.path.join(app.instance_path, 'seed_backup.json')
    if not _ose.path.exists(seed_path):
        return

    # Vérifier si la base est vraiment vide
    try:
        if BulletinData.query.count() > 0 or RecuPaiement.query.count() > 0:
            return  # Données déjà présentes — ne rien toucher
    except Exception:
        return

    _logger.info("[seed] Base vide détectée — restauration depuis seed_backup.json …")

    try:
        with open(seed_path, encoding='utf-8') as _f:
            bk = _jse.load(_f)
        td = bk.get('tables', {})
    except Exception as _e:
        _logger.error("[seed] Impossible de lire seed_backup.json : %s", _e)
        return

    def _dt(s):
        if not s: return None
        try: return datetime.fromisoformat(s)
        except: return None

    def _d(s):
        if not s: return None
        try: return date.fromisoformat(s)
        except: return None

    def _ins(obj):
        db.session.add(obj); db.session.flush()

    try:
        total = 0

        # 1. app_config
        for r in td.get('app_config', []):
            if not AppConfig.query.filter_by(key=r['key']).first():
                _ins(AppConfig(id=r.get('id'), key=r['key'], value=r.get('value', '')))
                total += 1

        # 2. page_contents
        for r in td.get('page_contents', []):
            if not PageContent.query.filter_by(page_name=r['page_name']).first():
                _ins(PageContent(id=r.get('id'), page_name=r['page_name'],
                                 content_json=r.get('content_json', '{}'),
                                 image_principale=r.get('image_principale'),
                                 date_modification=_dt(r.get('date_modification'))))
                total += 1

        # 3. liste_identifiants
        for r in td.get('liste_identifiants', []):
            _ins(ListeIdentifiants(
                id=r.get('id'), promotion=r.get('promotion',''), nom=r.get('nom',''),
                nom_norm=r.get('nom_norm',''), matricule=r.get('matricule',''),
                mot_de_passe=r.get('mot_de_passe'), date_import=_dt(r.get('date_import'))))
            total += 1

        # 4. etudiants
        for r in td.get('etudiants', []):
            if not Etudiant.query.filter_by(matricule=r['matricule']).first():
                _ins(Etudiant(id=r.get('id'), matricule=r['matricule'], nom=r['nom'],
                              postnom=r['postnom'], prenom=r.get('prenom',''),
                              sexe=r.get('sexe',''), telephone=r.get('telephone',''),
                              promotion=r.get('promotion',''), departement=r.get('departement',''),
                              photo=r.get('photo'), qrcode_path=r.get('qrcode_path'),
                              date_inscription=_dt(r.get('date_inscription')) or now_cat()))
                total += 1

        # 5. bulletin_sessions
        for r in td.get('bulletin_sessions', []):
            _ins(BulletinSession(
                id=r.get('id'), nom=r.get('nom'), annee=r.get('annee'),
                session_acad=r.get('session_acad'), semestre=r.get('semestre'),
                promotion=r.get('promotion'), montant_fc=r.get('montant_fc', 5000),
                departement=r.get('departement',''), texte_intro=r.get('texte_intro',''),
                type_grille=r.get('type_grille', 'initial'),
                date_import=_dt(r.get('date_import'))))
            total += 1

        # 6. bulletin_data
        for r in td.get('bulletin_data', []):
            _ins(BulletinData(
                id=r.get('id'), session_id=r.get('session_id'), matricule=r.get('matricule'),
                nom=r.get('nom'), sexe=r.get('sexe'), data_json=r.get('data_json'),
                numero_bulletin=r.get('numero_bulletin'), paye=r.get('paye', False),
                date_paiement=_dt(r.get('date_paiement')),
                methode_paiement=r.get('methode_paiement'),
                telephone_paiement=r.get('telephone_paiement'),
                reference_paiement=r.get('reference_paiement'),
                montant_paye=r.get('montant_paye', 0),
                nb_telechargements=r.get('nb_telechargements', 0),
                date_dernier_telechargement=_dt(r.get('date_dernier_telechargement'))))
            total += 1

        # 7. recus_paiement
        for r in td.get('recus_paiement', []):
            _ins(RecuPaiement(
                id=r.get('id'), numero=r['numero'], code_qr=r['code_qr'],
                lot=r.get('lot'), dept=r.get('dept'), annee=r.get('annee'),
                semestre=r.get('semestre'), montant=r.get('montant','5000 CDF'),
                montant_lettres=r.get('montant_lettres',''),
                motif=r.get('motif', ''),
                annee_complete=r.get('annee_complete','2025-2026'),
                type_recu=r.get('type_recu', 'bulletin'),
                date_creation=_dt(r.get('date_creation')),
                utilise=r.get('utilise', False),
                date_utilisation=_dt(r.get('date_utilisation')),
                matricule_etudiant=r.get('matricule_etudiant'),
                nom_etudiant=r.get('nom_etudiant'),
                bulletin_id=r.get('bulletin_id'),
                tentative_bulletin_id=r.get('tentative_bulletin_id'),
                tentative_matricule=r.get('tentative_matricule'),
                tentative_nom=r.get('tentative_nom'),
                date_tentative=_dt(r.get('date_tentative')),
                tentative_revue=r.get('tentative_revue', False)))
            total += 1

        # 8. paiement_audits
        for r in td.get('paiement_audits', []):
            _ins(PaiementAudit(
                id=r.get('id'), bulletin_id=r.get('bulletin_id'),
                old_montant=r.get('old_montant'), new_montant=r.get('new_montant'),
                old_methode=r.get('old_methode'), new_methode=r.get('new_methode'),
                old_reference=r.get('old_reference'), new_reference=r.get('new_reference'),
                old_date_paiement=_dt(r.get('old_date_paiement')),
                new_date_paiement=_dt(r.get('new_date_paiement')),
                date_modification=_dt(r.get('date_modification')) or now_cat(),
                operator_name=r.get('operator_name')))
            total += 1

        # 8b. administration_audits (optional for older backups)
        for r in td.get('administration_audits', []):
            _ins(AdministrationAudit(
                id=r.get('id'),
                action=r.get('action', ''),
                operator_name=r.get('operator_name'),
                details=r.get('details'),
                date_action=_dt(r.get('date_action')) or now_cat()))
            total += 1

        # 8c. releve_commandes (optional for older backups)
        for r in td.get('releve_commandes', []):
            _ins(ReleveCommande(
                id=r.get('id'),
                bulletin_id=r.get('bulletin_id'),
                recu_id=r.get('recu_id'),
                matricule=r.get('matricule', ''),
                nom_etudiant=r.get('nom_etudiant', ''),
                statut=r.get('statut', 'soumise'),
                date_commande=_dt(r.get('date_commande')) or now_cat(),
                date_traitement=_dt(r.get('date_traitement')),
                note_decanat=r.get('note_decanat'),
            ))
            total += 1

        # 9. recours
        for r in td.get('recours', []):
            _ins(Recours(
                id=r.get('id'), nom=r['nom'], postnom=r['postnom'],
                prenom=r.get('prenom',''), sexe=r.get('sexe'), telephone=r.get('telephone'),
                promotion=r['promotion'], filiere=r.get('filiere',''),
                reclamations_json=r.get('reclamations_json','[]'),
                recu_id=r.get('recu_id'), recu_numero=r.get('recu_numero'),
                date_soumission=_dt(r.get('date_soumission')) or now_cat(),
                statut=r.get('statut','soumis'), ip_soumission=r.get('ip_soumission'),
                carte_etudiant_filename=r.get('carte_etudiant_filename'),
                preuves_json=r.get('preuves_json','{}'),
                type_recu=r.get('type_recu', 'recours')))
            total += 1

        # 10. persistent uploaded files (optional for older backups)
        for r in backup.get('files', []):
            key = r.get('storage_key')
            encoded = r.get('data_b64')
            if not key or not encoded:
                continue
            if FileAsset.query.filter_by(storage_key=key).first():
                continue
            raw = base64.b64decode(encoded)
            _ins(FileAsset(
                storage_key=key,
                mime_type=r.get('mime_type') or 'application/octet-stream',
                data=raw,
                size_bytes=len(raw),
            ))
            total += 1

        db.session.commit()
        _logger.info("[seed] ✅ %d enregistrements restaurés depuis seed_backup.json.", total)

    except Exception as _e:
        db.session.rollback()
        _logger.error("[seed] ERREUR restauration : %s", _e)


with app.app_context():
    try:
        db.create_all()
        _run_migrations()
        _auto_seed_if_empty()
    except Exception as _db_init_err:
        # Ne pas crasher le worker gunicorn si la DB est temporairement inaccessible.
        # Les routes feront échouer les requêtes individuellement plutôt que de
        # tuer le processus entier au démarrage.
        import logging
        logging.getLogger(__name__).error(
            "DB init skipped at startup: %s", _db_init_err
        )
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['QRCODE_FOLDER'], exist_ok=True)
    os.makedirs(app.config['HORAIRES_FOLDER'], exist_ok=True)
    os.makedirs(app.config['ACTUALITES_FOLDER'], exist_ok=True)
    os.makedirs(app.config['COMMUNICATION_AUDIO_FOLDER'], exist_ok=True)
    try:
        _migrate_local_assets_to_db()
    except Exception as _asset_init_err:
        import logging
        logging.getLogger(__name__).error(
            "Persistent file migration skipped at startup: %s", _asset_init_err
        )

from sqlalchemy.exc import OperationalError as _SAOperationalError

@app.errorhandler(_SAOperationalError)
def handle_db_unavailable(e):
    """Affiche une page claire quand la base de données est inaccessible."""
    import logging
    logging.getLogger(__name__).error("DB unavailable: %s", e)
    return render_template('db_unavailable.html'), 503


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
